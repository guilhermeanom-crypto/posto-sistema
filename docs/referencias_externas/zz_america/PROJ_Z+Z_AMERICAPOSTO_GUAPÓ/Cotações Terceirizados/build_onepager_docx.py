#!/usr/bin/env python3
"""
Gerador DOCX do one-pager de cotacao COT-2026-001 (MCO + CONAMA 273).
Identidade visual Habilis: verde #1B5E20, Segoe UI / Consolas.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

VERDE = RGBColor(0x1B, 0x5E, 0x20)
VERDE_L = RGBColor(0x4C, 0xAF, 0x50)
VERDE_BG = "E8F5E9"
VERMELHO = RGBColor(0xB7, 0x1C, 0x1C)
VERMELHO_BG = "FFEBEE"
CINZA_9 = RGBColor(0x1C, 0x1C, 0x1E)
CINZA_7 = RGBColor(0x3A, 0x3A, 0x3C)
CINZA_5 = RGBColor(0x6C, 0x6C, 0x70)
CINZA_BG = "FAFAFA"
BRANCO = RGBColor(0xFF, 0xFF, 0xFF)
FONT_SANS = "Segoe UI"
FONT_MONO = "Consolas"


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_paragraph_border_left(p, size_pt, color_hex):
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), str(int(size_pt * 8)))
    left.set(qn('w:space'), '6'); left.set(qn('w:color'), color_hex)
    pbdr.append(left); p_pr.append(pbdr)


def styled_run(p, text, *, bold=False, italic=False, size=None, color=None, mono=False):
    run = p.add_run(text)
    font_name = FONT_MONO if mono else FONT_SANS
    run.font.name = font_name
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.find(qn('w:rFonts'))
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts'); r_pr.append(r_fonts)
    r_fonts.set(qn('w:ascii'), font_name); r_fonts.set(qn('w:hAnsi'), font_name); r_fonts.set(qn('w:cs'), font_name)
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    return run


def section_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '2'); bottom.set(qn('w:color'), '1B5E20')
    pbdr.append(bottom); p_pr.append(pbdr)
    styled_run(p, text.upper(), size=9, color=VERDE, bold=True, mono=True)


def body_p(doc, text, size=9.5, justify=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    styled_run(p, text, size=size, color=CINZA_7)
    return p


def kv_table(doc, rows, widths=(4.5, 12.0)):
    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for label, value in rows:
        r = t.add_row()
        c1, c2 = r.cells
        c1.width = Cm(widths[0]); c2.width = Cm(widths[1])
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(1); p1.paragraph_format.space_after = Pt(1)
        styled_run(p1, label, size=8.5, color=VERDE, bold=True, mono=True)
        p2 = c2.paragraphs[0]
        p2.paragraph_format.space_before = Pt(1); p2.paragraph_format.space_after = Pt(1)
        styled_run(p2, value, size=9.5, color=CINZA_9)
    return t


def timeline_table(doc, rows):
    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (marco, data, highlight) in enumerate(rows):
        r = t.add_row()
        c1, c2 = r.cells
        c1.width = Cm(11.0); c2.width = Cm(5.5)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(1); p1.paragraph_format.space_after = Pt(1)
        styled_run(p1, marco, size=9.5, color=CINZA_9, bold=highlight)
        p2 = c2.paragraphs[0]
        p2.paragraph_format.space_before = Pt(1); p2.paragraph_format.space_after = Pt(1)
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        styled_run(p2, data, size=9.5, color=VERDE if highlight else CINZA_9,
                   bold=highlight, mono=True)
        if highlight:
            set_cell_shading(c1, VERDE_BG); set_cell_shading(c2, VERDE_BG)
    return t


def callout_prazo(doc, text):
    t = doc.add_table(rows=1, cols=1)
    c = t.cell(0, 0)
    set_cell_shading(c, VERMELHO_BG)
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    styled_run(p, text, size=9.5, color=VERMELHO, bold=True)


def bullet_num(doc, num, text, size=9):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    styled_run(p, f"{num}. ", size=size, color=VERDE, bold=True, mono=True)
    styled_run(p, text, size=size, color=CINZA_7)


def build():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(1.5); s.bottom_margin = Cm(1.5)
    s.left_margin = Cm(1.8); s.right_margin = Cm(1.8)

    # Header
    h_p = s.header.paragraphs[0]
    styled_run(h_p, "Cotacao COT-2026-001 | Auto Posto America | Renovacao LAO 033/2020",
               size=8, color=CINZA_5, mono=True)

    # Footer
    f_p = s.footer.paragraphs[0]
    styled_run(f_p, "Habilis Licenciamento & Gestao Ambiental | guilherme.anom@gmail.com",
               size=8, color=CINZA_5, mono=True)

    # --- CABECALHO ---
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    styled_run(p, "COTACAO COT-2026-001", size=10, color=VERDE_L, bold=True, mono=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    styled_run(p, "Servico contratado: ", size=14, color=CINZA_9, bold=True)
    styled_run(p, "MCO + Anexos I e II CONAMA 273/2000", size=14, color=VERDE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    styled_run(p, "Habilis Licenciamento & Gestao Ambiental", size=10, color=CINZA_9, bold=True)
    styled_run(p, "  |  Coordenacao tecnica: Guilherme (RT Biologo, CRBio)", size=10, color=CINZA_7)
    styled_run(p, "  |  guilherme.anom@gmail.com", size=10, color=CINZA_7, mono=True)

    # Prazo resposta
    callout_prazo(doc, "Prazo para resposta da cotacao: 02/05/2026, ate as 18h. Cotacoes fora do prazo serao desconsideradas.")

    # --- OBJETO ---
    section_title(doc, "Objeto")
    body_p(doc,
        "Contratacao de profissional com registro no CREA (Civil / Ambiental / Mecanico / Quimico com "
        "atribuicao compativel) para emissao e assinatura, com ART propria, das pecas: "
        "(A) MCO - Memorial de Caracterizacao do Empreendimento (Anexo IX CEMAM-GO 029/2018); "
        "(B) Anexo I CONAMA 273/2000 - tabela descritiva; "
        "(C) Anexo II CONAMA 273/2000 - classificacao em classe de risco ambiental."
    )
    body_p(doc,
        "Finalidade: instrucao da renovacao da LAO 033/2020 do Auto Posto America junto a SEMMA Guapo. "
        "Protocolo-meta 26/06/2026 (vencimento 23/12/2026). Coordenacao tecnica integral pela Habilis "
        "(RT CRBio), com responsabilidade solidaria do profissional contratado pelas pecas emitidas."
    )

    # --- EMPREENDIMENTO ---
    section_title(doc, "Empreendimento")
    kv_table(doc, [
        ("Razao social", "AMERICA AUTO POSTO LTDA"),
        ("CNPJ", "34.144.804/0001-86"),
        ("Endereco", "Rodovia BR-060, km 203, S/N - Zona Rural - Guapo/GO"),
        ("Coordenadas (SIRGAS 2000)", "16°54'07,30\"S · 49°38'31,06\"O"),
        ("Atividade", "Posto revendedor varejista de combustiveis liquidos"),
        ("Area terreno / construida", "20.000 m² / 1.657 m²"),
        ("Ficha tecnica (tanques, SAO, tubulacoes)",
         "Enviada ao contratado ate 14/05/2026 (pos-inspecao Habilis 12/05)"),
    ])

    # --- PRAZOS ---
    section_title(doc, "Prazos (nao negociaveis)")
    timeline_table(doc, [
        ("Aceite e contratacao", "15/05/2026", False),
        ("Draft MCO para revisao", "25/05/2026", False),
        ("Entrega final MCO (PDF assinado + ART paga)", "28/05/2026", True),
        ("Entrega final Anexos I e II CONAMA 273", "29/05/2026", True),
        ("Protocolo SEMMA Guapo", "26/06/2026", False),
    ])

    # --- PAGAMENTO ---
    section_title(doc, "Pagamento")
    bullet_num(doc, "50%", "Na assinatura do aceite (15/05) via PIX no mesmo dia.")
    bullet_num(doc, "50%", "Na entrega final (29/05) via PIX no mesmo dia.")
    body_p(doc,
        "NF contra America Auto Posto Ltda. (CNPJ 34.144.804/0001-86) ou contra Habilis - a definir."
    )

    # --- DOCUMENTACAO ---
    section_title(doc, "Documentacao fornecida pela Habilis")
    body_p(doc,
        "LAO vigente | extrato de condicionantes | plantas (baixa, SASC, mecanico, piso) | mapa | "
        "RGI | MCE existente | laudos analiticos de agua e efluentes | teste de estanqueidade 10/2024 | "
        "dispensa de outorga | ficha tecnica dos tanques (pos-inspecao)."
    )

    # --- O QUE ENVIAR ---
    section_title(doc, "Sua cotacao precisa trazer")
    bullet_num(doc, "1", "Valor global do pacote (MCO + Anexos I e II) e valor por peca.")
    bullet_num(doc, "2", "ARTs inclusas no valor (indicar CREA-GO e atribuicao).")
    bullet_num(doc, "3", "Prazo de entrega a partir do aceite - deve caber em 28/05 e 29/05.")
    bullet_num(doc, "4", "Aceite do pagamento 50/50 via PIX ou contraproposta.")
    bullet_num(doc, "5", "Nome, CREA, especialidade e experiencia previa em postos.")
    bullet_num(doc, "6", "Numero de revisoes inclusas (recomendado: minimo 2).")
    bullet_num(doc, "7", "Suporte pos-entrega se SEMMA pedir complementacao.")
    bullet_num(doc, "8", "Validade da proposta (dias).")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    styled_run(p, "Enviar para: ", size=10, color=CINZA_9, bold=True)
    styled_run(p, "guilherme.anom@gmail.com", size=10, color=VERDE, bold=True, mono=True)

    # --- RODAPE INSTITUCIONAL ---
    section_title(doc, "Sobre a Habilis")
    body_p(doc,
        "Operacao recorrente em licenciamento ambiental de postos de combustiveis em Goias. "
        "Boas cotacoes evoluem para parceria com tabela fixa de precos e prioridade de atendimento "
        "em proximos projetos (MCO, CONAMA 273, SPDA, laudos estruturais, drenagem/SAO).",
        size=9
    )

    # Save
    out_dir = os.path.dirname(os.path.abspath(__file__))
    out_path = os.path.join(out_dir, "COT-2026-001_MCO_CONAMA273_AutoPostoAmerica_ONEPAGER.docx")
    doc.save(out_path)
    print(f"DOCX gerado: {out_path}")
    return out_path


if __name__ == "__main__":
    build()
