#!/usr/bin/env python3
"""
Pedidos formais de cotacao · 3 prestadores · Auto Posto America.
Documento interno Habilis · mensagens prontas para WhatsApp/Email.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

VERDE = RGBColor(0x1B, 0x5E, 0x20)
VERDE_L = RGBColor(0x4C, 0xAF, 0x50)
VERDE_BG = "E8F5E9"
AMBER = RGBColor(0xF5, 0x7F, 0x17)
AMBER_BG = "FFF8E1"
VERMELHO = RGBColor(0xB7, 0x1C, 0x1C)
VERMELHO_BG = "FFEBEE"
AZUL = RGBColor(0x0D, 0x47, 0xA1)
AZUL_BG = "E3F2FD"
CINZA_9 = RGBColor(0x1C, 0x1C, 0x1E)
CINZA_7 = RGBColor(0x3A, 0x3A, 0x3C)
CINZA_5 = RGBColor(0x6C, 0x6C, 0x70)
CINZA_BG = "FAFAFA"
CINZA_MSG = "F4F4F4"
BRANCO = RGBColor(0xFF, 0xFF, 0xFF)
FONT_SANS = "Segoe UI"
FONT_MONO = "Consolas"


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for name, val in edges.items():
        if val is None: continue
        size, color = val
        el = OxmlElement(f'w:{name}')
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(size)); el.set(qn('w:color'), color)
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def set_paragraph_shading(p, hex_color):
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    p_pr.append(shd)


def set_paragraph_border_left(p, size_pt, color_hex):
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), str(int(size_pt * 8)))
    left.set(qn('w:space'), '6'); left.set(qn('w:color'), color_hex)
    pbdr.append(left); p_pr.append(pbdr)


def styled_run(p, text, *, bold=False, italic=False, size=None, color=None, mono=False):
    run = p.add_run(text)
    font_name = FONT_MONO if mono else FONT_SANS
    run.font.name = font_name
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.find(qn('w:rFonts'))
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts'); r_pr.append(r_fonts)
    r_fonts.set(qn('w:ascii'), font_name); r_fonts.set(qn('w:hAnsi'), font_name); r_fonts.set(qn('w:cs'), font_name)
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    return run


def add_page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def heading_section(doc, num, title, space_before=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '14')
    bottom.set(qn('w:space'), '3'); bottom.set(qn('w:color'), '1B5E20')
    pbdr.append(bottom); p_pr.append(pbdr)
    styled_run(p, f"{num}  ", size=11, color=VERDE_L, mono=True)
    styled_run(p, title, size=15, color=VERDE, bold=True)


def label_mono(doc, text, space_before=6, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(2)
    styled_run(p, text.upper(), size=8.5, color=color or VERDE, bold=True, mono=True)


def body_p(doc, text, size=10, space_after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    styled_run(p, text, size=size, color=CINZA_7)


def mensagem_block(doc, text):
    """Bloco destacado para a mensagem copy/paste."""
    t = doc.add_table(rows=1, cols=1)
    c = t.cell(0, 0)
    set_cell_shading(c, CINZA_MSG)
    set_cell_border(c, left=(16, "1B5E20"), top=(4, "E0E0E0"),
                    right=(4, "E0E0E0"), bottom=(4, "E0E0E0"))
    p1 = c.paragraphs[0]
    p1.paragraph_format.space_after = Pt(3)
    styled_run(p1, "MENSAGEM · COPIAR E ENVIAR", size=8, color=CINZA_9, bold=True, mono=True)

    for linha in text.split("\n"):
        p = c.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        if linha.strip().startswith("•"):
            styled_run(p, linha, size=9.5, color=CINZA_9, mono=True)
        elif linha.strip().isupper() and len(linha.strip()) > 3 and not linha.startswith(" "):
            styled_run(p, linha, size=9, color=VERDE, bold=True, mono=True)
        else:
            styled_run(p, linha, size=9.5, color=CINZA_9, mono=True)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(2)


def contatos_block(doc, titulo, itens):
    """Bloco lista de contatos com padrao visual consistente."""
    label_mono(doc, titulo, space_before=8)
    for item in itens:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.4)
        styled_run(p, "• ", size=10, color=VERDE, bold=True)
        if isinstance(item, tuple):
            nome, contato = item
            styled_run(p, nome, size=10, color=CINZA_9, bold=True)
            styled_run(p, " · ", size=10, color=CINZA_5)
            styled_run(p, contato, size=9.5, color=CINZA_7, mono=True)
        else:
            styled_run(p, item, size=10, color=CINZA_7)


def table_hdr(doc, headers, widths_cm=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_shading(c, "1B5E20")
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        styled_run(p, h.upper(), size=8, color=BRANCO, bold=True, mono=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if widths_cm and i < len(widths_cm): c.width = Cm(widths_cm[i])
    return t


def add_row(t, values, zebra=False, bold_cols=None, mono_cols=None, is_total=False, size=9):
    row = t.add_row()
    for i, val in enumerate(values):
        c = row.cells[i]
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
        bold = (bold_cols and i in bold_cols) or is_total
        mono = mono_cols and i in mono_cols
        styled_run(p, val, size=size, color=CINZA_9 if bold else CINZA_7, bold=bold, mono=mono)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        if is_total: set_cell_shading(c, VERDE_BG)
        elif zebra: set_cell_shading(c, CINZA_BG)


def field_page(p):
    r = p.add_run()
    r.font.name = FONT_MONO; r.font.size = Pt(8); r.font.color.rgb = CINZA_5
    fld_b = OxmlElement('w:fldChar'); fld_b.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.text = 'PAGE'
    fld_e = OxmlElement('w:fldChar'); fld_e.set(qn('w:fldCharType'), 'end')
    r._r.append(fld_b); r._r.append(it); r._r.append(fld_e)


def field_numpages(p):
    r = p.add_run()
    r.font.name = FONT_MONO; r.font.size = Pt(8); r.font.color.rgb = CINZA_5
    fld_b = OxmlElement('w:fldChar'); fld_b.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.text = 'NUMPAGES'
    fld_e = OxmlElement('w:fldChar'); fld_e.set(qn('w:fldCharType'), 'end')
    r._r.append(fld_b); r._r.append(it); r._r.append(fld_e)


# ============ MENSAGENS (texto puro) ============
MSG_LAB = """Olá, [NOME], bom dia. Sou Guilherme, da Hábilis Consultoria.

Preciso de orçamento para coleta e análise laboratorial do Auto Posto América, em Guapó/GO (BR-060, km 203), dentro da renovação da LAO 033/2020 (protocolo-meta 26/06/2026).

ESCOPO PRINCIPAL

• Efluentes da SAO — 2 pontos (entrada e saída)
  Parâmetros: pH, temperatura, turbidez, óleos e graxas, sólidos sedimentáveis, DBO, DQO, OD, MBAS, metais (Pb, Zn, Cu, Ni, Cd, Ca, Ba) e TPH.
  Base: CONAMA 430/2011 + CEMAM 029/2018 Anexo XI.

• Água subterrânea — 3 poços de monitoramento já instalados (PM-01, PM-02, PM-03).
  Parâmetros: BTEX e PAH.
  Base: CONAMA 420/2009.

Data prevista da coleta: 22 de maio de 2026.
Prazo dos laudos: até 15 dias úteis após a coleta.

Requisitos do laboratório:
- Acreditação ISO/IEC 17025 vigente
- Credenciamento INMETRO vigente
- ART do químico responsável nos laudos
- Cadeia de custódia e coordenadas UTM das amostras

COTAR SEPARADO (CONTINGÊNCIA)

• Rodada complementar de coleta e análise de efluentes
  Para acionamento caso o primeiro laudo aponte parâmetros ainda fora de conformidade com CONAMA 430/2011 após intervenção na SAO. Cotar como "segunda rodada" com preço fechado.

INFORMAR NA PROPOSTA

- Valor global e desdobramento por parâmetro
- Valor da contingência (segunda rodada)
- Custo de deslocamento até Guapó
- Prazo dos laudos
- Comprovantes de ISO 17025 e INMETRO
- Validade da proposta

Consegue me retornar até 30/04/2026?

Obrigado.
Guilherme · Hábilis Consultoria"""


MSG_INVEST = """Olá, [NOME], bom dia. Sou Guilherme, da Hábilis Consultoria.

Preciso de orçamento para investigação ambiental confirmatória em posto de combustível — Auto Posto América, Guapó/GO (BR-060, km 203), dentro da renovação da LAO 033/2020 (protocolo-meta 26/06/2026).

ESCOPO PRINCIPAL (TURNKEY · TUDO INCLUSO)

• Sondagens SPT em no mínimo 3 furos em pontos estratégicos: tanques, SASC, bombas e área de descarga.
• Amostragem de solo e água subterrânea em cada furo.
• Análises químicas no laboratório parceiro (ISO 17025): VOC, BTEX, PAH, metais e TPH.
• Relatório de Investigação Confirmatória completo, com interpretação hidrogeológica, comparação com valores orientadores CONAMA 420/2009 e recomendação técnica.
• ART do geólogo ou engenheiro ambiental responsável.
• Dossiê fotográfico e georreferenciamento dos furos em UTM SIRGAS 2000.

Normas: NBR 15515-1/2/3, NBR 16209, CEMAM 029/2018, CONAMA 420/2009.

O posto já tem 3 poços de monitoramento instalados (PM-01, PM-02, PM-03), com análises de 07/2024 em conformidade com CONAMA 420. Não precisa perfurar novos poços.

Cronograma:
- Contratação: até 05/05/2026
- Mobilização em campo: 25/05/2026
- Entrega do laudo: 15/06/2026

COTAR SEPARADO (CONTINGÊNCIAS)

• Análise de Risco Ambiental — modalidade APP (Avaliação Preliminar de Perigo)
  Para acionamento se o laudo confirmatório indicar indícios de contaminação ou a SEMMA exigir.

• Análise de Risco Ambiental — modalidade ARQ (Análise Quantitativa de Risco)
  Para acionamento em caso de confirmação de passivo e exigência de avaliação de risco detalhada.

• Investigação Detalhada (NBR 15515-3)
  Para acionamento se o laudo confirmatório disparar necessidade de mais furos e caracterização aprofundada. Cotar por faixa estimada.

INFORMAR NA PROPOSTA

- Valor global da investigação turnkey
- Valor das 3 contingências acima, separadas
- Número de furos sugerido e justificativa técnica
- Profundidade máxima prevista
- Prazo de emissão da ART pelo CREA
- Portfolio e referências em postos de combustível
- Custo de deslocamento até Guapó
- Validade da proposta

Consegue me retornar até 30/04/2026?

Obrigado.
Guilherme · Hábilis Consultoria"""


MSG_ENG = """Olá, [NOME], bom dia. Sou Guilherme, da Hábilis Consultoria.

Preciso contratar profissional com atribuição CREA (Engenheiro Civil, Ambiental, Mecânico ou Químico) para duas peças técnicas do Auto Posto América, em Guapó/GO (BR-060, km 203), dentro da renovação da LAO 033/2020 (protocolo-meta 26/06/2026).

A coordenação geral do processo permanece com a Hábilis (RT Biólogo · CRBio). Você assina as peças com ART própria no CREA em composição multidisciplinar (Res. CONFEA 1.025/2009 + CONAMA 237/1997).

ESCOPO PRINCIPAL

• MCO — Memorial de Caracterização do Empreendimento conforme Anexo IX CEMAM 029/2018. Descrição técnica de tanques, SASC, tubulações, SAO, infraestrutura predial, fluxogramas e memoriais.

• Anexo I da CONAMA 273/2000 — tabela descritiva do empreendimento.

• Anexo II da CONAMA 273/2000 — classificação em classes de risco ambiental (1, 2 ou 3) com fundamentação.

A Hábilis fará inspeção técnica no local em 12/05 e compartilha o relatório com você como insumo das peças.

Cronograma:
- Contratação: até 15/05/2026
- Entrega do MCO: 28/05/2026
- Entrega dos Anexos I e II CONAMA 273: 29/05/2026

Entregáveis:
- Peças em DOCX editável e PDF
- Até 2 revisões inclusas
- ART do profissional registrada no CREA

Requisitos:
- Registro CREA ativo com atribuição compatível
- Experiência em postos de combustível (anexar portfolio)
- Emissão da ART em até 15 dias úteis

COTAR SEPARADO (CONTINGÊNCIAS)

• Projeto de reforma/adequação da SAO
  Para acionamento se diagnóstico operacional (feito pela Hábilis na inspeção 12/05) apontar necessidade de intervenção estrutural (redimensionamento, troca de compartimento, ajuste hidráulico com projeto de engenharia). Inclui memorial de cálculo, planta executiva, especificações técnicas e ART de projeto.

• Execução física da obra da SAO (turnkey, opcional)
  Para acionamento se a empresa trabalha também com execução direta da obra projetada acima (limpeza profunda não entra aqui — é serviço de mantenedor, não de eng.).

• Laudo técnico com ART para anuência de saneamento
  Para acionamento se a inspeção apontar fossa/sumidouro sem anuência da concessionária local. Serviço pequeno.

INFORMAR NA PROPOSTA

- Valor global do MCO + Anexos I e II
- Valor da ART e quantidade de revisões inclusas
- Valor de cada contingência acima, separado
- Registro CREA e atribuição específica
- Portfolio em postos de combustível
- Se requer visita ao local ou trabalha com dados da inspeção Hábilis de 12/05
- Validade da proposta

Consegue me retornar até 30/04/2026?

Obrigado.
Guilherme · Hábilis Consultoria"""


# ============ BUILD ============
def build():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(1.8); s.bottom_margin = Cm(1.8)
    s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)

    # Header
    h_p = s.header.paragraphs[0]
    styled_run(h_p, "Auto Posto América | Pedidos de Cotação · Terceiros", size=8, color=CINZA_5, mono=True)
    h_p.add_run("\t")
    styled_run(h_p, "Uso interno Hábilis", size=8, color=VERDE, bold=True, mono=True)

    # Footer
    f_p = s.footer.paragraphs[0]
    styled_run(f_p, "Hábilis · Guilherme · Abr/2026", size=8, color=CINZA_5, mono=True)
    f_p.add_run("\t")
    styled_run(f_p, "Página ", size=8, color=CINZA_5, mono=True)
    field_page(f_p)
    styled_run(f_p, " de ", size=8, color=CINZA_5, mono=True)
    field_numpages(f_p)

    # ============ CAPA ENXUTA ============
    p = doc.add_paragraph()
    styled_run(p, "H  Hábilis", size=14, color=VERDE, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    styled_run(p, "USO INTERNO · NÃO ENVIAR AO CLIENTE", size=9, color=VERMELHO, bold=True, mono=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    styled_run(p, "Pedidos de cotação · Terceiros", size=22, color=CINZA_9, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    styled_run(p, "Auto Posto América · Renovação LAO 033/2020", size=12, color=CINZA_7)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    styled_run(p, "Três mensagens prontas para envio via WhatsApp ou email, uma por tipo de "
               "prestador. Cada uma traz o ", size=10.5, color=CINZA_7)
    styled_run(p, "escopo principal ", size=10.5, color=CINZA_9, bold=True)
    styled_run(p, "+ ", size=10.5, color=CINZA_7)
    styled_run(p, "contingências cotadas em separado ", size=10.5, color=CINZA_9, bold=True)
    styled_run(p, "(para acionamento condicional). Ao final, tabela consolidada para registrar "
               "as cotações recebidas e escolher os vencedores.", size=10.5, color=CINZA_7)

    # Resumo dos 3 pedidos
    label_mono(doc, "Os 3 pedidos · visão geral", space_before=8)
    t = table_hdr(doc, ["#", "Prestador", "Escopo principal", "Contingências"],
                  [0.9, 3.8, 5.5, 6.8])
    add_row(t, ["01", "Laboratório analítico",
                "Efluentes SAO + água dos 3 poços · coleta 22/05",
                "Segunda rodada (se SAO não passar na 1ª)"],
            bold_cols=[0, 1], mono_cols=[0], size=9)
    add_row(t, ["02", "Empresa de investigação ambiental",
                "Turnkey · SPT + lab + laudo confirmatório + ART",
                "Análise de Risco (APP e ARQ) + Investigação Detalhada"],
            zebra=True, bold_cols=[0, 1], mono_cols=[0], size=9)
    add_row(t, ["03", "Profissional habilitado (CREA)",
                "MCO (Anexo IX CEMAM) + Anexo I e II CONAMA 273/2000",
                "Projeto de reforma SAO + execução turnkey + laudo saneamento"],
            bold_cols=[0, 1], mono_cols=[0], size=9)

    # O que fica com a Habilis
    label_mono(doc, "O que fica com a Hábilis (você · RT CRBio)", space_before=10)
    for item in [
        "Diagnóstico operacional da SAO após inspeção 12/05",
        "Plano de correção operacional (se não for estrutural)",
        "PGA/PCA atualizado (CP 4.14)",
        "RCA 2026 semestral (CP 4.4)",
        "PGRS · PAE · Plano das 28 condicionantes",
        "Relatório técnico consolidado · integração de todas as peças (19/06)",
        "Protocolo final no Processo SEMMA 22760/2020 (26/06)",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        styled_run(p, item, size=9.5, color=CINZA_7)

    add_page_break(doc)

    # ============ PEDIDO 01 · LABORATÓRIO ============
    heading_section(doc, "01", "Laboratório analítico · água e efluentes", space_before=0)

    body_p(doc,
           "Cotação focada em análise laboratorial rotineira (efluentes da SAO + água dos "
           "poços de monitoramento). Base para o RCA 2026 (CP 4.4) e para subsídio à "
           "verificação pós-intervenção na SAO.", size=10)

    contatos_block(doc, "Enviar para", [
        ("Vida Ambiental (Goiânia/GO)", "(62) 3261-1134 · comercial@aguaemeioambiente.com.br"),
        ("PROMATEC (Rio Claro/SP)", "(19) 3523-7455 · www.promatecambiental.com.br"),
        ("1 alternativa local a mapear", "buscar lab GO/DF com acreditação ISO 17025 e INMETRO"),
    ])

    mensagem_block(doc, MSG_LAB)

    add_page_break(doc)

    # ============ PEDIDO 02 · INVESTIGAÇÃO ============
    heading_section(doc, "02", "Empresa de investigação ambiental", space_before=0)

    body_p(doc,
           "Cotação turnkey (campo + laboratório + laudo + ART em pacote único). "
           "Correspondência direta com os itens 24 (mobilização), 33 (análise de risco) "
           "e 34 (laudo final) do checklist cronológico.", size=10)

    contatos_block(doc, "Enviar para", [
        "Sem referência no histórico · buscar 2-3 empresas com atuação em Goiás especializadas em passivo ambiental de postos de combustível.",
        "Sugestão: consultar colegas da área, grupos de RT ambiental em GO, ou buscar empresas cadastradas no CEMAM / listadas em processos SEMAD similares.",
    ])

    mensagem_block(doc, MSG_INVEST)

    add_page_break(doc)

    # ============ PEDIDO 03 · ENG CREA ============
    heading_section(doc, "03", "Profissional habilitado · MCO + Anexos CONAMA 273", space_before=0)

    body_p(doc,
           "Cotação focada exclusivamente em peças técnicas que exigem atribuição CREA "
           "(MCO + Anexos I e II da CONAMA 273/2000). Diagnóstico operacional da SAO e "
           "PGA/PCA ficam com a Hábilis. Projeto estrutural da SAO é contingente (só "
           "aciona se diagnóstico Hábilis apontar necessidade de reforma).", size=10)

    contatos_block(doc, "Enviar para", [
        ("Niveo Arquitetura e Engenharia", "(62) 98219-6796 · Natalie Reis · comercial@niveoarquitetura.com.br"),
        ("Ecotec Engenharia", "CREA-MG 78929/D · executora do teste de estanqueidade 2024 · confirmar atuação em GO"),
        ("1 alternativa local a mapear", "buscar eng. civil/ambiental com portfolio em postos de combustível em Goiás"),
    ])

    mensagem_block(doc, MSG_ENG)

    add_page_break(doc)

    # ============ CONSOLIDAÇÃO ============
    heading_section(doc, "R", "Consolidação das cotações · preencher", space_before=0)

    body_p(doc,
           "Tabela para registrar as cotações recebidas. 3 linhas por pedido para comparar "
           "até 3 prestadores. Marcar o escolhido em VENCEDOR.", size=10)

    t = table_hdr(doc,
                  ["#", "Prestador", "Contato", "Escopo principal (R$)", "Contingência (R$)", "Prazo", "Vencedor"],
                  [0.9, 2.8, 3.0, 2.6, 2.6, 1.5, 1.5])
    servicos_lista = [
        ("01", "Laboratório"),
        ("02", "Investigação"),
        ("03", "Eng. CREA"),
    ]
    for i, (num, serv) in enumerate(servicos_lista):
        for cot in range(3):
            zebra = (i % 2 == 1)
            first = (cot == 0)
            add_row(t,
                [num if first else "",
                 serv if first else "",
                 "", "", "", "", ""],
                zebra=zebra, bold_cols=[0, 1], mono_cols=[0], size=8.5)

    label_mono(doc, "Critérios de escolha", space_before=10)
    for item in [
        "Habilitação vigente confirmada (ISO 17025/INMETRO/CREA conforme o caso)",
        "Experiência comprovada em postos de combustível (portfolio anexado)",
        "Prazo de entrega compatível com o cronograma Hábilis",
        "Preço competitivo (não o mais baixo, o mais justo)",
        "Clareza na proposta (valores desdobrados, contingências precificadas)",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        styled_run(p, item, size=9.5, color=CINZA_7)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    styled_run(p, "Atenção: ", size=9.5, color=VERMELHO, bold=True)
    styled_run(p, "habilitação vencida (INMETRO, ISO, calibração) inviabiliza o dossiê mesmo "
               "com preço baixo. Conferir antes de contratar. A lição do teste de estanqueidade "
               "de 10/2024 (executado com calibração Suporty vencida desde 03/2023) é clara — "
               "laudo contestável custa muito mais do que a economia da contratação.", size=9.5, color=CINZA_7)

    # Salvar
    here = os.path.dirname(os.path.abspath(__file__))
    out = os.path.join(here, "Pedidos_Cotacao_Terceiros_AutoPostoAmerica_v1.docx")
    doc.save(out)
    print(f"OK -> {out}")
    return out


if __name__ == "__main__":
    build()
