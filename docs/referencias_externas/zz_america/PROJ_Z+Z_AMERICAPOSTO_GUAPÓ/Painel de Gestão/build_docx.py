#!/usr/bin/env python3
"""
Gerador DOCX nativo do Documento Consolidado Final · Auto Posto América.
Uso: python3 build_docx.py
Salva DOCX ao lado do HTML.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Paleta Hábilis
VERDE = RGBColor(0x1B, 0x5E, 0x20)
VERDE_L = RGBColor(0x4C, 0xAF, 0x50)
VERDE_BG = "E8F5E9"
AMBER = RGBColor(0xF5, 0x7F, 0x17)
AMBER_BG = "FFF8E1"
VERMELHO = RGBColor(0xB7, 0x1C, 0x1C)
VERMELHO_BG = "FFEBEE"
AZUL = RGBColor(0x0D, 0x47, 0xA1)
AZUL_BG = "E3F2FD"
CINZA_9 = RGBColor(0x1C, 0x1C, 0x1E)
CINZA_7 = RGBColor(0x3A, 0x3A, 0x3C)
CINZA_5 = RGBColor(0x6C, 0x6C, 0x70)
CINZA_BG = "FAFAFA"
BRANCO = RGBColor(0xFF, 0xFF, 0xFF)
FONT_SANS = "Segoe UI"
FONT_MONO = "Consolas"


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for name, val in edges.items():
        if val is None: continue
        size, color = val
        el = OxmlElement(f'w:{name}')
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(size)); el.set(qn('w:color'), color)
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def set_paragraph_shading(p, hex_color):
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    p_pr.append(shd)


def set_paragraph_border_left(p, size_pt, color_hex):
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), str(int(size_pt * 8)))
    left.set(qn('w:space'), '6'); left.set(qn('w:color'), color_hex)
    pbdr.append(left); p_pr.append(pbdr)


def styled_run(p, text, *, bold=False, italic=False, size=None, color=None, mono=False):
    run = p.add_run(text)
    font_name = FONT_MONO if mono else FONT_SANS
    run.font.name = font_name
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.find(qn('w:rFonts'))
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts'); r_pr.append(r_fonts)
    r_fonts.set(qn('w:ascii'), font_name); r_fonts.set(qn('w:hAnsi'), font_name); r_fonts.set(qn('w:cs'), font_name)
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    return run


def mixed_run(p, text, *, size=10.5, default_color=CINZA_7, bold_parts=None):
    """Renderiza parágrafo com partes em negrito/cor diferente."""
    if not bold_parts:
        styled_run(p, text, size=size, color=default_color)
        return
    remaining = text
    for target in bold_parts:
        idx = remaining.find(target)
        if idx == -1: continue
        if idx > 0: styled_run(p, remaining[:idx], size=size, color=default_color)
        styled_run(p, target, size=size, color=CINZA_9, bold=True)
        remaining = remaining[idx + len(target):]
    if remaining:
        styled_run(p, remaining, size=size, color=default_color)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def heading_section(doc, num, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '18')
    bottom.set(qn('w:space'), '4'); bottom.set(qn('w:color'), '1B5E20')
    pbdr.append(bottom); p_pr.append(pbdr)
    styled_run(p, f"{num}  ", size=13, color=VERDE_L, mono=True)
    styled_run(p, title, size=18, color=VERDE, bold=True)


def heading_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    set_paragraph_border_left(p, 2.5, "1B5E20")
    p.paragraph_format.left_indent = Cm(0.2)
    styled_run(p, text, size=12, color=CINZA_9, bold=True)


def heading_h4(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    styled_run(p, text.upper(), size=9.5, color=VERDE, bold=True, mono=True)


def body_p(doc, text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    mixed_run(p, text, bold_parts=bold_parts)


def lede(doc, text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_shading(p, VERDE_BG)
    set_paragraph_border_left(p, 2.5, "1B5E20")
    mixed_run(p, text, size=11, default_color=CINZA_9, bold_parts=bold_parts)


def callout(doc, title, body, kind="verde", bold_body_parts=None):
    bg = {"verde": VERDE_BG, "amber": AMBER_BG, "vermelho": VERMELHO_BG,
          "azul": AZUL_BG, "cinza": CINZA_BG}.get(kind, VERDE_BG)
    border = {"verde": "1B5E20", "amber": "F57F17", "vermelho": "B71C1C",
              "azul": "0D47A1", "cinza": "6C6C70"}.get(kind, "1B5E20")
    t = doc.add_table(rows=1, cols=1)
    c = t.cell(0, 0)
    set_cell_shading(c, bg)
    set_cell_border(c, left=(24, border))
    p1 = c.paragraphs[0]
    p1.paragraph_format.space_after = Pt(3)
    styled_run(p1, title.upper(), size=9.5, color=CINZA_9, bold=True, mono=True)
    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    mixed_run(p2, body, size=10, default_color=CINZA_7, bold_parts=bold_body_parts)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def table_hdr(doc, headers, widths_cm=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_shading(c, "1B5E20")
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        styled_run(p, h.upper(), size=8.5, color=BRANCO, bold=True, mono=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if widths_cm and i < len(widths_cm):
            c.width = Cm(widths_cm[i])
    return t


def add_row(t, values, zebra=False, bold_cols=None, mono_cols=None, is_total=False):
    row = t.add_row()
    for i, val in enumerate(values):
        c = row.cells[i]
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        bold = (bold_cols and i in bold_cols) or is_total
        mono = mono_cols and i in mono_cols
        styled_run(p, val, size=9.5, color=CINZA_9 if bold else CINZA_7, bold=bold, mono=mono)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        if is_total: set_cell_shading(c, VERDE_BG)
        elif zebra: set_cell_shading(c, CINZA_BG)


def field_page(p):
    """Insere campo PAGE."""
    r = p.add_run()
    r.font.name = FONT_MONO; r.font.size = Pt(8); r.font.color.rgb = CINZA_5
    fld_b = OxmlElement('w:fldChar'); fld_b.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.text = 'PAGE'
    fld_e = OxmlElement('w:fldChar'); fld_e.set(qn('w:fldCharType'), 'end')
    r._r.append(fld_b); r._r.append(it); r._r.append(fld_e)


def field_numpages(p):
    r = p.add_run()
    r.font.name = FONT_MONO; r.font.size = Pt(8); r.font.color.rgb = CINZA_5
    fld_b = OxmlElement('w:fldChar'); fld_b.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.text = 'NUMPAGES'
    fld_e = OxmlElement('w:fldChar'); fld_e.set(qn('w:fldCharType'), 'end')
    r._r.append(fld_b); r._r.append(it); r._r.append(fld_e)


def build():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(1.8); s.bottom_margin = Cm(1.8)
    s.left_margin = Cm(1.8); s.right_margin = Cm(1.8)

    # Header
    h_p = s.header.paragraphs[0]
    styled_run(h_p, "Auto Posto América · Regularização Ambiental", size=8, color=CINZA_5, mono=True)
    h_p.add_run("\t")
    styled_run(h_p, "Hábilis · Documento Consolidado", size=8, color=VERDE, bold=True, mono=True)

    # Footer
    f_p = s.footer.paragraphs[0]
    styled_run(f_p, "Hábilis · Abril 2026", size=8, color=CINZA_5, mono=True)
    f_p.add_run("\t")
    styled_run(f_p, "Página ", size=8, color=CINZA_5, mono=True)
    field_page(f_p)
    styled_run(f_p, " de ", size=8, color=CINZA_5, mono=True)
    field_numpages(f_p)

    # ===== CAPA =====
    p = doc.add_paragraph()
    styled_run(p, "H  ", size=16, color=VERDE, bold=True)
    styled_run(p, "Hábilis", size=16, color=VERDE, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(40)
    styled_run(p, "REGULARIZAÇÃO AMBIENTAL", size=9, color=CINZA_5, mono=True)
    for _ in range(4): doc.add_paragraph()

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    styled_run(p, "PROCESSO DE REGULARIZAÇÃO AMBIENTAL", size=10, color=VERDE, bold=True, mono=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    styled_run(p, "Auto Posto América", size=30, color=CINZA_9, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    styled_run(p, "Renovação LAO 033/2020 · SEMMA Guapó · Processo 22760/2020", size=13, color=CINZA_7)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    styled_run(p, "Documento único e consolidado para condução do processo de regularização, destinado ao ", size=11, color=CINZA_7)
    styled_run(p, "Grupo Z+Z", size=11, color=CINZA_9, bold=True)
    styled_run(p, ". Integra plano técnico contratado (3 frentes · 51 itens), racional de decisões, cronograma operacional (67 dias corridos · 4 fases), checklist cronológico e validação de auditoria. Apresenta, como bônus, a oportunidade de continuidade · Compliance Paralelo (§11 · 23 itens adicionais · fora do pacote atual). Protocolo-meta ", size=11, color=CINZA_7)
    styled_run(p, "26/06/2026 · 180 dias antes do vencimento da LAO.", size=11, color=CINZA_9, bold=True)

    tc = doc.add_table(rows=2, cols=4)
    for i, h in enumerate(["DESTINATÁRIO", "EMISSOR", "PROTOCOLO-META", "EMISSÃO"]):
        c = tc.cell(0, i)
        styled_run(c.paragraphs[0], h, size=8, color=CINZA_5, bold=True, mono=True)
        set_cell_border(c, top=(12, "1B5E20"))
    vals = ["Grupo Z+Z", "Hábilis", "26/06/2026", "19/04/2026"]
    for i, v in enumerate(vals):
        c = tc.cell(1, i)
        styled_run(c.paragraphs[0], v, size=10, color=CINZA_9, bold=True, mono=i >= 2)

    add_page_break(doc)

    # ===== §01 RESUMO =====
    heading_section(doc, "01", "Resumo Executivo")
    lede(doc, "O Auto Posto América opera sob a LAO 033/2020, vencimento 23/12/2026. "
         "O protocolo-meta acordado é de 180 dias antes do vencimento, o que fixa a submissão "
         "do dossiê em 26/06/2026. Janela de execução de 67 dias a partir desta data. Base "
         "documental suficiente para iniciar a estruturação.",
         ["LAO 033/2020", "23/12/2026", "180 dias antes do vencimento", "26/06/2026"])

    t = table_hdr(doc, ["Prazo", "Data", "Observação"], [4.8, 3.2, 9.5])
    add_row(t, ["Protocolo-meta · 180 dias antes", "26/06/2026",
                "Submissão do dossiê no Processo SEMMA 22760/2020. Janela de execução 67 dias."],
            bold_cols=[1], mono_cols=[1])
    add_row(t, ["Prazo regulatório · vencimento LAO", "23/12/2026",
                "Data-limite absoluta. 26/06 preserva 180 dias de folga."],
            zebra=True, bold_cols=[1], mono_cols=[1])
    add_row(t, ["Mínimo legal (120 dias)", "24/08/2026",
                "Piso legal · referência explicativa, não meta operacional."],
            bold_cols=[1], mono_cols=[1])
    doc.add_paragraph()

    heading_h3(doc, "Dados do empreendimento · conforme LAO 033/2020 (Série C2020033A)")
    t = table_hdr(doc, ["Campo", "Valor", "Observação"], [4.5, 6.5, 6.0])
    dados_lao = [
        ("Razão social", "AMÉRICA AUTO POSTO LTDA", "LAO atual emitida em 03/09/2025 em virtude de alteração da razão social anterior."),
        ("CNPJ", "34.144.804/0001-86", "Validar CNAE 'Comércio Varejista de Combustível e Lubrificantes' ativo."),
        ("Endereço", "Rodovia BR-060, km 203, S/N, sentido Posselândia/Guapó", "Microbacia Ribeirão das Posses · Rio dos Bois."),
        ("Zona", "Rural · Guapó/GO · CEP 75.350-000", "Confirmação libera atribuição Agronômica para Uso do Solo."),
        ("Coordenadas", "16°54'07,30\"S · 49°38'31,06\"O", "Base UTM SIRGAS 2000 (conferir em planta)."),
        ("Área do terreno", "20.000 m²", "Área construída: 1.657 m²."),
        ("Potencial poluidor", "2,0", "Classificação SEMMA."),
        ("Representante Legal", "Amarido Alves Pires · CPF 827.115.691-87 [VALIDAR]", "Checklist pede RG/CPF de Sérgio Marques e José Guilherme · confirmar quadro societário atual."),
        ("Interlocutor SEMMA", "Gleidson Nunes Ferreira · Secretário Municipal · Decreto 013/2025", "Assinatura da LAO atual · canal oficial para tramitação 2026."),
        ("RTs originais da LAO", "Edmilson Xavier de Souza · Osmar Roberto de Souza Filho", "Substituir por RTs Hábilis no dossiê de renovação (CP 4.1 atualizada)."),
    ]
    for i, row in enumerate(dados_lao):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0], mono_cols=[1])
    doc.add_paragraph()

    heading_h3(doc, "Cenário")
    body_p(doc, "A base documental recebida em 14/04/2026 cobre parte relevante do rol exigido "
           "pela SEMMA. Peças estruturais em mãos: LAO vigente, RGI atualizado, plantas técnicas, "
           "teste de estanqueidade e análises ambientais. Pendem atualização dos dados cadastrais "
           "(quadro societário), validação física em campo, novas coletas laboratoriais, "
           "atualização do MCO, inspeção técnica e consolidação do dossiê. Nenhum fato "
           "impeditivo identificado.", ["14/04/2026"])

    heading_h3(doc, "O que depende do Grupo Z+Z")
    pts = [
        ("Acesso via procuração para a Hábilis consolidar status no Sistema Veredas e protocolar novo pedido conforme IN 15/2026 SEMAD (DOE/GO 17/04/2026) — prazo de transição prorrogado para 15/07/2026; como o pedido será formalizado após 16/04/2026, o posto segue a regra do §15 (ingresso na fila após formalização, sem ordem cronológica de origem).",
         ["IN 15/2026 SEMAD", "15/07/2026", "§15"]),
        ("Contratação de serviços técnicos: laboratório, estanqueidade credenciada, investigação confirmatória e projetista.", []),
        ("Pagamento de taxas municipais, estaduais e de terceiros contratados.", []),
        ("Aprovação formal do relatório técnico antes do protocolo.", []),
    ]
    for txt, bolds in pts:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        mixed_run(p, txt, bold_parts=bolds)

    callout(doc, "Posição técnica",
            "Processo viável dentro da janela de 67 dias até o protocolo-meta. A consolidação do "
            "uso do solo, como pré-requisito da renovação, é o primeiro passo estruturante e está "
            "sob condução da Hábilis. O cronograma é factível desde que as quatro fases corram em "
            "sequência, sem gargalos de contratação de terceiros.",
            kind="verde", bold_body_parts=["67 dias", "uso do solo"])

    # ===== §02 ESTRUTURA =====
    heading_section(doc, "02", "Estrutura da Gestão Ambiental")
    lede(doc, "A gestão do Auto Posto América está estruturada em 3 frentes contratadas + 1 "
         "frente paralela de continuidade. O pacote formalizado da Renovação LAO totaliza 51 "
         "itens (11 Uso do Solo + 26 Renovação LAO + 14 Condicionantes). A Frente 04 · Compliance "
         "Paralelo (23 itens · ANP · IBAMA · NR-20 · PNRS · monitoramentos) foi mapeada como "
         "oportunidade de continuidade pós-protocolo, fora do pacote atual — apresentada no §11 "
         "para visualização do Grupo Z+Z.",
         ["3 frentes contratadas + 1 frente paralela de continuidade", "51 itens",
          "oportunidade de continuidade pós-protocolo, fora do pacote atual", "§11"])

    t = table_hdr(doc, ["Frente", "Itens", "Órgão / Norma", "Papel"], [4.2, 1.4, 4.2, 7.8])
    rows = [
        ("01 · Uso do Solo (pré-requisito)", "11", "Prefeitura de Guapó",
         "Destrava o protocolo. Sem certidão vigente, a SEMMA não recebe o dossiê."),
        ("02 · Renovação LAO (objeto central)", "26", "SEMMA Guapó · Proc. 22760/2020",
         "Dossiê técnico e relatório da Hábilis; peça principal do trabalho."),
        ("03 · Condicionantes", "14", "LAO 033/2020 · Extrato 05/2024",
         "Plano de atendimento das 28 CPs; SEMMA avalia o histórico antes de deferir."),
        ("04 · Compliance Paralelo", "23", "ANP · IBAMA · NR-20 · PNRS",
         "Frente permanente; autuação correlata contamina o processo principal."),
    ]
    for i, row in enumerate(rows):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0, 1], mono_cols=[1])
    add_row(t, ["Subtotal · pacote contratado", "51", "3 frentes", "11 + 26 + 14 = 51 itens contratados."],
            is_total=True, mono_cols=[1])
    add_row(t, ["Frente paralela · continuidade (§11)", "23", "Bônus · fora do pacote atual",
                "ANP + IBAMA + NR-20 + PNRS + Monitoramentos · oferta de gestão continuada pós-protocolo."],
            is_total=True, mono_cols=[1])
    add_row(t, ["Total mapeado", "74", "3 contratadas + 1 paralela",
                "51 contratados + 23 de continuidade."],
            is_total=True, mono_cols=[1])
    doc.add_paragraph()

    heading_h3(doc, "Como as frentes se conectam")
    body_p(doc, "Uso do Solo destrava a Renovação — sem certidão vigente, a SEMMA não recebe "
           "o protocolo. Renovação da LAO sustenta a operação licenciada e é o objeto central. "
           "As Condicionantes sustentam o deferimento (evidência técnica). Compliance Paralelo "
           "protege o conjunto: autuação em ANP, IBAMA, NR-20, PNRS ou monitoramentos gera "
           "passivo regulatório que contamina a análise do processo principal.",
           ["Uso do Solo destrava a Renovação", "objeto central", "Condicionantes",
            "Compliance Paralelo"])

    callout(doc, "Gestão efetiva, não apenas renovação",
            "A Hábilis conduz a gestão ambiental efetiva do empreendimento, não apenas a renovação "
            "documental. É a diferença entre cumprir um prazo e preservar a regularidade. A "
            "renovação segue cronograma de 67 dias. O compliance paralelo roda em cadência própria.",
            kind="verde", bold_body_parts=["gestão ambiental efetiva", "cronograma de 67 dias"])

    # ===== §03 PANORAMA =====
    heading_section(doc, "03", "Panorama Operacional")
    body_p(doc, "Três frentes contratadas + 1 frente paralela de continuidade (§11 · bônus). O uso do solo antecede a renovação. As condicionantes "
           "sustentam o dossiê técnico. O compliance paralelo roda em cadência própria, protegendo "
           "a integridade regulatória contra autuações correlatas.")

    heading_h3(doc, "3.1 · Uso do Solo — pré-requisito da renovação")
    t = doc.add_table(rows=2, cols=2)
    fr1 = [
        ("O que é", "Certidão de Uso e Ocupação do Solo emitida pela Prefeitura de Guapó, atualizada e vigente na data do protocolo."),
        ("Por que importa", "Sem a certidão vigente o protocolo de renovação não é recebido pela SEMMA. Base da regularidade locacional."),
        ("O que já existe", "RGI 2025 · documentação dominial · diagnóstico locacional compatível com uso atual."),
        ("O que falta", "Certidão de Uso e Ocupação vigente · requerimento (se inexistente) · certidão negativa municipal."),
    ]
    for i, (lbl, val) in enumerate(fr1):
        r, col = divmod(i, 2)
        c = t.cell(r, col)
        p1 = c.paragraphs[0]; p1.paragraph_format.space_after = Pt(2)
        styled_run(p1, lbl.upper(), size=8, color=AZUL, bold=True, mono=True)
        p2 = c.add_paragraph(); styled_run(p2, val, size=10, color=CINZA_7)
        set_cell_border(c, left=(18, "0D47A1"))
    doc.add_paragraph()
    callout(doc, "O que destrava a próxima etapa",
            "Certidão vigente em mãos, liberando a sequência do dossiê da LAO.",
            kind="azul", bold_body_parts=["Certidão vigente em mãos"])

    heading_h3(doc, "3.2 · Renovação da LAO 033/2020 — objeto central")
    body_p(doc, "Dossiê técnico no Processo SEMMA 22760/2020, instruído com peças obrigatórias "
           "e relatório técnico consolidado. Mantém o empreendimento licenciado e regular perante "
           "a SEMMA Guapó.", ["Processo SEMMA 22760/2020"])

    h = doc.add_paragraph()
    styled_run(h, "O QUE JÁ EXISTE (11 PEÇAS PRINCIPAIS)", size=9, color=CINZA_7, bold=True, mono=True)
    for item in [
        "LAO 033/2020 vigente (até 23/12/2026) · AVCB Cercon 161726 (válido até 20/10/2026)",
        "Planta Baixa, SASC e mecânica 05/2024 · Planta de Piso Nível 1 · Mapa SIRGAS 2000",
        "MCE 2020 (SCAN0035) · Memorial Fotográfico SASC 2023 · PGR 2024",
        "Teste de Estanqueidade 10/2024 (NBR 16795) · RGI 2025 · Despacho DNIT 4374748",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        styled_run(p, item, size=10, color=CINZA_7)

    h = doc.add_paragraph()
    styled_run(h, "O QUE FALTA", size=9, color=CINZA_7, bold=True, mono=True)
    for item in [
        "Requerimento modelo SEMMA · Publicação CONAMA 006/86",
        "MCO atualizado · Anexo IX CEMAM 029/2018 (+ART)",
        "Anexos I e II Res. CONAMA 273/2000 (+ART)",
        "Relatório técnico consolidado (+ART Hábilis)",
        "Laudo de Investigação Confirmatória (CEMAM 029/2018)",
        "RCA semestral · PGRS · PAE · Análise de Risco (se exigida)",
        "NFs de tanques, tubulações e dispositivos · Certificados CONAMA 319 + Portarias INMETRO",
        "Protocolo novo de outorga via Sistema Veredas · CNDs · DUAM",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        styled_run(p, item, size=10, color=CINZA_7)

    heading_h3(doc, "3.2.1 · Visita de Campo (12/05/2026) · 25 pontos em 5 eixos")
    body_p(doc, "Inspeção técnica de campo agendada para 12/05/2026, organizada em 5 eixos com "
           "25 pontos de verificação. A visita fecha as lacunas entre documentação histórica e "
           "realidade física, subsidiando o MCO atualizado, o novo RCA, o Laudo de Investigação "
           "Confirmatória, o Plano das 28 condicionantes, a validação das plantas e o levantamento "
           "hídrico que define o regime da outorga no Veredas (dispensa simplificada ou processo completo).",
           ["12/05/2026", "5 eixos com 25 pontos"])

    t = table_hdr(doc, ["Eixo", "Pontos", "Foco", "Subsidia"], [1.2, 1.3, 7.0, 8.0])
    eixos = [
        ("1", "4", "Implantação e imóvel (layout × plantas, SIRGAS 2000, RGI) + APP/RL/CAR + vegetação remanescente", "Validação de plantas · croqui · MCO · Uso do Solo"),
        ("2", "7", "Armazenamento, abastecimento + tubulações subterrâneas + levantamento hídrico (captação → regime outorga)", "MCO · NFs TR-16 · CONAMA 319 · Requerimento Veredas 13/05"),
        ("3", "5", "Efluentes, drenagem e esgoto — auditoria da reprovação RCA 2024", "RCA 03/06 · PGRS · anuência saneamento"),
        ("4", "4", "Resíduos e passivo + marcação física dos 3 furos + interferências enterradas", "Laudo Investigação · mobilização 25/05 · PGRS"),
        ("5", "5", "Segurança, documentos, rotina operacional + entrevista estruturada com operador (5 blocos: SAO/RCA · operações · incidentes · calibrações · gestão)", "PAE · Compliance Paralelo · plano de recomposição CPs · Risco 02"),
    ]
    for i, row in enumerate(eixos):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0], mono_cols=[0, 1])
    doc.add_paragraph()

    heading_h3(doc, "3.2.2 · Publicação em edital · Res. CONAMA 006/1986 (item 6 TR SEMMA)")
    body_p(doc, "A publicidade do pedido de renovação é exigência legal da Resolução CONAMA "
           "006/1986 e do item 6 do TR SEMMA Guapó. Garante o direito de informação da "
           "comunidade e permite manifestação pública antes da decisão do órgão. Não-cumprimento "
           "bloqueia o recebimento do dossiê. Data no cronograma: 04/05/2026 · responsável "
           "Hábilis · fase F1.", ["CONAMA 006/1986", "04/05/2026"])

    heading_h4(doc, "Duas publicações obrigatórias · custos estimados")
    t = table_hdr(doc, ["Veículo", "Para o Auto Posto América", "Custo", "Observação"],
                   [4.0, 5.5, 2.5, 5.0])
    rows_pub = [
        ("① Jornal oficial do estado", "DOE/GO · Diário Oficial de Goiás (ABC)",
         "R$ 400-800", "Cobrado por cm-coluna · anúncio padrão."),
        ("② Jornal de grande circulação", "O Popular · Diário da Manhã · ou regional certificado",
         "R$ 800-2.000", "Evitar jornais locais sem reconhecimento."),
        ("Total estimado", "Cliente paga · Hábilis contrata",
         "R$ 1.200-2.800", "Cabe em 'taxas e publicações' do CAPEX piso."),
    ]
    for i, row in enumerate(rows_pub):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0], mono_cols=[2])
    doc.add_paragraph()

    heading_h4(doc, "Texto modelo da publicação (pronto para usar)")
    callout(doc, "Modelo CONAMA 006/86 · Auto Posto América",
            "\"AMÉRICA AUTO POSTO LTDA, CNPJ 34.144.804/0001-86, torna público que requereu à "
            "Secretaria Municipal de Meio Ambiente — SEMMA de Guapó/GO a Renovação da Licença "
            "Ambiental de Operação — LAO nº 033/2020 para a atividade de Comércio Varejista de "
            "Combustíveis e Lubrificantes para Veículos Automotores e Similares, localizada na "
            "Rodovia BR-060, km 203, S/N, sentido Posselândia/Guapó, Zona Rural, Guapó-GO, "
            "CEP 75.350-000. Processo nº 22760/2020.\"",
            kind="azul",
            bold_body_parts=["AMÉRICA AUTO POSTO LTDA", "34.144.804/0001-86",
                             "SEMMA de Guapó/GO", "LAO nº 033/2020", "22760/2020"])

    heading_h4(doc, "Passos operacionais · 04/05 a 19/06")
    t = table_hdr(doc, ["#", "Passo", "Prazo", "Responsável"], [0.8, 8.0, 2.2, 6.0])
    rows_passos = [
        ("1", "Redigir texto da publicação (modelo CONAMA 006/86)", "até 02/05", "Hábilis · redação"),
        ("2", "Cliente aprova o texto (CNPJ · razão social · endereço)", "02/05", "Cliente · validação escrita"),
        ("3", "Contratar publicação no DOE/GO", "04/05", "Hábilis · pagamento via Cliente"),
        ("4", "Contratar publicação em jornal de grande circulação", "04-05/05", "Hábilis · preferir O Popular/Diário da Manhã"),
        ("5", "Receber e arquivar comprovantes das 2 publicações", "06-10/05", "Hábilis · pasta física + digital"),
        ("6", "Janela de contestação pública (30 dias)", "04/05 → 03/06", "—"),
        ("7", "Responder eventual contestação (raro)", "até 15/06", "Hábilis"),
        ("8", "Anexar os 2 comprovantes ao dossiê", "19/06", "Hábilis · Relatório Consolidado"),
    ]
    for i, row in enumerate(rows_passos):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[1], mono_cols=[0, 2])
    doc.add_paragraph()

    heading_h4(doc, "Riscos e cuidados operacionais")
    t = table_hdr(doc, ["Risco", "Causa", "Mitigação"], [4.5, 6.5, 5.5])
    rows_riscos_pub = [
        ("① Dados errados no anúncio",
         "CNPJ trocado · razão social desatualizada · endereço incompleto · processo errado",
         "Validação formal escrita do cliente antes de publicar · usar modelo deste documento."),
        ("② Razão social desatualizada",
         "LAO atual registra alteração societária; publicar com dados antigos gera divergência",
         "Resolver CNPJ atualizado (24/04) antes de publicar em 04/05."),
        ("③ Jornal sem 'grande circulação'",
         "Veículo local sem reconhecimento de circulação regional",
         "Consultar SEMMA Guapó sobre veículos aceitos · preferir jornais de Goiás."),
        ("④ Contestação pública (baixa probabilidade)",
         "Morador, ONG ou MP manifesta oposição",
         "Operação existente em zona rural · sem histórico de conflito. Hábilis responde em 15 dias se ocorrer."),
        ("⑤ Publicação adicional da outorga Veredas",
         "Se inspeção 12/05 apontar 'processo completo', SEMAD pode exigir publicidade adicional",
         "Publicidade da outorga normalmente automática pelo Veredas · confirmar com SEMAD."),
    ]
    for i, row in enumerate(rows_riscos_pub):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0])
    doc.add_paragraph()

    callout(doc, "Por que publicar em 04/05 e não mais tarde",
            "A janela de contestação pública de 30 dias (04/05 → 03/06) deve estar encerrada "
            "antes da consolidação do Relatório Técnico (19/06). Publicar após 15/05 atrasa o "
            "dossiê ou exige protocolo sem o período de publicidade cumprido — cenário em que "
            "a SEMMA rejeita o recebimento. Os 53 dias entre 04/05 e 26/06 são folga intencional, "
            "não desperdício.",
            kind="verde",
            bold_body_parts=["04/05 → 03/06", "antes da consolidação",
                             "folga intencional, não desperdício"])

    # ===== §3.2.3 Memorial de Caracterização da Obra (MCO) =====
    heading_h3(doc, "3.2.3 · Memorial de Caracterização da Obra (MCO) · Anexo IX CEMAM 029/2018")
    body_p(doc, "O MCO é a peça técnica central do dossiê — descreve o posto para a SEMMA e "
           "serve de base para as demais peças técnicas (Anexos I/II CONAMA 273, Relatório "
           "Consolidado, laudos). Se o MCO estiver desatualizado ou divergir da realidade "
           "física, todas as peças derivadas ficam em terreno movediço. Data no cronograma: "
           "28/05/2026 · produção posterior à inspeção de 12/05.",
           ["peça técnica central do dossiê", "28/05/2026"])

    heading_h4(doc, "Estrutura exigida pelo Anexo IX da Res. CEMAM 029/2018")
    t = table_hdr(doc, ["Capítulo do MCO", "Conteúdo obrigatório"], [5.5, 11.5])
    rows_mco = [
        ("1. Identificação", "Razão social · CNPJ · endereço · coordenadas UTM SIRGAS 2000 · área · potencial poluidor · representante legal · RT · ART"),
        ("2. Descrição do empreendimento", "Atividade · horário · funcionários · layout · histórico de ampliações/reformas"),
        ("3. Sistema de armazenamento", "Ficha técnica de cada tanque (fabricante · ano · capacidade · parede simples/dupla · produto · contenção secundária)"),
        ("4. Sistema de abastecimento (SASC)", "Bombas · ilhas · bacias · piso impermeável · drenagem · descarga selada · respiros · câmaras de contenção"),
        ("5. Tubulações subterrâneas", "Material · diâmetro · monitoramento intersticial · válvulas de retenção · válvulas de corte"),
        ("6. Sistemas auxiliares", "Lava-jato (vazão · destino) · trocador de óleo (contenção · destinação) · lanchonete · escritório"),
        ("7. Sistema de controle ambiental", "SAO · desarenador · poços de monitoramento · canaleta · segregação águas pluviais × servidas · fossa"),
        ("8. Caracterização ambiental do entorno", "APP · RL · vegetação · microbacia (Ribeirão das Posses/Rio dos Bois) · uso do solo adjacente"),
        ("9. Sistemas de emergência e segurança", "Kit de contenção · AVCB Cercon 161726 · plano de emergência · sinalização · NR-20"),
        ("10. Anexos cartográficos e fotográficos", "Planta baixa/situação · Quadro UTM · Memorial Fotográfico datado · .pdf + .dwg + shapefile (CP 3.12)"),
    ]
    for i, row in enumerate(rows_mco):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0])
    doc.add_paragraph()

    heading_h4(doc, "Status do acervo atual × o que falta")
    t = table_hdr(doc, ["Item", "Status", "Observação"], [5.0, 3.5, 8.5])
    rows_status_mco = [
        ("MCE 2020 (SCAN0035)", "Base, desatualizado",
         "Versão original · ponto de partida · exige atualização de 6 anos de operação."),
        ("Plantas 05/2024 (SASC + mecânica + piso)", "Em mãos",
         "Desenho mais recente · validar em campo (Eixo 1)."),
        ("Memorial Fotográfico SASC 2023", "Existe",
         "Atualizar com fotos datadas + UTM SIRGAS 2000 pós-inspeção 12/05."),
        ("Teste de Estanqueidade 10/2024", "Integra MCO",
         "Subsidia caracterização dos tanques (cap. 3)."),
        ("RGI 2025", "Em mãos",
         "Dados dominiais atualizados · entra no cap. 1."),
        ("Mapa SIRGAS 2000", "Em mãos",
         "Georreferenciamento · entra nos anexos (cap. 10)."),
        ("Ficha técnica detalhada dos tanques", "A levantar em campo",
         "Eixo 2 · parede, contenção, material das tubulações."),
        ("Descrição das tubulações subterrâneas", "A levantar em campo",
         "Eixo 2 · requerido pela CP 4.8 (Port. 084/2005 + CONAMA 273)."),
        ("Confirmação de lava-jato e trocador", "[VALIDAR em 12/05]",
         "Entrevista Bloco B · define cap. 6 e parâmetros do RCA."),
        ("ART do MCO", "Condicional",
         "Depende do parecer CREA/GO (30/04) · ver §6.1 e Risco 07."),
    ]
    for i, row in enumerate(rows_status_mco):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0])
    doc.add_paragraph()

    heading_h4(doc, "Insumos da inspeção 12/05 → capítulos do MCO")
    t = table_hdr(doc, ["Eixo", "Captura em campo", "Capítulo(s) do MCO"], [3.5, 6.0, 7.5])
    rows_insumos = [
        ("Eixo 1 · Implantação",
         "Layout × plantas · coordenadas · limites · APP/RL/CAR",
         "Cap. 1 Identificação · Cap. 2 Descrição · Cap. 8 Entorno · Cap. 10 Anexos"),
        ("Eixo 2 · Armazenamento + hídrico",
         "Ficha técnica tanques · tubulações · bombas/ilhas · descarga · levantamento hídrico",
         "Cap. 3 Armazenamento · Cap. 4 SASC · Cap. 5 Tubulações · Cap. 6 Auxiliares"),
        ("Eixo 3 · Efluentes/SAO",
         "SAO · diagnóstico · operações (lava-jato/trocador) · poços · drenagem/esgoto",
         "Cap. 6 Auxiliares · Cap. 7 Controle ambiental"),
        ("Eixo 4 · Resíduos/passivo",
         "Armazenamento resíduos perigosos · coleta · pontos críticos · marcação 3 furos",
         "Cap. 7 Controle ambiental · Cap. 10 Anexos (croqui)"),
        ("Eixo 5 · Segurança e rotina",
         "AVCB · licenças · contratos · calibrações · entrevista operador",
         "Cap. 9 Emergência e segurança · Cap. 2 Descrição"),
    ]
    for i, row in enumerate(rows_insumos):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0])
    doc.add_paragraph()

    heading_h4(doc, "Cronograma de produção do MCO")
    t = table_hdr(doc, ["Quando", "O quê", "Quem", "Entrega / insumo"], [2.0, 7.0, 3.0, 5.0])
    rows_prod = [
        ("12/05", "Inspeção de campo · captura de dados · Eixos 1-5", "Hábilis",
         "Formulário · fotos · medições"),
        ("13-15/05", "Digitalização dos dados · croqui atualizado · upload de fotos",
         "Hábilis · escritório", "Pasta do projeto organizada"),
        ("16-22/05", "Redação dos capítulos 1-10 do MCO conforme Anexo IX",
         "Hábilis · RT responsável", "Minuta do MCO"),
        ("23-27/05", "Integração com NFs (18/05) · CONAMA 319 · Portarias INMETRO · revisão",
         "Hábilis", "Versão final · ART emitida"),
        ("28/05", "Entrega do MCO consolidado", "Hábilis",
         "Peça do dossiê · integra Relatório 19/06"),
        ("29/05", "Anexos I e II CONAMA 273 · derivados do MCO", "Hábilis",
         "Peças complementares"),
    ]
    for i, row in enumerate(rows_prod):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[1], mono_cols=[0])
    doc.add_paragraph()

    heading_h4(doc, "Atribuição da ART · 2 cenários conforme parecer CREA/GO (30/04)")
    t = table_hdr(doc, ["Cenário", "Condição", "Custo adicional", "Prazo ART"],
                   [4.5, 6.5, 2.5, 4.0])
    rows_art = [
        ("A · RT Engenharia Agronômica interno (preferencial)",
         "CREA/GO aceita Res. 218/73 alínea c para posto em zona rural (LAO item 1.2)",
         "R$ 0", "Imediato · ART Hábilis em 27-28/05"),
        ("B · Terceiro Eng. Civil/Ambiental (contingência)",
         "CREA/GO interpreta de forma restritiva e nega cobertura agronômica",
         "R$ 8-15k", "5-15 dias úteis · contratação até 10/05"),
    ]
    for i, row in enumerate(rows_art):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0], mono_cols=[2, 3])
    doc.add_paragraph()

    heading_h4(doc, "Riscos específicos do MCO")
    t = table_hdr(doc, ["Risco", "Causa", "Mitigação"], [5.0, 6.0, 6.0])
    rows_riscos_mco = [
        ("① Divergência planta × realidade",
         "Alterações físicas não-declaradas desde 2020 (tanques, SASC, lava-jato)",
         "Eixo 1 da inspeção confere layout · MCO registra realidade + justificativa · pode gerar CP 3.5 retroativa."),
        ("② Atribuição CREA não resolvida",
         "Parecer CREA/GO atrasa ou sai restritivo (Risco 07)",
         "Lista de parceiros Eng. Civil/Ambiental pré-negociada até 05/05 · acionamento em 48h."),
        ("③ MCE 2020 como base incompleta",
         "Documento original com lacunas em tubulações, contenção ou sistemas auxiliares",
         "Inspeção 12/05 supre lacunas · validar capítulo por capítulo contra Anexo IX."),
        ("④ Lava-jato/trocador não declarados",
         "Se existirem fora do MCE 2020 · gera CP 3.5 retroativa + muda parâmetros do RCA",
         "Bloco B da entrevista fecha a questão · MCO traz cap. 6 completo · comunicação proativa à SEMMA."),
        ("⑤ Formato técnico não-conforme CP 3.12",
         "MCO entregue sem .dwg + shapefile + carimbo + UTM SIRGAS 2000",
         "Checklist final 24/06 valida formato · projetista contratado até 15/05 produz arquivos vetoriais."),
    ]
    for i, row in enumerate(rows_riscos_mco):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0])
    doc.add_paragraph()

    callout(doc, "Leitura prática · o MCO como peça ancorada",
            "O MCO não existe isoladamente — ele é ancorado em três dependências críticas: "
            "inspeção 12/05 (dados de campo), parecer CREA 30/04 (atribuição da ART) e NFs "
            "dos equipamentos até 18/05 (ficha técnica dos tanques). Se qualquer uma falhar, "
            "o MCO de 28/05 escorrega. Mitigação: monitoramento diário das três dependências "
            "na semana de 13-18/05. Uma vez entregue, o MCO destrava os Anexos I/II CONAMA "
            "273 (29/05 · §3.2.4) e alimenta o Relatório Consolidado (19/06).",
            kind="verde",
            bold_body_parts=["três dependências críticas", "12/05", "30/04", "18/05",
                             "13-18/05", "destrava os Anexos I/II CONAMA 273"])

    # ===== §3.2.4 Anexos I e II CONAMA 273/2000 =====
    heading_h3(doc, "3.2.4 · Anexos I e II da Res. CONAMA 273/2000 · classificação do empreendimento (item 23 TR)")
    body_p(doc, "Os Anexos I e II da Resolução CONAMA 273/2000 são formulários padronizados que "
           "acompanham o pedido de licença e são derivados diretos do MCO. O Anexo I traduz o "
           "MCO para o formato tabular; o Anexo II classifica o empreendimento em uma das 3 "
           "classes de risco ambiental — e essa classe define exigências técnicas adicionais. "
           "Data no cronograma: 29/05/2026 (dia seguinte ao MCO).",
           ["derivados diretos do MCO", "3 classes de risco ambiental", "29/05/2026"])

    heading_h4(doc, "Composição dos Anexos")
    t = table_hdr(doc, ["Peça", "Função", "Principais campos"], [3.5, 5.5, 8.0])
    rows_anexos = [
        ("Anexo I · Requerimento e informações",
         "Resumo tabular do MCO em formato CONAMA · a SEMMA consulta direto.",
         "Dados cadastrais · caracterização geral · SASC (tanques · produtos · capacidade) · bombas · tratamento (SAO · fossa · poços) · RT + ART."),
        ("Anexo II · Classificação do empreendimento",
         "Matriz de enquadramento de risco · define a classe do posto.",
         "Capacidade total · idade do SASC · parede simples/dupla · monitoramento intersticial · distância de poço (≥15m) · distância de curso d'água (≥30m) · tipo de produto · lava-jato/trocador."),
    ]
    for i, row in enumerate(rows_anexos):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0])
    doc.add_paragraph()

    heading_h4(doc, "As 3 classes de risco ambiental (Anexo II) e implicações")
    t = table_hdr(doc, ["Classe", "Critérios típicos", "Probab. p/ posto", "Implicações técnicas"],
                   [3.0, 5.0, 4.0, 5.0])
    rows_classes = [
        ("Classe 1 (baixo risco)",
         "Capacidade ≤15 m³ · parede dupla · monitoramento intersticial operacional · distâncias conforme",
         "Se parede dupla + monitoramento OK",
         "Exigências mínimas · cumprimento básico · menos peças no dossiê."),
        ("Classe 2 (médio risco)",
         "Capacidade intermediária · parede dupla com limitações · monitoramento parcial · distâncias limítrofes",
         "Mais provável dado potencial poluidor 2,0",
         "Exigências padrão · monitoramento reforçado · investigação confirmatória usual."),
        ("Classe 3 (alto risco)",
         "Capacidade elevada · parede simples · sem monitoramento intersticial · próximo a corpo hídrico",
         "Se tanques parede simples e/ou distâncias críticas",
         "Exigências adicionais · PGR obrigatório · monitoramento reforçado · auto de verificação · pode exigir nova LI (CP 4.12)."),
    ]
    for i, row in enumerate(rows_classes):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0])
    doc.add_paragraph()

    heading_h4(doc, "Dados decisivos para o Anexo II · validação na inspeção 12/05")
    t = table_hdr(doc, ["Variável", "O que a LAO traz", "O que validar em 12/05"],
                   [4.5, 5.5, 7.0])
    rows_decisivos = [
        ("Capacidade total", "Não explícita", "Soma dos tanques · Eixo 2"),
        ("Idade do SASC", "Licenciado desde 2020", "Instalação original pode ser anterior · Eixo 2"),
        ("Parede dos tanques (simples/dupla)", "Não informada", "Ponto decisivo da classe · Eixo 2 · ficha técnica + inspeção"),
        ("Monitoramento intersticial", "CP 4.8 exige", "Operacional ou só instalado? · Eixo 2"),
        ("Distância de poço de captação", "—", "Levantamento hídrico Eixo 2 · mínimo 15 m"),
        ("Distância de curso d'água", "Microbacia Ribeirão das Posses / Rio dos Bois", "Medir em campo · ≥30 m · pode puxar Classe 3 se próximo"),
        ("Lava-jato / trocador", "[VALIDAR]", "Bloco B da entrevista · Eixo 5"),
        ("Potencial poluidor", "2,0 (LAO item 2.1)", "Fixo · base da classificação"),
    ]
    for i, row in enumerate(rows_decisivos):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0])
    doc.add_paragraph()

    heading_h4(doc, "Riscos específicos dos Anexos I e II")
    t = table_hdr(doc, ["Risco", "Causa", "Mitigação"], [5.0, 6.0, 6.0])
    rows_riscos_anx = [
        ("① Divergência entre Anexo I e MCO",
         "Preenchimento paralelo gera inconsistências",
         "Anexos preenchidos a partir do MCO já consolidado (28/05) · revisão cruzada 24/06."),
        ("② Classe subestimada no Anexo II",
         "Metodologia defensável mas questionável pela SEMMA",
         "Classificar com critério honesto · se Classe 2 ou 3, apresentar com plano de controle desde o início."),
        ("③ Distâncias estimadas em vez de medidas",
         "Planilha exige valor numérico e o dado foi estimado visualmente",
         "Medição com trena em campo (Eixo 1-2) · registro fotográfico georreferenciado."),
        ("④ ART esquecida",
         "Comum confundir ART do MCO com a dos Anexos · ambos exigem ART",
         "Anexos têm ART própria · incluir no pedido CREA junto com a do MCO (27-28/05)."),
        ("⑤ Lava-jato/trocador não declarados afetam a classe",
         "Operações adicionais não no MCE 2020 mudam a classificação",
         "Bloco B da entrevista 12/05 resolve · classificação consistente MCO × Anexos."),
    ]
    for i, row in enumerate(rows_riscos_anx):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0])
    doc.add_paragraph()

    callout(doc, "Leitura prática · os Anexos como selo da classificação",
            "O Anexo II é mais importante que o Anexo I do ponto de vista regulatório — ele "
            "define a classe de risco ambiental do empreendimento, que por sua vez dispara "
            "(ou não) exigências técnicas adicionais. Posto Classe 1 tem dossiê mais enxuto; "
            "posto Classe 3 precisa de PGR robusto. A classificação correta vem da inspeção "
            "12/05, se apoia no MCO de 28/05, e se materializa nos Anexos de 29/05. Ordem: "
            "inspeção → MCO → Anexos I/II → dossiê consolidado.",
            kind="verde",
            bold_body_parts=["Anexo II é mais importante que o Anexo I",
                             "classe de risco ambiental", "inspeção → MCO → Anexos I/II → dossiê"])

    heading_h3(doc, "3.3 · Condicionantes — evidenciação técnica")
    body_p(doc, "Plano de atendimento e comprovação documental das 28 condicionantes da LAO "
           "033/2020 (12 Orientativas · itens 3.1 a 3.12 + 16 Específicas · itens 4.1 a 4.16). "
           "A SEMMA avalia o histórico antes de deferir a renovação. Classificação: "
           "10 regras declaratórias + 18 ações auditáveis (6 atendidas · 10 parciais · 2 pendentes).",
           ["28 condicionantes", "10 regras declaratórias", "18 ações auditáveis",
            "6 atendidas · 10 parciais · 2 pendentes"])

    heading_h3(doc, "3.3.1 · Quadro analítico das 28 condicionantes")

    heading_h4(doc, "Bloco 1 · Condicionantes Orientativas (12 · itens 3.1 a 3.12)")
    t = table_hdr(doc, ["#", "Texto sintético", "Tipo", "Status", "Evidência / Ação"],
                   [0.9, 6.8, 1.3, 2.0, 6.0])
    cps_orient = [
        ("3.1", "Veracidade dos documentos e projetos anexados ao processo 22760/2020.", "Ação", "✓ Atendida",
         "Acervo 14/04/2026 confere. Reafirmada no relatório consolidado."),
        ("3.2", "SEMMA pode modificar, suspender ou cancelar a licença.", "Regra", "— Declarat.",
         "Cláusula. Cumprida por aceite."),
        ("3.3", "Licença pode ser revogada por descumprimento.", "Regra", "— Declarat.",
         "Cláusula. Cumprida por aceite."),
        ("3.4", "Comunicação imediata à SEMMA em caso de acidentes.", "Ação", "✓ Atendida",
         "Sem incidentes. Procedimento no PAE (05/06)."),
        ("3.5", "Não autoriza modificações sem manifestação SEMMA.", "Ação", "✓ Atendida",
         "Conduta respeitada. Ampliação → LI (item 4.12)."),
        ("3.6", "Não dispensa outros alvarás federais/estaduais/municipais.", "Regra", "— Declarat.",
         "Ressalva jurídica. Cumprida por aceite."),
        ("3.7", "SEMMA reserva direito de novas exigências.", "Regra", "— Declarat.",
         "Cláusula de poder. Cumprida por aceite."),
        ("3.8", "Suspensão automática se demais licenças vencerem.", "Ação", "⚠ Parcial",
         "AVCB vigente 20/10/2026 ✓. LAI 008/2020 venceu 31/03/2026 · alvará sanitário [VALIDAR] · DURH002737 vencida 09/12/2023 (regularização Veredas)."),
        ("3.9", "Não produz efeitos de cessão ou aquisição.", "Regra", "— Declarat.",
         "Cláusula jurídica. Cumprida por aceite."),
        ("3.10", "Manter atualizado ANP e AVCB.", "Ação", "⚠ Parcial",
         "AVCB Cercon 161726 ✓. ANP cadastro + declarações [VALIDAR · Compliance Paralelo]."),
        ("3.11", "Renovação com antecedência mínima de 120 dias.", "Ação", "✓ Atendida",
         "Meta 26/06/2026 = 180 dias · 60 dias a mais que o mínimo."),
        ("3.12", "Premissas de relatórios: .jpg + ABNT · fotos datadas UTM SIRGAS 2000 · .pdf/.dwg · shapefile · RT + ART.", "Ação", "⚠ Parcial",
         "Atendido no Relatório Consolidado 19/06. Shapefile/.dwg via projetista (15/05)."),
    ]
    for i, row in enumerate(cps_orient):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0], mono_cols=[0])
    doc.add_paragraph()

    heading_h4(doc, "Bloco 2 · Condicionantes Específicas (16 · itens 4.1 a 4.16)")
    t = table_hdr(doc, ["#", "Texto sintético", "Período.", "Status", "Evidência / Ação"],
                   [0.9, 6.8, 1.3, 2.0, 6.0])
    cps_espec = [
        ("4.1", "Operação com RT habilitado.", "Contínua", "⚠ Parcial",
         "RTs originais: Edmilson Xavier e Osmar Roberto. Atualizar para Hábilis (procuração 23/04)."),
        ("4.2", "Seguir Res. CEMAM 29/2018.", "Contínua", "✓ Atendida",
         "Resolução é base do plano."),
        ("4.3", "Laudo estanqueidade bianual (NBR 13.784 cancelada · laudo 10/2024 já aplicou NBR 16795/2019).", "Bianual", "✓ Atendida (c/ informe)",
         "Laudo 10/2024 apresentado como está · cobre ciclo bianual até 10/2026. Opção B (próximo ciclo regular 10/2026) como plano principal. Opção A (teste de contingência em 30-60d) armada se SEMMA contestar por INMETRO NCC 21.08786 vencido 29/04/2024 ou calibração Suporty vencida 03/2023. Ver §3.3.3."),
        ("4.4", "Relatório Ambiental semestral (pH · turbidez · O&G · SS · DBO · DQO · OD · MBAS · metais · TPH · BTEX/PAH).", "Sem. fev+ago", "✗ Pendente",
         "Análise 2024 REPROVADA · 64 mg/L O&G. Novo RCA 03/06/2026 após adequação SAO. [VALIDAR] lava-jato (MBAS) e trocador (TPH)."),
        ("4.5", "Contrato + NFs semestrais manutenção SAO + coleta de resíduos oleosos.", "Sem. fev+ago", "⚠ Parcial",
         "Bauer + Limpmil + Ecofenix identificados. NFs a compilar 20/05. Validar licenciamento e certificação ANP."),
        ("4.6", "Óleo lubrificante → empresa autorizada ANP.", "Contínua", "✓ Atendida",
         "Ecofenix identificada. Validar cadastro ANP ativo."),
        ("4.7", "Certificados de Coleta disponíveis por 3 anos.", "Contínua", "⚠ Parcial",
         "Acervo [VALIDAR] · conferir no arquivo físico · inspeção 12/05 Eixo 5."),
        ("4.8", "CONAMA 273/2000 + Portaria 084/2005: descarga selada, câmaras contenção, válvulas retenção, parede dupla, monitoramento intersticial, impermeabilização, SAO.", "Estrutural", "⚠ Parcial",
         "Validação física · inspeção 12/05 · Eixo 2 (7 pts) + Eixo 3 (5 pts). Anexos I/II CONAMA 273 em 29/05."),
        ("4.9", "Plano de Gerenciamento de Risco: integridade · emergência · treinamento.", "Contínua", "⚠ Parcial",
         "PGR 2024 ✓. Validar emergência (PAE 05/06) e treinamento (inspeção Eixo 5). Cruzamento NR-20."),
        ("4.10", "Drenagem pluvial independente das águas servidas.", "Estrutural", "⚠ Parcial",
         "Validar em campo · inspeção 12/05 Eixo 3. Anuência saneamento condicional 25/05."),
        ("4.11", "Eficiência declarada é responsabilidade da empresa.", "Regra", "— Declarat.",
         "Cláusula. Cumprida por aceite."),
        ("4.12", "Ampliação/reforma exige nova LI.", "Regra", "— Declarat.",
         "Cláusula de controle. Cumprida por aceite."),
        ("4.13", "SEMMA promoverá avaliações de controle de poluição.", "Regra", "— Declarat.",
         "Cláusula de fiscalização. Cumprida por aceite."),
        ("4.14", "Atender Projeto Ambiental + PGA + PCA.", "Contínua", "✗ Pendente",
         "PGA e PCA [VALIDAR] no acervo. Se ausentes, recuperar ou produzir · integrar Relatório 19/06."),
        ("4.15", "Descumprimento causa suspensão/cassação.", "Regra", "— Declarat.",
         "Cláusula de penalidade. Cumprida por aceite."),
        ("4.16", "SEMMA reserva novas exigências (reiterada de 3.7).", "Regra", "— Declarat.",
         "Cláusula reiterada. Cumprida por aceite."),
    ]
    for i, row in enumerate(cps_espec):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0], mono_cols=[0])
    doc.add_paragraph()

    heading_h4(doc, "Bloco 3 · Cronograma periódico 2026 (3 CPs recorrentes · Extrato 05/2024)")
    t = table_hdr(doc, ["#", "CP periódica", "Periodicidade", "Entregas 2026 · absorção"],
                   [0.9, 7.1, 2.5, 6.5])
    cps_per = [
        ("4.3", "Laudo de estanqueidade (NBR 16795/2019)",
         "Bianual · jun", "Extrato: jun/2026. Plano: laudo 10/2024 apresentado como está (Opção B · plano principal) · próximo ciclo regular 10/2026. Contingência (Opção A) armada se SEMMA contestar. Ver §3.3.3."),
        ("4.4", "Relatório Ambiental semestral (efluentes + poços)",
         "Semestral · fev+ago", "Extrato: fev+ago/2026. Plano: 03/06/2026 (novo RCA absorve ciclo fev + corrige reprovação 2024); ago/2026 pós-protocolo."),
        ("4.5", "Contrato + NFs (Bauer/Limpmil/Ecofenix)",
         "Semestral · fev+ago", "Extrato: fev+ago/2026. Plano: compilação 20/05/2026 (F2 · Cliente)."),
    ]
    for i, row in enumerate(cps_per):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0], mono_cols=[0])
    doc.add_paragraph()

    callout(doc, "Leitura do quadro · síntese prática",
            "Das 28 CPs, 10 são regras declaratórias (cumpridas por aceite) e 18 são ações auditáveis. "
            "Das 18: 6 atendidas, 10 parciais, 2 pendentes — RCA 2024 reprovado (CP 4.4) e PGA/PCA "
            "a validar (CP 4.14). O protocolo de 26/06 depende de transformar as 10 parciais em "
            "atendidas e executar as 2 pendentes.",
            kind="verde",
            bold_body_parts=["10 são regras declaratórias", "18 são ações auditáveis",
                             "6 atendidas", "10 parciais", "2 pendentes"])

    heading_h3(doc, "3.3.2 · Status da consolidação do RCA 2026 (CP 4.4 · peça crítica)")
    body_p(doc, "O Relatório de Controle Ambiental semestral (RCA · Anexo XI CEMAM 029/2018) é a "
           "condicionante mais crítica do plano — hoje reprovada em 2024 (Óleos e Graxas 64 mg/L · "
           "3,2× o limite). A consolidação do novo RCA para 2026 depende de uma sequência "
           "obrigatória de 5 marcos entre 12/05 e 03/06. Em 21/04/2026, estamos no marco zero "
           "(pré-diagnóstico) — a inspeção de 12/05 é o ponto em que a consolidação começa de fato.",
           ["reprovada em 2024", "3,2× o limite", "5 marcos", "marco zero"])

    heading_h4(doc, "O que já temos")
    t = table_hdr(doc, ["Peça", "Status", "Observação"], [5.5, 3.0, 8.5])
    rows_temos = [
        ("Laudo RCA 2024 (Água Viva Ambiental)", "Reprovado",
         "O&G 64 mg/L · Turbidez fora do padrão · não pode ser reapresentado."),
        ("Contratos de manutenção SAO e coleta", "Identificados",
         "Bauer + Limpmil (SAO) + Ecofenix (óleo) · CP 4.5. NFs semestrais a compilar 20/05."),
        ("PGR 2024", "Em mãos",
         "Suporta CP 4.9 · validar atualização na inspeção 12/05."),
        ("Extrato de Condicionantes 05/2024", "Referência",
         "Confirma periodicidade semestral da CP 4.4 (fev+ago)."),
    ]
    for i, row in enumerate(rows_temos):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0])
    doc.add_paragraph()

    heading_h4(doc, "O que falta · em ordem de execução")
    t = table_hdr(doc, ["#", "Peça / Ação", "Prazo", "Bloqueio atual"],
                   [0.8, 4.5, 2.2, 9.5])
    rows_falta = [
        ("1", "Diagnóstico da SAO em campo", "12/05/2026",
         "Define Hipótese A/B/C. Sem isso, coleta em 22/05 pode reprovar. Eixo 3 da inspeção + Bloco A da entrevista."),
        ("2", "Confirmação de lava-jato e trocador de óleo", "12/05/2026",
         "Define se lab. coleta MBAS (lava-jato) e TPH (trocador). Bloco B da entrevista."),
        ("3", "Adequação física da SAO", "15-25/05/2026",
         "Depende do marco 1. Hip. A limpeza 1-2 sem · Hip. B obra 4-12 sem · Hip. C ajuste hidráulico 2-4 sem."),
        ("4", "Poços de monitoramento", "até 22/05",
         "CP 4.4 exige BTEX e PAH nos poços. Validar existência (Eixo 3). Se ausentes, instalar antes da coleta."),
        ("5", "Nova coleta de efluentes", "22/05/2026",
         "Terceiro · lab. acreditado. Só após adequação (marco 3)."),
        ("6", "Novo laudo RCA 2026 consolidado", "03/06/2026",
         "Peça final · RT com ART. Integra todos parâmetros da CP 4.4."),
        ("7", "NFs semestrais da CP 4.5 compiladas", "20/05/2026",
         "Cliente com contratos Bauer/Limpmil/Ecofenix."),
        ("8", "Certificados de Coleta (CP 4.7)", "12/05/2026",
         "Verificar arquivo físico · 3 anos em local de fácil acesso. Entrevista Eixo 5."),
    ]
    for i, row in enumerate(rows_falta):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[1], mono_cols=[0, 2])
    doc.add_paragraph()

    heading_h4(doc, "Linha do tempo até consolidar")
    callout(doc, "Sequência de 8 marcos · 21/04 → 26/06",
            "21/04 (HOJE pré-diagnóstico) → 12/05 (Inspeção SAO · decide Hip. A/B/C) → "
            "15-25/05 (Adequação SAO conforme hipótese) → 22/05 (Coleta de efluentes · só após adequação) → "
            "03/06 (Laudo RCA 2026 · peça final CP 4.4) → 19/06 (Relatório consolidado integra RCA) → "
            "26/06 (Protocolo SEMMA · dossiê completo).",
            kind="cinza",
            bold_body_parts=["21/04", "12/05", "15-25/05", "22/05", "03/06", "19/06", "26/06"])

    heading_h4(doc, "3 riscos reais que ainda pendem")
    t = table_hdr(doc, ["Risco", "Gatilho", "Consequência"], [4.5, 6.0, 6.5])
    rows_riscos = [
        ("① Hipótese B confirmada (dano estrutural)",
         "Inspeção 12/05 identifica trincas, chicana danificada, subdimensionamento",
         "Obra 4-12 semanas · RCA não consolidado até 26/06 · dossiê vai com plano corretivo anexo. SEMMA aceita com ressalva."),
        ("② Reprovação em sequência",
         "Coleta 22/05 antes de resolver SAO · reprovação 2026 se soma à de 2024",
         "SEMMA questiona viabilidade · risco de indeferimento ou nova LI (CP 4.12). Evitável: coleta só após adequação."),
        ("③ Parâmetros errados na coleta",
         "Coletar sem confirmar lava-jato (MBAS) ou trocador (TPH)",
         "Laudo incompleto · SEMMA pede complemento · atraso 30-60 dias. Evitável em 12/05 com entrevista."),
    ]
    for i, row in enumerate(rows_riscos):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0])
    doc.add_paragraph()

    # Essa chamada será depois desta nova subseção
    _placeholder_marker = None

    heading_h3(doc, "3.3.3 · Estanqueidade · estratégia de apresentação dupla (CP 4.3)")
    body_p(doc, "O Laudo de Estanqueidade de 10/2024 (conforme NBR 16795/2019) está em mãos e "
           "será apresentado à SEMMA como está, em linha com a periodicidade bianual da CP 4.3 "
           "(próximo ciclo regular em 10/2026). A estratégia substitui a execução de novo teste "
           "preventivo em 08/06/2026 por uma apresentação dupla: laudo atual + informe técnico "
           "do cenário + plano de contingência se contestado + programação da cadência bianual.",
           ["10/2024", "NBR 16795/2019", "apresentação dupla", "10/2026"])

    heading_h4(doc, "Informe técnico · status atual do laudo 10/2024")
    t = table_hdr(doc, ["Aspecto", "Situação", "Análise técnica"], [4.5, 3.5, 10.0])
    rows_informe = [
        ("Norma técnica aplicada", "NBR 16795/2019",
         "Norma vigente · substituta da NBR 13.784 cancelada. Adequação normativa cumprida. Sem passivo técnico."),
        ("Periodicidade da CP 4.3", "Bianual",
         "Laudo 10/2024 cobre tecnicamente até 10/2026 — ~4 meses após o protocolo de 26/06."),
        ("Credenciamento INMETRO da empresa", "NCC 21.08786 · venceu 29/04/2024",
         "Válido na data de execução do teste. Argumento a favor: não invalida retroativamente a prova pericial. Argumento contra possível: SEMMA pode exigir credenciamento vigente na data de apresentação."),
        ("Calibração do equipamento Suporty", "Venceu 03/2023",
         "Ponto mais sensível. Se o equipamento usado em 10/2024 era o Suporty com calibração de 2021-2022, rastreabilidade metrológica é válida. Validar com a empresa executora qual equipamento foi usado (inspeção 12/05 Eixo 5)."),
        ("Aproveitamento no dossiê", "Válido para renovação",
         "Laudo recente (~18 meses antes do protocolo), dentro da periodicidade bianual, norma atualizada."),
    ]
    for i, row in enumerate(rows_informe):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0])
    doc.add_paragraph()

    heading_h4(doc, "Opções de tratamento")
    t = table_hdr(doc, ["Opção", "Prazo", "Custo", "Quando ativar", "Racional técnico"],
                   [3.5, 2.5, 2.5, 3.5, 6.0])
    rows_opcoes = [
        ("A · Teste de contingência (plano B)",
         "30-60d pós-exigência", "R$ 15-30k (contingente)",
         "Só se SEMMA contestar formalmente o laudo 10/2024",
         "Acionamento reativo · empresa nova com INMETRO vigente e calibração ISO/IEC 17025 · lista de parceiros pré-negociada em F1."),
        ("B · Teste programado na cadência bianual (plano principal)",
         "10/2026", "R$ 15-30k (certo)",
         "Gestão continuada pós-protocolo · próximo ciclo regular da CP 4.3",
         "Cadência alinhada à periodicidade da condicionante · empresa pré-selecionada com credenciamento vigente · sem emergência."),
    ]
    for i, row in enumerate(rows_opcoes):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0], mono_cols=[1, 2])
    doc.add_paragraph()

    callout(doc, "Recomendação técnica",
            "Apresentar o laudo de 10/2024 como está + adotar Opção B como plano principal (teste "
            "programado na cadência bianual em 10/2026, gestão continuada). A Opção A fica armada "
            "como contingência: se a SEMMA contestar o laudo em resposta à triagem, a Hábilis "
            "aciona um parceiro com credenciamento vigente dentro de 30-60 dias, sem impacto no "
            "protocolo de 26/06 (certidão de tramitação da Veredas + dossiê já protocolado seguem válidos).",
            kind="verde",
            bold_body_parts=["laudo de 10/2024 como está", "Opção B como plano principal",
                             "Opção A fica armada como contingência"])

    callout(doc, "Ação imediata · validar equipamento do teste 10/2024",
            "Na visita de campo de 12/05 (Eixo 5 · entrevista estruturada + documentos físicos), "
            "confirmar com a empresa executora original (Suporty ou sucessora) qual equipamento "
            "foi utilizado e se a calibração estava vigente naquela data. Se sim, defesa técnica "
            "robusta; se não, vale antecipar a migração para Opção A proativamente.",
            kind="amber",
            bold_body_parts=["12/05 (Eixo 5"])

    # ===== §3.3.4 Anuência de saneamento =====
    heading_h3(doc, "3.3.4 · Anuência de saneamento · caracterização do esgoto sanitário (CP 4.10 · item 11 TR)")
    body_p(doc, "A anuência de saneamento atesta que o posto tem solução legal e técnica para "
           "o esgoto sanitário (banheiros, pias, vestiários), obrigatoriamente separado do "
           "efluente industrial (pista · SAO). Em zona rural (LAO item 1.2), SANEAGO tipicamente "
           "não atende a BR-060 no trecho do posto · solução esperada é fossa séptica + sumidouro. "
           "A visita de 12/05 (Eixo 3) confirma o cenário real em campo.",
           ["esgoto sanitário", "separado do efluente industrial", "fossa séptica + sumidouro"])

    heading_h4(doc, "Distinção fundamental · dois tipos de efluente que não podem se misturar")
    t = table_hdr(doc, ["Tipo", "Origem", "Destino obrigatório"], [3.0, 5.5, 8.5])
    rows_tipo = [
        ("Industrial", "Pista de abastecimento · troca de óleo · lavagem · descarga",
         "SAO (caixa separadora água-óleo) → rede pluvial / corpo d'água. Regulado pela CP 4.4 · monitorado no RCA semestral."),
        ("Sanitário", "Banheiros · pias · vestiários · lanchonete (se houver)",
         "Rede pública de esgoto OU fossa séptica + sumidouro · regulado pela CP 4.10 e item 11 do TR SEMMA."),
    ]
    for i, row in enumerate(rows_tipo):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0])
    p = doc.add_paragraph()
    styled_run(p, "Mistura cruzada entre os dois sistemas quebra a CP 4.10 e compromete "
               "simultaneamente o RCA (parâmetros sanitários não previstos em análise industrial) "
               "e a integridade da fossa (óleo rompe a colônia bacteriana anaeróbia).",
               size=9, color=CINZA_7, italic=True)
    doc.add_paragraph()

    heading_h4(doc, "Dois cenários para o Auto Posto América")
    t = table_hdr(doc, ["Cenário", "Probab.", "Ação", "Documento final", "Custo", "Prazo"],
                   [3.2, 1.8, 4.0, 4.0, 1.8, 2.2])
    rows_cenarios = [
        ("A · Rede pública (SANEAGO)", "Baixa",
         "Contato formal com SANEAGO · solicitar certidão de conexão + info da ETE",
         "Anuência em papel timbrado da concessionária",
         "Baixo/zero", "15-30 dias"),
        ("B · Fossa séptica + sumidouro", "Alta",
         "Laudo técnico com dimensionamento + memorial + fotos + ART",
         "Laudo NBR 7229/1993 + NBR 13969/1997",
         "R$ 3-8k", "5-10d pós-inspeção"),
    ]
    for i, row in enumerate(rows_cenarios):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0], mono_cols=[4, 5])
    doc.add_paragraph()

    heading_h4(doc, "O que o laudo precisa comprovar (Cenário B · fossa) · validado em 12/05")
    t = table_hdr(doc, ["Item", "Exigência normativa", "Como validar em campo"],
                   [3.8, 5.5, 7.7])
    rows_laudo = [
        ("Dimensionamento", "NBR 7229 · 1.000 L + 0,5 L/usuário/dia × nº usuários",
         "Medir fossa · estimar usuários (funcionários permanentes + clientes/dia + lanchonete)"),
        ("Distância mínima de poços", "≥ 15 m (NBR 7229)",
         "Medir até qualquer ponto de captação (poço, cisterna, mina)"),
        ("Distância mínima de cursos d'água", "≥ 30 m",
         "Medir até o córrego mais próximo · microbacia Ribeirão das Posses / Rio dos Bois"),
        ("Sumidouro/filtro anaeróbio a jusante", "NBR 13969 · tratamento complementar obrigatório",
         "Localizar estrutura · verificar funcionalidade (sem transbordamento, sem odor forte)"),
        ("Separação do efluente industrial (CP 4.10)", "Não pode receber lavagem, pátio, descarga",
         "Conferir tubulação · nenhuma conexão com SAO · Bloco A da entrevista"),
        ("Estado de manutenção", "Limpeza periódica · operação regular",
         "Entrevista operador · histórico de transbordamento · frequência de limpeza"),
    ]
    for i, row in enumerate(rows_laudo):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0])
    doc.add_paragraph()

    heading_h4(doc, "Três caminhos possíveis após a inspeção 12/05")
    t = table_hdr(doc, ["Caminho", "Descrição", "Entrega", "Impacto cronograma/CAPEX"],
                   [3.8, 6.0, 2.4, 4.8])
    rows_caminhos = [
        ("① Fossa conforme (MELHOR)",
         "Fossa existente dimensionada, distâncias atendidas, sumidouro funcional, sem conexão cruzada",
         "25/05/2026",
         "Sem impacto · laudo + ART · dentro do piso CAPEX."),
        ("② Ajuste simples (INTERMEDIÁRIO)",
         "Fossa existente mas com ajustes (limpeza profunda, sumidouro a construir, filtro anaeróbio a instalar)",
         "até 10/06/2026",
         "Impacto mínimo · Bauer/Limpmil · custo R$ 2-5k · cabe em F3."),
        ("③ Não-conforme ou ausente (PIOR)",
         "Sem fossa ou fora de norma (distâncias irregulares, sem sumidouro, dano estrutural)",
         "15-25/06/2026",
         "Projeto + obra · R$ 5-15k CAPEX · prazo apertado · pode exigir LI (CP 4.12)."),
    ]
    for i, row in enumerate(rows_caminhos):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0], mono_cols=[2])
    doc.add_paragraph()

    heading_h4(doc, "Riscos e cuidados operacionais")
    t = table_hdr(doc, ["Risco", "Causa", "Mitigação"], [4.5, 6.0, 6.5])
    rows_riscos_san = [
        ("① Fossa subdimensionada",
         "Posto ampliou operação (lanchonete, vestiário, mais funcionários) sem redimensionar a fossa",
         "Medir usuários reais × projetados na inspeção · se subdimensionada, Caminho ② ou ③."),
        ("② Distâncias irregulares",
         "Fossa próxima a poço (<15 m) ou curso d'água (<30 m)",
         "Relocar fossa · ou laudo geotécnico comprovando isolamento hidráulico · consultar SEMMA antes."),
        ("③ Conexão cruzada com SAO",
         "Tubulação mistura efluente da pista com sanitário (quebra CP 4.10)",
         "Vira Hip. C do Risco 02 (RCA reprovado) · ajuste hidráulico com ART · até 25/05."),
        ("④ Ausência de laudo anterior",
         "Nunca foi produzido laudo técnico · é primeira emissão",
         "Atestado técnico com ART da Hábilis · medição detalhada em campo · mais simples que anuência com concessionária."),
        ("⑤ Lanchonete/lavação/vestiário não declarados",
         "Operação real tem mais usuários sanitários que o projeto original considerou",
         "Entrevista estruturada (Bloco B do Eixo 5) captura realidade · dimensionamento revisado."),
    ]
    for i, row in enumerate(rows_riscos_san):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0])
    doc.add_paragraph()

    callout(doc, "Leitura prática · quem faz o quê e quando",
            "12/05 · Hábilis na inspeção (Eixo 3) · confirma tipo de solução (rede × fossa), mede, "
            "fotografa, valida distâncias e conexões. 15-20/05 · Hábilis no escritório · produz "
            "laudo técnico ou solicita anuência SANEAGO conforme o cenário. 25/05 · entrega para "
            "integração ao Relatório Consolidado (19/06). Se Caminho ② ou ③, cronograma estende "
            "até 10/06 ou 25/06 respectivamente. Todas as 3 saídas cabem dentro da janela F2/F3 "
            "sem empurrar o protocolo de 26/06.",
            kind="verde",
            bold_body_parts=["12/05", "15-20/05", "25/05", "19/06",
                             "sem empurrar o protocolo de 26/06"])

    # ===== §3.3.5 Plano de Atendimento a Emergências (PAE · CP 4.9) =====
    heading_h3(doc, "3.3.5 · Plano de Atendimento a Emergências (PAE · CP 4.9)")
    body_p(doc, "O PAE é a peça que RESPONDE quando os riscos mapeados pelo PGR se materializam. "
           "Exigido pela CP 4.9 da LAO, estrutura procedimentos operacionais, acionamento e "
           "comunicação em situações de vazamento, incêndio, derramamento e contaminação. Base "
           "normativa: ABNT NBR 14276/2020 (Brigada) + NBR 15219/2022 (Emergência contra Incêndio) "
           "+ NR-20 (Segurança com Inflamáveis). Data no cronograma: 05/06/2026 · F3 · Hábilis · "
           "RT Ambiental (CRBio).",
           ["CP 4.9 da LAO", "ABNT NBR 14276/2020", "NBR 15219/2022", "NR-20", "05/06/2026"])

    heading_h4(doc, "Estrutura obrigatória do PAE (9 capítulos)")
    t = table_hdr(doc, ["#", "Capítulo", "Conteúdo"], [1.0, 5.0, 11.0])
    rows_pae = [
        ("1", "Identificação", "Empreendimento · endereço · responsáveis · vigência · revisão"),
        ("2", "Análise preliminar de perigos (APP)", "Matriz probabilidade × severidade · alimentada pelo PGR 2024"),
        ("3", "Cenários de emergência", "Lista específica dos cenários esperados · não genérica"),
        ("4", "Matriz de resposta por cenário", "Procedimento operacional detalhado · fluxograma · ações nos primeiros 10/30/60 min"),
        ("5", "Recursos humanos", "Brigada NBR 14276 · responsabilidades · fluxo de acionamento"),
        ("6", "Recursos materiais", "Kit de contenção · EPIs · sistema de combate a incêndio (AVCB Cercon 161726) · sinalização"),
        ("7", "Comunicação externa", "Contatos · SEMMA · CBMGO · Defesa Civil · ANP · IBAMA · vizinhos"),
        ("8", "Treinamento e simulados", "Frequência · registro · avaliação · cruzamento com NR-20"),
        ("9", "Revisão e atualização", "Periodicidade anual · gatilhos de revisão"),
    ]
    for i, row in enumerate(rows_pae):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[1], mono_cols=[0])
    doc.add_paragraph()

    heading_h4(doc, "Cenários de emergência relevantes para o Auto Posto América")
    t = table_hdr(doc, ["Cenário", "Probab.", "Severidade", "Gatilho típico"],
                   [5.5, 2.5, 2.5, 6.5])
    rows_cenarios = [
        ("Pequeno derramamento na pista", "Alta (operacional)", "Baixa/média",
         "Falha no bico · transbordo do tanque do veículo"),
        ("Derramamento na descarga", "Média", "Média/alta",
         "Manobra do caminhão-tanque · falha de acoplamento"),
        ("Vazamento em tanque subterrâneo", "Média", "Crítica",
         "Corrosão · monitoramento intersticial falhou · relevante pelo RCA 2024"),
        ("Vazamento em tubulação subterrânea", "Média", "Crítica",
         "Corrosão · movimentação do solo · relacionado CP 4.8"),
        ("Incêndio no SASC ou pista", "Baixa", "Crítica",
         "Fagulha + vapor · AVCB Cercon 161726 é a primeira linha"),
        ("Explosão de tanque", "Muito baixa", "Catastrófica",
         "Acúmulo de vapor + ignição · cenário limite"),
        ("Contaminação de efluente (RCA)", "Alta (materializada 2024)", "Média",
         "Disfunção SAO · CP 4.4 · cruza com Risco 02"),
        ("Derramamento de óleo usado/resíduos", "Média", "Média",
         "Falha armazenamento · coleta Ecofenix atrasada"),
        ("Acidente rodoviário na BR-060 no acesso", "Alta (rural/rodovia)", "Média/alta",
         "Colisão no pátio · derrubada de bomba · Defesa Civil"),
        ("Contaminação do Ribeirão das Posses/Rio dos Bois", "Baixa/média", "Crítica",
         "Consequência propagada · microbacia registrada na LAO"),
    ]
    for i, row in enumerate(rows_cenarios):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0])
    doc.add_paragraph()

    heading_h4(doc, "Status do acervo atual × o que precisamos")
    t = table_hdr(doc, ["Peça", "Status", "Observação"], [6.0, 3.5, 7.5])
    rows_acervo_pae = [
        ("PGR 2024", "Em mãos",
         "Base técnica · APP alimenta cap. 2 · precisa revisão com cenários específicos."),
        ("AVCB Cercon 161726", "Vigente até 20/10/2026",
         "Sistema de combate a incêndio certificado · primeira linha · cap. 6."),
        ("Brigada + treinamento NR-20", "[VALIDAR em 12/05]",
         "Bloco D da entrevista · cruzamento Compliance · cap. 5 e 8."),
        ("Kit de contenção", "[VALIDAR em 12/05]",
         "Eixo 5 · material absorvente + sacos de areia + EPIs · cap. 6."),
        ("Sinalização + rotas de fuga", "[VALIDAR em 12/05]",
         "Eixo 5 · extintores · placas · cap. 6."),
        ("Registro de incidentes anteriores", "[VALIDAR em 12/05]",
         "Bloco C da entrevista · histórico · APP empírica."),
        ("Contatos oficiais atualizados", "A produzir",
         "SEMMA (Gleidson Nunes Ferreira) · CBMGO · Defesa Civil · ANP · IBAMA · vizinhos · cap. 7."),
    ]
    for i, row in enumerate(rows_acervo_pae):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0])
    doc.add_paragraph()

    heading_h4(doc, "Diferença PGR × PAE (e integração)")
    t = table_hdr(doc, ["Peça", "Função", "Relação"], [2.5, 7.0, 7.5])
    rows_pgr_pae = [
        ("PGR", "Mapeia riscos e define medidas preventivas (APP · matriz · controles).",
         "Funciona ANTES do evento. Alimenta cap. 2 do PAE. PGR 2024 em mãos serve de base · precisa revisão."),
        ("PAE", "Responde quando o risco se materializa (procedimento · acionamento · comunicação).",
         "Funciona DURANTE e APÓS o evento. Depende do PGR. Se PGR tem lacunas, PAE completa e força revisão."),
    ]
    for i, row in enumerate(rows_pgr_pae):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0])
    doc.add_paragraph()

    heading_h4(doc, "Insumos da inspeção 12/05 → capítulos do PAE")
    t = table_hdr(doc, ["Eixo", "Captura em campo", "Capítulo(s) do PAE"], [3.5, 6.5, 7.0])
    rows_pae_insumos = [
        ("Eixo 1 · Implantação", "Layout · rotas de fuga · distância do Ribeirão",
         "Cap. 2 APP · Cap. 5 Recursos humanos (pontos de encontro)"),
        ("Eixo 2 · Armazenamento + hídrico", "Tanques · tubulações · bombas · lava-jato/trocador",
         "Cap. 3 Cenários · Cap. 4 Matriz de resposta"),
        ("Eixo 3 · Efluentes/SAO", "SAO · drenagem · poços · disfunção RCA 2024",
         "Cap. 3 (contaminação) · Cap. 4 · cruza com Risco 02"),
        ("Eixo 4 · Resíduos/passivo", "Armazenamento · pontos críticos",
         "Cap. 3 (derramamento) · Cap. 6 (kit dedicado)"),
        ("Eixo 5 · Segurança e rotina", "AVCB · brigada · kit · NR-20 · sinalização",
         "Cap. 5 · Cap. 6 · Cap. 8"),
    ]
    for i, row in enumerate(rows_pae_insumos):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0])
    doc.add_paragraph()

    heading_h4(doc, "Riscos específicos do PAE")
    t = table_hdr(doc, ["Risco", "Causa", "Mitigação"], [5.0, 6.0, 6.0])
    rows_riscos_pae = [
        ("① PAE genérico/copiado",
         "Plano padrão não reflete posto específico (sem microbacia, sem BR-060, sem lava-jato/trocador)",
         "Cenários construídos a partir da inspeção 12/05 · 10 cenários específicos · APP empírica do PGR 2024."),
        ("② Cenários faltando",
         "Não contempla acidente rodoviário BR-060, lava-jato, trocador, microbacia próxima",
         "Tabela de cenários (acima) · 4 cenários de severidade crítica priorizados."),
        ("③ Sem registro de treinamento NR-20",
         "Brigada não treinada · CP 4.9 exige",
         "Bloco D da entrevista 12/05 · se ausente, capacitação até 31/05 (antes do PAE 05/06)."),
        ("④ Desalinhamento com PGR 2024",
         "PAE e PGR com cenários/recursos diferentes",
         "PAE construído SOBRE o PGR 2024 · se PGR incompleto, revisão anexa · consolidação 19/06."),
        ("⑤ Contatos desatualizados",
         "Telefones de Bombeiros, SEMMA, Defesa Civil, ANP fora de data",
         "Cap. 7 com lista validada em maio/2026 · Gleidson Nunes Ferreira (Decreto 013/2025) · revisão anual."),
        ("⑥ Kit de contenção apenas listado, não comprovado",
         "PAE afirma 'temos kit' mas inspeção revela ausência/validade vencida",
         "Eixo 5 valida fisicamente · se ausente, reposição antes de 05/06 · custo R$ 500-2k (piso)."),
    ]
    for i, row in enumerate(rows_riscos_pae):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0])
    doc.add_paragraph()

    callout(doc, "Leitura prática · o PAE como ponte entre prevenção e resposta",
            "O PAE depende de 4 insumos que precisam estar validados antes de 05/06: PGR 2024 "
            "revisado (base · cap. 2), inspeção 12/05 (cenários · recursos · brigada · kit), "
            "contatos oficiais atualizados (cap. 7) e registros de treinamento NR-20 (cap. 8). "
            "Uma vez entregue, o PAE cruza com o Risco 02 (Hip. B · obra estrutural da SAO pode "
            "disparar cenário de emergência ambiental) e com o Compliance Paralelo (NR-20 · "
            "treinamento contínuo). É peça viva: revisão anual + após incidente real + após "
            "mudança operacional.",
            kind="verde",
            bold_body_parts=["4 insumos", "05/06", "Risco 02", "Hip. B",
                             "Compliance Paralelo", "peça viva"])

    callout(doc, "Conclusão objetiva · status do RCA 2026",
            "O RCA não está consolidado e nem poderia estar ainda — a ordem técnica obriga a "
            "sequência diagnóstico → adequação → coleta → laudo. Em 21/04/2026, estamos no marco "
            "zero: as ações da próxima semana (procuração 23/04 · CNPJ 24/04 · contratação "
            "investigação 05/05) destravam a visita 12/05, que é o ponto em que a consolidação "
            "começa de fato. Cronograma até 03/06 é factível se Hip. A ou C. Se Hip. B (obra "
            "estrutural), RCA consolidado escorrega para depois do protocolo e o dossiê é "
            "submetido com plano corretivo em anexo — cenário transparente, não inviabilizante.",
            kind="verde",
            bold_body_parts=["não está consolidado", "marco zero", "12/05", "Hip. B",
                             "plano corretivo em anexo"])

    callout(doc, "Alerta técnico · RCA 2024 reprovado",
            "A análise qualitativa de efluentes 2024 registra não conformidade em Óleos e Graxas "
            "(64 mg/L · limite 20 mg/L) e Turbidez. Indica disfunção do Sistema de Tratamento de "
            "Efluentes Oleosos. Se o novo RCA (03/06) for coletado sem adequação prévia do sistema, "
            "há risco de nova reprovação. Recomenda-se auditoria do SAO durante a inspeção de "
            "campo (12/05) e eventual adequação física antes da coleta do novo RCA.",
            kind="vermelho",
            bold_body_parts=["Óleos e Graxas (64 mg/L · limite 20 mg/L)", "auditoria do SAO"])

    heading_h3(doc, "3.4 · Compliance Paralelo — frente paralela permanente")
    body_p(doc, "23 itens em frentes regulatórias correlatas: ANP (cadastro e declarações), "
           "IBAMA (CTF/APP + RAPP), NR-20 (PGR + capacitação), PNRS (MTR) e monitoramentos "
           "contínuos. Cadência própria, independente do cronograma da renovação.",
           ["23 itens", "Cadência própria"])

    # ===== §04 CRONOGRAMA =====
    heading_section(doc, "04", "Cronograma Operacional")
    body_p(doc, "Quatro fases sequenciais dentro da janela de execução de 67 dias corridos "
           "(20/04/2026 a 26/06/2026). O protocolo-meta em 26/06 corresponde a 180 dias antes "
           "do vencimento da LAO.",
           ["67 dias corridos", "20/04/2026 a 26/06/2026", "180 dias antes"])

    t = table_hdr(doc, ["Fase", "Período", "Dias", "Objetivo e entregáveis"], [3.0, 3.0, 1.5, 10.0])
    fases = [
        ("F1 · Estruturação", "20/04 → 08/05", "19",
         "Consolidar uso do solo e estruturar base técnica. Entregáveis: verificação da certidão, verificação de outorga, requerimento municipal, contratação da investigação confirmatória, análise documental."),
        ("F2 · Consolidação", "09/05 → 30/05", "22",
         "Validar realidade física e repor base técnica. Entregáveis: inspeção de campo (12/05), protocolo Veredas (11/05), novas análises água e efluentes, mobilização da investigação, MCO atualizado, Plano das 28 CPs."),
        ("F3 · Execução e Evidência", "31/05 → 19/06", "20",
         "Executar atendimento das condicionantes e produzir relatório. Entregáveis: PGRS, PAE, Análise de Risco, teste de estanqueidade (08/06), laudo investigação (15/06), relatório consolidado (19/06)."),
        ("F4 · Protocolo", "20/06 → 26/06", "7",
         "Aprovar dossiê, recolher taxas e protocolar no SEMMA. Entregáveis: aprovação do relatório (22/06), CNDs (23/06), checklist final (24/06), taxas (25/06), protocolo (26/06)."),
    ]
    for i, row in enumerate(fases):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0], mono_cols=[1, 2])
    doc.add_paragraph()

    callout(doc, "Cenário contingente · passivo ambiental",
            "Caso o Laudo de Investigação Confirmatória (15/06) identifique solo ou água "
            "subterrânea com concentrações acima dos valores de investigação (CONAMA 420/2009), "
            "abre-se escopo contingente não coberto pela janela de 67 dias: Investigação Detalhada "
            "(NBR 16210 · 60-90 dias) + Plano de Intervenção (NBR 16784 · 3-18 meses) + "
            "Monitoramento pós-remediação. A Hábilis apresentaria plano revisado em 10 dias úteis, "
            "preservando a margem até 23/12 para protocolo emergencial.",
            kind="vermelho",
            bold_body_parts=["Laudo de Investigação Confirmatória (15/06)", "10 dias úteis"])

    callout(doc, "Pontos de atenção · riscos mapeados · outorga via Veredas",
            "IN 15/2026 SEMAD (DOE/GO 17/04/2026) prorrogou o prazo de transição Web Outorga → "
            "Veredas para 15/07/2026, com §14 (pedidos até 16/04: ordem de origem) e §15 (após: "
            "fila pós-formalização). Auto Posto América cai no §15. Sequência invertida: inspeção "
            "12/05 precede o protocolo, que passa para 13/05, para que o levantamento hídrico "
            "defina o regime do pedido (dispensa simplificada vs. processo completo) antes da "
            "formalização. Aproveita a mesma viagem a Guapó e produz memorial técnico com dados "
            "frescos. Os 2 dias de atraso na fila §15 são ruído frente aos 60-120 dias de análise. "
            "Protocolo em 13/05 gera certidão de tramitação suficiente para o dossiê da LAO; "
            "análise SEMAD corre em paralelo sem bloquear 26/06.",
            kind="amber",
            bold_body_parts=["IN 15/2026 SEMAD", "15/07/2026", "§15", "12/05", "13/05",
                             "dispensa simplificada vs. processo completo"])

    # ===== §05 CHECKLIST =====
    add_page_break(doc)
    heading_section(doc, "05", "Checklist Cronológico · 40 itens")
    body_p(doc, "Itens ordenados pela data-alvo. Coluna Responsável identifica a parte executora "
           "(Cliente · Hábilis · Terceiro). Detalhamento dos 40 itens abaixo.", ["40 itens"])

    t = table_hdr(doc, ["Prazo", "Item", "Resp.", "Fase", "Observação"], [1.8, 5.3, 1.5, 1.0, 7.9])
    checklist = [
        ("22/04", "Coleta de RG/CPF dos sócios (Sérgio Marques, José Guilherme)", "Cliente", "F1",
         "Exigência dos dois processos. Não localizado no acervo histórico."),
        ("22/04", "Validação do AVCB · Cercon 161726", "Hábilis", "F1",
         "Válido até 20/10/2026. Em mãos."),
        ("23/04", "Procuração pública para a Hábilis (SEMMA/Prefeitura/SEMAD/Veredas)", "Cliente", "F1",
         "Histórica é em nome da A3 Engenharia. Precisa substituir."),
        ("24/04", "Cartão CNPJ atualizado + comprovante de endereço", "Cliente", "F1",
         "Emissão Receita Federal · comprovantes dos últimos 3 meses."),
        ("25/04", "Verificação formal da Certidão de Uso do Solo", "Cliente", "F1",
         "Confirmar existência, número, data. Apoio da Hábilis."),
        ("27/04", "Enquadramento do imóvel + CAR (confirmado rural conforme LAO item 1.2)", "Hábilis", "F1",
         "Imóvel classificado como Zona Rural (LAO 033/2020 item 1.2). Resolve: (1) CAR obrigatório · validar inscrição; (2) ART Uso do Solo via RT Agronômica (atribuição nata · Res. CONFEA 218/73)."),
        ("27/04", "ART específica para Uso do Solo", "Hábilis · RT Agronômica", "F1",
         "ART pela Hábilis via RT Engenharia Agronômica · atribuição nata em zona rural (Res. CONFEA 218/73 Art. 5º alínea c)."),
        ("28/04", "Verificação e diagnóstico de outorga no Sistema Veredas", "Hábilis", "F1",
         "DURH002737 venceu 09/12/2023. IN 15/2026 SEMAD prorrogou transição para 15/07/2026; pedido pós-16/04 segue §15 (fila após formalização, sem ordem cronológica de origem)."),
        ("30/04", "Protocolo do requerimento da Certidão (condicional)", "Hábilis", "F1",
         "Somente se verificação indicar ausência ou desatualização."),
        ("30/04", "Consulta CREA/GO · atribuição para MCO e Anexos CONAMA 273", "Hábilis", "F1",
         "Parecer formal do CREA/GO sobre atribuição Eng. Agrônomo para posto urbano. Resposta 3-7d. Define se MCO/Anexos são assinados internamente (RT Agronômica) ou se contrata Terceiro Eng. Civil/Ambiental. Impacto potencial: +R$ 8-15k CAPEX e +5-15d prazo ART."),
        ("04/05", "Publicação do pedido de licenciamento (CONAMA 006/1986)", "Hábilis", "F1",
         "Jornal local + Diário Oficial. 15-30 dias para contestação pública."),
        ("05/05", "Contratação da empresa de investigação confirmatória", "Cliente", "F1",
         "NBR 15515-2/15515-3 e NBR 16209. ART do geólogo/eng. ambiental."),
        ("08/05", "Análise documental consolidada", "Hábilis", "F1",
         "Fechamento da triagem do lote 14/04 e identificação de lacunas."),
        ("10/05", "Mapeamento das 8 condicionantes pendentes", "Hábilis", "F1",
         "Valida se o prazo 15/06 é factível ou exige reabertura."),
        ("12/05", "Inspeção técnica de campo · 20 pontos em 5 eixos", "Hábilis", "F2",
         "Conteúdo programático detalhado na subseção 3.2.1."),
        ("12/05", "Validação das plantas em campo", "Hábilis", "F2",
         "Quadro de Áreas, UTM SIRGAS 2000, APP, Reserva Legal."),
        ("12/05", "Levantamento hídrico · define regime da outorga", "Hábilis", "F2",
         "Eixos 1-2: identifica captação direta de recurso hídrico. Sem captação → dispensa simplificada; com captação (poço/mina/cisterna) → processo completo com memorial técnico e provável visita SEMAD."),
        ("13/05", "Protocolo novo de outorga via Sistema Veredas", "Hábilis", "F2",
         "Dia útil seguinte à inspeção de 12/05, aproveitando dados frescos para o memorial. Regime (dispensa ou completo) definido pelo levantamento hídrico. Dentro do prazo 15/07/2026 (IN 15/2026 SEMAD). §15: fila pós-formalização, sem ordem de origem. Gera certidão de tramitação para o dossiê LAO. SEMAD analisa mérito em 60-120d."),
        ("15/05", "Contratação de laboratório, estanqueidade e projetista", "Cliente", "F2",
         "Indicações da Hábilis. Investigação confirmatória já contratada em 05/05."),
        ("18/05", "NFs e certificados dos equipamentos (tanques, tubulações, detecção)", "Cliente", "F2",
         "CONAMA 319/2002 + Portarias INMETRO 37·109·110·111·009/2005."),
        ("20/05", "Contratos de manutenção, limpeza e destinação (Ecofenix, Bauer, Limpmil)", "Cliente", "F2",
         "Validar vigência do licenciamento e certificação ANP antes de anexar."),
        ("22/05", "Nova análise de efluentes", "Terceiro · Lab", "F2",
         "Hábilis valida o laudo. ART do químico."),
        ("22/05", "Nova análise de água", "Terceiro · Lab", "F2",
         "Conjugada com efluentes. ART do químico."),
        ("25/05", "Anuência de saneamento ou laudo de fossa (condicional)", "Hábilis", "F2",
         "Conforme diagnóstico de 12/05 (item 11 TR)."),
        ("25/05", "Mobilização da investigação confirmatória (3 furos mín.)", "Terceiro · Invest.", "F2",
         "VOC, BTEX, PAH, metais, TPH. CEMAM 029/2018 + CONAMA 420/2009."),
        ("28/05", "Atualização do MCO · Anexo IX CEMAM 029/2018", "Hábilis ou Terceiro", "F2",
         "Posterior à inspeção. Atribuição ART depende do parecer CREA/GO de 30/04: RT Engenharia Agronômica se aceito, senão Terceiro Eng. Civil/Ambiental contratado. Base de dados: Eixo 2 da inspeção (ficha tanques + tubulações subterrâneas)."),
        ("29/05", "Anexos I e II da Resolução CONAMA 273/2000 (+ART)", "Hábilis ou Terceiro", "F2",
         "Item 23 TR. Mesma lógica de atribuição do MCO (validação CREA/GO pendente)."),
        ("30/05", "Plano de atendimento das 28 condicionantes", "Hábilis", "F2",
         "Responsável, evidência e prazo por condicionante."),
        ("02/06", "PGRS · Plano de Gerenciamento de Resíduos", "Hábilis", "F3",
         "Caracterização, rotas de destinação e responsáveis. ART via RT Ambiental (CRBio) ou RT Agronômica (CREA) — ambas compatíveis."),
        ("03/06", "RCA semestral · Anexo XI CEMAM 029/2018", "Hábilis", "F3",
         "Análises SAO: pH, turbidez, O&G, sólidos, DBO, DQO, MBAS, metais, TPH, BTEX/PAH. ART via RT Ambiental (CRBio) — atribuição nata Lei 6.684/79."),
        ("05/06", "PAE · Plano de Atendimento a Emergências", "Hábilis", "F3",
         "Cenários (vazamentos, incêndios, derramamento). ART via RT Ambiental (CRBio); acopla Eng. Segurança se SEMMA exigir escopo NR-20/NR-23."),
        ("Conting.", "Teste de estanqueidade · contingência (Opção A)", "Terceiro · Credenciada", "F3",
         "Substituído pelo laudo 10/2024 no dossiê (plano principal · §3.3.3). Acionado apenas se SEMMA contestar o laudo (INMETRO NCC 21.08786 vencido 29/04/2024 · calibração Suporty vencida 03/2023). Prazo 30-60d pós-exigência. Próximo ciclo regular (Opção B) em 10/2026."),
        ("10/06", "Análise de Risco (condicional · contingência)", "Terceiro · Eng. Químico/Ambiental", "F3",
         "Item mantido no plano como contingência armada, não como entrega obrigatória. Peça SEM citação literal na LAO 033/2020 ou TR SEMMA (ver nota de transparência no §6.1). Aplicável se exigida pela SEMMA em resposta à triagem (via itens 3.7 e 4.16 LAO). APP 15-30d R$ 5-15k · ARQ 60-90d R$ 30-80k. Indicação Hábilis · contratação Cliente."),
        ("12/06", "Recomposição das 10 condicionantes parciais", "Hábilis", "F3",
         "Recompostas documentalmente conforme Plano de Atendimento de 30/05 · das 28 CPs totais (6 atendidas / 10 parciais / 2 pendentes / 10 regras declaratórias)."),
        ("15/06", "Laudo de Investigação Confirmatória · MARCO DECISIVO", "Terceiro · Invest.", "F3",
         "Dispara cenário contingente se P-01 positivo (CONAMA 420/09)."),
        ("15/06", "Execução das 8 condicionantes pendentes", "Terceiro · Contratados", "F3",
         "Serviços físicos ou ensaios."),
        ("19/06", "Relatório técnico consolidado", "Hábilis", "F3",
         "Peça central. ART do RT coordenador da Hábilis (atribuição segue definição CREA/GO de 30/04). Laudos específicos mantêm ART dos RTs próprios."),
        ("22/06", "Aprovação formal do relatório técnico", "Cliente", "F4",
         "Aprovação escrita antes do protocolo."),
        ("23/06", "CND Municipal (Guapó) + CND Estadual (SEFAZ/GO)", "Cliente", "F4",
         "Validade típica 90 dias. Emitir próximo ao protocolo."),
        ("24/06", "Checklist final cruzado", "Hábilis", "F4",
         "Conferência de peças, assinaturas, ARTs e taxas."),
        ("25/06", "Pagamento das taxas de protocolo", "Cliente", "F4",
         "SEMMA, municipais e demais."),
        ("26/06", "PROTOCOLO no Processo SEMMA 22760/2020", "Hábilis", "F4",
         "Submissão do dossiê completo. Marco final."),
    ]
    for i, row in enumerate(checklist):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[1], mono_cols=[0, 3])
    doc.add_paragraph()

    # ===== §06 RESPONSABILIDADES =====
    heading_section(doc, "06", "Responsabilidades")
    body_p(doc, "Atribuição operacional distribuída em quatro camadas. Alinha quem responde por "
           "cada entrega em qualquer ponto do cronograma.")

    resp_blocks = [
        ("1 · Hábilis · execução técnica direta", VERDE_BG, "1B5E20",
         ["Análise documental e gestão do checklist",
          "Requerimento e acompanhamento da Certidão de Uso do Solo",
          "Inspeção técnica de campo e parecer interno",
          "Atualização do MCO",
          "Plano de atendimento das 28 condicionantes",
          "Relatório técnico consolidado",
          "Montagem do dossiê e protocolo no Processo SEMMA 22760/2020",
          "Interlocução com SEMMA e acompanhamento pós-protocolo"]),
        ("2 · Cliente (Grupo Z+Z) · decisão e liberação", AZUL_BG, "0D47A1",
         ["Confirmação formal da certidão de uso do solo existente",
          "Autorização e pagamento da investigação confirmatória",
          "Contratação de laboratório, estanqueidade e projetista",
          "Assinaturas em requerimentos, procurações e formulários",
          "Pagamento de taxas municipais, SEMMA, laboratórios e terceiros",
          "Aprovação formal do relatório técnico antes do protocolo",
          "Retorno em até 5 dias úteis às solicitações formais da Hábilis"]),
        ("3 · Terceiros · execução obrigatória", CINZA_BG, "6C6C70",
         ["SEMMA Guapó, órgão licenciador",
          "Prefeitura de Guapó, Certidão de Uso e Ocupação",
          "SEMARH/GO, dispensa de outorga",
          "RTs executores (CREA/CAU), ART por peça técnica",
          "Laboratórios acreditados, análises de água e efluentes",
          "Empresa credenciada, teste de estanqueidade",
          "Empresa especializada, investigação ambiental confirmatória",
          "Projetista, revisão de plantas (se aplicável)"]),
        ("4 · Hábilis · consultoria e validação", AMBER_BG, "F57F17",
         ["Indicação de laboratórios e profissionais de mercado",
          "Definição do escopo técnico dos serviços de terceiros",
          "Validação técnica de laudos, ARTs e entregáveis terceirizados",
          "Orientação em pendências técnicas complexas",
          "Interface com terceiros para aderência ao dossiê da SEMMA"]),
    ]
    for titulo, bg, border, items in resp_blocks:
        heading_h3(doc, titulo)
        t = doc.add_table(rows=1, cols=1)
        c = t.cell(0, 0)
        set_cell_shading(c, bg)
        set_cell_border(c, left=(18, border))
        for item in items:
            pi = c.add_paragraph()
            pi.paragraph_format.space_after = Pt(1)
            pi.paragraph_format.left_indent = Cm(0.2)
            styled_run(pi, "→ ", size=10, color=VERDE, bold=True, mono=True)
            styled_run(pi, item, size=10, color=CINZA_7)
        c.paragraphs[0]._p.getparent().remove(c.paragraphs[0]._p)

    callout(doc, "Delimitação da atuação da Hábilis",
            "A Hábilis executa apenas atividades dentro da sua capacidade técnica direta. "
            "Serviços que exigem profissionais específicos não são executados pela Hábilis, "
            "cabendo ao cliente sua contratação. A Hábilis atua na orientação e validação técnica.",
            kind="amber",
            bold_body_parts=["capacidade técnica direta", "orientação e validação técnica"])

    # ===== §07 CAPEX =====
    heading_section(doc, "07", "Investimento · CAPEX em faixa")
    lede(doc, "O investimento total é apresentado como faixa, não número único, porque o "
         "principal fator de variação — passivo ambiental no solo — só se revela com o Laudo "
         "de Investigação Confirmatória em 15/06/2026, 11 dias antes do protocolo-meta. Piso e "
         "teto refletem, respectivamente, cenário sem passivo detectado e cenário com investigação "
         "detalhada disparada.",
         ["faixa, não número único", "15/06/2026", "Piso e teto"])

    tc = doc.add_table(rows=1, cols=2)
    c = tc.cell(0, 0); set_cell_shading(c, CINZA_BG); set_cell_border(c, left=(18, "1B5E20"))
    styled_run(c.paragraphs[0], "PISO · CENÁRIO LIMPO", size=8, color=CINZA_5, bold=True, mono=True)
    styled_run(c.add_paragraph(), "R$ 1.398k", size=20, color=VERDE, bold=True, mono=True)
    styled_run(c.add_paragraph(),
               "Laudo confirmatório negativo. RCA dentro do limite após limpeza do SAO. Sem "
               "alvará sanitário adicional. Cobre taxas oficiais, laboratórios, estanqueidade "
               "(contingência), investigação 1ª fase, publicação CONAMA 006/86 e ajustes "
               "operacionais. Honorários Hábilis em contrato separado · fora deste CAPEX.",
               size=9, color=CINZA_7)

    c = tc.cell(0, 1); set_cell_shading(c, CINZA_BG); set_cell_border(c, left=(18, "B71C1C"))
    styled_run(c.paragraphs[0], "TETO · CENÁRIO CONTINGENTE", size=8, color=CINZA_5, bold=True, mono=True)
    styled_run(c.add_paragraph(), "R$ 2.567k", size=20, color=VERMELHO, bold=True, mono=True)
    styled_run(c.add_paragraph(),
               "Laudo dispara Investigação Detalhada (NBR 16210 · 60-90d) + Plano de Intervenção "
               "(NBR 16784 · 3-18 meses). Obra SAO. Monitoramento pós-remediação. Diferença: R$ 1.169k.",
               size=9, color=CINZA_7)
    doc.add_paragraph()

    heading_h3(doc, "7.1 · Composição detalhada do CAPEX piso (sem honorários Hábilis)")
    body_p(doc, "O piso cobre 4 categorias. Honorários Hábilis seguem em contrato separado, "
           "não computados aqui. Valores são estimativas de mercado 2026 · cotações formais "
           "no §7.3.",
           ["4 categorias", "contrato separado"])

    t = table_hdr(doc, ["Categoria", "Faixa estimada", "Composição"], [4.5, 3.0, 9.5])
    rows_cat = [
        ("① Taxas e emolumentos oficiais", "R$ 40k-80k",
         "DUAM SEMMA · taxa Renovação LAO · Prefeitura (Certidão Uso do Solo) · SEMAD/GO (Veredas) · CND Municipal + Estadual · cartório (procuração + firmas)."),
        ("② Publicações CONAMA 006/86", "R$ 1,2k-2,8k",
         "DOE/GO + jornal de grande circulação (O Popular · Diário da Manhã ou regional). Detalhamento no §3.2.2."),
        ("③ Serviços de terceiros", "R$ 180k-450k",
         "Investigação confirmatória · laboratórios · manutenção SAO (Bauer/Limpmil/Ecofenix) · projetista · estanqueidade (contingência) · NR-20. Abertura em §7.2."),
        ("④ Ajustes operacionais + reserva técnica", "R$ 50k-250k",
         "Limpeza profunda SAO · kit de contenção · sinalização · poços de monitoramento · ajustes hidráulicos leves · pequenas adequações."),
        ("Total piso · faixa estimada", "R$ 271k-782k",
         "Valor de referência da Hábilis: R$ 1.398k (inclui margem para adequações não mapeadas + cotações formais ainda a obter)."),
    ]
    for i, row in enumerate(rows_cat):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0], mono_cols=[1])
    doc.add_paragraph()

    callout(doc, "Observação sobre a diferença entre faixa estimada e valor de referência",
            "A faixa estimada (R$ 271k-782k) baseia-se em valores de mercado 2026 sem cotações "
            "formais. O valor de referência da Hábilis (R$ 1.398k) preserva margem para: "
            "(a) adequações que só aparecem na inspeção 12/05, (b) cotações específicas do "
            "interior de Goiás (sobretaxa rural BR-060 km 203), (c) contratações emergenciais. "
            "As cotações formais do §7.2 refinam essa margem ao longo de F1-F2.",
            kind="amber",
            bold_body_parts=["R$ 271k-782k", "R$ 1.398k", "refinam essa margem"])

    heading_h3(doc, "Itens que mais movem a faixa")
    for i, item in enumerate([
        "Investigação detalhada, disparada se P-01 der acima dos valores de investigação em "
        "qualquer dos 3 furos. Abre 60-90 dias adicionais.",
        "Obra de adequação do SAO, se a nova análise de efluentes (22/05) não normalizar após "
        "limpeza Bauer/Limpmil. Pode extrapolar F3.",
        "Análise de Risco nível APP/ARQ, se exigida pela SEMMA. Passa a terceiro especializado "
        "com prazo dilatado e custo dedicado.",
    ], 1):
        p = doc.add_paragraph(style="List Number")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        styled_run(p, item, size=10.5, color=CINZA_7)

    callout(doc, "Delimitação do CAPEX · dentro e fora",
            "DENTRO: taxas SEMMA/Prefeitura/SEMAD/cartório · publicação CONAMA 006/86 · "
            "laboratórios · estanqueidade (contingência) · investigação 1ª fase · ajustes "
            "operacionais leves. FORA: investigação detalhada (se passivo confirmado) · "
            "remediação · obra do SAO (Hip. B Risco 02) · passivos Compliance Paralelo · "
            "alvará sanitário · Análise de Risco. TAMBÉM FORA: honorários Hábilis em contrato "
            "separado, não computados neste CAPEX.",
            kind="verde", bold_body_parts=["DENTRO:", "FORA:", "TAMBÉM FORA:"])

    heading_h3(doc, "7.2 · Cotações a serem indicadas pela Hábilis")
    body_p(doc, "Para cada serviço terceirizado, a Hábilis indica ao Grupo Z+Z 2-3 fornecedores "
           "com cotação formal · o Cliente decide a contratação. Faixas de preço são estimativas "
           "de mercado 2026.",
           ["2-3 fornecedores"])

    t = table_hdr(doc, ["#", "Serviço", "Escopo", "Norma", "Cotações", "Faixa", "Observação"],
                   [0.6, 2.6, 3.5, 2.2, 1.6, 1.8, 4.7])
    rows_cotacao = [
        ("A", "Investigação Ambiental Confirmatória",
         "3 furos · sondagens · amostragem solo e água · VOC/BTEX/PAH/metais/TPH",
         "NBR 15515-2/3 · CONAMA 420/09",
         "2-3 empresas", "R$ 80k-180k",
         "ART do geólogo ou eng. ambiental · experiência em postos"),
        ("B", "Laboratório acreditado para RCA",
         "Análise efluente + água poços · parâmetros CP 4.4",
         "INMETRO · Anexo XI CEMAM 029",
         "2-3 laboratórios", "R$ 8k-18k",
         "Acreditação INMETRO vigente · ART químico"),
        ("C", "Teste de Estanqueidade (contingência Opção A)",
         "Ensaio NBR 16795/2019 · ART eng. mecânico",
         "INMETRO + ISO/IEC 17025",
         "2-3 empresas", "R$ 15k-30k",
         "Pré-negociado até 10/05 · acionado só se SEMMA contestar 10/2024"),
        ("D", "Manutenção SAO + coleta resíduos",
         "Limpeza SAO · coleta areias oleosas · borra · óleo",
         "CONAMA 362/05 · CP 4.5/4.6",
         "Bauer + Limpmil + Ecofenix", "R$ 15k-40k",
         "Validar licenciamento + certificação ANP · cotação atualizada"),
        ("E", "Projetista · revisão de plantas",
         "Plantas + shapefile + .dwg · Quadro UTM (CP 3.12)",
         "ABNT · SIRGAS 2000",
         "2-3 profissionais", "R$ 10k-25k",
         "ART eng. civil/arquiteto · até F2 antes do MCO 28/05"),
        ("F", "RT Externo Eng. Civil/Ambiental (contingência)",
         "ART MCO + Anexos CONAMA 273 se CREA restringir agronômica",
         "Res. CONFEA 218/73 + 1010/05",
         "2-3 parceiros", "R$ 8k-15k + ART",
         "Pré-negociado 05/05 · acionamento 48h se CREA restritivo"),
        ("G", "Capacitação Brigada NR-20",
         "Treinamento + registro · integra PAE cap. 8",
         "NR-20 · NBR 14276/2020",
         "1-2 capacitadores", "R$ 3k-8k",
         "Bloco D entrevista 12/05 · entrega até 31/05"),
        ("H", "Eng. Químico/Ambiental (contingência AR)",
         "Análise de Risco APP (15-30d) ou ARQ (60-90d)",
         "ABNT · prática SEMMA",
         "1-2 especialistas", "APP R$ 5-15k · ARQ R$ 30-80k",
         "Peça sem citação literal na LAO · só se SEMMA exigir"),
        ("I", "Adequação SAO (contingência Hip. B)",
         "Obra civil · reforma/substituição SAO",
         "NBR 14605 + CONAMA 273",
         "2-3 empresas de obra", "R$ 30k-150k",
         "Acionamento 15/05 se Hip. B · pode exigir LI (CP 4.12)"),
        ("J", "Poços de monitoramento (se ausentes)",
         "2-3 poços · CP 4.4 (BTEX + PAH)",
         "NBR 15495 + CEMAM 029",
         "1-2 empresas", "R$ 20k-60k",
         "Validar em 12/05 · se ausentes, antes do RCA 22/05"),
    ]
    for i, row in enumerate(rows_cotacao):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[1], mono_cols=[0, 4, 5])
    doc.add_paragraph()

    callout(doc, "Orientação de boa prática · 3 cotações por serviço",
            "A Hábilis indica ao Grupo Z+Z no mínimo 3 fornecedores para cada serviço acima "
            "(exceto contingências com disponibilidade restrita). Cotações apresentadas em "
            "formato comparativo (escopo · preço · prazo · certificações · experiência) "
            "para decisão com critério objetivo. Honorários Hábilis pela coordenação das "
            "cotações estão em contrato separado.",
            kind="verde",
            bold_body_parts=["3 fornecedores", "formato comparativo", "contrato separado"])

    heading_h3(doc, "7.3 · Cronograma de contratações")
    body_p(doc, "Calendário dos gatilhos · cada data representa o prazo máximo para cotações em "
           "mãos do Cliente. Contratação efetiva no dia útil seguinte.")

    t = table_hdr(doc, ["Data-limite", "Ação", "Responsável", "Insumo/Dependência"],
                   [2.5, 6.0, 3.0, 5.5])
    rows_cron = [
        ("até 05/05", "Cotações Investigação Confirmatória (A) + lista RTs externos (F)",
         "Hábilis · indicação", "Pré-negociação em F1 · 2-3 cada"),
        ("05/05", "Contratação Investigação Confirmatória",
         "Cliente · decisão", "Depende cotações Hábilis"),
        ("até 10/05", "Lista empresas estanqueidade Opção A (C) + RT externo MCO (F)",
         "Hábilis · indicação", "Só se gatilhos (Risco 07/08) se materializarem"),
        ("até 15/05", "Cotações laboratório RCA (B) + projetista (E)",
         "Hábilis · indicação", "Antes da inspeção 12/05"),
        ("15/05", "Contratação laboratório + projetista",
         "Cliente · decisão", "Mobilização RCA 22/05 e MCO 28/05"),
        ("até 20/05", "Validação contratos Bauer/Limpmil/Ecofenix (D)",
         "Cliente", "Verificar licenciamento + ANP"),
        ("até 22/05", "Contratação poços de monitoramento (J · se ausentes)",
         "Cliente · decisão", "Depende Eixo 3 inspeção 12/05"),
        ("até 31/05", "Capacitação Brigada NR-20 (G · se necessário)",
         "Cliente · decisão", "Bloco D entrevista · integra PAE 05/06"),
        ("Contingente", "Obra SAO (I) · AR (H) · RT externo MCO (F)",
         "Cliente · decisão", "Acionamento 48h após materialização do gatilho"),
    ]
    for i, row in enumerate(rows_cron):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[1], mono_cols=[0])
    doc.add_paragraph()

    callout(doc, "Ponto crítico · lista de cotações da Hábilis em 05/05",
            "O grande marco de contratação é 05/05/2026 · data em que o Cliente contrata a "
            "Investigação Confirmatória. Para viabilizar a decisão, a Hábilis deve entregar "
            "até 30/04 o dossiê com 3 cotações comparativas (escopo · preço · prazo · "
            "certificações · experiência). Mesmo fluxo em 15/05 para laboratório + projetista.",
            kind="amber",
            bold_body_parts=["05/05/2026", "até 30/04", "3 cotações comparativas"])

    # ===== §08 RISCOS =====
    heading_section(doc, "08", "Riscos e Contingências")
    lede(doc, "Oito cenários podem empurrar o protocolo ou abrir passivo pós-protocolo. Todos "
         "previsíveis, todos com gatilho de contingência definido. Identificar antes permite "
         "ação nas primeiras 48 horas do evento em vez das últimas 48 horas do prazo.",
         ["Oito cenários", "previsíveis", "gatilho de contingência"])

    t = table_hdr(doc, ["#", "Risco", "Gatilho", "Impacto", "Contingência"],
                  [0.8, 4.2, 3.6, 4.2, 4.2])
    riscos = [
        ("01", "Laudo Confirmatório positivo",
         "Detecção acima CONAMA 420/09 em qualquer dos 3 furos · 15/06",
         "Protocolo 26/06 INVIÁVEL. Abre Investigação Detalhada + Plano de Intervenção.",
         "Hábilis apresenta plano revisado em 10 dias úteis. Margem até 23/12."),
        ("02", "RCA 2026 reprovado · disfunção SAO (3 hipóteses de causa)",
         "Nova análise efluentes com O&G >20 mg/L · histórico 2024 = 64 mg/L (3,2× limite). Diagnóstico causal via inspeção 12/05 Eixo 3.",
         "Hip. A (saturação por falta de limpeza): correção 1-2 semanas, sem impacto cronograma · Hip. B (dano estrutural/subdimensionamento): OBRA 4-12 semanas · extrapola F3 · dossiê com plano corretivo · Hip. C (conexão cruzada pluvial/lava-jato/trocador direto): adequação hidráulica 2-4 semanas.",
         "Diagnóstico prévio 12/05 via Eixo 3 + Eixo 5 (entrevista operador) define a hipótese · Hip. A: Bauer/Limpmil limpeza até 20/05 · Hip. B: projetista/obra até 15/05 + plano corretivo no dossiê · Hip. C: ajuste hidráulico com ART até 25/05. Coleta do RCA só após adequação."),
        ("03", "Certidão Uso do Solo > 45d",
         "Prefeitura não emite até 10/06 (45d após 25/04)",
         "Bloqueio do protocolo SEMMA. Certidão é pré-requisito.",
         "Protocolo municipal antecipado 30/04. Acompanhamento direto Hábilis."),
        ("04", "Cliente atrasa contratação",
         "Investigação não contratada até 05/05 ou demais até 15/05",
         "Mobilização 25/05 inviável. Laudos empilham em F3.",
         "Retorno em 5 dias úteis (§06). Alerta D-5 de cada contratação."),
        ("05", "AVCB exige renovação",
         "SEMMA pede AVCB vigente pós-20/10/2026",
         "Resposta pós 20/10 exige novo AVCB. Frente CBMGO fora do escopo.",
         "Verificação AVCB 12/05 (Eixo 5). Se renovação, frente jun-set."),
        ("06", "Outorga rebaixada para processo completo",
         "Inspeção 12/05 (Eixos 1-2) identifica captação direta (poço/mina/cisterna)",
         "Pedido Veredas deixa de ser dispensa e passa a processo completo: memorial expandido, planta hidráulica, ANA/CNARH, provável visita técnica SEMAD 15-45d pós-protocolo.",
         "Sequência 12/05 → 13/05 já antecipa esta hipótese. Kit de docs hídricos pré-validado entre 13-20/05 para eventual complementação."),
        ("07", "CREA/GO pode restringir Agrônomo para MCO/Anexos 273 (MITIGADO)",
         "Parecer CREA/GO à consulta 30/04 com interpretação restritiva da Res. 218/73, apesar da zona rural",
         "MCO (28/05) e Anexos CONAMA 273 (29/05) precisariam de Terceiro Eng. Civil/Ambiental: +R$ 8-15k +5-15d ART. Risco reduzido: zona rural da LAO fortalece atribuição agronômica.",
         "Consulta mantida como precaução. Lista de parceiros Eng. Civil/Ambiental até 05/05 como plano B. ART Uso do Solo segue firme com RT Agronômica independente do parecer."),
        ("08", "SEMMA contesta Laudo de Estanqueidade 10/2024",
         "SEMMA, na triagem do dossiê (jul-ago/2026), exige novo laudo por INMETRO NCC 21.08786 vencido 29/04/2024 ou calibração Suporty vencida 03/2023",
         "Laudo apresentado pode ser insuficiente. Exigência de complemento. Não bloqueia protocolo de 26/06, mas abre passivo respondível em 30-60 dias.",
         "Opção A armada em F1. Lista de empresas com INMETRO vigente + equipamento ISO/IEC 17025 pré-negociada até 10/05. Validação em 12/05 (Eixo 5) sobre equipamento usado no teste original · ver §3.3.3."),
    ]
    for i, row in enumerate(riscos):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[1], mono_cols=[0])
    doc.add_paragraph()

    callout(doc, "Feriados na janela · 3 dias úteis a menos",
            "Dentro da janela de 67 dias corridos, 3 feriados caem em dias úteis: 21/04 Tiradentes "
            "· 01/05 Dia do Trabalho · 04/06 Corpus Christi.",
            kind="cinza", bold_body_parts=["21/04", "01/05", "04/06"])

    callout(doc, "Nota de transparência · Análise de Risco no plano",
            "A Análise de Risco está listada como contingência armada, não como peça obrigatória. "
            "Não tem citação literal na LAO 033/2020 nem no TR SEMMA Guapó — o que existe é a "
            "CP 4.9 citando Plano de Gerenciamento de Risco (PGR · peça distinta) e os itens 3.7 "
            "e 4.16 da LAO que reservam à SEMMA o direito de fazer 'novas exigências'. A NR-20 "
            "exige APP formal apenas para Classe III (acima de 5.000 m³); postos padrão Classe II "
            "não se enquadram. A CONAMA 398/2008 aplica a terminais/refinarias, não varejistas. "
            "Na prática, AR pode ser exigida em 5 gatilhos: Classe 3 no Anexo II CONAMA 273 · "
            "proximidade com corpo hídrico sensível · histórico de não-conformidade · laudo "
            "confirmatório positivo · mudança operacional relevante. O posto tem probabilidade "
            "média de ativação (RCA 2024 reprovado + microbacia Ribeirão das Posses). Decisão: "
            "manter no plano como contingência, sem mobilizar CAPEX preventivo; acionar em 48h "
            "se SEMMA exigir.",
            kind="amber",
            bold_body_parts=["contingência armada", "sem citação literal na LAO 033/2020",
                             "CP 4.9", "itens 3.7 e 4.16", "5 gatilhos",
                             "manter no plano como contingência"])

    body_p(doc, "Mesmo que o pior cenário se materialize, a janela entre 26/06 (protocolo-meta) "
           "e 23/12 (vencimento LAO) reserva 180 dias de margem regulatória. É por isso que 26/06 "
           "foi acordado em vez do mínimo legal de 24/08 (120 dias).",
           ["180 dias de margem regulatória", "26/06"])

    # ===== §09 VALIDAÇÃO =====
    heading_section(doc, "09", "Validação Final · Auditoria")
    lede(doc, "Antes do envio do dossiê à SEMMA e da aprovação pelo Grupo Z+Z, este checklist "
         "aplica quatro testes independentes — ancoragem normativa, rastreabilidade de itens, "
         "dependências explicitadas e margem regulatória suficiente.",
         ["quatro testes independentes"])

    heading_h3(doc, "17 Datas-âncora oficiais")
    t = table_hdr(doc, ["Âncora", "Data", "Documento / Norma de origem"], [7.5, 2.3, 7.2])
    ancoras = [
        ("Calibração Suporty vencida", "03/2023", "Certificado do equipamento · ISO/IEC 17025"),
        ("Dispensa DURH002737 vencida", "09/12/2023", "Declaração de Uso de Recursos Hídricos · SEMAD/GO"),
        ("Credenciamento INMETRO NCC 21.08786 vencido", "29/04/2024", "Empresa de estanqueidade (anterior) · INMETRO"),
        ("LAI 008/2020 · vencimento", "31/03/2026", "Licença Ambiental de Instalação · validar status residual"),
        ("Marco divisor da ordem cronológica no Veredas", "16/04/2026", "IN 15/2026 SEMAD · §14 (até 16/04: ordem de origem) vs §15 (após: fila pós-formalização)"),
        ("Prazo transição Web Outorga → Veredas (prorrogado)", "15/07/2026", "IN 15/2026 SEMAD (altera IN 4/2025) · DOE/GO 17/04/2026 · pena de arquivamento"),
        ("Início da janela de execução", "20/04/2026", "Data após triagem do lote recebido 14/04/2026"),
        ("Verificação de outorga (F1 · Hábilis)", "28/04/2026", "Consulta ao Sistema Veredas · SEMAD/GO"),
        ("Consulta CREA/GO · atribuição MCO (F1)", "30/04/2026", "Parecer formal sobre Eng. Agrônomo para MCO/Anexos CONAMA 273 em posto · define RT interno ou externo"),
        ("Inspeção técnica de campo (F2)", "12/05/2026", "25 pontos em 5 eixos · Seção 3.2.1"),
        ("Protocolo Veredas (F2 · Hábilis · pós-inspeção)", "13/05/2026", "Requerimento + memorial + plantas no Veredas, dia útil seguinte à inspeção 12/05"),
        ("Laudo de Investigação Confirmatória (F3)", "15/06/2026", "CEMAM 029/2018 + CONAMA 420/2009 · marco do CAPEX"),
        ("Laudo de Estanqueidade · ciclo regular seguinte", "10/2026", "NBR 16795/2019 · gestão continuada (Opção B · CP 4.3) · laudo 10/2024 apresentado no dossiê"),
        ("Protocolo SEMMA · META", "26/06/2026", "Processo SEMMA 22760/2020 · 180 dias antes"),
        ("Protocolo mínimo legal (120d)", "24/08/2026", "Limite inferior legal · referência explicativa"),
        ("Vencimento AVCB Cercon 161726", "20/10/2026", "CBMGO · Certificado nº 161726 · emissão 07/11/2025"),
        ("Vencimento LAO 033/2020", "23/12/2026", "Deadline regulatório absoluto · Processo 22760/2020"),
    ]
    for i, row in enumerate(ancoras):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0], mono_cols=[1])
    doc.add_paragraph()

    heading_h3(doc, "Checklist de Auditoria · 4 testes")
    for titulo, texto in [
        ("Teste 1 · Ancoragem normativa",
         "Toda data citada está ancorada em documento oficial verificável? Toda norma citada "
         "(CONAMA 006/86, 273/2000, 319/2002, 362/2005, 420/2009 · CEMAM 029/2018 · NBR 16795/2019, "
         "16210 · Portarias INMETRO · IN 15/2026 SEMAD) tem aplicação justificada?"),
        ("Teste 2 · Rastreabilidade dos 51 itens contratados (+ 23 de continuidade)",
         "Os 11 itens de Uso do Solo mapeiam contra o checklist da Prefeitura? Os 26 itens da LAO "
         "contra o TR SEMMA? Os 14 de Condicionantes cobrem 10 parciais + 2 pendentes + 6 atendidas + 10 regras declaratórias das 28 CPs totais? "
         "Os 23 de Compliance cobrem ANP, IBAMA, NR-20, PNRS e monitoramentos?"),
        ("Teste 3 · Dependências explicitadas",
         "Cada marco tem dependência a montante e a jusante? Os 8 gatilhos de contingência (§08) "
         "têm responsável nomeado e prazo de acionamento?"),
        ("Teste 4 · Margem regulatória suficiente",
         "A janela 26/06 → 23/12 (180 dias) cobre o pior cenário? O cenário contingente tem plano "
         "alternativo dentro dos 180 dias? O CAPEX contempla piso e teto auditados? As cadências "
         "de Compliance Paralelo estão acordadas?"),
    ]:
        callout(doc, titulo, texto, kind="verde")

    # ===== §10 CONCLUSÃO =====
    heading_section(doc, "10", "Conclusão Técnica")
    lede(doc, "Processo viável dentro da janela de 67 dias corridos até o protocolo-meta em "
         "26/06/2026, fixado em 180 dias antes do vencimento da LAO. Escopo acordado organizado "
         "em 3 frentes contratadas · 51 itens · 4 fases (+ Compliance Paralelo como continuidade · 23 itens · §11 · fora do pacote atual). Cumprimento depende de organização "
         "documental e decisão tempestiva do cliente.",
         ["67 dias corridos", "26/06/2026", "180 dias antes", "51 itens · 4 fases"])

    heading_h3(doc, "1 · Gestão em 3 frentes contratadas + 1 de continuidade")
    body_p(doc, "A regularização está tratada como gestão ambiental efetiva, não apenas como "
           "renovação documental. As três primeiras frentes convergem para o protocolo de 26/06. "
           "A quarta (Compliance Paralelo) roda em cadência própria e permanente.")

    heading_h3(doc, "2 · Protocolo-meta em 180 dias antes do vencimento")
    body_p(doc, "Protocolo fixado em 26/06/2026 preserva 180 dias de margem antes do vencimento "
           "em 23/12/2026. Reduz exposição a exigências complementares.")

    heading_h3(doc, "3 · Janela de execução tight mas factível")
    body_p(doc, "67 dias corridos para quatro fases em sequência. Base documental suficiente para "
           "estruturar a F1 imediatamente. Sem folgas para paralisação, sem gargalos técnicos.")

    heading_h3(doc, "4 · Condução depende de organização e decisão")
    body_p(doc, "Quatro pontos concentram o risco: outorga via Veredas e AVCB, contratação da "
           "investigação confirmatória em F1, contratação dos demais terceiros em meados de maio, "
           "retorno em 5 dias úteis. Risco exógeno principal: passivo ambiental (Risco 01 · §08).")

    callout(doc, "Síntese final",
            "O Auto Posto América tem processo em condução ordenada, base documental funcional "
            "e escopo acordado bem delimitado. A antecipação de 180 dias antes do vencimento é "
            "a postura mais segura do ponto de vista regulatório. Meta operacional: protocolo "
            "em 26/06/2026.",
            kind="verde",
            bold_body_parts=["condução ordenada", "180 dias antes do vencimento", "26/06/2026"])

    heading_h3(doc, "Próximas entregas imediatas")
    t = table_hdr(doc, ["Entrega", "Responsável", "Prazo"], [10.5, 3.5, 3.0])
    for i, row in enumerate([
        ("Verificação formal da Certidão de Uso do Solo", "Cliente · apoio Hábilis", "25/04/2026"),
        ("Verificação de outorga no Sistema Veredas (conforme IN 15/2026)", "Hábilis", "28/04/2026"),
        ("Consulta CREA/GO · atribuição MCO (RT interno ou externo)", "Hábilis", "30/04/2026"),
        ("Inspeção de campo · 25 pontos · 5 eixos + levantamento hídrico", "Hábilis", "12/05/2026"),
        ("Protocolo novo via Sistema Veredas (pós-inspeção · §15)", "Hábilis", "13/05/2026"),
        ("Contratação da investigação confirmatória", "Cliente", "05/05/2026"),
        ("Mapeamento técnico das 8 condicionantes pendentes", "Hábilis", "10/05/2026"),
        ("Contratação de laboratório, estanqueidade e projetista", "Cliente", "15/05/2026"),
        ("Plano de atendimento das 28 condicionantes", "Hábilis", "30/05/2026"),
        ("PGRS, PAE e (se exigida) Análise de Risco", "Hábilis", "02–10/06/2026"),
        ("Laudo de Investigação Confirmatória · marco decisivo", "Terceiro · Invest.", "15/06/2026"),
        ("Estanqueidade · apresentar laudo 10/2024 (contingência se contestado)", "Hábilis", "19/06/2026"),
        ("Relatório técnico consolidado", "Hábilis", "19/06/2026"),
        ("Aprovação formal do relatório", "Cliente", "22/06/2026"),
        ("Protocolo no Processo SEMMA 22760/2020", "Hábilis", "26/06/2026"),
    ]):
        add_row(t, list(row), zebra=(i % 2 == 1), mono_cols=[2])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr'); top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single'); top.set(qn('w:sz'), '6'); top.set(qn('w:color'), 'E5E5EA')
    pbdr.append(top); p_pr.append(pbdr)
    styled_run(p, "Documento emitido por  ", size=8, color=CINZA_5, mono=True)
    styled_run(doc.add_paragraph(), "Hábilis Consultoria", size=13, color=VERDE, bold=True)
    styled_run(doc.add_paragraph(), "Regularização Ambiental · Documento Consolidado Final · v1",
               size=9, color=CINZA_5)
    styled_run(doc.add_paragraph(),
               "Versão Consolidada 1.0 · Confidencial · Emitido em 19/04/2026 · Destinatário: Grupo Z+Z",
               size=8.5, color=CINZA_5, mono=True)

    # ===== §11 ANEXOS =====
    add_page_break(doc)
    # ===== §11 · CONTINUIDADE DA GESTÃO · COMPLIANCE PARALELO (BÔNUS · FORA DO PACOTE) =====
    add_page_break(doc)
    heading_section(doc, "11", "Continuidade da Gestão · Compliance Paralelo (bônus · fora do pacote contratado)")

    callout(doc, "Capítulo adicional · oportunidade de continuidade",
            "Esta seção apresenta o Compliance Paralelo mapeado pela Hábilis como oferta de "
            "gestão continuada pós-protocolo da renovação LAO. Não integra o pacote técnico "
            "contratado atual (3 frentes · 51 itens · R$ 1.398k-2.567k). Está aqui para dar "
            "visibilidade ao Grupo Z+Z do que seguirá pendente após 26/06/2026 e do valor "
            "adicional da gestão efetiva quando conduzida de forma contínua. A formalização "
            "como serviço contratado é decisão do cliente.",
            kind="azul",
            bold_body_parts=["gestão continuada pós-protocolo", "Não integra o pacote",
                             "51 itens · R$ 1.398k-2.567k", "decisão do cliente"])

    heading_h3(doc, "11.1 · O que é Compliance Paralelo")
    body_p(doc, "A Renovação da LAO é um evento episódico · ocorre a cada 6 anos (próxima em "
           "23/12/2032). O Compliance Paralelo é permanente · engloba as obrigações regulatórias "
           "que rodam continuamente enquanto o posto opera. Sem a gestão dessas obrigações "
           "paralelas, a LAO renovada pode ser suspensa por descumprimento correlato em ANP, "
           "IBAMA, MTE (NR-20) ou PNRS · mesmo que o dossiê da SEMMA tenha sido protocolado "
           "perfeitamente.",
           ["episódico", "permanente", "suspensa por descumprimento correlato"])
    body_p(doc, "A Hábilis posiciona-se como consultoria de gestão ambiental efetiva, não apenas "
           "de renovação. Por isso mapeou o Compliance Paralelo junto com as outras frentes — "
           "para que o Grupo Z+Z visualize o escopo total da regularidade ambiental mesmo que "
           "a formalização desta frente como serviço contratado fique para depois da renovação.",
           ["gestão ambiental efetiva", "escopo total da regularidade ambiental"])

    heading_h3(doc, "11.2 · Frente ANP · Agência Nacional do Petróleo")
    t = table_hdr(doc, ["Dimensão", "Detalhamento"], [4.0, 13.0])
    rows_anp = [
        ("Base normativa", "Res. ANP 41/2013 · Res. ANP 58/2014 · Res. ANP 20/2009"),
        ("Obrigações", "Cadastro de revendedor · declaração anual de vendas · placa ao consumidor · certificado de coleta de óleo usado · aprovações de tanques/bombas"),
        ("Cadência", "Anual · contínua · pontual"),
        ("Custo anual estimado", "R$ 2k-5k"),
        ("Responsável", "Cliente (execução) + Hábilis (validação · se contratada)"),
        ("Consequência", "Auto ANP · multa · suspensão da autorização de revenda · bloqueio de NFs"),
    ]
    for i, row in enumerate(rows_anp):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0])
    doc.add_paragraph()

    heading_h3(doc, "11.3 · Frente IBAMA · CTF/APP + RAPP + TCFA")
    t = table_hdr(doc, ["Dimensão", "Detalhamento"], [4.0, 13.0])
    rows_ibama = [
        ("Base normativa", "Lei 6.938/1981 · Lei 10.165/2000 (TCFA) · IN IBAMA 06/2013 (CTF/APP e RAPP)"),
        ("Obrigações", "CTF/APP ativo · RAPP anual até 31/03 · TCFA trimestral paga"),
        ("Cadência", "Anual (RAPP) · trimestral (TCFA) · contínua (CTF)"),
        ("Custo anual estimado", "R$ 3k-8k (TCFA ~R$ 2k-4k + assessoria RAPP)"),
        ("Responsável", "Cliente (pagamento · dados) + Hábilis (RAPP · CTF · se contratada)"),
        ("Consequência", "Auto federal · multa (R$ 500-100k) · bloqueio de licenças · cadastro de devedores"),
    ]
    for i, row in enumerate(rows_ibama):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0])
    doc.add_paragraph()

    heading_h3(doc, "11.4 · Frente NR-20 · Segurança com Inflamáveis")
    t = table_hdr(doc, ["Dimensão", "Detalhamento"], [4.0, 13.0])
    rows_nr20 = [
        ("Base normativa", "NR-20 MTE · ABNT NBR 14276/2020 · NBR 15219/2022"),
        ("Classificação do posto", "Classe II (10-5.000 m³ · perfil típico varejista) · validar na inspeção 12/05"),
        ("Obrigações", "PGR atualizado · PPE · capacitação por cargo (Básica 16h · Intermediária 20h · Avançada 40h) · treinamento periódico bienal/trienal · brigada treinada · procedimentos operacionais"),
        ("Cadência", "Contínua + bienal/trienal"),
        ("Custo anual estimado", "R$ 5k-15k (treinamentos + revisão PGR + assessoria)"),
        ("Responsável", "Cliente + capacitador credenciado + Hábilis/Eng. Seg. Trabalho (se contratado)"),
        ("Consequência", "Auto MTE · multa · embargo · responsabilização criminal · impacto CP 4.9 e 3.10"),
    ]
    for i, row in enumerate(rows_nr20):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0])
    doc.add_paragraph()

    heading_h3(doc, "11.5 · Frente PNRS · Política Nacional de Resíduos Sólidos")
    t = table_hdr(doc, ["Dimensão", "Detalhamento"], [4.0, 13.0])
    rows_pnrs = [
        ("Base normativa", "Lei 12.305/2010 · CONAMA 362/2005 · Decreto 10.936/2022 + legislação estadual GO"),
        ("Obrigações", "PGRS revisado anualmente · MTR por destinação · certificados de destinação final · logística reversa do óleo · rastreabilidade"),
        ("Cadência", "Contínua · mensal/trimestral · revisão anual do PGRS"),
        ("Custo anual estimado", "R$ 2k-8k (PGRS + rastreabilidade · custos da destinação estão no contrato Bauer/Limpmil/Ecofenix)"),
        ("Responsável", "Cliente + Ecofenix/Bauer/Limpmil + Hábilis (PGRS · se contratada)"),
        ("Consequência", "Auto ambiental · multa · obrigação de recolher passivo · cadastro de irregulares · impacto CP 4.6/4.7"),
    ]
    for i, row in enumerate(rows_pnrs):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0])
    doc.add_paragraph()

    heading_h3(doc, "11.6 · Monitoramentos periódicos pós-LAO")
    t = table_hdr(doc, ["Dimensão", "Detalhamento"], [4.0, 13.0])
    rows_mon = [
        ("Base normativa", "CP 4.3 estanqueidade · CP 4.4 RCA · CP 4.5 contratos · CP 4.7 certificados · CEMAM 029/2018"),
        ("Obrigações", "RCA semestral (fev+ago) · Laudo Estanqueidade bianual (NBR 16795/2019) · contratos e NFs semestrais · certificados por 3 anos · CNDs a cada 90 dias"),
        ("Cadência", "Semestral · bianual · contínua · trimestral"),
        ("Custo anual estimado", "R$ 25k-60k (laboratórios + estanqueidade em anos pares + assessoria)"),
        ("Responsável", "Cliente + Laboratórios + Empresa INMETRO + Hábilis (validação · se contratada)"),
        ("Consequência", "Não-conformidade no Extrato · impacto na próxima renovação (2032) · auto ambiental se reprovação recorrente"),
    ]
    for i, row in enumerate(rows_mon):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0])
    doc.add_paragraph()

    heading_h3(doc, "11.7 · Cronograma anual integrado · matriz 12 meses × 5 frentes")
    t = table_hdr(doc, ["Frente / obrigação", "Jan", "Fev", "Mar", "Abr", "Mai", "Jun",
                         "Jul", "Ago", "Set", "Out", "Nov", "Dez"],
                   [4.5, 1.0, 1.0, 1.0, 1.0, 1.0, 1.0, 1.0, 1.0, 1.0, 1.0, 1.0, 1.0])
    rows_cron_anual = [
        ("ANP · declaração anual", "", "", "X", "", "", "", "", "", "", "", "", ""),
        ("IBAMA · RAPP (até 31/03)", "", "X", "X", "", "", "", "", "", "", "", "", ""),
        ("IBAMA · TCFA trimestral", "X", "", "", "X", "", "", "X", "", "", "X", "", ""),
        ("NR-20 · PGR/treinamento", "", "", "", "", "X", "", "", "", "", "", "X", ""),
        ("PNRS · MTR mensal", "X", "X", "X", "X", "X", "X", "X", "X", "X", "X", "X", "X"),
        ("PNRS · PGRS anual", "", "", "", "", "", "", "", "", "", "", "", "X"),
        ("RCA semestral · fev+ago", "", "X", "", "", "", "", "", "X", "", "", "", ""),
        ("Contratos Bauer/Limpmil/Ecofenix", "", "X", "", "", "", "", "", "X", "", "", "", ""),
        ("Estanqueidade · bianual (anos pares)", "", "", "", "", "", "X", "", "", "", "", "", ""),
        ("CND Municipal + Estadual · 90 dias", "X", "", "", "X", "", "", "X", "", "", "X", "", ""),
    ]
    for i, row in enumerate(rows_cron_anual):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0], mono_cols=[1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12])
    doc.add_paragraph()

    heading_h3(doc, "11.8 · Custo anual estimado · consolidado")
    t = table_hdr(doc, ["Frente", "Piso anual", "Teto anual", "Observação"],
                   [6.0, 3.0, 3.0, 5.0])
    rows_custo = [
        ("ANP (§11.2)", "R$ 2k", "R$ 5k", ""),
        ("IBAMA · CTF/APP + RAPP + TCFA (§11.3)", "R$ 3k", "R$ 8k", ""),
        ("NR-20 (§11.4)", "R$ 5k", "R$ 15k", "Bienal pode elevar ~R$ 10k"),
        ("PNRS (§11.5)", "R$ 2k", "R$ 8k", ""),
        ("Monitoramentos pós-LAO (§11.6)", "R$ 25k", "R$ 60k", "Estanqueidade em anos pares"),
        ("TOTAL ANUAL CONSOLIDADO", "R$ 37k", "R$ 96k", "Fora do CAPEX da renovação"),
    ]
    for i, row in enumerate(rows_custo):
        is_total = (i == len(rows_custo) - 1)
        add_row(t, list(row), zebra=(i % 2 == 1 and not is_total), bold_cols=[0],
                mono_cols=[1, 2], is_total=is_total)
    doc.add_paragraph()

    heading_h3(doc, "11.9 · Matriz de responsabilidades · Cliente × Hábilis × Terceiros")
    t = table_hdr(doc, ["Frente", "Cliente", "Hábilis (se contratada)", "Terceiros"],
                   [3.5, 4.5, 4.5, 4.5])
    rows_resp = [
        ("ANP", "Pagamento taxas · dados · placa", "Validação · calendário · acompanhamento", "—"),
        ("IBAMA", "Pagamento TCFA · dados", "RAPP · CTF · alertas", "—"),
        ("NR-20", "Aplicação + registro · brigada", "Atualização PGR · auditoria · PAE", "Capacitador credenciado"),
        ("PNRS", "Geração + armazenamento", "PGRS · MTR · revisão anual", "Bauer · Limpmil · Ecofenix"),
        ("Monitoramentos", "Pagamento · acesso", "Validação · protocolação · alertas", "Laboratórios · empresa INMETRO"),
    ]
    for i, row in enumerate(rows_resp):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0])
    doc.add_paragraph()

    heading_h3(doc, "11.10 · Como formalizar a continuidade")
    body_p(doc, "Se o Grupo Z+Z desejar formalizar o Compliance Paralelo como serviço continuado "
           "pós-protocolo, a Hábilis apresenta aditivo contratual específico com:")
    for item in [
        "Escopo por frente (marcar frentes de interesse · não precisa ser tudo)",
        "Honorários Hábilis anuais para coordenação e alertas (contrato separado do CAPEX de terceiros)",
        "Indicação de parceiros para execução (capacitador NR-20 · laboratórios · destinadores)",
        "Painel de gestão (reutilização do Dashboard_LAO_v5 ampliado com Compliance) · acompanhamento",
        "Cronograma do primeiro ano de continuidade (ancorado em 26/06/2026 como marco inicial)",
        "Cláusulas de revisão e atualização conforme mudanças normativas",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        styled_run(p, item, size=10, color=CINZA_7)
    body_p(doc, "Início sugerido: agosto/2026 · logo após o ciclo pós-protocolo ser estabilizado. "
           "Primeiro marco seria a coleta do RCA semestral em ago/2026, seguido da revisão anual "
           "do PGRS em dez/2026.",
           ["agosto/2026", "ago/2026", "dez/2026"])

    callout(doc, "Por que esta seção está aqui",
            "Este capítulo adicional existe para que o Grupo Z+Z visualize o escopo total da "
            "regularidade ambiental · não apenas o que está dentro do pacote contratado da "
            "renovação LAO (R$ 1.398k-2.567k · 3 frentes · 51 itens · 67 dias). A renovação "
            "resolve o problema imediato de 23/12/2026. O Compliance Paralelo resolve o problema "
            "eterno de manter a regularidade que sustentará as próximas 3 renovações "
            "(2032 · 2038 · 2044). A Hábilis oferece visibilidade aqui · a contratação fica para "
            "a decisão estratégica do cliente após 26/06.",
            kind="azul",
            bold_body_parts=["escopo total da regularidade ambiental", "eterno",
                             "2032 · 2038 · 2044", "decisão estratégica"])

    # ===== §12 ANEXOS (renumeração) =====
    add_page_break(doc)
    heading_section(doc, "12", "Anexos · Material de apoio operacional")
    lede(doc, "Dois anexos acompanham este relatório como material de apoio operacional: "
         "o Anexo A é uma cartilha em linguagem direta para o Grupo Z+Z acompanhar o processo "
         "da outorga no Sistema Veredas; o Anexo B é o roteiro operacional da auditoria da "
         "Caixa Separadora Água-Óleo (SAO) na inspeção de campo de 12/05/2026.",
         ["Anexo A", "Anexo B", "12/05/2026"])

    # ===== ANEXO A · OUTORGA =====
    heading_h3(doc, "Anexo A · Outorga do Auto Posto América no Sistema Veredas (cartilha para o Grupo Z+Z)")

    heading_h4(doc, "A.1 · O que é a outorga e por que precisamos")
    body_p(doc, "Todo empreendimento que usa ou interfere em recurso hídrico (poço, nascente, "
           "captação, lançamento) precisa de outorga do Estado · autorização formal da SEMAD/GO. "
           "O posto pode cair em: (a) dispensa simplificada, se usa apenas rede pública de água; "
           "(b) processo completo, se tem captação própria. O regime correto é decidido pelo "
           "levantamento hídrico na inspeção de 12/05/2026 (Eixo 2).",
           ["dispensa simplificada", "processo completo", "12/05/2026"])

    heading_h4(doc, "A.2 · Onde estamos hoje")
    t = table_hdr(doc, ["Fato", "Situação", "O que significa"], [5.0, 3.0, 9.0])
    rows_hoje = [
        ("Dispensa anterior · DURH002737", "Venceu 09/12/2023",
         "Posto opera sem dispensa há 2+ anos. Regularização de inconformidade, não atualização."),
        ("Sistema anterior · Web Outorga", "Descontinuado",
         "SEMAD migrou para o Sistema Veredas."),
        ("IN 15/2026 SEMAD · DOE 17/04/2026", "Prazo prorrogado 15/07/2026",
         "Prorrogou a transição MAS criou 2 filas: §14 até 16/04 com prioridade · §15 após 16/04 sem prioridade."),
        ("Situação do Auto Posto América", "§15 da IN",
         "Protocolo em 13/05/2026 cai no §15. Entra na fila sem herdar ordem cronológica. Análise SEMAD leva 60-120 dias de qualquer forma."),
    ]
    for i, row in enumerate(rows_hoje):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0], mono_cols=[1])
    doc.add_paragraph()

    heading_h4(doc, "A.3 · Timeline · passo a passo")
    t = table_hdr(doc, ["Quando", "O quê", "Quem", "Grupo Z+Z precisa"], [2.2, 5.0, 2.3, 7.5])
    rows_tl = [
        ("23/04/2026", "Procuração pública para a Hábilis (SEMMA, Prefeitura, SEMAD, Veredas)",
         "Cliente · cartório", "Assinar e reconhecer firma. Modelo pronto."),
        ("28/04/2026", "Consulta no Sistema Veredas · diagnóstico", "Hábilis",
         "Aguardar resultado (1-3 dias úteis)."),
        ("12/05/2026", "Inspeção de campo · levantamento hídrico (define regime)", "Hábilis",
         "Liberar acesso ao posto · ponto focal no local · documentação física."),
        ("13/05/2026", "Protocolo do pedido no Veredas", "Hábilis",
         "Pagamento da taxa SEMAD (DUAM)."),
        ("13-14/05/2026", "Recebimento da Certidão de Tramitação", "Veredas",
         "Documento imediato pós-protocolo · compõe dossiê da LAO (26/06)."),
        ("Jul-Set/2026", "Análise de mérito SEMAD (60-120d)", "SEMAD/GO",
         "Responder complementos em 5 dias úteis. Se processo completo, receber visita técnica."),
        ("Depois", "Deferimento da outorga ou dispensa", "SEMAD/GO",
         "Receber documento oficial · arquivar · cumprir condicionantes."),
    ]
    for i, row in enumerate(rows_tl):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0], mono_cols=[0])
    doc.add_paragraph()

    heading_h4(doc, "A.4 · Checklist · o que o Grupo Z+Z precisa providenciar (até 05/05)")
    for item in [
        "Procuração pública com firma reconhecida · SEMMA Guapó · Prefeitura · SEMAD/GO · Veredas · em nome da Hábilis (a de 2023 em nome da A3 Engenharia NÃO serve).",
        "Cartão CNPJ atualizado · Receita Federal online · CNAE 'Comércio Varejista de Combustíveis' ativo.",
        "RG/CPF do representante legal · se houve mudança do quadro societário (Amarido × Sérgio/José), trazer ato societário atualizado.",
        "Documentos do imóvel · RGI 2025 · IPTU/ITR vigente · certidão municipal de enquadramento (confirma zona rural).",
        "Liberação da visita 12/05 · autorização formal · ponto focal no local · acesso a documentos físicos.",
        "Pagamento de taxa DUAM/SEMAD · valor confirmado após montagem do requerimento.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        styled_run(p, item, size=10, color=CINZA_7)

    heading_h4(doc, "A.5 · O que esperar após o protocolo 13/05")
    t = table_hdr(doc, ["Cenário", "Significado", "Grupo Z+Z precisa fazer"], [4.0, 5.5, 7.5])
    rows_esp = [
        ("① Dispensa simplificada", "Posto usa só rede pública. Análise SEMAD 60-120d.",
         "Aguardar. Responder complemento em 5 dias úteis."),
        ("② Processo completo", "Posto tem captação. Análise + provável visita técnica SEMAD 15-45d.",
         "Receber visita SEMAD com ponto focal. Responder em 5 dias úteis. Hábilis orienta peças adicionais."),
        ("③ Exigência complementar", "SEMAD solicita ajuste antes de decidir.",
         "Hábilis monta a resposta. Cliente assina se necessário. Prazo 5 dias úteis."),
        ("④ Indeferimento (improvável)", "SEMAD nega o pedido.",
         "Hábilis analisa e prepara recurso/novo pedido. Não afeta LAO (certidão de tramitação basta)."),
    ]
    for i, row in enumerate(rows_esp):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0])
    doc.add_paragraph()

    heading_h4(doc, "A.6 · Glossário rápido")
    t = table_hdr(doc, ["Termo", "Explicação"], [3.0, 14.0])
    rows_glos = [
        ("DURH", "Declaração de Uso de Recursos Hídricos · documento anterior que o posto tinha (DURH002737, venceu 09/12/2023)."),
        ("Veredas", "Sistema oficial da SEMAD/GO para outorgas desde 2025. Substituiu o Web Outorga."),
        ("SEMAD", "Secretaria Estadual de Meio Ambiente e Desenvolvimento Sustentável (GO). Diferente da SEMMA Guapó (municipal)."),
        ("IN 15/2026 SEMAD", "Instrução Normativa de 17/04/2026. Prorrogou transição para 15/07/2026 e criou §14 (com prioridade) e §15 (sem prioridade)."),
        ("Certidão de Tramitação", "Comprovante automático pós-protocolo. Suficiente para dossiê da LAO antes do deferimento."),
        ("ANA · CNARH", "Agência Nacional de Águas · Cadastro Nacional de Usuários de Recursos Hídricos · necessário se há captação própria."),
        ("DUAM", "Documento Único de Arrecadação · boleto para taxas SEMAD."),
    ]
    for i, row in enumerate(rows_glos):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0], mono_cols=[0])
    doc.add_paragraph()

    # ===== ANEXO B · ROTEIRO SAO =====
    add_page_break(doc)
    heading_h3(doc, "Anexo B · Roteiro operacional · Auditoria da SAO na inspeção 12/05/2026")
    lede(doc, "Protocolo passo a passo para uso em campo. Objetivo: separar em campo as 3 "
         "hipóteses de causa do RCA 2024 reprovado (Hip. A saturação · Hip. B dano estrutural · "
         "Hip. C conexão cruzada) e produzir diagnóstico antes da coleta do novo RCA em 03/06/2026. "
         "Formato prancheta · imprimir e levar para o campo.",
         ["3 hipóteses de causa", "Hip. A", "Hip. B", "Hip. C"])

    heading_h4(doc, "B.1 · Preparação antes de sair")
    for item in [
        "Câmera/celular com GPS · registro datado",
        "Trena 5m + vara de medição com régua",
        "Prancheta + este formulário impresso",
        "Caneta + marcador permanente",
        "EPI: luvas nitrílicas + óculos + botina",
        "Lanterna de alta intensidade",
        "Cópia do Contrato Bauer/Limpmil + NFs semestrais",
        "Cópia do laudo RCA 2024 reprovado (Água Viva)",
        "Planta baixa 05/2024 · ART Hábilis emitida (27/04)",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        styled_run(p, item, size=10, color=CINZA_7)

    # Passos em callouts azuis
    passos = [
        ("PASSO 1 · Localização e identificação",
         ["Confirmar localização vs. planta 05/2024 · GPS registrado",
          "Tipo de construção (alvenaria · pré-moldada · polietileno · outro)",
          "Operador identificado e acesso liberado",
          "Fotos: localização externa · tampas · entorno"]),
        ("PASSO 2 · Inspeção externa (antes de abrir)",
         ["Sinais de vazamento ao redor? (manchas, odor, vegetação morta)",
          "Tampas íntegras e vedadas?",
          "Piso de concreto sem trincas ou afundamentos?",
          "Canaleta de entrada limpa e funcional?",
          "Odor perceptível sem abrir? (nenhum / leve / forte)"]),
        ("PASSO 3 · Inspeção interna dos 3 compartimentos",
         ["A SAO tem 3 compartimentos: (1) Desarenador · (2) Separador · (3) Caixa final",
          "C1 · Desarenador: profundidade do sedimento (cm) · limpo/parcial/saturado",
          "C2 · Separador: espessura da camada de óleo boiando (cm) · odor leve/forte",
          "C2 · Chicana divisora no lugar e íntegra?",
          "C3 · Caixa final: cor do efluente · transparência (clara/turva/opaca)",
          "Parede interna: trincas visíveis? (fotografar)",
          "Fotos internas: C1 · C2 · C3 · chicanas · trincas"]),
        ("PASSO 4 · Vazão e operação em tempo real",
         ["Com bombas operando, vazão de entrada (gotejando/fluxo pequeno/alto/nenhum)",
          "Saída C3 (efluente final) flui? (sim/não/intermitente)",
          "Para onde vai o efluente final (rede pluvial, corpo d'água, valo)"]),
        ("PASSO 5 · Segregação de águas pluviais (CP 4.10)",
         ["Pista de abastecimento drena para SAO?",
          "Pátio externo drena para rede pluvial SEPARADA?",
          "Telhado e áreas cobertas independentes da SAO?",
          "Se lava-jato: água vai para SAO ou direto ao esgoto?",
          "Se trocador de óleo: para onde vai a água de derramamento/limpeza?"]),
        ("PASSO 6 · Entrevista focada com operador",
         ["Tempo de trabalho no posto",
          "Frequência real da limpeza da SAO (mensal/trimestral/semestral/quando entope)",
          "Última limpeza: quando, quem, tem NF",
          "Histórico de entupimento/transbordamento/extravasamento",
          "Posto opera lava-jato? Quantos carros/dia?",
          "Posto opera trocador de óleo? Desde quando?",
          "Vazamento de combustível: histórico, ação tomada, reportado?",
          "RT atual que assina documentos",
          "Representante legal atual (Amarido ainda ou outro?)"]),
        ("PASSO 7 · Marcação prévia dos 3 pontos de sondagem (Eixo 4)",
         ["Definir 3 pontos estratégicos para Investigação Confirmatória 25/05",
          "Marcar com estaca/pintura no chão",
          "Conferir redes enterradas com operador (elétrica, hidráulica, gás)",
          "Fotografar cada ponto com GPS",
          "Autorização formal do proprietário para perfuração"]),
    ]
    for titulo, items in passos:
        callout(doc, titulo, " · ".join(items), kind="azul")

    heading_h4(doc, "B.3 · Matriz de decisão em campo · separa Hip. A / B / C")
    t = table_hdr(doc, ["Hipótese", "Sinais que confirmam", "Ação imediata pós-visita"],
                   [5.0, 6.0, 6.0])
    rows_mtx = [
        ("Hip. A · Saturação por falta de limpeza (MELHOR CENÁRIO)",
         "Desarenador saturado · óleo grosso >10 cm · operador diz limpeza 'quando entope' · compartimentos íntegros · sem trincas",
         "Bauer/Limpmil limpeza profunda até 15/05. Reset de contrato com nova frequência. RCA 03/06 com expectativa positiva."),
        ("Hip. B · Dano estrutural ou subdimensionamento (PIOR CENÁRIO)",
         "Trincas visíveis · chicana danificada/ausente · paredes deterioradas · volume menor que vazão · compartimentos não funcionais",
         "Projetista até 15/05 + plano corretivo no dossiê · aditivo CAPEX · pode exigir nova LI (CP 4.12)."),
        ("Hip. C · Conexão cruzada (CENÁRIO INTERMEDIÁRIO)",
         "Pluvial entrando na SAO · lava-jato/trocador direto sem passar pela SAO · efluente estranho (detergente espumoso, cor não-oleosa)",
         "Ajuste hidráulico com ART até 25/05. Anuência de saneamento pode ser necessária. Refazer hidráulica do posto."),
    ]
    for i, row in enumerate(rows_mtx):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0])
    doc.add_paragraph()

    callout(doc, "B.4 · Ação imediata pós-visita (mesmo dia, volta a Guapó)",
            "Baixar fotos em pasta datada · Preencher formulário eletrônico · Definir hipótese "
            "(A/B/C) e acionar parceiro · Reportar ao Grupo Z+Z em até 24 h · Atualizar o Plano "
            "das 28 CPs com a hipótese confirmada · Se Hipótese B, montar pedido de LI em paralelo "
            "à renovação (CP 4.12).",
            kind="amber",
            bold_body_parts=["24 h", "Hipótese B"])

    heading_h4(doc, "B.5 · Identificação do executor da auditoria")
    t = table_hdr(doc, ["Campo", "Preenchimento"], [5.0, 12.0])
    rows_id = [
        ("Data da auditoria", "____/____/2026"),
        ("Hora início / término", "____:____ / ____:____"),
        ("Executor (Hábilis)", "Nome · CRBio/CREA · ART nº"),
        ("Ponto focal no local", "Nome · cargo · telefone"),
        ("Condições climáticas", "Tempo · temperatura · chuva 24h?"),
        ("Hipótese confirmada", "A / B / C (circular)"),
        ("Próxima ação definida", "________________________________"),
        ("Assinatura do executor", "________________________________"),
    ]
    for i, row in enumerate(rows_id):
        add_row(t, list(row), zebra=(i % 2 == 1), bold_cols=[0])

    out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       "Documento_Consolidado_Final_AutoPostoAmerica_v1.docx")
    doc.save(out)
    print(f"DOCX salvo: {out}")
    print(f"Tamanho: {os.path.getsize(out) / 1024:.1f} KB")


if __name__ == "__main__":
    build()
