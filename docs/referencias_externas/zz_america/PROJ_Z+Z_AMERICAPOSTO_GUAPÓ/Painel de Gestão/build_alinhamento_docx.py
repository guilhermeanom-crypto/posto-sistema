#!/usr/bin/env python3
"""
Gerador DOCX · Alinhamento Executivo · Auto Posto América.
Peça enxuta (3 paginas + 1 de Compliance Paralelo) para validacao do proprietario.
Uso: python3 build_alinhamento_docx.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Paleta Habilis
VERDE = RGBColor(0x1B, 0x5E, 0x20)
VERDE_L = RGBColor(0x4C, 0xAF, 0x50)
VERDE_BG = "E8F5E9"
AMBER = RGBColor(0xF5, 0x7F, 0x17)
AMBER_BG = "FFF8E1"
VERMELHO = RGBColor(0xB7, 0x1C, 0x1C)
VERMELHO_BG = "FFEBEE"
AZUL = RGBColor(0x0D, 0x47, 0xA1)
AZUL_BG = "E3F2FD"
CINZA_9 = RGBColor(0x1C, 0x1C, 0x1E)
CINZA_7 = RGBColor(0x3A, 0x3A, 0x3C)
CINZA_5 = RGBColor(0x6C, 0x6C, 0x70)
CINZA_BG = "FAFAFA"
BRANCO = RGBColor(0xFF, 0xFF, 0xFF)
FONT_SANS = "Segoe UI"
FONT_MONO = "Consolas"


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for name, val in edges.items():
        if val is None: continue
        size, color = val
        el = OxmlElement(f'w:{name}')
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(size)); el.set(qn('w:color'), color)
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def set_paragraph_shading(p, hex_color):
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    p_pr.append(shd)


def set_paragraph_border_left(p, size_pt, color_hex):
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), str(int(size_pt * 8)))
    left.set(qn('w:space'), '6'); left.set(qn('w:color'), color_hex)
    pbdr.append(left); p_pr.append(pbdr)


def styled_run(p, text, *, bold=False, italic=False, size=None, color=None, mono=False):
    run = p.add_run(text)
    font_name = FONT_MONO if mono else FONT_SANS
    run.font.name = font_name
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.find(qn('w:rFonts'))
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts'); r_pr.append(r_fonts)
    r_fonts.set(qn('w:ascii'), font_name); r_fonts.set(qn('w:hAnsi'), font_name); r_fonts.set(qn('w:cs'), font_name)
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    return run


def mixed_run(p, text, *, size=10, default_color=CINZA_7, bold_parts=None):
    if not bold_parts:
        styled_run(p, text, size=size, color=default_color); return
    remaining = text
    for target in bold_parts:
        idx = remaining.find(target)
        if idx == -1: continue
        if idx > 0: styled_run(p, remaining[:idx], size=size, color=default_color)
        styled_run(p, target, size=size, color=CINZA_9, bold=True)
        remaining = remaining[idx + len(target):]
    if remaining:
        styled_run(p, remaining, size=size, color=default_color)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def heading_section(doc, num, title, space_before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '14')
    bottom.set(qn('w:space'), '3'); bottom.set(qn('w:color'), '1B5E20')
    pbdr.append(bottom); p_pr.append(pbdr)
    styled_run(p, f"{num}  ", size=11, color=VERDE_L, mono=True)
    styled_run(p, title, size=15, color=VERDE, bold=True)


def heading_h3(doc, text, space_before=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    set_paragraph_border_left(p, 2, "1B5E20")
    p.paragraph_format.left_indent = Cm(0.2)
    styled_run(p, text, size=10.5, color=CINZA_9, bold=True)


def body_p(doc, text, bold_parts=None, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    mixed_run(p, text, bold_parts=bold_parts, size=size)


def lede(doc, text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.25)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_shading(p, VERDE_BG)
    set_paragraph_border_left(p, 2.5, "1B5E20")
    mixed_run(p, text, size=10.5, default_color=CINZA_9, bold_parts=bold_parts)


def callout(doc, title, body, kind="verde", bold_body_parts=None):
    bg = {"verde": VERDE_BG, "amber": AMBER_BG, "vermelho": VERMELHO_BG,
          "azul": AZUL_BG, "cinza": CINZA_BG}.get(kind, VERDE_BG)
    border = {"verde": "1B5E20", "amber": "F57F17", "vermelho": "B71C1C",
              "azul": "0D47A1", "cinza": "6C6C70"}.get(kind, "1B5E20")
    t = doc.add_table(rows=1, cols=1)
    c = t.cell(0, 0)
    set_cell_shading(c, bg)
    set_cell_border(c, left=(24, border))
    p1 = c.paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    styled_run(p1, title.upper(), size=9, color=CINZA_9, bold=True, mono=True)
    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    mixed_run(p2, body, size=9.5, default_color=CINZA_7, bold_parts=bold_body_parts)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(2)


def table_hdr(doc, headers, widths_cm=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_shading(c, "1B5E20")
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        styled_run(p, h.upper(), size=8, color=BRANCO, bold=True, mono=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if widths_cm and i < len(widths_cm): c.width = Cm(widths_cm[i])
    return t


def add_row(t, values, zebra=False, bold_cols=None, mono_cols=None, is_total=False, size=9):
    row = t.add_row()
    for i, val in enumerate(values):
        c = row.cells[i]
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
        bold = (bold_cols and i in bold_cols) or is_total
        mono = mono_cols and i in mono_cols
        styled_run(p, val, size=size, color=CINZA_9 if bold else CINZA_7, bold=bold, mono=mono)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        if is_total: set_cell_shading(c, VERDE_BG)
        elif zebra: set_cell_shading(c, CINZA_BG)


def bullet(doc, text, bold_parts=None, size=10):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    mixed_run(p, text, bold_parts=bold_parts, size=size)


# =========== Cores por responsavel (checklist cronologico) ===========
RESP_COLORS = {
    "cliente":  "0D47A1",   # azul
    "habilis":  "1B5E20",   # verde
    "terceiro": "6C6C70",   # cinza
    "misto":    "4A148C",   # roxo (Habilis ou Terceiro)
}


# =========== Checklist tipado por chip ===========
CHIP_STYLE = {
    "DOC":     ("0D47A1", "Documento"),
    "DECISAO": ("1B5E20", "Decisao"),
    "PAGAR":   ("F57F17", "Pagamento"),
    "ACESSO":  ("6C6C70", "Acesso"),
    "ASSINAR": ("2E7D32", "Assinatura"),
}


def set_tbl_borders(t, color="E0E0E0", size=4):
    tbl_pr = t._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(size))
        el.set(qn('w:color'), color)
        borders.append(el)
    tbl_pr.append(borders)


def phase_checklist(doc, title, period, days, items):
    """Header da fase + tabela (chip tipo | descricao)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    set_paragraph_border_left(p, 2.5, "1B5E20")
    p.paragraph_format.left_indent = Cm(0.2)
    styled_run(p, f"{title}  ", size=11, color=CINZA_9, bold=True)
    styled_run(p, f"{period}  ", size=9.5, color=VERDE, mono=True)
    styled_run(p, f"({days} dias)", size=9, color=CINZA_5, mono=True)

    # Contadores por tipo
    counts = {}
    for kind, _ in items:
        counts[kind] = counts.get(kind, 0) + 1
    mini = " | ".join(f"{v} {CHIP_STYLE[k][1].lower()}" for k, v in counts.items())
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(3)
    p2.paragraph_format.left_indent = Cm(0.25)
    styled_run(p2, mini, size=8.5, color=CINZA_5, mono=True)

    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_tbl_borders(t, "F0F0F0", 4)
    for kind, text in items:
        row = t.add_row()
        c_chip, c_body = row.cells[0], row.cells[1]
        color_hex, label = CHIP_STYLE[kind]
        set_cell_shading(c_chip, color_hex)
        pc = c_chip.paragraphs[0]
        pc.paragraph_format.space_before = Pt(2); pc.paragraph_format.space_after = Pt(2)
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(pc, label.upper(), size=7.5, color=BRANCO, bold=True, mono=True)
        c_chip.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c_chip.width = Cm(2.4)

        pb = c_body.paragraphs[0]
        pb.paragraph_format.space_before = Pt(2); pb.paragraph_format.space_after = Pt(2)
        pb.paragraph_format.left_indent = Cm(0.15)
        styled_run(pb, text, size=9.5, color=CINZA_9)
        c_body.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c_body.width = Cm(14.4)


# =========== Matriz visual (cards horizontais) ===========
def matriz_cards(doc, cards, accent_hex="1B5E20"):
    """
    cards: list of dicts { 'title', 'sub', 'stats': [(label, value, color), ...], 'next': str }
    Gera uma tabela N colunas com mini-cards visuais.
    """
    n = len(cards)
    t = doc.add_table(rows=1, cols=n)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, card in enumerate(cards):
        c = t.rows[0].cells[i]
        set_cell_shading(c, "FAFAFA")
        set_cell_border(c, top=(24, accent_hex), left=(4, "E0E0E0"),
                        right=(4, "E0E0E0"), bottom=(4, "E0E0E0"))
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # Título
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(1)
        styled_run(p, card["title"], size=10.5, color=CINZA_9, bold=True)

        # Subtítulo
        p = c.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(5)
        styled_run(p, card["sub"], size=8.5, color=CINZA_5, mono=True)

        # Stats
        for label, val, color in card.get("stats", []):
            p = c.add_paragraph()
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(1)
            styled_run(p, f"{val} ", size=13, color=color, bold=True, mono=True)
            styled_run(p, label, size=7.5, color=CINZA_5, bold=True, mono=True)

        # Próximo
        if card.get("next"):
            p = c.add_paragraph()
            p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(3)
            set_paragraph_border_left(p, 1.5, accent_hex)
            p.paragraph_format.left_indent = Cm(0.15)
            styled_run(p, "PROXIMA ACAO", size=7, color=CINZA_5, bold=True, mono=True)
            p = c.add_paragraph()
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(4)
            set_paragraph_border_left(p, 1.5, accent_hex)
            p.paragraph_format.left_indent = Cm(0.15)
            styled_run(p, card["next"], size=9, color=CINZA_9, bold=True)


def field_page(p):
    r = p.add_run()
    r.font.name = FONT_MONO; r.font.size = Pt(8); r.font.color.rgb = CINZA_5
    fld_b = OxmlElement('w:fldChar'); fld_b.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.text = 'PAGE'
    fld_e = OxmlElement('w:fldChar'); fld_e.set(qn('w:fldCharType'), 'end')
    r._r.append(fld_b); r._r.append(it); r._r.append(fld_e)


def field_numpages(p):
    r = p.add_run()
    r.font.name = FONT_MONO; r.font.size = Pt(8); r.font.color.rgb = CINZA_5
    fld_b = OxmlElement('w:fldChar'); fld_b.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.text = 'NUMPAGES'
    fld_e = OxmlElement('w:fldChar'); fld_e.set(qn('w:fldCharType'), 'end')
    r._r.append(fld_b); r._r.append(it); r._r.append(fld_e)


def build():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(1.5); s.bottom_margin = Cm(1.5)
    s.left_margin = Cm(1.8); s.right_margin = Cm(1.8)

    # Header
    h_p = s.header.paragraphs[0]
    styled_run(h_p, "Auto Posto América | Alinhamento Executivo", size=8, color=CINZA_5, mono=True)
    h_p.add_run("\t")
    styled_run(h_p, "Hábilis | Abr/2026", size=8, color=VERDE, bold=True, mono=True)

    # Footer
    f_p = s.footer.paragraphs[0]
    styled_run(f_p, "Hábilis Regularização Ambiental", size=8, color=CINZA_5, mono=True)
    f_p.add_run("\t")
    styled_run(f_p, "Pagina ", size=8, color=CINZA_5, mono=True)
    field_page(f_p)
    styled_run(f_p, " de ", size=8, color=CINZA_5, mono=True)
    field_numpages(f_p)

    # ====================================================================
    # CAPA
    # ====================================================================
    p = doc.add_paragraph()
    styled_run(p, "H  ", size=16, color=VERDE, bold=True)
    styled_run(p, "Hábilis", size=16, color=VERDE, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(40)
    styled_run(p, "REGULARIZAÇÃO AMBIENTAL", size=9, color=CINZA_5, mono=True)
    for _ in range(4): doc.add_paragraph()

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(p, "PROCESSO DE REGULARIZAÇÃO AMBIENTAL", size=10, color=VERDE, bold=True, mono=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(p, "Auto Posto América", size=30, color=CINZA_9, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(p, "Renovação LAO 033/2020 | SEMMA Guapó | Processo 22760/2020",
               size=13, color=CINZA_7)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    styled_run(p, "Alinhamento executivo para validação do Grupo Z+Z. Resume o plano "
               "técnico em execução pela Hábilis: 4 frentes (Uso do Solo, Renovação LAO, "
               "Condicionantes, Compliance Paralelo) totalizando 74 itens, cronograma "
               "operacional em 4 fases (67 dias corridos) e checklist cronológico "
               "integral com 41 entregas. Protocolo-meta ", size=11, color=CINZA_7)
    styled_run(p, "26/06/2026 | 180 dias antes do vencimento da LAO.", size=11, color=CINZA_9, bold=True)

    for _ in range(2): doc.add_paragraph()

    tc = doc.add_table(rows=2, cols=4)
    for i, h in enumerate(["DESTINATÁRIO", "EMISSOR", "PROTOCOLO-META", "EMISSÃO"]):
        c = tc.cell(0, i)
        styled_run(c.paragraphs[0], h, size=8, color=CINZA_5, bold=True, mono=True)
        set_cell_border(c, top=(12, "1B5E20"))
    vals = ["Grupo Z+Z", "Hábilis", "26/06/2026", "22/04/2026"]
    for i, v in enumerate(vals):
        c = tc.cell(1, i)
        styled_run(c.paragraphs[0], v, size=10, color=CINZA_9, bold=True, mono=i >= 2)

    add_page_break(doc)

    heading_section(doc, "01", "Panorama", space_before=0)

    lede(doc,
         "A LAO 033/2020 do Auto Posto América vence em 23/12/2026. A Hábilis "
         "conduz a renovação com protocolo-meta em 26/06/2026 (180 dias antes do "
         "vencimento), preservando folga regulatória para eventuais exigências "
         "complementares da SEMMA. A base documental recebida em 14/04/2026 "
         "suporta o início da execução sem fato impeditivo identificado.",
         ["23/12/2026", "26/06/2026", "180 dias antes", "14/04/2026"])

    heading_h3(doc, "Quatro frentes integradas | 74 itens")
    t = table_hdr(doc, ["Frente", "Itens", "Órgão", "Escopo"], [4.0, 1.5, 3.2, 8.5])
    add_row(t, ["Uso do Solo", "11", "Prefeitura Guapó", "Pré-requisito para protocolo da renovação."], bold_cols=[0, 1], mono_cols=[1])
    add_row(t, ["Renovação LAO", "26", "SEMMA (Proc. 22760/2020)", "Dossiê técnico principal do processo."], zebra=True, bold_cols=[0, 1], mono_cols=[1])
    add_row(t, ["Condicionantes LAO 033/2020", "14", "SEMMA", "15 atendidas | 3 parciais | 8 pendentes."], bold_cols=[0, 1], mono_cols=[1])
    add_row(t, ["Compliance Paralelo", "23", "ANP | IBAMA | NR-20 | PNRS", "Cadência própria (detalhe no anexo)."], zebra=True, bold_cols=[0, 1], mono_cols=[1])
    add_row(t, ["Total", "74", "4 órgãos/frentes", "Plano único, execução coordenada."], bold_cols=[0, 1, 2, 3], mono_cols=[1], is_total=True)

    # ====================================================================
    # SECAO 02 · CRONOGRAMA (fluxo continuo)
    # ====================================================================

    heading_section(doc, "02", "Cronograma | 4 fases em 67 dias", space_before=12)

    body_p(doc,
           "Janela de execução de 67 dias, de 20/04/2026 a 26/06/2026. Cada fase tem "
           "entregáveis próprios e depende de ações do Grupo Z+Z (seção 03). Atrasos em "
           "F1 comprimem F3 e F4 proporcionalmente.",
           ["67 dias", "20/04/2026", "26/06/2026"], size=9.5)

    t = table_hdr(doc, ["Fase", "Período", "Dias", "Entregáveis-chave"], [3.3, 3.4, 1.2, 9.3])
    add_row(t,
            ["F1 | Estruturação", "20/04 -> 08/05", "19",
             "Procuração, dados societários, contratação de terceiros, Uso do Solo protocolado, status outorga consolidado."],
            bold_cols=[0, 1, 2], mono_cols=[1, 2], size=8.5)
    add_row(t,
            ["F2 | Consolidação", "09/05 -> 30/05", "22",
             "Inspeção em campo, novas coletas laboratoriais, estanqueidade com INMETRO vigente, MCO atualizado."],
            zebra=True, bold_cols=[0, 1, 2], mono_cols=[1, 2], size=8.5)
    add_row(t,
            ["F3 | Execução e Evidência", "31/05 -> 19/06", "20",
             "Laudos técnicos assinados, condicionantes pendentes encerradas, dossiê fechado."],
            bold_cols=[0, 1, 2], mono_cols=[1, 2], size=8.5)
    add_row(t,
            ["F4 | Protocolo", "20/06 -> 26/06", "7",
             "Aprovação final, taxa paga, protocolo no Processo 22760/2020 | SEMMA."],
            zebra=True, bold_cols=[0, 1, 2], mono_cols=[1, 2], size=8.5)

    # ====================================================================
    # SECAO 03 · CHECKLIST CRONOLOGICO (fluxo continuo)
    # ====================================================================

    heading_section(doc, "03", "Checklist cronológico | 41 itens", space_before=12)

    body_p(doc,
           "Itens em ordem cronológica conforme documento apresentado em reunião. "
           "Responsável identificado por cor: Cliente em azul, Hábilis em verde e "
           "Terceiro em cinza. MCO e Anexo I/II CONAMA 273 são executados por Terceiro "
           "especializado (Eng. Civil/Ambiental) sob coordenação da Hábilis.",
           ["Cliente", "Hábilis", "Terceiro"], size=9.5)

    # Legenda de responsaveis (3 cores: sem mais 'misto')
    legenda = doc.add_table(rows=1, cols=3)
    legenda.alignment = WD_TABLE_ALIGNMENT.LEFT
    legenda_data = [
        ("Cliente", "cliente"), ("Hábilis", "habilis"), ("Terceiro", "terceiro"),
    ]
    for i, (label, kind) in enumerate(legenda_data):
        c = legenda.rows[0].cells[i]
        set_cell_shading(c, RESP_COLORS[kind])
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(p, label.upper(), size=8, color=BRANCO, bold=True, mono=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Tabela cronologica (sem checkbox, versao 'proposto')
    t = table_hdr(doc, ["Prazo", "Item", "Responsável", "Fase", "Observação"],
                  [1.5, 5.2, 2.8, 0.9, 7.0])

    itens_chrono = [
        ("22/04", "Coleta de RG/CPF dos sócios", "Cliente", "cliente", "F1",
         "RG e CPF de Sérgio Marques e José Guilherme. Exigência dos dois processos (Uso do Solo + LAO)."),
        ("23/04", "Procuração pública para a Hábilis", "Cliente", "cliente", "F1",
         "Firma reconhecida. Poderes para SEMMA, Prefeitura, SEMAD/GO e Veredas. Substitui a procuração A3 de 2023."),
        ("24/04", "Cartão CNPJ + Comprovante de Endereço", "Cliente", "cliente", "F1",
         "CNPJ ativo conferindo CNAE/endereço. Comprovante de endereço dos sócios nos últimos 3 meses."),
        ("25/04", "Verificação formal da Certidão de Uso do Solo", "Cliente", "cliente", "F1",
         "Confirmar existência, número, data e disponibilidade. Apoio da Hábilis."),
        ("27/04", "Enquadramento do imóvel + CAR", "Hábilis", "habilis", "F1",
         "Zona rural confirmada via LAO. CAR obrigatório. ART de Uso do Solo via Eng. Agronômica (CREA)."),
        ("27/04", "ART específica para Uso do Solo", "Hábilis", "habilis", "F1",
         "ART distinta da Renovação LAO. Assinada pela Hábilis via RT Eng. Agronômica."),
        ("28/04", "Diagnóstico de outorga no Web Outorga", "Hábilis", "habilis", "F1",
         "Consulta ao Sistema Web Outorga para histórico da DURH002737 (processos legados ainda não migrados ao Veredas). Protocolo novo será no Veredas em 13/05 (item 19). DURH002737 vencida 09/12/2023."),
        ("30/04", "Protocolo do requerimento da Certidão", "Hábilis", "habilis", "F1",
         "Somente se verificação indicar ausência ou desatualização."),
        ("05/05", "Contratação: investigação confirmatória", "Cliente", "cliente", "F1",
         "Empresa NBR 15515-2/3 e NBR 16209. ART geólogo/eng. ambiental. Indicação Hábilis."),
        ("08/05", "Análise documental consolidada", "Hábilis", "habilis", "F1",
         "Fechamento da triagem do lote de 14/04 e identificação de lacunas."),
        ("04/05", "Publicação do pedido de licenciamento", "Hábilis", "habilis", "F1",
         "Jornal local + Diário Oficial (CONAMA 006/1986). Aguardar 15-30 dias para contestação pública."),
        ("10/05", "Mapeamento técnico das 12 CPs com ação", "Hábilis", "habilis", "F1",
         "Identificação do tipo de intervenção por CP pendente (documental, física, contratação)."),
        ("22/04", "Validação do AVCB Cercon 161726", "Hábilis", "habilis", "F1",
         "Certificado CBMGO 161726, emitido 07/11/2025, válido até 20/10/2026. Em mãos."),
        ("12/05", "Inspeção técnica de campo", "Hábilis", "habilis", "F2",
         "25 pontos em 5 eixos: implantação/APP/RL/CAR, armazenamento, efluentes/SAO, resíduos, segurança."),
        ("25/05", "Anuência de saneamento ou laudo de fossa", "Hábilis", "habilis", "F2",
         "Rede pública: anuência da concessionária. Fossa/sumidouro: laudo técnico com ART."),
        ("12/05", "Validação das plantas em campo", "Hábilis", "habilis", "F2",
         "Plantas 05/2024 vs configuração atual. Quadro de Áreas + UTM SIRGAS 2000, APP e RL."),
        ("12/05", "Levantamento hídrico: define regime Veredas", "Hábilis", "habilis", "F2",
         "Rede pública: dispensa simplificada. Captação direta: processo completo com memorial técnico."),
        ("13/05", "Protocolo novo de outorga via Veredas", "Hábilis", "habilis", "F2",
         "Dia útil seguinte à inspeção. Gera certidão de tramitação que já compõe o dossiê SEMMA. IN 15/2026 §15."),
        ("15/05", "Contratação dos serviços de terceiros", "Cliente", "cliente", "F2",
         "Laboratório, estanqueidade, projetista e Eng. Civil/Ambiental (para MCO + Anexos CONAMA 273). Indicações da Hábilis."),
        ("18/05", "NFs e certificados dos equipamentos", "Cliente", "cliente", "F2",
         "NFs de tanques/tubulações + certificados CONAMA 319/2002 e INMETRO 37/109/110/111/009."),
        ("20/05", "Contratos de manutenção, limpeza e destinação", "Cliente", "cliente", "F2",
         "Ecofenix, Bauer, Limpmil. Validar ANP + licenciamento + autorização de transporte."),
        ("22/05", "Nova análise de efluentes", "Terceiro", "terceiro", "F2",
         "Laboratório indicado pela Hábilis. ART do químico responsável acompanha o laudo."),
        ("22/05", "Nova análise de água", "Terceiro", "terceiro", "F2",
         "Conjugada com análise de efluentes para otimizar coleta. ART químico."),
        ("25/05", "Mobilização da investigação confirmatória", "Terceiro", "terceiro", "F2",
         "Sondagens SPT, mínimo 3 furos. VOC, BTEX, PAH, metais, TPH. CEMAM 029/2018 + CONAMA 420/2009."),
        ("28/05", "Atualização do MCO", "Terceiro", "terceiro", "F2",
         "Peça central do dossiê (Anexo IX CEMAM 029/2018). Eng. Civil/Ambiental subcontratado pela Hábilis. ART do profissional contratado."),
        ("29/05", "Anexo I e II da CONAMA 273/2000", "Terceiro", "terceiro", "F2",
         "Anexo I tabular + Anexo II classifica risco (3 classes). ART própria (item 23 TR). Eng. Civil/Ambiental subcontratado pela Hábilis."),
        ("30/05", "Plano das 28 condicionantes", "Hábilis", "habilis", "F2",
         "Documento interno: responsável, evidência e prazo por condicionante."),
        ("Conting.", "Teste de estanqueidade (contingência)", "Terceiro", "terceiro", "F3",
         "Acionado apenas se SEMMA contestar laudo 10/2024. NBR 16795/2019. Empresa INMETRO vigente."),
        ("12/06", "Recomposição das 10 CPs parciais", "Hábilis", "habilis", "F3",
         "Recomposição documental conforme Plano de Atendimento de 30/05."),
        ("02/06", "PGRS · Gerenciamento de Resíduos", "Hábilis", "habilis", "F3",
         "Caracterização, rotas de destinação e responsáveis. ART RT Ambiental (CRBio) ou Agronômica."),
        ("03/06", "RCA · Relatório de Controle Ambiental", "Hábilis", "habilis", "F3",
         "Anexo XI CEMAM 029/2018. Análises SAO + poços. ART RT Ambiental (CRBio)."),
        ("05/06", "PAE · Atendimento a Emergências", "Hábilis", "habilis", "F3",
         "NBR 14276 + 15219. 9 capítulos, 10 cenários específicos. ART RT Ambiental."),
        ("10/06", "Análise de Risco (contingência)", "Terceiro", "terceiro", "F3",
         "Acionada se exigida pela SEMMA. APP: 15-30d. ARQ: 60-90d."),
        ("15/06", "Laudo de Investigação Confirmatória", "Terceiro", "terceiro", "F3",
         "CEMAM 029/2018 + CONAMA 420/2009. VOC, BTEX, PAH, metais, TPH + análise de riscos."),
        ("15/06", "Execução das 2 CPs pendentes", "Terceiro", "terceiro", "F3",
         "Serviços físicos ou ensaios: CP 4.4 (RCA semestral) + CP 4.14 (PGA/PCA a validar)."),
        ("19/06", "Relatório técnico consolidado", "Hábilis", "habilis", "F3",
         "Peça central do dossiê. Integra MCO, laudos, plantas e plano de condicionantes. ART do RT coordenador CRBio."),
        ("22/06", "Aprovação formal do relatório", "Cliente", "cliente", "F4",
         "Aprovação escrita do Grupo Z+Z antes do protocolo."),
        ("23/06", "CND Municipal + CND Estadual", "Cliente", "cliente", "F4",
         "Certidões Negativas de Débitos emitidas próximas ao protocolo. Validade 90 dias."),
        ("24/06", "Checklist final cruzado", "Hábilis", "habilis", "F4",
         "Conferência de peças, assinaturas, ARTs e taxas antes da submissão."),
        ("25/06", "Pagamento das taxas de protocolo", "Cliente", "cliente", "F4",
         "Taxas SEMMA, municipais e demais custos da submissão."),
        ("26/06", "Protocolo no Processo SEMMA 22760/2020", "Hábilis", "habilis", "F4",
         "Submissão do dossiê completo. Acompanhamento pós-protocolo se inicia."),
    ]

    for prazo, item, resp, resp_kind, fase, obs in itens_chrono:
        row = t.add_row()
        # Prazo
        c = row.cells[0]
        pc = c.paragraphs[0]
        pc.paragraph_format.space_before = Pt(1); pc.paragraph_format.space_after = Pt(1)
        styled_run(pc, prazo, size=8, color=CINZA_9, bold=True, mono=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        # Item
        c = row.cells[1]
        pc = c.paragraphs[0]
        pc.paragraph_format.space_before = Pt(1); pc.paragraph_format.space_after = Pt(1)
        styled_run(pc, item, size=8.5, color=CINZA_9, bold=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        # Responsavel (celula colorida)
        c = row.cells[2]
        set_cell_shading(c, RESP_COLORS[resp_kind])
        pc = c.paragraphs[0]
        pc.paragraph_format.space_before = Pt(1); pc.paragraph_format.space_after = Pt(1)
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(pc, resp.upper(), size=7.5, color=BRANCO, bold=True, mono=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Fase
        c = row.cells[3]
        pc = c.paragraphs[0]
        pc.paragraph_format.space_before = Pt(1); pc.paragraph_format.space_after = Pt(1)
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(pc, fase, size=8, color=CINZA_9, bold=True, mono=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Observacao
        c = row.cells[4]
        pc = c.paragraphs[0]
        pc.paragraph_format.space_before = Pt(1); pc.paragraph_format.space_after = Pt(1)
        styled_run(pc, obs, size=8, color=CINZA_7)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # ====================================================================
    # ANEXO A1 · COMPLIANCE PARALELO (fluxo continuo)
    # ====================================================================

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(0)
    styled_run(p, "ANEXO | CADÊNCIA PRÓPRIA", size=9, color=VERDE, bold=True, mono=True)

    heading_section(doc, "A1", "Compliance Paralelo | 23 itens", space_before=4)

    lede(doc,
         "Frente complementar à renovação da LAO, com cadência própria. Itens aqui não "
         "integram o dossiê SEMMA, mas falhas nesta frente contaminam o processo principal "
         "(ex.: RAPP IBAMA em atraso impede emissão de certidões exigidas pela SEMMA; NR-20 "
         "irregular bloqueia operação). Execução em paralelo reduz risco e evita retrabalho.",
         ["cadência própria", "contaminam o processo principal"])

    heading_h3(doc, "Sumário por órgão | 23 itens", space_before=4)
    t = table_hdr(doc, ["Órgão | Norma", "Itens", "Natureza"], [5.5, 1.5, 10.2])
    add_row(t, ["ANP", "7", "Cadastro, revenda, livro de movimentação, relatórios comerciais, comprovação de distribuidora ativa."], bold_cols=[0, 1], mono_cols=[1])
    add_row(t, ["IBAMA | CTF/APP + RAPP", "4", "Cadastro Técnico Federal ativo e Relatório Anual de Atividades Potencialmente Poluidoras entregue."], zebra=True, bold_cols=[0, 1], mono_cols=[1])
    add_row(t, ["NR-20 | Segurança do Trabalho", "7", "PGR, PCMSO, treinamento NR-20 válido por função, inspeção SESMT e plano de emergência."], bold_cols=[0, 1], mono_cols=[1])
    add_row(t, ["PNRS | Resíduos Sólidos", "5", "Plano de Gerenciamento de Resíduos, MTR, contrato com transportador licenciado e destino final."], zebra=True, bold_cols=[0, 1], mono_cols=[1])
    add_row(t, ["Total", "23", "Rodando em paralelo a F1-F4. Diagnóstico individual em documento próprio."], bold_cols=[0, 1, 2], mono_cols=[1], is_total=True)

    heading_h3(doc, "O que depende do Grupo Z+Z | compliance", space_before=8)
    bullet(doc, "Acesso completo aos cadastros ANP (login/senha ou procuração com poderes específicos).", ["cadastros ANP"])
    bullet(doc, "Acesso ao CTF/APP IBAMA (senha ou nomeação de preposto com poderes de consulta e entrega).", ["CTF/APP IBAMA"])
    bullet(doc, "Documentos do SESMT ou contrato com serviço terceirizado de segurança do trabalho: PGR, PCMSO, ASO, ficha de EPI.", ["PGR", "PCMSO"])
    bullet(doc, "Contratos e MTRs vigentes com transportadores e destinadores de resíduos licenciados.", ["MTRs"])
    bullet(doc, "Liberação orçamentária para treinamentos NR-20 faltantes e eventual regularização de PGR/PCMSO.", ["NR-20"])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_shading(p, AZUL_BG)
    set_paragraph_border_left(p, 2, "0D47A1")
    p.paragraph_format.left_indent = Cm(0.25)
    styled_run(p, "Duas decisões requeridas neste anexo: ", size=9, color=CINZA_9, bold=True)
    styled_run(p, "aprovação do escopo de 23 itens + liberação do diagnóstico inicial, em paralelo a F1.", size=9, color=CINZA_7)

    here = os.path.dirname(os.path.abspath(__file__))
    out = os.path.join(here, "Alinhamento_Executivo_AutoPostoAmerica_v1.docx")
    doc.save(out)
    print(f"OK -> {out}")
    return out


if __name__ == "__main__":
    build()
