#!/usr/bin/env python3
"""
Briefing para Cotacao de Terceiros v3 · Auto Posto America.
Peca exclusiva para direcionar a pessoa responsavel pelas cotacoes.
Contem: mensagem pronta + tabela geral v3 + fichas tecnicas por item.
Alinhada ao modelo de 4 camadas do Alinhamento Executivo v3.
Uso: python3 build_briefing_cotacao_v3.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Paleta Habilis
VERDE = RGBColor(0x1B, 0x5E, 0x20)
VERDE_L = RGBColor(0x4C, 0xAF, 0x50)
VERDE_BG = "E8F5E9"
AMBER = RGBColor(0xF5, 0x7F, 0x17)
AMBER_BG = "FFF8E1"
VERMELHO = RGBColor(0xB7, 0x1C, 0x1C)
VERMELHO_BG = "FFEBEE"
AZUL = RGBColor(0x0D, 0x47, 0xA1)
AZUL_BG = "E3F2FD"
ROXO = RGBColor(0x4A, 0x14, 0x8C)
ROXO_BG = "F3E5F5"
CINZA_9 = RGBColor(0x1C, 0x1C, 0x1E)
CINZA_7 = RGBColor(0x3A, 0x3A, 0x3C)
CINZA_5 = RGBColor(0x6C, 0x6C, 0x70)
CINZA_BG = "FAFAFA"
CINZA_MSG = "F4F4F4"
BRANCO = RGBColor(0xFF, 0xFF, 0xFF)
FONT_SANS = "Segoe UI"
FONT_MONO = "Consolas"

# Cores das classificacoes
CLASSIF_COLORS = {
    "obrig":   "1B5E20",  # verde
    "prov":    "F57F17",  # amber
    "cont":    "F57F17",  # amber
    "opc":     "0D47A1",  # azul
}
CLASSIF_LABEL = {
    "obrig":   "OBRIG. IMEDIATO",
    "prov":    "PROVÁVEL",
    "cont":    "CONTINGENTE",
    "opc":     "OPCIONAL",
}


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for name, val in edges.items():
        if val is None: continue
        size, color = val
        el = OxmlElement(f'w:{name}')
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(size)); el.set(qn('w:color'), color)
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def set_paragraph_shading(p, hex_color):
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    p_pr.append(shd)


def set_paragraph_border_left(p, size_pt, color_hex):
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), str(int(size_pt * 8)))
    left.set(qn('w:space'), '6'); left.set(qn('w:color'), color_hex)
    pbdr.append(left); p_pr.append(pbdr)


def styled_run(p, text, *, bold=False, italic=False, size=None, color=None, mono=False):
    run = p.add_run(text)
    font_name = FONT_MONO if mono else FONT_SANS
    run.font.name = font_name
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.find(qn('w:rFonts'))
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts'); r_pr.append(r_fonts)
    r_fonts.set(qn('w:ascii'), font_name); r_fonts.set(qn('w:hAnsi'), font_name); r_fonts.set(qn('w:cs'), font_name)
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    return run


def mixed_run(p, text, *, size=10, default_color=CINZA_7, bold_parts=None):
    if not bold_parts:
        styled_run(p, text, size=size, color=default_color); return
    remaining = text
    for target in bold_parts:
        idx = remaining.find(target)
        if idx == -1: continue
        if idx > 0: styled_run(p, remaining[:idx], size=size, color=default_color)
        styled_run(p, target, size=size, color=CINZA_9, bold=True)
        remaining = remaining[idx + len(target):]
    if remaining:
        styled_run(p, remaining, size=size, color=default_color)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def heading_section(doc, num, title, space_before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '14')
    bottom.set(qn('w:space'), '3'); bottom.set(qn('w:color'), '1B5E20')
    pbdr.append(bottom); p_pr.append(pbdr)
    styled_run(p, f"{num}  ", size=11, color=VERDE_L, mono=True)
    styled_run(p, title, size=15, color=VERDE, bold=True)


def heading_h3(doc, text, space_before=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    set_paragraph_border_left(p, 2, "1B5E20")
    p.paragraph_format.left_indent = Cm(0.2)
    styled_run(p, text, size=10.5, color=CINZA_9, bold=True)


def body_p(doc, text, bold_parts=None, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    mixed_run(p, text, bold_parts=bold_parts, size=size)


def table_hdr(doc, headers, widths_cm=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_shading(c, "1B5E20")
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        styled_run(p, h.upper(), size=8, color=BRANCO, bold=True, mono=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if widths_cm and i < len(widths_cm): c.width = Cm(widths_cm[i])
    return t


def add_row(t, values, zebra=False, bold_cols=None, mono_cols=None, is_total=False, size=9):
    row = t.add_row()
    for i, val in enumerate(values):
        c = row.cells[i]
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
        bold = (bold_cols and i in bold_cols) or is_total
        mono = mono_cols and i in mono_cols
        styled_run(p, val, size=size, color=CINZA_9 if bold else CINZA_7, bold=bold, mono=mono)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        if is_total: set_cell_shading(c, VERDE_BG)
        elif zebra: set_cell_shading(c, CINZA_BG)


def bullet(doc, text, bold_parts=None, size=10):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    mixed_run(p, text, bold_parts=bold_parts, size=size)


def field_page(p):
    r = p.add_run()
    r.font.name = FONT_MONO; r.font.size = Pt(8); r.font.color.rgb = CINZA_5
    fld_b = OxmlElement('w:fldChar'); fld_b.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.text = 'PAGE'
    fld_e = OxmlElement('w:fldChar'); fld_e.set(qn('w:fldCharType'), 'end')
    r._r.append(fld_b); r._r.append(it); r._r.append(fld_e)


def field_numpages(p):
    r = p.add_run()
    r.font.name = FONT_MONO; r.font.size = Pt(8); r.font.color.rgb = CINZA_5
    fld_b = OxmlElement('w:fldChar'); fld_b.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.text = 'NUMPAGES'
    fld_e = OxmlElement('w:fldChar'); fld_e.set(qn('w:fldCharType'), 'end')
    r._r.append(fld_b); r._r.append(it); r._r.append(fld_e)


def mensagem_pronta_box(doc):
    """Caixa destacada com mensagem pronta para copiar e enviar."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
    styled_run(p, "MENSAGEM PRONTA · COPIAR E ENVIAR", size=9, color=VERDE, bold=True, mono=True)

    tb = doc.add_table(rows=1, cols=1)
    c = tb.cell(0, 0)
    set_cell_shading(c, CINZA_MSG)
    set_cell_border(c, left=(32, "1B5E20"), top=(6, "E0E0E0"),
                    right=(6, "E0E0E0"), bottom=(6, "E0E0E0"))

    linhas = [
        "Oi, tudo bem?",
        "",
        "Preciso da sua ajuda com cotações para o processo de renovação da "
        "Licença Ambiental de Operação (LAO) do Auto Posto América, em Guapó-GO. "
        "Processo SEMMA 22760/2020 · protocolo-meta 26/06/2026.",
        "",
        "Montei um briefing técnico com 14 itens que precisam ser cotados junto a "
        "terceiros. Cada item tem ficha própria com: fundamento normativo, escopo "
        "da entrega, perfil do prestador exigido, prazo-alvo, quantidades/parâmetros "
        "e observações para cotação. Os valores ficam em branco — é o que você vai "
        "preencher conforme for recebendo as propostas.",
        "",
        "Os itens estão classificados em 4 gatilhos:",
        "· OBRIG. IMEDIATO — contrata já (MCO, Anexos CONAMA 273, publicação).",
        "· PROVÁVEL — contrata em seguida, muito provavelmente será acionado (RCA, análises).",
        "· CONTINGENTE — cotar mesmo assim, mas só aciona se houver gatilho "
        "formal (investigação confirmatória, análise de risco, estanqueidade extra, "
        "adequação física).",
        "· OPCIONAL — cotar por garantia, aciona só se o órgão pedir.",
        "",
        "Como preciso que volte:",
        "1. Tabela geral preenchida (R$ + nome do prestador + validade da proposta).",
        "2. Mínimo 2 propostas por item sempre que possível (3 para os contingentes).",
        "3. Se algum prestador recusar ou não tiver acreditação exigida, me avisa.",
        "4. Prazos de resposta curtos: tenho que fechar a contratação dos "
        "OBRIG. IMEDIATO em F1 (até 08/05).",
        "",
        "Qualquer dúvida técnica sobre escopo ou norma, me chama direto — o "
        "fundamento está na ficha de cada item.",
        "",
        "Obrigado.",
    ]

    first = True
    for linha in linhas:
        if first:
            p = c.paragraphs[0]
            first = False
        else:
            p = c.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
        if linha == "":
            continue
        if linha.startswith("·") or (len(linha) > 2 and linha[0].isdigit() and linha[1] == "."):
            p.paragraph_format.left_indent = Cm(0.4)
        styled_run(p, linha, size=10, color=CINZA_9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def ficha_tecnica(doc, num, titulo, classif, dados):
    """
    Ficha tecnica de um item de cotacao.
    dados: dict com chaves:
      - fundamento (str): norma/CP que obriga
      - acionamento (str): quando acionar
      - escopo (str): o que o terceiro entrega
      - perfil (str): perfil tecnico do prestador
      - prazo (str): data ou janela alvo
      - quantidades (str): volumes/parametros (pode ser None)
      - docs_fornecidos (str): o que a Habilis fornece para o prestador
      - observacoes (str): notas para cotacao
    """
    # Titulo da ficha
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    set_paragraph_border_left(p, 2.5, "1B5E20")
    p.paragraph_format.left_indent = Cm(0.2)
    styled_run(p, f"Ficha {num:02d}  ", size=10, color=VERDE_L, bold=True, mono=True)
    styled_run(p, titulo, size=12, color=CINZA_9, bold=True)

    # Chip de classificacao
    chip_tbl = doc.add_table(rows=1, cols=1)
    chip_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cc = chip_tbl.cell(0, 0)
    set_cell_shading(cc, CLASSIF_COLORS[classif])
    cc.width = Cm(4.5)
    pc = cc.paragraphs[0]
    pc.paragraph_format.space_before = Pt(1); pc.paragraph_format.space_after = Pt(1)
    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(pc, CLASSIF_LABEL[classif], size=8, color=BRANCO, bold=True, mono=True)
    cc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

    # Tabela de dados (key-value)
    campos = [
        ("Fundamento normativo", dados["fundamento"]),
        ("Quando acionar", dados["acionamento"]),
        ("Escopo da entrega", dados["escopo"]),
        ("Perfil exigido do prestador", dados["perfil"]),
        ("Prazo-alvo", dados["prazo"]),
    ]
    if dados.get("quantidades"):
        campos.append(("Quantidades e parâmetros", dados["quantidades"]))
    if dados.get("docs_fornecidos"):
        campos.append(("Documentos que a Hábilis fornece", dados["docs_fornecidos"]))
    campos.append(("Observações para cotação", dados["observacoes"]))

    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for label, value in campos:
        row = t.add_row()
        c_l = row.cells[0]
        c_v = row.cells[1]
        c_l.width = Cm(4.8)
        c_v.width = Cm(12.6)
        set_cell_shading(c_l, "E8F5E9")
        set_cell_border(c_l, right=(4, "E0E0E0"), bottom=(4, "F0F0F0"))
        set_cell_border(c_v, bottom=(4, "F0F0F0"))
        pl = c_l.paragraphs[0]
        pl.paragraph_format.space_before = Pt(2); pl.paragraph_format.space_after = Pt(2)
        styled_run(pl, label.upper(), size=8, color=CINZA_9, bold=True, mono=True)
        c_l.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        pv = c_v.paragraphs[0]
        pv.paragraph_format.space_before = Pt(2); pv.paragraph_format.space_after = Pt(2)
        pv.paragraph_format.left_indent = Cm(0.15)
        styled_run(pv, value, size=9.5, color=CINZA_7)
        c_v.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Campos em branco para preencher
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(0)
    styled_run(p, "CAMPOS A PREENCHER APÓS COTAÇÃO", size=8, color=CINZA_5, bold=True, mono=True)

    fill = doc.add_table(rows=0, cols=2)
    fill.alignment = WD_TABLE_ALIGNMENT.LEFT
    for label in ["Prestador 1  ·  Nome / CNPJ",
                  "Prestador 1  ·  Valor (R$)",
                  "Prestador 1  ·  Validade da proposta",
                  "Prestador 2  ·  Nome / CNPJ",
                  "Prestador 2  ·  Valor (R$)",
                  "Prestador 2  ·  Validade da proposta"]:
        row = fill.add_row()
        c_l = row.cells[0]; c_v = row.cells[1]
        c_l.width = Cm(5.8); c_v.width = Cm(11.6)
        set_cell_shading(c_l, CINZA_BG)
        set_cell_border(c_l, right=(4, "E0E0E0"), bottom=(4, "F0F0F0"))
        set_cell_border(c_v, bottom=(4, "E0E0E0"))
        pl = c_l.paragraphs[0]
        pl.paragraph_format.space_before = Pt(1); pl.paragraph_format.space_after = Pt(1)
        styled_run(pl, label, size=8.5, color=CINZA_7, mono=True)
        pv = c_v.paragraphs[0]
        pv.paragraph_format.space_before = Pt(1); pv.paragraph_format.space_after = Pt(1)
        styled_run(pv, "                                                                        ",
                   size=9.5, color=CINZA_5)


def build():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(1.5); s.bottom_margin = Cm(1.5)
    s.left_margin = Cm(1.8); s.right_margin = Cm(1.8)

    # Header
    h_p = s.header.paragraphs[0]
    styled_run(h_p, "Auto Posto América | Briefing de Cotação v3 · USO INTERNO", size=8, color=CINZA_5, mono=True)
    h_p.add_run("\t")
    styled_run(h_p, "Hábilis | Abr/2026", size=8, color=VERDE, bold=True, mono=True)

    # Footer
    f_p = s.footer.paragraphs[0]
    styled_run(f_p, "Hábilis Regularização Ambiental · uso interno · não enviar ao cliente", size=8, color=CINZA_5, mono=True)
    f_p.add_run("\t")
    styled_run(f_p, "Pagina ", size=8, color=CINZA_5, mono=True)
    field_page(f_p)
    styled_run(f_p, " de ", size=8, color=CINZA_5, mono=True)
    field_numpages(f_p)

    # ====================================================================
    # CAPA
    # ====================================================================
    p = doc.add_paragraph()
    styled_run(p, "H  ", size=16, color=VERDE, bold=True)
    styled_run(p, "Hábilis", size=16, color=VERDE, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(40)
    styled_run(p, "REGULARIZAÇÃO AMBIENTAL · USO INTERNO", size=9, color=CINZA_5, mono=True)
    for _ in range(3): doc.add_paragraph()

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(p, "BRIEFING PARA COTAÇÃO DE TERCEIROS", size=10, color=VERDE, bold=True, mono=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(p, "Auto Posto América", size=28, color=CINZA_9, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(p, "14 itens · 4 gatilhos · campos em branco para preenchimento",
               size=12, color=CINZA_7)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    styled_run(p,
               "Peça de uso exclusivamente interno Hábilis. Destinada à pessoa responsável "
               "por levantar cotações junto a prestadores de serviço para a renovação da LAO "
               "033/2020 do Auto Posto América (Processo SEMMA 22760/2020). Contém: ",
               size=11, color=CINZA_7)
    styled_run(p, "mensagem pronta para envio · tabela geral dos 14 itens · 14 fichas "
               "técnicas com fundamento, escopo e perfil do prestador · campos em branco "
               "para preenchimento dos valores e prestadores.",
               size=11, color=CINZA_9, bold=True)

    for _ in range(1): doc.add_paragraph()

    tc = doc.add_table(rows=2, cols=4)
    for i, h in enumerate(["DESTINATÁRIO", "EMISSOR", "PROTOCOLO-META", "EMISSÃO"]):
        c = tc.cell(0, i)
        styled_run(c.paragraphs[0], h, size=8, color=CINZA_5, bold=True, mono=True)
        set_cell_border(c, top=(12, "1B5E20"))
    vals = ["Cotador interno", "Hábilis", "26/06/2026", "23/04/2026"]
    for i, v in enumerate(vals):
        c = tc.cell(1, i)
        styled_run(c.paragraphs[0], v, size=10, color=CINZA_9, bold=True, mono=i >= 2)

    # Aviso de confidencialidade
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_shading(p, AMBER_BG)
    set_paragraph_border_left(p, 2.5, "F57F17")
    p.paragraph_format.left_indent = Cm(0.25)
    styled_run(p, "Uso interno · ", size=9.5, color=CINZA_9, bold=True)
    styled_run(p,
               "Este briefing não deve ser enviado ao Grupo Z+Z nem a prestadores. "
               "É ferramenta de trabalho do cotador. A mensagem da próxima página é o "
               "único conteúdo pensado para ser repassado (à pessoa de cotação, não a "
               "prestadores diretos).",
               size=9.5, color=CINZA_7)

    add_page_break(doc)

    # ====================================================================
    # MENSAGEM PRONTA
    # ====================================================================
    heading_section(doc, "01", "Mensagem pronta", space_before=0)

    body_p(doc,
           "Texto pronto para copiar e enviar à pessoa responsável pelas cotações (WhatsApp, "
           "e-mail ou chat interno). Fecha o contexto, define o formato de retorno e delega "
           "com clareza. Edite o cumprimento e o fechamento conforme a relação.",
           ["copiar e enviar", "formato de retorno"], size=10)

    mensagem_pronta_box(doc)

    add_page_break(doc)

    # ====================================================================
    # TABELA GERAL
    # ====================================================================
    heading_section(doc, "02", "Tabela geral · 14 itens", space_before=0)

    body_p(doc,
           "Mesma tabela do Alinhamento Executivo v3, seção 03. Visão condensada para rápida "
           "consulta. Detalhe completo de cada item nas fichas técnicas da seção 03.",
           ["14 itens", "fichas técnicas"], size=10)

    heading_h3(doc, "Legenda de classificação", space_before=4)
    legenda = doc.add_table(rows=1, cols=4)
    legenda.alignment = WD_TABLE_ALIGNMENT.LEFT
    legenda_data = [
        ("Obrigatório imediato", CLASSIF_COLORS["obrig"]),
        ("Provável", CLASSIF_COLORS["prov"]),
        ("Contingente", CLASSIF_COLORS["cont"]),
        ("Opcional", CLASSIF_COLORS["opc"]),
    ]
    for i, (label, color) in enumerate(legenda_data):
        c = legenda.rows[0].cells[i]
        set_cell_shading(c, color)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(p, label.upper(), size=7.5, color=BRANCO, bold=True, mono=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    t = table_hdr(doc,
                  ["#", "Item", "Classificação", "Faixa de Custo", "Condição de acionamento", "Fase"],
                  [0.7, 5.0, 2.4, 2.4, 5.2, 0.9])

    itens_geral = [
        ("01", "Eng. Civil/Ambiental · ART do MCO", "obrig",
         "Literal TR SEMMA. Hábilis subcontrata sob coordenação CRBio.", "F2"),
        ("02", "Eng. Civil/Ambiental · ART Anexos I/II CONAMA 273", "obrig",
         "Item 23 TR SEMMA. Mesmo prestador do MCO.", "F2"),
        ("03", "Laboratório INMETRO · RCA semestral", "prov",
         "Obrigatório para CP 4.4. Contratação após diagnóstico 12/05.", "F2"),
        ("04", "Laboratório · análises de efluentes/água", "prov",
         "Após inspeção 12/05 validar conformidade SAO.", "F2"),
        ("05", "Publicação CONAMA 006/86", "obrig",
         "Item 6 TR SEMMA. Veículo oficial + jornal regional.", "F1"),
        ("06", "Investigação confirmatória · NBR 15515/16209", "cont",
         "Se contestação pública ou passivo identificado.", "F2/F3"),
        ("07", "Sondagens SPT (mín. 3 furos)", "cont",
         "Componente da investigação confirmatória.", "F2/F3"),
        ("08", "Análise de Risco (APP ou ARQ)", "cont",
         "Se exigida pela SEMMA.", "F3"),
        ("09", "Teste de estanqueidade extra · NBR 16795", "cont",
         "Se SEMMA contestar laudo 10/2024.", "F3"),
        ("10", "Adequação física da SAO", "cont",
         "Se inspeção 12/05 apontar não-conformidade.", "F2/F3"),
        ("11", "Laudo de saneamento ou fossa/sumidouro", "cont",
         "Rede pública: anuência (sem custo). Fossa: laudo + ART.", "F2"),
        ("12", "Recontrato Bauer, Limpmil, Ecofenix", "prov",
         "Validação de contratos vigentes · CP 4.5/4.6/4.7.", "F2"),
        ("13", "Levantamento topográfico complementar", "opc",
         "Se Veredas exigir memorial adicional.", "F2"),
        ("14", "Projeto de adequação paisagística/APP", "opc",
         "Se inspeção 12/05 identificar APP comprometida.", "F3"),
    ]

    for num, item, klass, cond, fase in itens_geral:
        row = t.add_row()
        # numero
        c = row.cells[0]
        pc = c.paragraphs[0]
        pc.paragraph_format.space_before = Pt(1); pc.paragraph_format.space_after = Pt(1)
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(pc, num, size=8.5, color=CINZA_9, bold=True, mono=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # item
        c = row.cells[1]
        pc = c.paragraphs[0]
        pc.paragraph_format.space_before = Pt(1); pc.paragraph_format.space_after = Pt(1)
        styled_run(pc, item, size=8.5, color=CINZA_9, bold=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        # classificacao chip
        c = row.cells[2]
        set_cell_shading(c, CLASSIF_COLORS[klass])
        pc = c.paragraphs[0]
        pc.paragraph_format.space_before = Pt(1); pc.paragraph_format.space_after = Pt(1)
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(pc, CLASSIF_LABEL[klass], size=7, color=BRANCO, bold=True, mono=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # faixa custo
        c = row.cells[3]
        set_cell_shading(c, CINZA_BG)
        pc = c.paragraphs[0]
        pc.paragraph_format.space_before = Pt(1); pc.paragraph_format.space_after = Pt(1)
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(pc, "R$ ______________", size=8, color=CINZA_5, mono=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # condicao
        c = row.cells[4]
        pc = c.paragraphs[0]
        pc.paragraph_format.space_before = Pt(1); pc.paragraph_format.space_after = Pt(1)
        styled_run(pc, cond, size=8, color=CINZA_7)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        # fase
        c = row.cells[5]
        pc = c.paragraphs[0]
        pc.paragraph_format.space_before = Pt(1); pc.paragraph_format.space_after = Pt(1)
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(pc, fase, size=8, color=CINZA_9, bold=True, mono=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    add_page_break(doc)

    # ====================================================================
    # FICHAS TECNICAS (14)
    # ====================================================================
    heading_section(doc, "03", "Fichas técnicas · especificação para cotação", space_before=0)

    body_p(doc,
           "Cada ficha traz fundamento normativo, escopo da entrega, perfil exigido do "
           "prestador, prazo-alvo, quantidades/parâmetros e observações para cotação. "
           "Campos em branco ao final de cada ficha para preenchimento de até 2 propostas "
           "(Prestador 1 e Prestador 2). Para itens CONTINGENTE e OPCIONAL, cote normalmente: "
           "o valor fica pronto caso haja gatilho de acionamento.",
           ["fundamento normativo", "até 2 propostas", "cote normalmente"], size=10)

    # Ficha 01 · MCO
    ficha_tecnica(doc, 1,
                  "Eng. Civil/Ambiental · ART do MCO (Memorial de Caracterização do Empreendimento)",
                  "obrig",
                  {
                      "fundamento": "Anexo IX CEMAM 029/2018 · Termo de Referência SEMMA Guapó "
                                    "para renovação de LAO de posto de combustível.",
                      "acionamento": "Imediato. É peça central do dossiê SEMMA. Sem MCO, não "
                                     "há protocolo.",
                      "escopo": "Elaboração do Memorial de Caracterização do Empreendimento "
                                "(MCO) completo, conforme modelo Anexo IX CEMAM 029/2018: "
                                "identificação do empreendimento, caracterização da atividade, "
                                "descrição dos processos, equipamentos, armazenamento, fluxos "
                                "de efluentes/resíduos, sistema SAO, poços de monitoramento, "
                                "plantas e croquis. Emissão da ART CREA vinculada. Arquivo "
                                "final entregue em .docx + .pdf + ART assinada.",
                      "perfil": "Engenheiro(a) Civil ou Ambiental com CREA-GO ativo e "
                                "experiência comprovada em MCO para postos de combustível. "
                                "Trabalha sob coordenação da CRBio responsável técnica da "
                                "Hábilis (modelo de subcontratação padrão).",
                      "prazo": "Entrega até 28/05/2026. Contratação efetivada até 08/05/2026 "
                               "(F1).",
                      "quantidades": "1 peça. Uma revisão pós-inspeção de campo 12/05 "
                                     "incluída no escopo.",
                      "docs_fornecidos": "LAO 033/2020 e anexos · plantas .dwg 05/2024 · "
                                         "levantamento documental lote 14/04 · relatório de "
                                         "inspeção técnica 12/05 · laudo de estanqueidade "
                                         "10/2024 · histórico RCA.",
                      "observacoes": "Pacote com a Ficha 02 (Anexos I/II CONAMA 273) "
                                     "tende a ter desconto. Solicitar cotação em pacote "
                                     "único e em separado para comparação. Verificar se o "
                                     "prestador emite ART individualizada por peça ou uma "
                                     "ART consolidada (ambas aceitas pela SEMMA).",
                  })

    # Ficha 02 · Anexos CONAMA 273
    ficha_tecnica(doc, 2,
                  "Eng. Civil/Ambiental · ART dos Anexos I e II da Resolução CONAMA 273/2000",
                  "obrig",
                  {
                      "fundamento": "Resolução CONAMA 273/2000 · Item 23 do Termo de "
                                    "Referência SEMMA.",
                      "acionamento": "Imediato. Item expresso no TR. Mesmo prestador da Ficha 01.",
                      "escopo": "Elaboração do Anexo I (formulário tabular de caracterização) "
                                "e do Anexo II (classificação de risco em 3 classes conforme "
                                "critério CONAMA 273). Emissão da ART específica. Entrega em "
                                ".docx + .pdf + ART.",
                      "perfil": "Mesmo da Ficha 01. Engenheiro(a) Civil ou Ambiental CREA-GO "
                                "ativo.",
                      "prazo": "Entrega até 29/05/2026.",
                      "quantidades": "2 anexos (I e II). ART pode ser única para o conjunto "
                                     "MCO + Anexos, ou uma por peça.",
                      "docs_fornecidos": "Os mesmos da Ficha 01. Peças se apoiam no MCO.",
                      "observacoes": "Fechar em pacote com Ficha 01. O Anexo II demanda "
                                     "avaliação técnica de risco baseada em idade dos tanques, "
                                     "histórico operacional e sensibilidade ambiental do "
                                     "entorno — o prestador precisa ter dados da Ficha 01 "
                                     "antes de finalizar.",
                  })

    # Ficha 03 · Laboratório RCA
    ficha_tecnica(doc, 3,
                  "Laboratório acreditado INMETRO · RCA semestral",
                  "prov",
                  {
                      "fundamento": "CP 4.4 LAO 033/2020 · Anexo XI CEMAM 029/2018 · "
                                    "ABNT NBR ISO/IEC 17025.",
                      "acionamento": "Provável. O diagnóstico de campo de 12/05 define o "
                                     "escopo exato de parâmetros. Contratação após 12/05.",
                      "escopo": "Coleta e análise físico-química e cromatográfica de amostras "
                                "de: (a) efluente do sistema separador água-óleo (SAO) na "
                                "saída; (b) poços de monitoramento existentes. Laudo final com "
                                "ART do químico responsável, comparação com padrões CONAMA "
                                "430/2011 (efluentes) e Portaria MS 888/2021 (água). "
                                "Metodologia rastreada e certificado de acreditação anexados.",
                      "perfil": "Laboratório com acreditação ABNT NBR ISO/IEC 17025 vigente "
                                "pelo INMETRO, escopo contendo os parâmetros listados. "
                                "Coleta realizada por técnico do próprio laboratório (não "
                                "aceitar amostra terceirizada sem cadeia de custódia).",
                      "prazo": "Coleta em 22/05/2026. Laudo entregue até 03/06/2026 para "
                               "consolidar no RCA.",
                      "quantidades": "Parâmetros em SAO: pH, turbidez, óleos e graxas (O&G), "
                                     "sólidos suspensos, DBO5, DQO, OD, surfactantes (MBAS), "
                                     "metais pesados (Pb, Zn, Cr, Cd, Ni, Cu, Fe), TPH total "
                                     "e frações. Parâmetros em poços de monitoramento: BTEX "
                                     "(benzeno, tolueno, etilbenzeno, xilenos), PAH "
                                     "(hidrocarbonetos policíclicos aromáticos), TPH, metais.",
                      "docs_fornecidos": "Plantas com localização dos poços e ponto de coleta "
                                         "SAO · histórico de laudos anteriores · Anexo XI "
                                         "CEMAM 029/2018.",
                      "observacoes": "Cotar separadamente: (a) pacote completo (SAO + 3 poços) "
                                     "e (b) unitário por ponto, para flexibilidade se um poço "
                                     "estiver obstruído. Solicitar prazo de emissão do laudo "
                                     "a partir da data da coleta (crítico para cumprir 03/06).",
                  })

    # Ficha 04 · Análises efluentes/água complementares
    ficha_tecnica(doc, 4,
                  "Laboratório · nova análise de efluentes e água (subsidia RCA)",
                  "prov",
                  {
                      "fundamento": "Apoia a elaboração do RCA (CP 4.4) e o diagnóstico de "
                                    "adequação do sistema SAO identificado na inspeção 12/05.",
                      "acionamento": "Provável. Coleta em 22/05 (conjugada com a coleta do "
                                     "RCA, Ficha 03, para otimizar). Só se prestador do RCA "
                                     "não cobrir escopo complementar.",
                      "escopo": "Coleta complementar de efluentes e água subterrânea para "
                                "validar conformidade da SAO pós-diagnóstico 12/05. Laudo "
                                "com ART do químico.",
                      "perfil": "Mesmo perfil da Ficha 03. Preferencialmente o mesmo "
                                "laboratório (otimiza custo e logística de coleta).",
                      "prazo": "Coleta 22/05. Laudo até 30/05.",
                      "quantidades": "A definir pós-inspeção 12/05. Estimativa: até 2 pontos "
                                     "adicionais, mesmos parâmetros da Ficha 03.",
                      "docs_fornecidos": "Relatório de inspeção 12/05 · plantas.",
                      "observacoes": "Item pode ser absorvido no pacote da Ficha 03. Cotar "
                                     "como adicional.",
                  })

    # Ficha 05 · Publicação CONAMA 006/86
    ficha_tecnica(doc, 5,
                  "Publicação do pedido de licenciamento · CONAMA 006/1986",
                  "obrig",
                  {
                      "fundamento": "Resolução CONAMA 006/1986 · Item 6 do Termo de "
                                    "Referência SEMMA Guapó.",
                      "acionamento": "Imediato. Obrigatório em F1. Abre janela de 15-30 dias "
                                     "para contestação pública, que é crítica para gatilhos "
                                     "de itens contingentes (Ficha 06).",
                      "escopo": "Publicação do aviso padrão de requerimento de licença "
                                "ambiental em: (a) Diário Oficial do Estado de Goiás (DOE/GO) "
                                "e (b) jornal de circulação regional em Guapó-GO. Texto "
                                "conforme modelo oficial. Entrega das publicações originais "
                                "(recortes digitalizados e exemplares físicos).",
                      "perfil": "Serviço contratado diretamente com o Diário Oficial de "
                                "Goiás e com o jornal escolhido. Agência de publicações "
                                "legais pode intermediar (acelera e padroniza).",
                      "prazo": "Publicação em 04/05/2026 simultânea nos dois veículos. "
                               "Entregas físicas/digitais até 10/05.",
                      "docs_fornecidos": "Texto padrão do aviso (Hábilis fornece) · dados do "
                                         "empreendimento · nº do processo SEMMA.",
                      "observacoes": "Custos separados: (a) taxa DOE/GO (tabelada) e (b) jornal "
                                     "(variável · cotar 2-3 jornais de Guapó ou região). "
                                     "Priorizar veículo com maior circulação comprovada na "
                                     "zona rural do município.",
                  })

    # Ficha 06 · Investigação confirmatória
    ficha_tecnica(doc, 6,
                  "Investigação confirmatória · NBR 15515-2/3 + NBR 16209",
                  "cont",
                  {
                      "fundamento": "ABNT NBR 15515-2 (Avaliação Preliminar) · NBR 15515-3 "
                                    "(Investigação Confirmatória) · NBR 16209 (Avaliação de "
                                    "Risco à Saúde Humana) · CEMAM 029/2018 · CONAMA 420/2009.",
                      "acionamento": "CONTINGENTE. Aciona se: (a) contestação pública na "
                                     "janela CONAMA 006/86 (04/05–04/06); (b) análise "
                                     "documental 08/05 identificar passivo; (c) RCA 03/06 "
                                     "apresentar parâmetro fora; (d) SEMMA exigir formalmente.",
                      "escopo": "Avaliação preliminar (fase documental e visual) + "
                                "investigação confirmatória (sondagens SPT + coletas de solo "
                                "e água subterrânea) + avaliação de risco preliminar. Laudo "
                                "completo com ART de geólogo ou engenheiro ambiental. "
                                "Interpretação cruzada com a LAO e recomendações.",
                      "perfil": "Empresa especializada em passivo ambiental com: (i) "
                                "geólogo(a) ou eng. ambiental com CREA ativo, (ii) "
                                "equipamento de sondagem próprio ou terceirizado com "
                                "certificação, (iii) acordo com laboratório acreditado "
                                "INMETRO para análise química, (iv) portfólio em postos de "
                                "combustível ou áreas contaminadas.",
                      "prazo": "Se acionada, execução entre 25/05 (mobilização) e 15/06 "
                               "(laudo). Janela de 21 dias corridos.",
                      "quantidades": "Mínimo 3 sondagens SPT (ver Ficha 07). Coletas em "
                                     "solo superficial, subsuperfície e água subterrânea. "
                                     "Parâmetros cromatográficos na Ficha 07.",
                      "docs_fornecidos": "Laudos anteriores · plantas · histórico operacional "
                                         "do posto · relatórios de estanqueidade · MCO "
                                         "atualizado (Ficha 01).",
                      "observacoes": "Ficha 06 e Ficha 07 (sondagens) geralmente cotadas em "
                                     "pacote único. Solicitar cotação combinada e em "
                                     "separado para visibilidade. Se houver possibilidade de "
                                     "investigação detalhada (NBR 15515-4) após o preliminar, "
                                     "pedir cotação extra. Prestadores em Goiânia: pedir "
                                     "pelo menos 3 propostas por causa do custo.",
                  })

    # Ficha 07 · Sondagens SPT
    ficha_tecnica(doc, 7,
                  "Sondagens SPT · mínimo 3 furos · análise química cromatográfica",
                  "cont",
                  {
                      "fundamento": "ABNT NBR 15515-3 · ABNT NBR 6484 (SPT) · CEMAM "
                                    "029/2018 · CONAMA 420/2009.",
                      "acionamento": "CONTINGENTE. Componente da Ficha 06. Aciona junto.",
                      "escopo": "Execução de no mínimo 3 furos de sondagem SPT com "
                                "amostragem contínua, até profundidade mínima de 6m ou até "
                                "atingir o nível freático. Coleta de amostras de solo em 3 "
                                "profundidades por furo + 1 amostra de água subterrânea por "
                                "furo (se atingido). Análise laboratorial dos parâmetros "
                                "listados. Laudo com resultados + mapa de locação dos furos.",
                      "perfil": "Empresa de sondagem com equipamento calibrado + laboratório "
                                "acreditado INMETRO para análise química. Normalmente "
                                "fornecido em pacote pela mesma empresa da Ficha 06.",
                      "prazo": "Execução em 25/05/2026 (se acionada). Laudo analítico até "
                               "10/06.",
                      "quantidades": "Parâmetros cromatográficos por amostra: VOC (compostos "
                                     "orgânicos voláteis), BTEX, PAH (16 compostos EPA), "
                                     "metais pesados (Pb, Cr, Cd, Zn, Cu, Ni, Fe, Mn), TPH "
                                     "total + frações C9–C36. Solo: acrescentar granulometria "
                                     "e matéria orgânica. Água: pH, condutividade, sólidos.",
                      "docs_fornecidos": "Plantas com sugestão de locação dos furos "
                                         "(áreas críticas: tanques, bombas, SAO, fronteira "
                                         "com APP/poços). Laudo de estanqueidade.",
                      "observacoes": "Locação dos furos deve ser validada pela Hábilis antes "
                                     "da execução. Pedir custo unitário por furo adicional "
                                     "(pode precisar de 4º furo conforme achados).",
                  })

    # Ficha 08 · Análise de Risco
    ficha_tecnica(doc, 8,
                  "Análise de Risco · APP (Preliminar) ou ARQ (Quantitativa)",
                  "cont",
                  {
                      "fundamento": "CEMAM 029/2018 · CONAMA 420/2009 · metodologias ASTM "
                                    "RBCA (Risk-Based Corrective Action).",
                      "acionamento": "CONTINGENTE. Aciona só se SEMMA exigir formalmente, "
                                     "tipicamente após resultado da investigação "
                                     "confirmatória (Ficha 06) apontar concentrações acima "
                                     "do VI (Valor de Investigação).",
                      "escopo": "APP (Análise de Risco Preliminar): avaliação qualitativa "
                                "de cenários de exposição, sem modelagem numérica. Entrega "
                                "em 15–30 dias. ARQ (Análise de Risco Quantitativa): "
                                "modelagem completa com software (RBCA Tool Kit, RISC ou "
                                "similar), cálculo de concentrações-meta, entrega em 60–90 "
                                "dias. Laudo com ART de profissional habilitado.",
                      "perfil": "Empresa com eng. ambiental/químico com formação específica "
                                "em análise de risco. Portfólio comprovado em áreas "
                                "contaminadas avaliadas pela SEMAD ou IBAMA.",
                      "prazo": "APP: 15–30 dias após acionamento. ARQ: 60–90 dias. "
                               "Preferência pela APP primeiro (etapa menor).",
                      "docs_fornecidos": "Laudo da investigação confirmatória (Ficha 06) · "
                                         "dados da vizinhança (poços cadastrados, uso do solo "
                                         "no raio de 500m) · plantas.",
                      "observacoes": "Pedir cotação dos dois formatos (APP e ARQ) separados. "
                                     "Algumas empresas incluem APP no pacote da Ficha 06 sem "
                                     "custo adicional — confirmar.",
                  })

    # Ficha 09 · Estanqueidade extra
    ficha_tecnica(doc, 9,
                  "Teste de estanqueidade extra · NBR 16795/2019",
                  "cont",
                  {
                      "fundamento": "ABNT NBR 16795/2019 · CP 3 da LAO 033/2020 · "
                                    "Resolução INMETRO 37/2009.",
                      "acionamento": "CONTINGENTE. Aciona só se SEMMA contestar o laudo de "
                                     "10/2024 em mãos (validade bienal · próximo legal em "
                                     "10/2026). Improvável antes do protocolo.",
                      "escopo": "Teste de estanqueidade em tanques subterrâneos e tubulações "
                                "de produto, conforme NBR 16795/2019. Emissão de laudo "
                                "aprovado ou reprovado, com ART. Se reprovar, apontar ponto "
                                "de vazamento.",
                      "perfil": "Empresa com acreditação INMETRO vigente para teste de "
                                "estanqueidade em SASC (Sistema de Armazenamento "
                                "Subterrâneo de Combustíveis). Verificar validade da "
                                "acreditação na data da contratação.",
                      "prazo": "Execução em até 15 dias corridos após acionamento.",
                      "quantidades": "Tanques e linhas do Auto Posto América · verificar "
                                     "quantidade exata no MCO 2024.",
                      "docs_fornecidos": "Laudo 10/2024 em mãos · MCO · cadastro dos "
                                         "tanques.",
                      "observacoes": "Como o laudo atual está vigente, cotar apenas como "
                                     "backup (valor registrado, sem mobilização). Custo "
                                     "típico em Goiás: faixa de mercado cotada para "
                                     "comparação histórica.",
                  })

    # Ficha 10 · Adequação SAO
    ficha_tecnica(doc, 10,
                  "Adequação física do sistema SAO · serviços de obra civil",
                  "cont",
                  {
                      "fundamento": "Conformidade com CONAMA 273/2000 + padrões CONAMA "
                                    "430/2011 para efluentes · diagnóstico da inspeção de "
                                    "12/05 ou da análise de efluentes de 22/05.",
                      "acionamento": "CONTINGENTE. Aciona só se inspeção 12/05 ou análise "
                                     "de efluentes apontar não-conformidade (subdimensionamento, "
                                     "fissuras, manutenção atrasada, falta de declividade).",
                      "escopo": "Serviços de obra civil para adequação do separador "
                                "água-óleo: pode envolver impermeabilização, troca de "
                                "componentes, ampliação, instalação de novo SAO, "
                                "adequação de canaletas e grelhas. Escopo definido após "
                                "diagnóstico técnico.",
                      "perfil": "Empreiteira com experiência em postos de combustível · "
                                "aceitável: prestador regular de manutenção de SAO em postos "
                                "da região de Goiânia.",
                      "prazo": "Execução entre 20/05 e 12/06, se acionado. Dependente do "
                               "escopo.",
                      "docs_fornecidos": "Diagnóstico da inspeção 12/05 · plantas do SAO · "
                                         "laudo de efluentes.",
                      "observacoes": "Cotar duas faixas: (a) manutenção corretiva simples "
                                     "(limpeza reforçada, troca de selos) e (b) obra de "
                                     "adequação estrutural (mais cara, caso fissuras ou "
                                     "subdimensionamento). Incluir custo de descarte do "
                                     "material removido (resíduo classe I).",
                  })

    # Ficha 11 · Laudo saneamento/fossa
    ficha_tecnica(doc, 11,
                  "Laudo de anuência de saneamento ou de fossa/sumidouro",
                  "cont",
                  {
                      "fundamento": "Exigência do TR SEMMA sobre destinação de efluentes "
                                    "sanitários.",
                      "acionamento": "CONTINGENTE. Depende de como o posto destina o efluente "
                                     "sanitário: (a) rede pública → anuência da "
                                     "concessionária (sem custo de terceiro); (b) fossa/"
                                     "sumidouro → laudo técnico com ART.",
                      "escopo": "Laudo técnico sobre o sistema de fossa/sumidouro "
                                "(dimensionamento, estanqueidade, distância de poços), com "
                                "ART. Se for rede pública, apenas requerimento de anuência "
                                "junto à concessionária local.",
                      "perfil": "Engenheiro(a) civil ou sanitarista com CREA ativo. Para "
                                "rede pública: contato direto com concessionária de Guapó.",
                      "prazo": "Até 25/05/2026.",
                      "docs_fornecidos": "Plantas · dados do sistema sanitário · memorial "
                                         "de construção do sumidouro (se disponível).",
                      "observacoes": "Verificar primeiro qual é o regime de destinação "
                                     "(pergunta para o Grupo Z+Z). Se for rede pública, "
                                     "cotar apenas custo de anuência/taxa municipal — sem "
                                     "prestador técnico. Se for fossa, cotar laudo com ART.",
                  })

    # Ficha 12 · Recontrato Bauer/Limpmil/Ecofenix
    ficha_tecnica(doc, 12,
                  "Recontrato de serviços terceirizados · Bauer, Limpmil, Ecofenix",
                  "prov",
                  {
                      "fundamento": "CP 4.5 / 4.6 / 4.7 da LAO 033/2020 · exigência de "
                                    "contratos vigentes e regulares com prestadores "
                                    "licenciados.",
                      "acionamento": "Provável. Validação imediata dos contratos vigentes. "
                                     "Recontrato só se algum prestador estiver irregular (ANP, "
                                     "licenciamento ambiental ou autorização de transporte "
                                     "vencidos).",
                      "escopo": "Validação de: (Bauer) limpeza de tanques e SAO; (Limpmil) "
                                "coleta e transporte de resíduos; (Ecofenix) destinação "
                                "final de óleo lubrificante usado. Verificar licenças "
                                "ambientais, cadastro IBAMA CTF/APP, autorização ANTT "
                                "(transporte), certificados de destinação.",
                      "perfil": "Os próprios prestadores atuais ou substitutos de mesmo "
                                "porte. Opção alternativa: empresa regional com licenças "
                                "vigentes e frota própria.",
                      "prazo": "Validação até 20/05/2026. Recontrato (se necessário) até "
                               "30/05.",
                      "docs_fornecidos": "Contratos vigentes · NFs dos últimos 3 anos · "
                                         "certificados de destinação.",
                      "observacoes": "Este item tende a não gerar custo novo. Cotar apenas "
                                     "caso irregularidade exija substituição. Solicitar "
                                     "também cópias atualizadas de: licença ambiental, "
                                     "cadastro ANP (distribuidora) e certificado de "
                                     "destinação dos últimos 3 anos.",
                  })

    # Ficha 13 · Topografia
    ficha_tecnica(doc, 13,
                  "Levantamento topográfico complementar · UTM SIRGAS 2000",
                  "opc",
                  {
                      "fundamento": "Potencial exigência do Veredas ao receber o protocolo "
                                    "de outorga em 13/05 · CP 3.12 LAO (georreferenciamento).",
                      "acionamento": "OPCIONAL. Aciona só se o Veredas exigir memorial "
                                     "topográfico adicional ou se as plantas 05/2024 não "
                                     "tiverem precisão suficiente para o dossiê SEMMA.",
                      "escopo": "Levantamento topográfico planialtimétrico com coordenadas "
                                "UTM SIRGAS 2000 do perímetro do posto + APP + corpo "
                                "hídrico próximo (se houver). Entrega em .dwg + memorial "
                                "descritivo + ART.",
                      "perfil": "Eng. agrimensor ou técnico em topografia com CREA/CFT "
                                "ativo · equipamento GPS/RTK.",
                      "prazo": "Até 10 dias corridos após acionamento.",
                      "docs_fornecidos": "Plantas 05/2024 · CAR.",
                      "observacoes": "Cotar como reserva. Preferencialmente a mesma empresa "
                                     "que fez as plantas 05/2024.",
                  })

    # Ficha 14 · Paisagismo/APP
    ficha_tecnica(doc, 14,
                  "Projeto de adequação paisagística ou recuperação de APP",
                  "opc",
                  {
                      "fundamento": "Código Florestal (Lei 12.651/2012) · CAR · diagnóstico "
                                    "da inspeção 12/05.",
                      "acionamento": "OPCIONAL. Aciona só se inspeção 12/05 identificar APP "
                                     "degradada ou descumprimento de faixa mínima. Não há "
                                     "evidência preliminar.",
                      "escopo": "Projeto de Recuperação de Área Degradada (PRAD) simples ou "
                                "projeto paisagístico de adequação. Inclui escolha de "
                                "espécies nativas, plantio, cronograma, ART.",
                      "perfil": "Eng. agrônomo(a) ou florestal com CREA ativo. "
                                "Preferencialmente com experiência em PRAD no Cerrado.",
                      "prazo": "Projeto em 15 dias. Execução definida pelo projeto.",
                      "docs_fornecidos": "CAR · imagens de satélite · inspeção 12/05.",
                      "observacoes": "Item só entra em cena em cenário específico. Cotar "
                                     "apenas projeto (não execução) como reserva de "
                                     "precificação.",
                  })

    # ====================================================================
    # SECAO 04 · ORIENTACOES DE RETORNO
    # ====================================================================
    add_page_break(doc)
    heading_section(doc, "04", "Orientações para retorno das cotações", space_before=0)

    body_p(doc,
           "Padrão de retorno esperado do cotador. Quanto mais estruturado, mais rápido a "
           "Hábilis aprova e contrata.",
           ["padrão de retorno"], size=10)

    heading_h3(doc, "O que devolver", space_before=4)
    bullet(doc, "Tabela geral (seção 02) preenchida com: valor R$ · nome do prestador · "
                "validade da proposta · telefone/e-mail do contato · CNPJ.",
           ["Tabela geral", "validade da proposta"])
    bullet(doc, "Propostas originais recebidas dos prestadores (PDF, e-mail, orçamento "
                "digitalizado) anexadas ao retorno.",
           ["Propostas originais"])
    bullet(doc, "Para cada item, mínimo 2 propostas comparáveis. Para itens CONTINGENTE, "
                "ideal 3 propostas (maior variação de preço).",
           ["mínimo 2 propostas", "CONTINGENTE"])
    bullet(doc, "Sinalizar prestadores que recusaram, não têm acreditação exigida ou não "
                "responderam no prazo.",
           ["Sinalizar prestadores que recusaram"])
    bullet(doc, "Destacar se alguma proposta veio em pacote (ex: Fichas 01+02 juntas) com "
                "desconto. Comparar com o unitário.",
           ["pacote"])

    heading_h3(doc, "Ordem de prioridade de cotação", space_before=8)
    bullet(doc, "Prioridade 1 (até 05/05/2026) · Fichas 01, 02, 05 — itens OBRIG. "
                "IMEDIATO. Sem estas cotações, a contratação de F1 atrasa.",
           ["Prioridade 1", "OBRIG. IMEDIATO"])
    bullet(doc, "Prioridade 2 (até 15/05/2026) · Fichas 03, 04, 11, 12 — itens PROVÁVEL. "
                "Coleta depende destes fechamentos.",
           ["Prioridade 2", "PROVÁVEL"])
    bullet(doc, "Prioridade 3 (até 25/05/2026) · Fichas 06, 07, 08, 09, 10 — itens "
                "CONTINGENTE. Cotados mesmo sem acionamento, ficam pré-aprovados.",
           ["Prioridade 3", "CONTINGENTE"])
    bullet(doc, "Prioridade 4 (até 30/05/2026) · Fichas 13, 14 — itens OPCIONAL. "
                "Reserva de precificação.",
           ["Prioridade 4", "OPCIONAL"])

    heading_h3(doc, "Canal e formato de retorno", space_before=8)
    bullet(doc, "Enviar pacote consolidado para o e-mail Hábilis de referência com "
                "assunto: \"Cotações Auto Posto América · [data]\".",
           ["e-mail Hábilis de referência", "[data]"])
    bullet(doc, "Arquivos: (1) este documento preenchido em .docx ou PDF anotado; "
                "(2) pasta com as propostas originais nomeadas \"F[NN]_[prestador].pdf\".",
           ["F[NN]_[prestador].pdf"])
    bullet(doc, "Dúvidas técnicas sobre o escopo de cada ficha: WhatsApp ou ligação direta "
                "com o responsável técnico Hábilis antes de mandar pergunta para o prestador.",
           ["WhatsApp ou ligação direta"])

    # Callout final
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_shading(p, VERDE_BG)
    set_paragraph_border_left(p, 2.5, "1B5E20")
    p.paragraph_format.left_indent = Cm(0.25)
    styled_run(p, "Regra de ouro · ", size=10, color=CINZA_9, bold=True)
    styled_run(p,
               "prestador sem acreditação específica (INMETRO para laboratórios, CREA ativo "
               "para ART, INMETRO para estanqueidade) não entra na tabela, mesmo com preço "
               "menor. A ART ou o laudo sem cobertura formal inviabiliza o item no dossiê "
               "SEMMA — custo zero com dano máximo.",
               size=10, color=CINZA_7)

    here = os.path.dirname(os.path.abspath(__file__))
    out = os.path.join(here, "Briefing_Cotacao_Terceiros_AutoPostoAmerica_v3.docx")
    doc.save(out)
    print(f"OK -> {out}")
    return out


if __name__ == "__main__":
    build()
