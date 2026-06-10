#!/usr/bin/env python3
"""
Briefing para Cotacao de Terceiros - Auto Posto America.
Documento interno para Guilherme levantar orcamentos.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

VERDE = RGBColor(0x1B, 0x5E, 0x20)
VERDE_L = RGBColor(0x4C, 0xAF, 0x50)
VERDE_BG = "E8F5E9"
AMBER = RGBColor(0xF5, 0x7F, 0x17)
AMBER_BG = "FFF8E1"
VERMELHO = RGBColor(0xB7, 0x1C, 0x1C)
VERMELHO_BG = "FFEBEE"
AZUL = RGBColor(0x0D, 0x47, 0xA1)
AZUL_BG = "E3F2FD"
CINZA_9 = RGBColor(0x1C, 0x1C, 0x1E)
CINZA_7 = RGBColor(0x3A, 0x3A, 0x3C)
CINZA_5 = RGBColor(0x6C, 0x6C, 0x70)
CINZA_BG = "FAFAFA"
BRANCO = RGBColor(0xFF, 0xFF, 0xFF)
FONT_SANS = "Segoe UI"
FONT_MONO = "Consolas"


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for name, val in edges.items():
        if val is None: continue
        size, color = val
        el = OxmlElement(f'w:{name}')
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(size)); el.set(qn('w:color'), color)
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def set_paragraph_shading(p, hex_color):
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    p_pr.append(shd)


def set_paragraph_border_left(p, size_pt, color_hex):
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), str(int(size_pt * 8)))
    left.set(qn('w:space'), '6'); left.set(qn('w:color'), color_hex)
    pbdr.append(left); p_pr.append(pbdr)


def styled_run(p, text, *, bold=False, italic=False, size=None, color=None, mono=False):
    run = p.add_run(text)
    font_name = FONT_MONO if mono else FONT_SANS
    run.font.name = font_name
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.find(qn('w:rFonts'))
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts'); r_pr.append(r_fonts)
    r_fonts.set(qn('w:ascii'), font_name); r_fonts.set(qn('w:hAnsi'), font_name); r_fonts.set(qn('w:cs'), font_name)
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    return run


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def heading_section(doc, num, title, space_before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '14')
    bottom.set(qn('w:space'), '3'); bottom.set(qn('w:color'), '1B5E20')
    pbdr.append(bottom); p_pr.append(pbdr)
    styled_run(p, f"{num}  ", size=11, color=VERDE_L, mono=True)
    styled_run(p, title, size=15, color=VERDE, bold=True)


def heading_h3(doc, text, space_before=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    set_paragraph_border_left(p, 2, "1B5E20")
    p.paragraph_format.left_indent = Cm(0.2)
    styled_run(p, text, size=10.5, color=CINZA_9, bold=True)


def label_mono(doc, text, space_before=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(2)
    styled_run(p, text.upper(), size=8.5, color=VERDE, bold=True, mono=True)


def body_p(doc, text, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    styled_run(p, text, size=size, color=CINZA_7)


def bullet(doc, text, size=9.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    styled_run(p, text, size=size, color=CINZA_7)


def template_block(doc, title, text):
    """Bloco destacado com template de mensagem para copy/paste."""
    t = doc.add_table(rows=1, cols=1)
    c = t.cell(0, 0)
    set_cell_shading(c, "F5F5F5")
    set_cell_border(c, left=(16, "6C6C70"), top=(4, "E0E0E0"),
                    right=(4, "E0E0E0"), bottom=(4, "E0E0E0"))
    p1 = c.paragraphs[0]
    p1.paragraph_format.space_after = Pt(4)
    styled_run(p1, title.upper(), size=8.5, color=CINZA_9, bold=True, mono=True)
    p2 = c.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    styled_run(p2, text, size=9.5, color=CINZA_7, mono=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def table_hdr(doc, headers, widths_cm=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_shading(c, "1B5E20")
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        styled_run(p, h.upper(), size=8, color=BRANCO, bold=True, mono=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if widths_cm and i < len(widths_cm): c.width = Cm(widths_cm[i])
    return t


def add_row(t, values, zebra=False, bold_cols=None, mono_cols=None, is_total=False, size=9):
    row = t.add_row()
    for i, val in enumerate(values):
        c = row.cells[i]
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
        bold = (bold_cols and i in bold_cols) or is_total
        mono = mono_cols and i in mono_cols
        styled_run(p, val, size=size, color=CINZA_9 if bold else CINZA_7, bold=bold, mono=mono)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        if is_total: set_cell_shading(c, VERDE_BG)
        elif zebra: set_cell_shading(c, CINZA_BG)


def field_page(p):
    r = p.add_run()
    r.font.name = FONT_MONO; r.font.size = Pt(8); r.font.color.rgb = CINZA_5
    fld_b = OxmlElement('w:fldChar'); fld_b.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.text = 'PAGE'
    fld_e = OxmlElement('w:fldChar'); fld_e.set(qn('w:fldCharType'), 'end')
    r._r.append(fld_b); r._r.append(it); r._r.append(fld_e)


def field_numpages(p):
    r = p.add_run()
    r.font.name = FONT_MONO; r.font.size = Pt(8); r.font.color.rgb = CINZA_5
    fld_b = OxmlElement('w:fldChar'); fld_b.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.text = 'NUMPAGES'
    fld_e = OxmlElement('w:fldChar'); fld_e.set(qn('w:fldCharType'), 'end')
    r._r.append(fld_b); r._r.append(it); r._r.append(fld_e)


# ============ SERVICO BLOCK =============
def servico(doc, numero, titulo, quando, escopo, normas, entregaveis, habilitacao, pontos_orcamento, template_msg):
    heading_section(doc, numero, titulo, space_before=14)

    # Linha de metadados
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    styled_run(p, "QUANDO: ", size=8.5, color=VERDE, bold=True, mono=True)
    styled_run(p, quando, size=9.5, color=CINZA_9, bold=True, mono=True)

    label_mono(doc, "Escopo")
    body_p(doc, escopo)

    label_mono(doc, "Normas tecnicas aplicaveis")
    for n in normas:
        bullet(doc, n)

    label_mono(doc, "Entregaveis obrigatorios")
    for e in entregaveis:
        bullet(doc, e)

    label_mono(doc, "Habilitacao exigida do prestador")
    for h in habilitacao:
        bullet(doc, h)

    label_mono(doc, "Pontos que o orcamento precisa detalhar")
    for po in pontos_orcamento:
        bullet(doc, po)

    template_block(doc, "Template de mensagem (copiar e enviar)", template_msg)


def build():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(1.8); s.bottom_margin = Cm(1.8)
    s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)

    # Header
    h_p = s.header.paragraphs[0]
    styled_run(h_p, "Auto Posto America | Briefing para Cotacao de Terceiros", size=8, color=CINZA_5, mono=True)
    h_p.add_run("\t")
    styled_run(h_p, "Uso interno Habilis", size=8, color=VERDE, bold=True, mono=True)

    # Footer
    f_p = s.footer.paragraphs[0]
    styled_run(f_p, "Habilis | Guilherme | Abr/2026", size=8, color=CINZA_5, mono=True)
    f_p.add_run("\t")
    styled_run(f_p, "Pagina ", size=8, color=CINZA_5, mono=True)
    field_page(f_p)
    styled_run(f_p, " de ", size=8, color=CINZA_5, mono=True)
    field_numpages(f_p)

    # ============ CAPA ENXUTA ============
    p = doc.add_paragraph()
    styled_run(p, "H  Habilis", size=14, color=VERDE, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    styled_run(p, "USO INTERNO | NAO ENVIAR AO CLIENTE", size=9, color=VERMELHO, bold=True, mono=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    styled_run(p, "Briefing para cotacao de terceiros", size=22, color=CINZA_9, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    styled_run(p, "Auto Posto America | Renovacao LAO 033/2020", size=12, color=CINZA_7)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    styled_run(p, "Documento de apoio para o levantamento de orcamentos dos servicos contratados "
               "de terceiros previstos no plano de regularizacao. Cada bloco contem: escopo tecnico, "
               "normas aplicaveis, entregaveis exigidos, requisitos de habilitacao do prestador "
               "(credenciamentos, ART, conselho) e template de mensagem pronta para envio via "
               "WhatsApp/email. Ao final, tabela consolidada para preencher com cotacoes recebidas.",
               size=10.5, color=CINZA_7)

    # Mini sumario
    heading_h3(doc, "Servicos a cotar - visao geral", space_before=6)
    t = table_hdr(doc, ["#", "Servico", "Quando", "Prioridade"], [1.0, 8.0, 3.0, 4.0])
    add_row(t, ["01", "Laboratorio de analises ambientais", "A contratar ate 15/05", "Alta"], bold_cols=[0, 1, 3], mono_cols=[0, 2])
    add_row(t, ["02", "Investigacao ambiental (SPT + amostragem + laudo)", "Contratar ate 05/05", "Alta"], zebra=True, bold_cols=[0, 1, 3], mono_cols=[0, 2])
    add_row(t, ["03", "Eng. Civil/Ambiental - MCO + Anexos CONAMA 273", "Contratar ate 15/05", "Alta"], bold_cols=[0, 1, 3], mono_cols=[0, 2])
    add_row(t, ["04", "Empresa de estanqueidade (contingencia)", "Mapear ja, acionar se SEMMA contestar", "Media"], zebra=True, bold_cols=[0, 1, 3], mono_cols=[0, 2])
    add_row(t, ["05", "Servicos fisicos para CPs pendentes (CP 4.4 + 4.14)", "Contratar ate 01/06", "Media"], bold_cols=[0, 1, 3], mono_cols=[0, 2])

    add_page_break(doc)

    # ============ 01 - LABORATORIO ============
    servico(doc,
        numero="01",
        titulo="Laboratorio de analises ambientais",
        quando="Contratar ate 15/05 | coleta 22/05 | laudo em F2",
        escopo="Coleta e analise laboratorial de amostras de efluentes e agua do Auto Posto America, "
               "conjugadas na mesma visita para otimizar deslocamento. Agua coletada dos pocos de monitoramento "
               "e do abastecimento (se houver captacao). Efluentes coletados na saida da SAO (Sistema Separador "
               "de Agua e Oleo) e em pontos de descarte.",
        normas=[
            "CONAMA 357/2005 (qualidade de agua) e 430/2011 (efluentes).",
            "CEMAM 029/2018 Anexo XI (RCA semestral - parametros SAO).",
            "NBR 9898 (preservacao e tecnicas de amostragem).",
            "ISO/IEC 17025 (acreditacao laboratorial).",
        ],
        entregaveis=[
            "Laudo tecnico com resultados e comparacao com limites normativos.",
            "ART do quimico responsavel emitindo assinatura do laudo (CRQ).",
            "Cadeia de custodia das amostras, fotografias da coleta, coordenadas UTM SIRGAS 2000.",
            "Certificado de calibracao dos equipamentos utilizados.",
        ],
        habilitacao=[
            "Laboratorio com acreditacao ISO/IEC 17025 vigente (RBC ou equivalente).",
            "Credenciamento INMETRO vigente na data da coleta.",
            "Quimico responsavel registrado no CRQ com ART ativa.",
            "Historico em postos de combustiveis (preferencial, nao eliminatorio).",
        ],
        pontos_orcamento=[
            "Parametros a analisar em agua: pH, turbidez, O&G, solidos sedimentaveis, DBO, DQO, OD, MBAS (se houver lava-jato), metais (Pb, Zn, Cu, Ni, Cd, Ca, Ba), TPH (se trocador de oleo), BTEX/PAH nos pocos de monitoramento.",
            "Parametros a analisar em efluentes: mesmos acima + especificos da SAO.",
            "Numero de pontos de coleta (tipicamente 3-5 pocos de monitoramento + entrada/saida SAO).",
            "Prazo de entrega do laudo apos a coleta (buscar <= 15 dias uteis).",
            "Custo de deslocamento ate Guapo/GO (BR-060, km 203).",
            "Desdobramento do preco por parametro (se possivel), para comparacao entre cotacoes.",
            "Validade da proposta comercial.",
        ],
        template_msg=(
            "Ola, [NOME]. Preciso de um orcamento para servico de coleta e analise laboratorial de "
            "efluentes e agua no Auto Posto America, em Guapo/GO. Os parametros sao os do Anexo XI da "
            "CEMAM 029/2018 e dos limites da CONAMA 357/2005 e 430/2011. A coleta esta prevista para "
            "22/05/2026. Preciso de laudo com ART do quimico responsavel e amostragem conforme NBR 9898. "
            "O laboratorio precisa ter ISO 17025 e INMETRO vigentes. Total estimado: 3-5 pontos de "
            "coleta (pocos de monitoramento + SAO). Pode me mandar o orcamento com prazo de entrega e "
            "detalhamento por parametro?"
        ),
    )

    add_page_break(doc)

    # ============ 02 - INVESTIGACAO AMBIENTAL ============
    servico(doc,
        numero="02",
        titulo="Investigacao ambiental (confirmatoria)",
        quando="Contratacao ate 05/05 | mobilizacao 25/05 | laudo 15/06",
        escopo="Sondagens SPT, amostragem de solo e agua subterranea e emissao do laudo de Investigacao "
               "Confirmatoria para o Auto Posto America. Minimo 3 furos em pontos estrategicos (tanques, SASC, "
               "bombas, area de descarga). O mesmo prestador pode entregar a mobilizacao em 25/05 e o laudo final "
               "em 15/06. Se necessaria Analise de Risco na sequencia (contingencia F3), pode ser cotada tambem.",
        normas=[
            "NBR 15515-1/2/3 (investigacao de passivo ambiental em solo e agua subterranea).",
            "NBR 16209 (avaliacao de risco).",
            "CEMAM 029/2018 (diretrizes estaduais GO).",
            "CONAMA 420/2009 (VRQ e VI para solos).",
        ],
        entregaveis=[
            "Relatorio de Investigacao Confirmatoria com resultados e interpretacao.",
            "ART do geologo ou engenheiro ambiental responsavel tecnico.",
            "Locacao georreferenciada dos furos (UTM SIRGAS 2000).",
            "Laudo com VOC, BTEX, PAH, metais e TPH nas amostras coletadas.",
            "Recomendacao sobre necessidade ou nao de investigacao detalhada/remediacao.",
        ],
        habilitacao=[
            "Empresa com historico em postos de combustiveis.",
            "Equipe com geologo ou engenheiro ambiental registrado no CREA.",
            "Equipamento de sondagem SPT proprio ou contratado (confirmar).",
            "Laboratorio parceiro acreditado ISO 17025 (confirmar se esta incluso).",
        ],
        pontos_orcamento=[
            "Numero minimo de furos: 3 (tanques, SASC, bombas). Indicar se ha sugestao de ampliacao.",
            "Profundidade maxima de sondagem prevista.",
            "Se o laboratorio para analises esta incluso ou e contratacao separada.",
            "Prazo total: contratacao (05/05) -> mobilizacao em campo (25/05) -> laudo final (15/06).",
            "Cotacao separada para Analise de Risco (contingencia F3): APP 15-30d / ARQ 60-90d.",
            "Cobranca de deslocamento ate Guapo/GO (BR-060, km 203).",
            "Se ART e de emissao imediata ou exige prazo do CREA.",
        ],
        template_msg=(
            "Ola, [NOME]. Preciso de orcamento para Investigacao Confirmatoria em posto de combustivel "
            "(Auto Posto America, Guapo/GO). Escopo: sondagens SPT, amostragem de solo e agua subterranea "
            "com minimo 3 furos, analises de VOC/BTEX/PAH/metais/TPH e emissao de Relatorio de Investigacao "
            "Confirmatoria conforme CEMAM 029/2018 + CONAMA 420/2009 + NBR 15515-2/3 e NBR 16209. "
            "Contratacao em 05/05/2026, mobilizacao 25/05 e laudo em 15/06. Preciso de ART de geologo/eng. "
            "ambiental. Por favor cote tambem, separadamente, a Analise de Risco (opcional/contingencia), e "
            "inclua no orcamento o deslocamento ate Guapo."
        ),
    )

    add_page_break(doc)

    # ============ 03 - ENG CIVIL/AMBIENTAL MCO+273 ============
    servico(doc,
        numero="03",
        titulo="Eng. Civil/Ambiental | MCO + Anexos CONAMA 273",
        quando="Contratar ate 15/05 | MCO 28/05 | Anexos 29/05",
        escopo="Subcontratacao de Engenheiro Civil, Ambiental ou Mecanico para elaborar e assinar com "
               "ART propria as pecas tecnicas que a Habilis nao tem atribuicao interna para emitir. Escopo: "
               "(1) MCO - Memorial de Caracterizacao do Empreendimento (Anexo IX CEMAM 029/2018); "
               "(2) Anexos I e II da Resolucao CONAMA 273/2000. Coordenacao integral pela Habilis; o "
               "profissional entrega as pecas com ART propria e eventuais suportes de consulta tecnica.",
        normas=[
            "CEMAM 029/2018 Anexo IX (MCO).",
            "Resolucao CONAMA 273/2000 Anexos I e II.",
            "NBR da area de engenharia conforme equipamentos (NBR 13781 SASC, NBR 16161 tubulacoes, NBR 14605 SAO, etc).",
            "Res. CONFEA 218/73 (atribuicoes profissionais).",
        ],
        entregaveis=[
            "MCO completo com descricao tecnica: tanques, SASC, tubulacoes, SAO, infraestrutura predial, fluxogramas e memoriais.",
            "Anexo I CONAMA 273/2000 (tabela descritiva do empreendimento).",
            "Anexo II CONAMA 273/2000 (classificacao em classes de risco ambiental 1/2/3).",
            "ART propria (CREA) do profissional responsavel.",
            "Arquivo editavel (.docx ou .pdf) das pecas para eventual revisao.",
        ],
        habilitacao=[
            "Registro ativo no CREA (civil, ambiental, mecanico ou quimico com atribuicao compativel).",
            "Experiencia em memoriais tecnicos para licenciamento de postos de combustiveis.",
            "ART pode ser emitida em ate 15 dias uteis apos inicio do servico.",
            "Disponibilidade para entrega ate 29/05 (prazo nao negociavel por risco ao protocolo 26/06).",
        ],
        pontos_orcamento=[
            "Valor global pelo conjunto (MCO + Anexo I + Anexo II) versus valor por peca.",
            "Numero de revisoes inclusas (recomendado: 2 revisoes).",
            "Custo da ART (se incluso no valor ou cobrado a parte).",
            "Se precisa de visita tecnica ao posto ou se trabalha com dados fornecidos pela Habilis (inspecao 12/05).",
            "Responsabilidade solidaria: confirmacao que aceita composicao com RT coordenador CRBio (Habilis).",
            "Suporte pos-entrega se SEMMA pedir complementacao.",
        ],
        template_msg=(
            "Ola, [NOME]. A Habilis precisa contratar um Eng. Civil/Ambiental/Mecanico para emitir duas pecas "
            "tecnicas de licenciamento ambiental em posto de combustivel (Auto Posto America, Guapo/GO). Sao: "
            "(1) MCO conforme Anexo IX CEMAM 029/2018 e (2) Anexos I e II da CONAMA 273/2000, ambas com ART "
            "propria no CREA. Prazo: contratacao 15/05, entregas 28/05 e 29/05. A coordenacao do processo e "
            "da Habilis (RT Biologo/CRBio); sua atuacao e pontual nas pecas. Preciso de orcamento global, "
            "valor da ART e numero de revisoes inclusas."
        ),
    )

    add_page_break(doc)

    # ============ 04 - ESTANQUEIDADE (CONTINGENCIA) ============
    servico(doc,
        numero="04",
        titulo="Teste de estanqueidade | contingencia",
        quando="Mapear agora | acionar apenas se SEMMA contestar laudo 10/2024",
        escopo="Teste de estanqueidade em tanques, tubulacoes e linhas do Auto Posto America conforme NBR "
               "16795/2019. Inclusao no plano como contingencia armada: sera acionado apenas se a SEMMA "
               "contestar o laudo atual de 10/2024 em razao do credenciamento INMETRO (NCC 21.08786) vencido "
               "em 29/04/2024 ou da calibracao Suporty vencida em 03/2023. Objetivo do orcamento e ter "
               "prestador qualificado ja mapeado para mobilizacao em 30-60 dias se a exigencia vier.",
        normas=[
            "NBR 16795/2019 (teste de estanqueidade em tanques e tubulacoes de postos).",
            "Portarias INMETRO aplicaveis (37, 109, 110, 111, 009/2005).",
            "CEMAM 029/2018 (diretrizes estaduais).",
        ],
        entregaveis=[
            "Laudo de estanqueidade dentro da NBR 16795/2019.",
            "ART do engenheiro mecanico responsavel.",
            "Certificado de calibracao ISO/IEC 17025 vigente dos equipamentos utilizados.",
            "Credenciamento INMETRO da empresa na data do teste (vigente).",
        ],
        habilitacao=[
            "Empresa com credenciamento INMETRO VIGENTE na data da contratacao (condicao eliminatoria).",
            "Equipamentos com calibracao ISO 17025 vigente (condicao eliminatoria).",
            "Engenheiro mecanico com ART no CREA.",
            "Historico de atuacao em postos de combustiveis em Goias (preferencial).",
        ],
        pontos_orcamento=[
            "Confirmacao expressa de que credenciamento INMETRO esta vigente.",
            "Prazo maximo entre contratacao e execucao do teste em campo (buscar <= 15 dias).",
            "Numero de tanques e tubulacoes que serao testados (levantar na inspecao 12/05).",
            "Custo por tanque ou pacote fechado.",
            "Validade da proposta (minimo 90 dias, ja que acionamento e contingente).",
            "Inclusao de retorno ao local se houver reprovacao em algum teste.",
        ],
        template_msg=(
            "Ola, [NOME]. Estou mapeando prestadores para eventual teste de estanqueidade em tanques e "
            "tubulacoes do Auto Posto America (Guapo/GO) conforme NBR 16795/2019. O acionamento e contingente "
            "(30-60 dias apos eventual exigencia da SEMMA), mas preciso ja ter orcamento e confirmacao de "
            "que o credenciamento INMETRO da empresa e a calibracao dos equipamentos estao vigentes. Pode "
            "me passar proposta com validade de 90 dias e custo por tanque ou pacote?"
        ),
    )

    add_page_break(doc)

    # ============ 05 - SERVICOS FISICOS CPs ============
    servico(doc,
        numero="05",
        titulo="Servicos fisicos para CPs pendentes (CP 4.4 + CP 4.14)",
        quando="Contratar ate 01/06 | execucao ate 15/06",
        escopo="Execucao de servicos fisicos ou ensaios para atendimento das 2 condicionantes pendentes "
               "(CPs) da LAO 033/2020 do Auto Posto America. CP 4.4 (Relatorio Ambiental semestral novo) e "
               "CP 4.14 (PGA/PCA a validar). O detalhamento exato dos servicos so sera definido apos o "
               "mapeamento tecnico da Habilis em 10/05, mas cabe comecar sondagem de prestadores ja.",
        normas=[
            "LAO 033/2020 - condicionantes 4.4 e 4.14.",
            "CEMAM 029/2018 Anexo XI (RCA semestral - subsidia CP 4.4).",
            "PGA/PCA (Plano de Gestao Ambiental / Plano de Controle Ambiental) conforme TR SEMMA.",
        ],
        entregaveis=[
            "Servicos/ensaios conforme diagnostico Habilis 10/05.",
            "Relatorio descritivo da execucao com fotos e coordenadas UTM.",
            "ART do profissional responsavel pela execucao.",
            "Cruzamento entre os servicos executados e os textos literais das CPs 4.4 e 4.14.",
        ],
        habilitacao=[
            "A definir conforme diagnostico Habilis 10/05 (pode ser engenharia civil, ambiental ou servicos especializados).",
            "Registro profissional no conselho competente.",
            "Historico em atendimento de condicionantes em licencas ambientais.",
        ],
        pontos_orcamento=[
            "Aguardar fechamento do diagnostico 10/05 antes de detalhar o escopo.",
            "Preparar prestadores de 2-3 perfis diferentes (ensaios, servicos de SAO, reparos, projetistas).",
            "Prazo entre contratacao (01/06) e execucao (15/06): 14 dias corridos.",
            "Recomendavel ter flexibilidade para ajuste de escopo pos-diagnostico.",
            "Cotacao valor/hora ou valor por servico especifico.",
        ],
        template_msg=(
            "Ola, [NOME]. A Habilis esta mapeando prestadores para execucao de servicos fisicos para "
            "atendimento de condicionantes ambientais do Auto Posto America (Guapo/GO). O escopo exato sera "
            "definido ate 10/05 (estamos em fase de diagnostico). Trata-se de CP 4.4 (RCA semestral) e CP 4.14 "
            "(PGA/PCA). Contratacao prevista para 01/06 e execucao em 15/06. Pode me passar, por enquanto, "
            "seu portfolio de servicos em licenciamento ambiental e tabela de valores/hora ou por servico?"
        ),
    )

    add_page_break(doc)

    # ============ TABELA CONSOLIDADA ============
    heading_section(doc, "R", "Consolidacao de cotacoes | preencher conforme respostas", space_before=0)

    body_p(doc, "Tabela para preencher com as cotacoes recebidas. Campos em branco para anotacao manual.")

    t = table_hdr(doc, ["#", "Servico", "Prestador", "Telefone/email", "Valor", "Prazo", "ART ok?", "Obs"],
                  [0.9, 3.2, 2.8, 3.0, 1.8, 1.5, 1.5, 2.3])
    servicos_lista = [
        ("01", "Laboratorio analises"),
        ("02", "Investigacao ambiental"),
        ("03", "Eng. Civil/Amb. MCO+273"),
        ("04", "Estanqueidade"),
        ("05", "Servicos fisicos CPs"),
    ]
    # 3 linhas por servico para anotacao de ate 3 cotacoes
    for i, (num, serv) in enumerate(servicos_lista):
        for cot in range(3):
            zebra = (i % 2 == 1)
            first = (cot == 0)
            add_row(t,
                [num if first else "",
                 serv if first else "",
                 "", "", "", "", "", ""],
                zebra=zebra, bold_cols=[0, 1], mono_cols=[0], size=8.5)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    styled_run(p, "Criterios para escolha final: ", size=9.5, color=CINZA_9, bold=True)
    styled_run(p, "menor prazo entre orcamentos tecnicamente equivalentes > habilitacao vigente "
               "confirmada > preco. ", size=9.5, color=CINZA_7)
    styled_run(p, "Nao escolher so por preco; habilitacao vencida custa mais do que a diferenca de orcamento.",
               size=9.5, color=VERMELHO, bold=True)

    # Salvar
    here = os.path.dirname(os.path.abspath(__file__))
    out = os.path.join(here, "Briefing_Cotacao_Terceiros_AutoPostoAmerica_v1.docx")
    doc.save(out)
    print(f"OK -> {out}")
    return out


if __name__ == "__main__":
    build()
