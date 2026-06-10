#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Tabela única para o dono do posto · Auto Posto América · Guapó/GO
Estrutura:
  Bloco A · Uso do Solo (1 linha · Status OK Z+Z)
  Bloco B · LAO · Checklist conforme CEMAM 029/2018 (peças do dossiê)
  Bloco C · Condicionantes LAO 033/2020 (28 · íntegra)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

VERDE   = RGBColor(0x1B, 0x5E, 0x20)
CINZA_7 = RGBColor(0x42, 0x42, 0x42)
CINZA_5 = RGBColor(0x75, 0x75, 0x75)
BRANCO  = RGBColor(0xFF, 0xFF, 0xFF)
ZEBRA   = "F5F5F5"
HEADER  = "1B5E20"
BLOCO_A = "E8F5E9"
BLOCO_B = "E3F2FD"
BLOCO_C = "FFF8E1"

FONTE_CORPO  = "DM Sans"
FONTE_FALL   = "Segoe UI"
FONTE_MONO   = "DM Mono"
FONTE_MONO_F = "Consolas"


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color="BDBDBD", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        borders.append(b)
    tc_pr.append(borders)


def styled(p, text, size=9, color=CINZA_7, bold=False, mono=False):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.bold = bold
    r.font.name = FONTE_MONO if mono else FONTE_CORPO
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    fam = FONTE_MONO if mono else FONTE_CORPO
    fall = FONTE_MONO_F if mono else FONTE_FALL
    rFonts.set(qn("w:ascii"), fam)
    rFonts.set(qn("w:hAnsi"), fam)
    rFonts.set(qn("w:cs"), fall)
    return r


def set_col_widths(table, widths_cm):
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[idx])


def add_header_row(table, headers):
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr[i]
        set_cell_shading(cell, HEADER)
        set_cell_borders(cell, color="0D3911", size="8")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        styled(p, h.upper(), size=8, color=BRANCO, bold=True, mono=True)


def add_data_row(table, values, fill=None, mono_cols=None, bold_cols=None):
    mono_cols = mono_cols or []
    bold_cols = bold_cols or []
    cells = table.add_row().cells
    for i, v in enumerate(values):
        cell = cells[i]
        if fill:
            set_cell_shading(cell, fill)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        styled(p, v, size=8.5, color=CINZA_7,
               bold=(i in bold_cols), mono=(i in mono_cols))
    return cells


def add_section_row(table, label, fill):
    cells = table.add_row().cells
    a = cells[0]
    a.merge(cells[-1])
    set_cell_shading(a, fill)
    set_cell_borders(a, color="9E9E9E", size="6")
    p = a.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    styled(p, label, size=9.5, color=VERDE, bold=True, mono=True)


def build():
    doc = Document()

    # ===== Página A4 paisagem =====
    sec = doc.sections[0]
    sec.page_width  = Cm(29.7)
    sec.page_height = Cm(21.0)
    sec.top_margin    = Cm(1.2)
    sec.bottom_margin = Cm(1.2)
    sec.left_margin   = Cm(1.2)
    sec.right_margin  = Cm(1.2)

    # ===== Cabeçalho do documento =====
    h = doc.add_paragraph()
    h.paragraph_format.space_after = Pt(2)
    styled(h, "AUTO POSTO AMÉRICA · Guapó/GO · Renovação LAO 033/2020 · Proc. SEMMA 22760/2020",
           size=11, color=VERDE, bold=True, mono=True)
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(8)
    styled(sub, "Quadro consolidado · Uso do Solo · Checklist LAO (CEMAM 029/2018) · 28 Condicionantes · Hábilis · 23/04/2026",
           size=9, color=CINZA_5, mono=True)

    # ===== Tabela única · 5 colunas =====
    # Bloco | # | Item | O que tem (status atual) | O que precisa providenciar / Serviço vinculado
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    add_header_row(table, [
        "Bloco",
        "#",
        "Item / Texto literal",
        "O que tem (status atual)",
        "O que precisa providenciar · Serviço vinculado",
    ])
    set_col_widths(table, [2.4, 1.4, 8.2, 7.6, 7.7])

    # ===== BLOCO A · USO DO SOLO =====
    add_section_row(table, "A · USO DO SOLO · pré-requisito da renovação · Status: OK pelo Grupo Z+Z",
                    BLOCO_A)
    add_data_row(table, [
        "Uso do Solo",
        "A.1",
        "Certidão de Uso e Ocupação do Solo · Prefeitura de Guapó · vigente na data do protocolo (LAO 033/2020 item 1.2 · zona rural).",
        "RGI 2025 em mãos · documentação dominial · diagnóstico locacional compatível com o uso atual · enquadramento Zona Rural confirmado pelo Grupo Z+Z (status OK).",
        "Verificação formal da Certidão vigente (25/04 · Cliente) · ART Uso do Solo via RT Agronômica (27/04 · Hábilis) · validar inscrição CAR. Sem serviço de terceiro adicional.",
    ], fill=None, mono_cols=[1])

    # ===== BLOCO B · LAO · CHECKLIST CEMAM 029/2018 =====
    add_section_row(table,
                    "B · LAO · CHECKLIST DO DOSSIÊ DE RENOVAÇÃO · CEMAM 029/2018 · Proc. SEMMA 22760/2020",
                    BLOCO_B)

    lao_itens = [
        ("B.1",
         "LAO 033/2020 vigente · Licença de Operação Ambiental.",
         "Em mãos · vigente até 23/12/2026.",
         "Protocolar renovação até 26/06/2026 (180 d · meta) — mínimo legal CP 3.11: 24/08/2026 (120 d). Hábilis."),
        ("B.2",
         "Requerimento padrão SEMMA Guapó · formulário oficial de renovação.",
         "Não emitido.",
         "Preencher e protocolar junto ao dossiê. Hábilis · Fase F4 (20-26/06/2026)."),
        ("B.3",
         "Procuração pública · SEMMA · Prefeitura · SEMAD · Veredas (em nome da Hábilis).",
         "Procuração 2023 em nome de A3 Engenharia NÃO serve.",
         "Procuração nova com firma reconhecida (23/04/2026 · Cliente · cartório). Sem terceiro."),
        ("B.4",
         "Publicação em edital · Resolução CONAMA 006/1986 (item 6 do TR SEMMA).",
         "Não publicada.",
         "2 publicações (DOE/GO + jornal local) em 04/05/2026 · Hábilis · Fase F1. Serviço: agência de publicação legal."),
        ("B.5",
         "MCO · Memorial de Caracterização do Empreendimento atualizado + Anexo IX CEMAM 029/2018.",
         "MCE 2020 (SCAN0035) em mãos como referência. MCO atualizado não emitido.",
         "Produzir MCO + Anexo IX com ART. Entrega 28/05/2026. **Serviço vinculado: terceiro Eng. Civil/Ambiental + ART** (Hábilis subcontrata · CRBio coordena)."),
        ("B.6",
         "Anexos I e II da Resolução CONAMA 273/2000 (caracterização técnica de SASC + sistemas de proteção).",
         "Validação física pendente · inspeção 12/05/2026 (Eixo 2 · 7 pontos).",
         "Preencher Anexos I/II com ART. Entrega 29/05/2026. **Serviço vinculado: terceiro Eng. Civil/Ambiental + ART** (Hábilis subcontrata)."),
        ("B.7",
         "Plantas técnicas · Planta Baixa, SASC, mecânica · Piso Nível 1 · Mapa SIRGAS 2000.",
         "Plantas 05/2024 em mãos · .dwg disponível.",
         "Complementação de shapefile + quadro UTM SIRGAS 2000 · memorial fotográfico datado (CP 3.12). Hábilis."),
        ("B.8",
         "Laudo de Estanqueidade dos tanques · NBR 16795/2019 · periodicidade bianual (CP 4.3).",
         "Laudo 10/2024 em mãos (NBR 16795/2019) · cobre ciclo bianual até 10/2026.",
         "Apresentar laudo 10/2024 + informe técnico (Opção B · plano principal). Contingência (Opção A): novo teste em 30-60 d se SEMMA contestar INMETRO NCC 21.08786 vencido 29/04/2024 ou calibração Suporty venc. 03/2023. **Serviço vinculado (contingência): empresa de teste com INMETRO vigente.**"),
        ("B.9",
         "AVCB · Auto de Vistoria do Corpo de Bombeiros (CP 3.10).",
         "AVCB Cercon 161726 vigente até 20/10/2026 (pós-protocolo).",
         "Nada a providenciar para o protocolo. Renovação em frente própria após 20/10/2026."),
        ("B.10",
         "Contratos + NFs semestrais de manutenção da SAO + coleta de resíduos oleosos (CP 4.5/4.6/4.7).",
         "Bauer + Limpmil (SAO) + Ecofenix (óleo) identificados.",
         "Compilar contratos + NFs (fev+ago) até 20/05/2026. Cliente · validar licenciamento e cadastro ANP dos prestadores. **Serviço vinculado: Bauer · Limpmil · Ecofenix.**"),
        ("B.11",
         "Certificados de Coleta de resíduos · disponíveis por 3 anos (CP 4.7).",
         "Acervo a validar · arquivo físico do posto.",
         "Verificar arquivo físico na inspeção 12/05/2026 (Eixo 5). Se ausentes, solicitar 2ª via aos prestadores. Cliente + Hábilis."),
        ("B.12",
         "RCA · Relatório de Controle Ambiental semestral (Anexo XI CEMAM 029/2018 · CP 4.4) · pH · turbidez · O&G · SS · DBO · DQO · OD · MBAS · metais · TPH · BTEX/PAH.",
         "Laudo RCA 2024 (Água Viva Ambiental) REPROVADO · O&G 64 mg/L (3,2× o limite de 20).",
         "Sequência obrigatória: diagnóstico SAO 12/05 → adequação 15-25/05 (Hipótese A/B/C) → coleta 22/05 → laudo 03/06/2026. **Serviços vinculados: laboratório acreditado INMETRO + RT químico + adequação SAO (Bauer/Limpmil).**"),
        ("B.13",
         "PAE · Plano de Atendimento a Emergências (CP 4.9 · NBR 14276/2020 · NBR 15219/2022).",
         "PGR 2024 em mãos como base.",
         "Produzir PAE entregando 05/06/2026 · validar inspeção 12/05 Eixo 5 (brigada · kit · equipamentos · contatos atualizados). Hábilis."),
        ("B.14",
         "PGA / PCA · Plano de Gestão Ambiental + Plano de Controle Ambiental (CP 4.14).",
         "A validar no acervo do Cliente.",
         "Se presentes: revisar e integrar ao Relatório 19/06. Se ausentes: produzir novos com ART. Hábilis · prazo 19/06/2026."),
        ("B.15",
         "Drenagem pluvial independente das águas servidas (CP 4.10) · anuência de saneamento.",
         "Validar em campo · inspeção 12/05/2026 · Eixo 3.",
         "Anuência saneamento (SANEAGO ou laudo NBR 7229/1993 + 13969/1997 para fossa séptica) até 25/05/2026. **Serviço vinculado (se fossa): engenheiro civil para laudo NBR · R$ 3-8k.**"),
        ("B.16",
         "Outorga de captação de água · DURH002737.",
         "DURH002737 venceu 09/12/2023 · posto opera sem outorga vigente há 2+ anos.",
         "Novo protocolo via Sistema Veredas em 13/05/2026 (após levantamento hídrico na inspeção 12/05 · Eixo 2). Hábilis. Cai no §15 da IN 15/2026 SEMAD (sem ordem cronológica preservada)."),
        ("B.17",
         "Relatório Técnico Consolidado · peça-síntese do dossiê · ART Hábilis CRBio.",
         "Não produzido (peça final).",
         "Consolidar todas as peças anteriores. Entrega 19/06/2026. Hábilis · ART CRBio."),
        ("B.18",
         "CNDs · Certidões Negativas de Débitos · Municipal + Estadual.",
         "A emitir.",
         "Solicitação online · prazo F4. Cliente."),
        ("B.19",
         "DUAM · Documento Único de Arrecadação Municipal · taxa SEMMA + taxa SEMAD/Veredas.",
         "A emitir/pagar.",
         "Emissão e pagamento na fase F4 (20-26/06/2026). Cliente."),
    ]
    for it in lao_itens:
        add_data_row(table, ["LAO · Checklist", it[0], it[1], it[2], it[3]],
                     fill=None, mono_cols=[1], bold_cols=[1])

    # ===== BLOCO C · 28 CONDICIONANTES =====
    add_section_row(table,
                    "C · CONDICIONANTES DA LAO 033/2020 · 28 itens (íntegra) · Bloco 1 Orientativas (3.1-3.12) + Bloco 2 Específicas (4.1-4.16)",
                    BLOCO_C)

    cps = [
        # ----- Bloco 1 · Orientativas (3.1 a 3.12)
        ("3.1", "Veracidade dos documentos e projetos anexados ao processo 22760/2020.",
         "Atendida · acervo 14/04/2026 confere.",
         "Reafirmar no Relatório Consolidado (19/06). Hábilis."),
        ("3.2", "SEMMA pode modificar, suspender ou cancelar a licença.",
         "Cláusula declaratória.",
         "Cumprida por aceite · sem ação."),
        ("3.3", "Licença pode ser revogada por descumprimento.",
         "Cláusula declaratória.",
         "Cumprida por aceite · sem ação."),
        ("3.4", "Comunicação imediata à SEMMA em caso de acidentes.",
         "Sem incidentes registrados.",
         "Procedimento formalizado no PAE (entrega 05/06). Hábilis."),
        ("3.5", "Não autoriza modificações sem manifestação SEMMA.",
         "Conduta respeitada · sem reformas/ampliações irregulares declaradas.",
         "Validar em campo (inspeção 12/05 · Eixo 1). Ampliação futura → nova LI (CP 4.12)."),
        ("3.6", "Não dispensa outros alvarás federais/estaduais/municipais.",
         "Cláusula declaratória.",
         "Cumprida por aceite · sem ação direta. Reforça relevância do Compliance Paralelo."),
        ("3.7", "SEMMA reserva direito de novas exigências.",
         "Cláusula declaratória.",
         "Cumprida por aceite · sem ação."),
        ("3.8", "Suspensão automática se demais licenças vencerem.",
         "AVCB Cercon 161726 vigente até 20/10/2026 ✓ · DURH002737 vencida 09/12/2023 · Alvará Sanitário a validar.",
         "Regularizar outorga via Veredas (13/05) · validar Alvará Sanitário em frente paralela. Hábilis."),
        ("3.9", "Não produz efeitos de cessão ou aquisição.",
         "Cláusula declaratória.",
         "Cumprida por aceite · sem ação."),
        ("3.10", "Manter atualizado ANP e AVCB.",
         "AVCB Cercon 161726 ✓ vigente. Cadastro ANP do revendedor: tratado em Compliance Paralelo (frente separada).",
         "Manter AVCB vigente. Cadastro ANP correlato (Res. ANP 41/2013 · 58/2014 · 20/2009): atualizar no Compliance Paralelo · não é peça do dossiê."),
        ("3.11", "Renovação com antecedência mínima de 120 dias.",
         "Atendida · meta 26/06/2026 = 180 dias antes do vencimento (60 d a mais que o piso).",
         "Cumprir o protocolo até 26/06/2026. Hábilis."),
        ("3.12", "Premissas dos relatórios: .jpg + ABNT · fotos datadas UTM SIRGAS 2000 · .pdf/.dwg · shapefile · RT + ART.",
         "Plantas 05/2024 em .dwg ✓.",
         "Complementar shapefile + quadro UTM SIRGAS 2000 · memorial fotográfico datado · ART em todas as peças. Hábilis · 19/06."),
        # ----- Bloco 2 · Específicas (4.1 a 4.16)
        ("4.1", "Operação com Responsável Técnico habilitado.",
         "RTs originais: Edmilson Xavier de Souza · Osmar Roberto de Souza Filho.",
         "Substituir por RTs Hábilis no dossiê de renovação · ART Hábilis CRBio · procuração 23/04. Hábilis."),
        ("4.2", "Seguir Resolução CEMAM 029/2018.",
         "Resolução é a base normativa do plano.",
         "Cumprida por aderência ao plano. Sem ação adicional."),
        ("4.3", "Laudo de estanqueidade bianual (NBR 13.784 cancelada · NBR 16795/2019 vigente).",
         "Laudo 10/2024 (NBR 16795/2019) em mãos · cobre ciclo bianual até 10/2026.",
         "Apresentar laudo 10/2024 como está (Opção B · plano principal). Contingência (Opção A): novo teste 30-60 d. **Serviço vinculado (contingência): empresa de teste com INMETRO vigente.**"),
        ("4.4", "Relatório Ambiental semestral (efluentes · pH · turbidez · O&G · SS · DBO · DQO · OD · MBAS · metais · TPH · BTEX/PAH) · periodicidade fev+ago.",
         "RCA 2024 REPROVADO (O&G 64 mg/L · 3,2× o limite). Análise indeferida.",
         "Diagnóstico SAO 12/05 → adequação 15-25/05 → coleta 22/05 → laudo 03/06/2026. **Serviços vinculados: laboratório INMETRO + RT químico + adequação SAO (Bauer/Limpmil).**"),
        ("4.5", "Contrato + NFs semestrais de manutenção da SAO + coleta de resíduos oleosos · fev+ago.",
         "Bauer + Limpmil + Ecofenix identificados como prestadores.",
         "Compilar NFs semestrais até 20/05/2026 · validar licenciamento e cadastro ANP dos prestadores. Cliente. **Serviços vinculados: Bauer · Limpmil · Ecofenix.**"),
        ("4.6", "Óleo lubrificante usado destinado a empresa autorizada ANP.",
         "Ecofenix identificada · cadastro ANP a validar.",
         "Confirmar autorização ANP ativa da Ecofenix · juntar Certificado de Coleta. Cliente."),
        ("4.7", "Certificados de Coleta disponíveis por 3 anos.",
         "Acervo a validar · arquivo físico do posto.",
         "Verificar arquivo físico (inspeção 12/05 · Eixo 5). 2ª via aos prestadores se ausente. Cliente + Hábilis."),
        ("4.8", "CONAMA 273/2000 + Portaria 084/2005: descarga selada · câmaras de contenção · válvulas de retenção · parede dupla · monitoramento intersticial · impermeabilização · SAO.",
         "Validação física pendente · inspeção 12/05 · Eixo 2 (7 pts) + Eixo 3 (5 pts).",
         "Preencher Anexos I/II CONAMA 273 em 29/05. **Serviço vinculado: terceiro Eng. Civil/Ambiental + ART** (Hábilis subcontrata)."),
        ("4.9", "Plano de Gerenciamento de Risco · integridade · emergência · treinamento.",
         "PGR 2024 em mãos.",
         "Validar emergência via PAE 05/06 + treinamento via inspeção (Eixo 5). Capacitação NR-20 em Compliance Paralelo. Hábilis."),
        ("4.10", "Drenagem pluvial independente das águas servidas.",
         "Validar em campo · inspeção 12/05 · Eixo 3.",
         "Anuência de saneamento (SANEAGO ou laudo fossa NBR 7229/1993 + 13969/1997) até 25/05/2026. **Serviço vinculado (se fossa): engenheiro civil para laudo NBR · R$ 3-8k.**"),
        ("4.11", "Eficiência declarada é responsabilidade da empresa.",
         "Cláusula declaratória.",
         "Cumprida por aceite · sem ação."),
        ("4.12", "Ampliação/reforma exige nova LI (Licença de Instalação).",
         "Cláusula declaratória de controle.",
         "Cumprida por aceite · sem ação. Validar inexistência de obras informais (inspeção 12/05 · Eixo 1)."),
        ("4.13", "SEMMA promoverá avaliações de controle de poluição.",
         "Cláusula declaratória de fiscalização.",
         "Cumprida por aceite · sem ação."),
        ("4.14", "Atender Projeto Ambiental + PGA + PCA.",
         "PGA e PCA a validar no acervo do Cliente.",
         "Se presentes: revisar e integrar ao Relatório 19/06. Se ausentes: produzir novos com ART. Hábilis."),
        ("4.15", "Descumprimento causa suspensão/cassação.",
         "Cláusula declaratória de penalidade.",
         "Cumprida por aceite · sem ação."),
        ("4.16", "SEMMA reserva novas exigências (reiteração da CP 3.7).",
         "Cláusula declaratória reiterada.",
         "Cumprida por aceite · sem ação."),
    ]
    for cp in cps:
        add_data_row(table,
                     ["Condicionante", f"CP {cp[0]}", cp[1], cp[2], cp[3]],
                     fill=None, mono_cols=[1], bold_cols=[1])

    # ===== Aplica zebra (linha sim, linha não) ignorando header e seções merged =====
    bloco_fills = {BLOCO_A, BLOCO_B, BLOCO_C}
    data_idx = 0
    for r_idx, row in enumerate(table.rows):
        if r_idx == 0:
            continue
        # detecta seção merged (1ª célula = última)
        if row.cells[0]._tc is row.cells[-1]._tc:
            continue
        if data_idx % 2 == 1:
            for c in row.cells:
                # respeita células sem fill explícito
                set_cell_shading(c, ZEBRA)
        data_idx += 1

    # ===== Rodapé =====
    foot = doc.add_paragraph()
    foot.paragraph_format.space_before = Pt(8)
    styled(foot,
           "Hábilis Regularização Ambiental · CRBio · Auto Posto América · Guapó/GO · Quadro consolidado em 23/04/2026 · "
           "Fontes literais: LAO 033/2020 · Extrato de Condicionantes 05/2024 · Resolução CEMAM 029/2018 · TR SEMMA Guapó.",
           size=7.5, color=CINZA_5, mono=True)

    out = "/home/guilherme/Projetos VS CODE/Z+Z - América/PROJ_Z+Z_AMERICAPOSTO_GUAPÓ/Painel de Gestão/Quadro_Unico_AutoPostoAmerica_v1.docx"
    doc.save(out)
    print(f"OK · {out}")


if __name__ == "__main__":
    build()
