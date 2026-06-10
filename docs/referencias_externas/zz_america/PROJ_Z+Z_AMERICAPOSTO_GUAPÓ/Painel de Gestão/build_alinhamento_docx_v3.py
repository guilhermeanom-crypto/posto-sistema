#!/usr/bin/env python3
"""
Gerador DOCX · Alinhamento Executivo v3 · Auto Posto América.
Base: v1 oficial (74 itens, 41 entregas cronológicas, 4 frentes, Compliance Paralelo).
Refinamento v3: modelo de escopo em 4 camadas classificadas (Base · Contingente ·
Opcional Futura · Bastidor) + quadro central de Terceiros e Contingentes + discurso
de proteção de caixa. Não reduz escopo. Adiciona disciplina de acionamento.
Uso: python3 build_alinhamento_docx_v3.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Paleta Habilis
VERDE = RGBColor(0x1B, 0x5E, 0x20)
VERDE_L = RGBColor(0x4C, 0xAF, 0x50)
VERDE_BG = "E8F5E9"
AMBER = RGBColor(0xF5, 0x7F, 0x17)
AMBER_BG = "FFF8E1"
VERMELHO = RGBColor(0xB7, 0x1C, 0x1C)
VERMELHO_BG = "FFEBEE"
AZUL = RGBColor(0x0D, 0x47, 0xA1)
AZUL_BG = "E3F2FD"
ROXO = RGBColor(0x4A, 0x14, 0x8C)
ROXO_BG = "F3E5F5"
CINZA_9 = RGBColor(0x1C, 0x1C, 0x1E)
CINZA_7 = RGBColor(0x3A, 0x3A, 0x3C)
CINZA_5 = RGBColor(0x6C, 0x6C, 0x70)
CINZA_BG = "FAFAFA"
BRANCO = RGBColor(0xFF, 0xFF, 0xFF)
FONT_SANS = "Segoe UI"
FONT_MONO = "Consolas"

# Cores das 4 camadas
CAMADA_COLORS = {
    "base":      "1B5E20",  # verde · Escopo-Base
    "conting":   "F57F17",  # amber · Contingente
    "opcional":  "0D47A1",  # azul · Opcional Futura
    "bastidor":  "4A148C",  # roxo · Bastidor Estrategico
}
CAMADA_LABEL = {
    "base":     "BASE",
    "conting":  "CONT",
    "opcional": "OPC",
    "bastidor": "BAST",
}

# Cores por responsavel
RESP_COLORS = {
    "cliente":  "0D47A1",
    "habilis":  "1B5E20",
    "terceiro": "6C6C70",
}


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for name, val in edges.items():
        if val is None: continue
        size, color = val
        el = OxmlElement(f'w:{name}')
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(size)); el.set(qn('w:color'), color)
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def set_paragraph_shading(p, hex_color):
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    p_pr.append(shd)


def set_paragraph_border_left(p, size_pt, color_hex):
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), str(int(size_pt * 8)))
    left.set(qn('w:space'), '6'); left.set(qn('w:color'), color_hex)
    pbdr.append(left); p_pr.append(pbdr)


def styled_run(p, text, *, bold=False, italic=False, size=None, color=None, mono=False):
    run = p.add_run(text)
    font_name = FONT_MONO if mono else FONT_SANS
    run.font.name = font_name
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.find(qn('w:rFonts'))
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts'); r_pr.append(r_fonts)
    r_fonts.set(qn('w:ascii'), font_name); r_fonts.set(qn('w:hAnsi'), font_name); r_fonts.set(qn('w:cs'), font_name)
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    return run


def mixed_run(p, text, *, size=10, default_color=CINZA_7, bold_parts=None):
    if not bold_parts:
        styled_run(p, text, size=size, color=default_color); return
    remaining = text
    for target in bold_parts:
        idx = remaining.find(target)
        if idx == -1: continue
        if idx > 0: styled_run(p, remaining[:idx], size=size, color=default_color)
        styled_run(p, target, size=size, color=CINZA_9, bold=True)
        remaining = remaining[idx + len(target):]
    if remaining:
        styled_run(p, remaining, size=size, color=default_color)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def heading_section(doc, num, title, space_before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '14')
    bottom.set(qn('w:space'), '3'); bottom.set(qn('w:color'), '1B5E20')
    pbdr.append(bottom); p_pr.append(pbdr)
    styled_run(p, f"{num}  ", size=11, color=VERDE_L, mono=True)
    styled_run(p, title, size=15, color=VERDE, bold=True)


def heading_h3(doc, text, space_before=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    set_paragraph_border_left(p, 2, "1B5E20")
    p.paragraph_format.left_indent = Cm(0.2)
    styled_run(p, text, size=10.5, color=CINZA_9, bold=True)


def body_p(doc, text, bold_parts=None, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    mixed_run(p, text, bold_parts=bold_parts, size=size)


def lede(doc, text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.25)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_shading(p, VERDE_BG)
    set_paragraph_border_left(p, 2.5, "1B5E20")
    mixed_run(p, text, size=10.5, default_color=CINZA_9, bold_parts=bold_parts)


def callout(doc, title, body, kind="verde", bold_body_parts=None):
    bg = {"verde": VERDE_BG, "amber": AMBER_BG, "vermelho": VERMELHO_BG,
          "azul": AZUL_BG, "cinza": CINZA_BG, "roxo": ROXO_BG}.get(kind, VERDE_BG)
    border = {"verde": "1B5E20", "amber": "F57F17", "vermelho": "B71C1C",
              "azul": "0D47A1", "cinza": "6C6C70", "roxo": "4A148C"}.get(kind, "1B5E20")
    t = doc.add_table(rows=1, cols=1)
    c = t.cell(0, 0)
    set_cell_shading(c, bg)
    set_cell_border(c, left=(24, border))
    p1 = c.paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    styled_run(p1, title.upper(), size=9, color=CINZA_9, bold=True, mono=True)
    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    mixed_run(p2, body, size=9.5, default_color=CINZA_7, bold_parts=bold_body_parts)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(2)


def table_hdr(doc, headers, widths_cm=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_shading(c, "1B5E20")
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        styled_run(p, h.upper(), size=8, color=BRANCO, bold=True, mono=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if widths_cm and i < len(widths_cm): c.width = Cm(widths_cm[i])
    return t


def add_row(t, values, zebra=False, bold_cols=None, mono_cols=None, is_total=False, size=9):
    row = t.add_row()
    for i, val in enumerate(values):
        c = row.cells[i]
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
        bold = (bold_cols and i in bold_cols) or is_total
        mono = mono_cols and i in mono_cols
        styled_run(p, val, size=size, color=CINZA_9 if bold else CINZA_7, bold=bold, mono=mono)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        if is_total: set_cell_shading(c, VERDE_BG)
        elif zebra: set_cell_shading(c, CINZA_BG)


def bullet(doc, text, bold_parts=None, size=10):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    mixed_run(p, text, bold_parts=bold_parts, size=size)


def set_tbl_borders(t, color="E0E0E0", size=4):
    tbl_pr = t._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(size))
        el.set(qn('w:color'), color)
        borders.append(el)
    tbl_pr.append(borders)


def field_page(p):
    r = p.add_run()
    r.font.name = FONT_MONO; r.font.size = Pt(8); r.font.color.rgb = CINZA_5
    fld_b = OxmlElement('w:fldChar'); fld_b.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.text = 'PAGE'
    fld_e = OxmlElement('w:fldChar'); fld_e.set(qn('w:fldCharType'), 'end')
    r._r.append(fld_b); r._r.append(it); r._r.append(fld_e)


def field_numpages(p):
    r = p.add_run()
    r.font.name = FONT_MONO; r.font.size = Pt(8); r.font.color.rgb = CINZA_5
    fld_b = OxmlElement('w:fldChar'); fld_b.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.text = 'NUMPAGES'
    fld_e = OxmlElement('w:fldChar'); fld_e.set(qn('w:fldCharType'), 'end')
    r._r.append(fld_b); r._r.append(it); r._r.append(fld_e)


def camada_chip(cell, camada_kind):
    """Preenche uma celula com chip da camada."""
    set_cell_shading(cell, CAMADA_COLORS[camada_kind])
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(p, CAMADA_LABEL[camada_kind], size=7, color=BRANCO, bold=True, mono=True)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def build():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(1.5); s.bottom_margin = Cm(1.5)
    s.left_margin = Cm(1.8); s.right_margin = Cm(1.8)

    # Header
    h_p = s.header.paragraphs[0]
    styled_run(h_p, "Auto Posto América | Alinhamento Executivo v3", size=8, color=CINZA_5, mono=True)
    h_p.add_run("\t")
    styled_run(h_p, "Hábilis | Abr/2026", size=8, color=VERDE, bold=True, mono=True)

    # Footer
    f_p = s.footer.paragraphs[0]
    styled_run(f_p, "Hábilis Regularização Ambiental", size=8, color=CINZA_5, mono=True)
    f_p.add_run("\t")
    styled_run(f_p, "Pagina ", size=8, color=CINZA_5, mono=True)
    field_page(f_p)
    styled_run(f_p, " de ", size=8, color=CINZA_5, mono=True)
    field_numpages(f_p)

    # ====================================================================
    # CAPA
    # ====================================================================
    p = doc.add_paragraph()
    styled_run(p, "H  ", size=16, color=VERDE, bold=True)
    styled_run(p, "Hábilis", size=16, color=VERDE, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(40)
    styled_run(p, "REGULARIZAÇÃO AMBIENTAL", size=9, color=CINZA_5, mono=True)
    for _ in range(4): doc.add_paragraph()

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(p, "PROCESSO DE REGULARIZAÇÃO AMBIENTAL", size=10, color=VERDE, bold=True, mono=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(p, "Auto Posto América", size=30, color=CINZA_9, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(p, "Renovação LAO 033/2020 | SEMMA Guapó | Processo 22760/2020",
               size=13, color=CINZA_7)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    styled_run(p,
               "Alinhamento executivo v3 para validação do Grupo Z+Z. Após leitura técnica "
               "completa do empreendimento, refinamos a condução para um modelo de ",
               size=11, color=CINZA_7)
    styled_run(p, "4 camadas classificadas", size=11, color=CINZA_9, bold=True)
    styled_run(p,
               " (Escopo-Base, Contingente, Opcional Futura, Bastidor Estratégico): foco no "
               "núcleo necessário à renovação, com acionamento pontual de terceiros e execução "
               "complementar apenas quando houver exigência formal, bloqueio regulatório ou "
               "deliberação do cliente. Preserva as ",
               size=11, color=CINZA_7)
    styled_run(p, "4 frentes integradas (74 itens), o cronograma de 4 fases em 67 dias e o "
               "checklist cronológico de 41 entregas", size=11, color=CINZA_9, bold=True)
    styled_run(p, " do plano oficial, agora com disciplina de acionamento e proteção de caixa.",
               size=11, color=CINZA_7)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    styled_run(p, "Protocolo-meta ", size=11, color=CINZA_7)
    styled_run(p, "26/06/2026 (180 dias antes do vencimento da LAO). Mínimo legal CP 3.11: "
               "24/08/2026 (120 dias).", size=11, color=CINZA_9, bold=True)

    for _ in range(1): doc.add_paragraph()

    tc = doc.add_table(rows=2, cols=4)
    for i, h in enumerate(["DESTINATÁRIO", "EMISSOR", "PROTOCOLO-META", "EMISSÃO"]):
        c = tc.cell(0, i)
        styled_run(c.paragraphs[0], h, size=8, color=CINZA_5, bold=True, mono=True)
        set_cell_border(c, top=(12, "1B5E20"))
    vals = ["Grupo Z+Z", "Hábilis", "26/06/2026", "22/04/2026"]
    for i, v in enumerate(vals):
        c = tc.cell(1, i)
        styled_run(c.paragraphs[0], v, size=10, color=CINZA_9, bold=True, mono=i >= 2)

    add_page_break(doc)

    # ====================================================================
    # SECAO 01 · PANORAMA
    # ====================================================================
    heading_section(doc, "01", "Panorama", space_before=0)

    lede(doc,
         "A LAO 033/2020 do Auto Posto América vence em 23/12/2026. A Hábilis conduz a "
         "renovação com protocolo-meta em 26/06/2026 (180 dias antes do vencimento), "
         "preservando folga regulatória para eventuais exigências complementares da SEMMA. "
         "A base documental recebida em 14/04/2026 suporta o início da execução sem fato "
         "impeditivo identificado. O escopo técnico é o do plano oficial: nenhum item foi "
         "removido. O refinamento v3 introduz a disciplina de 4 camadas — a execução permanece "
         "completa, o acionamento se torna cirúrgico.",
         ["23/12/2026", "26/06/2026", "180 dias antes", "14/04/2026",
          "nenhum item foi removido", "4 camadas", "acionamento se torna cirúrgico"])

    heading_h3(doc, "Quatro frentes integradas | 74 itens")
    t = table_hdr(doc, ["Frente", "Itens", "Órgão", "Escopo"], [4.0, 1.5, 3.2, 8.5])
    add_row(t, ["Uso do Solo", "11", "Prefeitura Guapó",
                "Pré-requisito para protocolo da renovação."],
            bold_cols=[0, 1], mono_cols=[1])
    add_row(t, ["Renovação LAO", "26", "SEMMA (Proc. 22760/2020)",
                "Dossiê técnico principal do processo."],
            zebra=True, bold_cols=[0, 1], mono_cols=[1])
    add_row(t, ["Condicionantes LAO 033/2020", "14", "SEMMA",
                "15 atendidas | 3 parciais | 8 pendentes."],
            bold_cols=[0, 1], mono_cols=[1])
    add_row(t, ["Compliance Paralelo", "23", "ANP | IBAMA | NR-20 | PNRS",
                "Cadência própria (detalhe no anexo)."],
            zebra=True, bold_cols=[0, 1], mono_cols=[1])
    add_row(t, ["Total", "74", "4 órgãos/frentes",
                "Plano único, execução coordenada."],
            bold_cols=[0, 1, 2, 3], mono_cols=[1], is_total=True)

    # ====================================================================
    # SECAO 02 · MODELO DE ESCOPO EM 4 CAMADAS (NOVO)
    # ====================================================================
    heading_section(doc, "02", "Modelo de escopo em 4 camadas", space_before=12)

    body_p(doc,
           "A condução v3 classifica cada item do plano em uma de quatro camadas. Essa disciplina "
           "evita superdimensionamento sem reduzir escopo técnico: a Hábilis mantém a leitura "
           "completa do empreendimento e expõe, para cada peça, qual é o gatilho de acionamento. "
           "O cliente enxerga o núcleo fechado e previsível separado das frentes condicionadas, "
           "opcionais e do bastidor técnico.",
           ["gatilho de acionamento", "núcleo fechado e previsível"], size=10)

    # Tabela das 4 camadas
    t = table_hdr(doc, ["Camada", "O que entra", "Orçamento", "Gatilho de acionamento"],
                  [3.0, 7.0, 2.5, 5.0])
    # Linha 1 - Base
    row = t.add_row()
    camada_chip(row.cells[0], "base")
    # Override: texto maior na coluna camada
    row.cells[0].paragraphs[0].clear()
    p = row.cells[0].paragraphs[0]
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(p, "1 · ESCOPO-BASE", size=8.5, color=BRANCO, bold=True, mono=True)
    styled_run(row.cells[1].paragraphs[0],
               "Indispensável para protocolar. Núcleo TR SEMMA + peças literais da LAO 033/2020 + "
               "28 condicionantes + Uso do Solo.",
               size=9, color=CINZA_7)
    styled_run(row.cells[2].paragraphs[0], "Fechado · previsível",
               size=9, color=CINZA_9, bold=True, mono=True)
    styled_run(row.cells[3].paragraphs[0], "Imediato · contratado",
               size=9, color=CINZA_7, mono=True)

    # Linha 2 - Contingente
    row = t.add_row()
    set_cell_shading(row.cells[0], CAMADA_COLORS["conting"])
    row.cells[0].paragraphs[0].clear()
    p = row.cells[0].paragraphs[0]
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(p, "2 · CONTINGENTE", size=8.5, color=BRANCO, bold=True, mono=True)
    row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_shading(row.cells[1], CINZA_BG)
    set_cell_shading(row.cells[2], CINZA_BG)
    set_cell_shading(row.cells[3], CINZA_BG)
    styled_run(row.cells[1].paragraphs[0],
               "Pode ser exigido, mas ainda não é certo. Investigação confirmatória, sondagens SPT, "
               "análise de risco, estanqueidade extra, adequações físicas.",
               size=9, color=CINZA_7)
    styled_run(row.cells[2].paragraphs[0], "Separado · forecast",
               size=9, color=CINZA_9, bold=True, mono=True)
    styled_run(row.cells[3].paragraphs[0], "Exigência formal · deliberação",
               size=9, color=CINZA_7, mono=True)

    # Linha 3 - Opcional Futura
    row = t.add_row()
    set_cell_shading(row.cells[0], CAMADA_COLORS["opcional"])
    row.cells[0].paragraphs[0].clear()
    p = row.cells[0].paragraphs[0]
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(p, "3 · OPCIONAL FUTURA", size=8.5, color=BRANCO, bold=True, mono=True)
    row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    styled_run(row.cells[1].paragraphs[0],
               "Importante, mas não trava a renovação agora. Compliance Paralelo (ANP, IBAMA, "
               "NR-20, PNRS, alvará sanitário).",
               size=9, color=CINZA_7)
    styled_run(row.cells[2].paragraphs[0], "Aditivo anual",
               size=9, color=CINZA_9, bold=True, mono=True)
    styled_run(row.cells[3].paragraphs[0], "Cadência própria · paralela",
               size=9, color=CINZA_7, mono=True)

    # Linha 4 - Bastidor
    row = t.add_row()
    set_cell_shading(row.cells[0], CAMADA_COLORS["bastidor"])
    row.cells[0].paragraphs[0].clear()
    p = row.cells[0].paragraphs[0]
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(p, "4 · BASTIDOR", size=8.5, color=BRANCO, bold=True, mono=True)
    row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_shading(row.cells[1], CINZA_BG)
    set_cell_shading(row.cells[2], CINZA_BG)
    set_cell_shading(row.cells[3], CINZA_BG)
    styled_run(row.cells[1].paragraphs[0],
               "Inteligência técnica de consultoria. Monitoramento de hipóteses, leitura de "
               "passivo, posturas defensivas, cenários de fiscalização — não vira obrigação.",
               size=9, color=CINZA_7)
    styled_run(row.cells[2].paragraphs[0], "Sem custo adicional",
               size=9, color=CINZA_9, bold=True, mono=True)
    styled_run(row.cells[3].paragraphs[0], "Interno Hábilis",
               size=9, color=CINZA_7, mono=True)

    # Callout de proteção de caixa (preview - versão completa no final)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_shading(p, AMBER_BG)
    set_paragraph_border_left(p, 2.5, "F57F17")
    p.paragraph_format.left_indent = Cm(0.25)
    styled_run(p, "Proteção de caixa · ", size=9.5, color=CINZA_9, bold=True)
    styled_run(p,
               "Camadas 2, 3 e 4 não são custo comprometido. São forecast controlado. Nenhum "
               "item dessas camadas é acionado sem exigência formal do órgão, bloqueio "
               "regulatório concreto ou deliberação explícita do Grupo Z+Z. O escopo-base é "
               "fechado. O restante é gatilho.",
               size=9.5, color=CINZA_7)

    add_page_break(doc)

    # ====================================================================
    # SECAO 03 · QUADRO DE TERCEIROS E CONTINGENTES (NOVO · PECA CENTRAL)
    # ====================================================================
    heading_section(doc, "03", "Quadro de Terceiros e Contingentes", space_before=0)

    body_p(doc,
           "Todos os itens que envolvem terceiros ou peças condicionadas, classificados por "
           "gatilho de acionamento. A coluna Faixa de Custo fica em branco para preenchimento "
           "após cotação formal (Hábilis provê indicações). Nada é contratado de ofício: cada "
           "linha explicita a base técnica ou regulatória do item e a condição objetiva que "
           "destrava o acionamento, para evidenciar que não se trata de levantamento arbitrário.",
           ["gatilho de acionamento", "em branco para preenchimento",
            "base técnica ou regulatória", "não se trata de levantamento arbitrário"], size=10)

    heading_h3(doc, "Classificação por gatilho", space_before=4)

    # Mini-legenda das 4 classificações
    legenda = doc.add_table(rows=1, cols=4)
    legenda.alignment = WD_TABLE_ALIGNMENT.LEFT
    legenda_data = [
        ("Obrigatório imediato", CAMADA_COLORS["base"]),
        ("Provável", CAMADA_COLORS["conting"]),
        ("Contingente", CAMADA_COLORS["conting"]),
        ("Opcional", CAMADA_COLORS["opcional"]),
    ]
    for i, (label, color) in enumerate(legenda_data):
        c = legenda.rows[0].cells[i]
        set_cell_shading(c, color)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(p, label.upper(), size=7.5, color=BRANCO, bold=True, mono=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # TABELA PRINCIPAL: Item | Classificação | Faixa Custo | Base objetiva do acionamento | Fase
    t = table_hdr(doc,
                  ["Item", "Classificação", "Faixa de Custo", "Base objetiva do acionamento", "Fase"],
                  [5.0, 2.4, 2.4, 6.0, 0.9])

    CLASSIF_COLORS = {
        "obrig":   CAMADA_COLORS["base"],
        "prov":    CAMADA_COLORS["conting"],
        "cont":    CAMADA_COLORS["conting"],
        "opc":     CAMADA_COLORS["opcional"],
    }
    CLASSIF_LABEL = {
        "obrig":   "OBRIG. IMEDIATO",
        "prov":    "PROVÁVEL",
        "cont":    "CONTINGENTE",
        "opc":     "OPCIONAL",
    }

    terceiros = [
        # (item, classificação, base_objetiva, fase)
        ("Eng. Civil/Ambiental · ART do MCO (Anexo IX CEMAM 029/2018)",
         "obrig",
         "Base: o MCO integra o dossiê principal da renovação e exige responsabilidade técnica compatível. "
         "Aciona porque a Hábilis não possui Eng. Civil/Ambiental interno para essa ART; portanto a execução "
         "ocorre por subcontratação sob coordenação CRBio.",
         "F2"),
        ("Eng. Civil/Ambiental · ART dos Anexos I e II CONAMA 273/2000",
         "obrig",
         "Base: item 23 do TR SEMMA exige os Anexos I e II com responsabilidade técnica própria. "
         "Aciona porque essa peça não pode ser absorvida pela ART do MCO; usa-se o mesmo terceiro para "
         "manter coerência técnica e reduzir retrabalho.",
         "F2"),
        ("Laboratório acreditado INMETRO · RCA semestral (CP 4.4 LAO)",
         "prov",
         "Base: a CP 4.4 da LAO exige RCA semestral com parâmetros analíticos compatíveis. "
         "Aciona quando o diagnóstico de 12/05 fechar o escopo da coleta e confirmar quais parâmetros "
         "precisam ser analisados por laboratório acreditado.",
         "F2"),
        ("Laboratório · nova coleta de efluentes e água (subsidia RCA)",
         "prov",
         "Base: o RCA não pode ser sustentado apenas por acervo antigo se a condição operacional atual "
         "precisa ser demonstrada. Aciona quando a inspeção de 12/05 validar a SAO e indicar nova coleta "
         "como evidência atualizada; o laudo sai com ART do químico responsável.",
         "F2"),
        ("Publicação CONAMA 006/86 · DOE/GO + jornal de grande circulação",
         "obrig",
         "Base: item 6 do TR SEMMA exige publicação formal do pedido de licenciamento. "
         "Aciona porque essa etapa depende de veículo oficial habilitado e não de produção interna.",
         "F1"),
        ("Investigação confirmatória · NBR 15515-2/3 + NBR 16209",
         "cont",
         "Base: investigação confirmatória só se justifica quando houver indício objetivo de passivo. "
         "Aciona apenas se a SEMMA formalizar exigência após a publicação CONAMA 006/86 ou se a análise "
         "documental/técnica apontar evidência consistente de contaminação.",
         "F2/F3"),
        ("Sondagens SPT (mín. 3 furos) · VOC, BTEX, PAH, metais, TPH",
         "cont",
         "Base: as sondagens não são frente autônoma; elas existem apenas como etapa executiva da investigação "
         "confirmatória. Aciona em conjunto com ela, quando houver necessidade de comprovação amostral em campo.",
         "F2/F3"),
        ("Análise de Risco Preliminar (APP) ou Quantitativa (ARQ)",
         "cont",
         "Base: APP/ARQ é medida de aprofundamento, não peça ordinária da renovação. "
         "Aciona somente se a SEMMA exigir tecnicamente esse nível adicional de avaliação, nos termos da "
         "CEMAM 029/2018 e CONAMA 420/2009.",
         "F3"),
        ("Teste de estanqueidade extra · NBR 16795/2019 (empresa INMETRO)",
         "cont",
         "Base: já existe laudo de estanqueidade 10/2024 em mãos. "
         "Aciona apenas se a SEMMA contestar sua suficiência, validade, cadeia metrológica ou aderência formal.",
         "F3"),
        ("Adequação física da SAO · serviços de obra civil",
         "cont",
         "Base: obra civil só se sustenta por não conformidade material, e não por cautela genérica. "
         "Aciona somente se a inspeção de 12/05 ou os resultados analíticos apontarem falha física na SAO.",
         "F2/F3"),
        ("Laudo de anuência de saneamento ou fossa/sumidouro",
         "cont",
         "Base: a solução depende do regime real de esgotamento. "
         "Se houver rede pública, basta anuência da concessionária, sem custo de terceiro; se houver "
         "fossa/sumidouro, aciona laudo técnico com ART para comprovação formal.",
         "F2"),
        ("Recontrato de serviços terceirizados (Bauer, Limpmil, Ecofenix)",
         "prov",
         "Base: as CPs 4.5, 4.6 e 4.7 exigem continuidade documental e operacional dos serviços. "
         "Aciona apenas se a validação dos contratos vigentes identificar irregularidade, lacuna de licença, "
         "falha de rastreabilidade ou insuficiência de evidência.",
         "F2"),
        ("Levantamento topográfico complementar (UTM SIRGAS 2000)",
         "opc",
         "Base: o protocolo parte das plantas e referências já existentes. "
         "Aciona somente se o Veredas exigir memorial topográfico complementar ou se a validação de campo "
         "mostrar insuficiência do material atual.",
         "F2"),
        ("Projeto de adequação paisagística / APP (se exigido)",
         "opc",
         "Base: não há evidência preliminar de passivo paisagístico ou APP comprometida. "
         "Aciona apenas se a inspeção de 12/05 identificar intervenção, conflito locacional ou exigência "
         "de recomposição específica.",
         "F3"),
    ]

    for item, klass, cond, fase in terceiros:
        row = t.add_row()
        # Item
        c = row.cells[0]
        pc = c.paragraphs[0]
        pc.paragraph_format.space_before = Pt(1); pc.paragraph_format.space_after = Pt(1)
        styled_run(pc, item, size=8.5, color=CINZA_9, bold=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        # Classificação (chip colorido)
        c = row.cells[1]
        set_cell_shading(c, CLASSIF_COLORS[klass])
        pc = c.paragraphs[0]
        pc.paragraph_format.space_before = Pt(1); pc.paragraph_format.space_after = Pt(1)
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(pc, CLASSIF_LABEL[klass], size=7, color=BRANCO, bold=True, mono=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Faixa de custo (em branco)
        c = row.cells[2]
        set_cell_shading(c, CINZA_BG)
        pc = c.paragraphs[0]
        pc.paragraph_format.space_before = Pt(1); pc.paragraph_format.space_after = Pt(1)
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(pc, "R$ ______________", size=8, color=CINZA_5, mono=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Condição de acionamento
        c = row.cells[3]
        pc = c.paragraphs[0]
        pc.paragraph_format.space_before = Pt(1); pc.paragraph_format.space_after = Pt(1)
        styled_run(pc, cond, size=8, color=CINZA_7)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        # Fase
        c = row.cells[4]
        pc = c.paragraphs[0]
        pc.paragraph_format.space_before = Pt(1); pc.paragraph_format.space_after = Pt(1)
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(pc, fase, size=8, color=CINZA_9, bold=True, mono=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Callout sobre preenchimento
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_shading(p, AZUL_BG)
    set_paragraph_border_left(p, 2, "0D47A1")
    p.paragraph_format.left_indent = Cm(0.25)
    styled_run(p, "Preenchimento das faixas de custo · ", size=9, color=CINZA_9, bold=True)
    styled_run(p,
               "Coluna em branco intencionalmente. As cotações formais serão incorporadas após "
               "fechamento do briefing com terceiros indicados (indicação Hábilis). Esta peça "
               "entrega ao Grupo Z+Z a visão completa de quais itens existem, em qual camada "
               "estão e sob qual gatilho são acionados — preservando decisão orçamentária no "
               "momento certo.",
               size=9, color=CINZA_7)

    add_page_break(doc)

    # ====================================================================
    # SECAO 04 · CRONOGRAMA
    # ====================================================================
    heading_section(doc, "04", "Cronograma | 4 fases em 67 dias", space_before=0)

    body_p(doc,
           "Janela de execução de 67 dias, de 20/04/2026 a 26/06/2026. Cada fase tem "
           "entregáveis próprios e depende de ações do Grupo Z+Z (seção 05). Atrasos em "
           "F1 comprimem F3 e F4 proporcionalmente.",
           ["67 dias", "20/04/2026", "26/06/2026"], size=9.5)

    t = table_hdr(doc, ["Fase", "Período", "Dias", "Entregáveis-chave"], [3.3, 3.4, 1.2, 9.3])
    add_row(t,
            ["F1 | Estruturação", "20/04 -> 08/05", "19",
             "Procuração, dados societários, contratação de terceiros da Camada 1 (Eng. Civil/Ambiental), Uso do Solo protocolado, status outorga consolidado, publicação CONAMA 006/86."],
            bold_cols=[0, 1, 2], mono_cols=[1, 2], size=8.5)
    add_row(t,
            ["F2 | Consolidação", "09/05 -> 30/05", "22",
             "Inspeção em campo (gatilho de itens da Camada 2), novas coletas laboratoriais, MCO atualizado, Anexos I/II CONAMA 273, Plano das 28 CPs."],
            zebra=True, bold_cols=[0, 1, 2], mono_cols=[1, 2], size=8.5)
    add_row(t,
            ["F3 | Execução e Evidência", "31/05 -> 19/06", "20",
             "RCA, PAE, PGRS, recomposição das CPs parciais, Relatório Consolidado. Contingentes da Camada 2 só se exigência formal SEMMA."],
            bold_cols=[0, 1, 2], mono_cols=[1, 2], size=8.5)
    add_row(t,
            ["F4 | Protocolo", "20/06 -> 26/06", "7",
             "Aprovação final, CNDs, checklist cruzado, taxa paga, protocolo no Processo 22760/2020 | SEMMA."],
            zebra=True, bold_cols=[0, 1, 2], mono_cols=[1, 2], size=8.5)

    # ====================================================================
    # SECAO 05 · CHECKLIST CRONOLOGICO (41 itens + coluna Camada)
    # ====================================================================
    heading_section(doc, "05", "Checklist cronológico | 41 itens", space_before=12)

    body_p(doc,
           "Itens em ordem cronológica conforme documento oficial. Responsável identificado por "
           "cor (Cliente azul, Hábilis verde, Terceiro cinza). Nova coluna Camada indica a "
           "classificação v3: Base (verde) é imediato; Contingente (amber) depende de gatilho; "
           "Opcional (azul) fica no Compliance Paralelo. MCO e Anexos I/II CONAMA 273 executados "
           "por terceiro Eng. Civil/Ambiental sob coordenação CRBio da Hábilis.",
           ["Cliente", "Hábilis", "Terceiro", "Camada",
            "Base", "Contingente", "Opcional"], size=9.5)

    # Legenda de responsaveis
    legenda = doc.add_table(rows=1, cols=3)
    legenda.alignment = WD_TABLE_ALIGNMENT.LEFT
    legenda_data = [
        ("Cliente", "cliente"), ("Hábilis", "habilis"), ("Terceiro", "terceiro"),
    ]
    for i, (label, kind) in enumerate(legenda_data):
        c = legenda.rows[0].cells[i]
        set_cell_shading(c, RESP_COLORS[kind])
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(p, label.upper(), size=8, color=BRANCO, bold=True, mono=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Tabela cronologica COM coluna Camada
    t = table_hdr(doc, ["Prazo", "Item", "Resp.", "Fase", "Camada", "Observação"],
                  [1.4, 4.7, 2.3, 0.8, 1.0, 6.5])

    # itens: (prazo, item, resp, resp_kind, fase, camada_kind, obs)
    itens_chrono = [
        ("22/04", "Coleta de RG/CPF dos sócios", "Cliente", "cliente", "F1", "base",
         "RG e CPF de Sérgio Marques e José Guilherme. Exigência dos dois processos (Uso do Solo + LAO)."),
        ("23/04", "Procuração pública para a Hábilis", "Cliente", "cliente", "F1", "base",
         "Firma reconhecida. Poderes para SEMMA, Prefeitura, SEMAD/GO e Veredas. Substitui a procuração A3 de 2023."),
        ("24/04", "Cartão CNPJ + Comprovante de Endereço", "Cliente", "cliente", "F1", "base",
         "CNPJ ativo conferindo CNAE/endereço. Comprovante de endereço dos sócios nos últimos 3 meses."),
        ("25/04", "Verificação formal da Certidão de Uso do Solo", "Cliente", "cliente", "F1", "base",
         "Confirmar existência, número, data e disponibilidade. Apoio da Hábilis."),
        ("27/04", "Enquadramento do imóvel + CAR", "Hábilis", "habilis", "F1", "base",
         "Zona rural confirmada via LAO. CAR obrigatório. ART de Uso do Solo via Eng. Agronômica (CREA)."),
        ("27/04", "ART específica para Uso do Solo", "Hábilis", "habilis", "F1", "base",
         "ART distinta da Renovação LAO. Assinada pela Hábilis via RT Eng. Agronômica."),
        ("28/04", "Diagnóstico de outorga no Web Outorga", "Hábilis", "habilis", "F1", "base",
         "Consulta ao Sistema Web Outorga para histórico da DURH002737 (processos legados ainda não migrados ao Veredas). Protocolo novo no Veredas em 13/05. DURH002737 vencida 09/12/2023."),
        ("30/04", "Protocolo do requerimento da Certidão", "Hábilis", "habilis", "F1", "base",
         "Somente se verificação indicar ausência ou desatualização."),
        ("05/05", "Contratação: investigação confirmatória", "Cliente", "cliente", "F1", "conting",
         "Empresa NBR 15515-2/3 e NBR 16209. ART geólogo/eng. ambiental. Indicação Hábilis. Acionamento condicionado (seção 03)."),
        ("08/05", "Análise documental consolidada", "Hábilis", "habilis", "F1", "base",
         "Fechamento da triagem do lote de 14/04 e identificação de lacunas."),
        ("04/05", "Publicação do pedido de licenciamento (CONAMA 006/86)", "Hábilis", "habilis", "F1", "base",
         "Jornal local + Diário Oficial. Aguardar 15-30 dias para contestação pública."),
        ("10/05", "Mapeamento técnico das 12 CPs com ação", "Hábilis", "habilis", "F1", "base",
         "Identificação do tipo de intervenção por CP pendente (documental, física, contratação)."),
        ("22/04", "Validação do AVCB Cercon 161726", "Hábilis", "habilis", "F1", "base",
         "Certificado CBMGO 161726, emitido 07/11/2025, válido até 20/10/2026. Em mãos."),
        ("12/05", "Inspeção técnica de campo", "Hábilis", "habilis", "F2", "base",
         "25 pontos em 5 eixos: implantação/APP/RL/CAR, armazenamento, efluentes/SAO, resíduos, segurança. É o gatilho que confirma ou dispensa itens da Camada 2."),
        ("25/05", "Anuência de saneamento ou laudo de fossa", "Hábilis", "habilis", "F2", "base",
         "Rede pública: anuência da concessionária. Fossa/sumidouro: laudo técnico com ART."),
        ("12/05", "Validação das plantas em campo", "Hábilis", "habilis", "F2", "base",
         "Plantas 05/2024 vs configuração atual. Quadro de Áreas + UTM SIRGAS 2000, APP e RL."),
        ("12/05", "Levantamento hídrico: define regime Veredas", "Hábilis", "habilis", "F2", "base",
         "Rede pública: dispensa simplificada. Captação direta: processo completo com memorial técnico."),
        ("13/05", "Protocolo novo de outorga via Veredas", "Hábilis", "habilis", "F2", "base",
         "Dia útil seguinte à inspeção. Gera certidão de tramitação que já compõe o dossiê SEMMA. IN 15/2026 §15."),
        ("15/05", "Contratação dos serviços de terceiros", "Cliente", "cliente", "F2", "base",
         "Laboratório, estanqueidade, projetista e Eng. Civil/Ambiental (MCO + Anexos CONAMA 273). Indicações da Hábilis."),
        ("18/05", "NFs e certificados dos equipamentos", "Cliente", "cliente", "F2", "base",
         "NFs de tanques/tubulações + certificados CONAMA 319/2002 e INMETRO 37/109/110/111/009."),
        ("20/05", "Contratos de manutenção, limpeza e destinação", "Cliente", "cliente", "F2", "base",
         "Ecofenix, Bauer, Limpmil. Validar ANP + licenciamento + autorização de transporte."),
        ("22/05", "Nova análise de efluentes", "Terceiro", "terceiro", "F2", "base",
         "Laboratório indicado pela Hábilis. ART do químico responsável acompanha o laudo."),
        ("22/05", "Nova análise de água", "Terceiro", "terceiro", "F2", "base",
         "Conjugada com análise de efluentes para otimizar coleta. ART químico."),
        ("25/05", "Mobilização da investigação confirmatória", "Terceiro", "terceiro", "F2", "conting",
         "Sondagens SPT, mínimo 3 furos. VOC, BTEX, PAH, metais, TPH. CEMAM 029/2018 + CONAMA 420/2009. Condicionada a gatilho (seção 03)."),
        ("28/05", "Atualização do MCO", "Terceiro", "terceiro", "F2", "base",
         "Peça central do dossiê (Anexo IX CEMAM 029/2018). Eng. Civil/Ambiental subcontratado sob coordenação CRBio. ART do profissional contratado."),
        ("29/05", "Anexo I e II da CONAMA 273/2000", "Terceiro", "terceiro", "F2", "base",
         "Anexo I tabular + Anexo II classifica risco (3 classes). ART própria (item 23 TR)."),
        ("30/05", "Plano das 28 condicionantes", "Hábilis", "habilis", "F2", "base",
         "Documento interno: responsável, evidência e prazo por condicionante."),
        ("Conting.", "Teste de estanqueidade (contingência)", "Terceiro", "terceiro", "F3", "conting",
         "Acionado apenas se SEMMA contestar laudo 10/2024. NBR 16795/2019. Empresa INMETRO vigente."),
        ("12/06", "Recomposição das 10 CPs parciais", "Hábilis", "habilis", "F3", "base",
         "Recomposição documental conforme Plano de Atendimento de 30/05."),
        ("02/06", "PGRS · Gerenciamento de Resíduos", "Hábilis", "habilis", "F3", "base",
         "Caracterização, rotas de destinação e responsáveis. ART RT Ambiental (CRBio) ou Agronômica."),
        ("03/06", "RCA · Relatório de Controle Ambiental", "Hábilis", "habilis", "F3", "base",
         "Anexo XI CEMAM 029/2018. CP 4.4. Análises SAO + poços. ART RT Ambiental (CRBio)."),
        ("05/06", "PAE · Atendimento a Emergências", "Hábilis", "habilis", "F3", "base",
         "NBR 14276 + 15219. CP 4.9. 9 capítulos, 10 cenários específicos. ART RT Ambiental."),
        ("10/06", "Análise de Risco (contingência)", "Terceiro", "terceiro", "F3", "conting",
         "Acionada se exigida pela SEMMA. APP ou ARQ, conforme classificação CONAMA 420/2009."),
        ("15/06", "Laudo de Investigação Confirmatória", "Terceiro", "terceiro", "F3", "conting",
         "CEMAM 029/2018 + CONAMA 420/2009. VOC, BTEX, PAH, metais, TPH + análise de riscos. Condicionada."),
        ("15/06", "Execução das 2 CPs pendentes", "Terceiro", "terceiro", "F3", "base",
         "Serviços físicos ou ensaios: CP 4.4 (RCA semestral) + CP 4.14 (PGA/PCA a validar)."),
        ("19/06", "Relatório técnico consolidado", "Hábilis", "habilis", "F3", "base",
         "Peça integradora: MCO, laudos, plantas, plano de CPs, evidências. ART do RT coordenador CRBio."),
        ("22/06", "Aprovação formal do relatório", "Cliente", "cliente", "F4", "base",
         "Aprovação escrita do Grupo Z+Z antes do protocolo."),
        ("23/06", "CND Municipal + CND Estadual", "Cliente", "cliente", "F4", "base",
         "Certidões Negativas de Débitos emitidas próximas ao protocolo. Validade 90 dias."),
        ("24/06", "Checklist final cruzado", "Hábilis", "habilis", "F4", "base",
         "Conferência de peças, assinaturas, ARTs e taxas antes da submissão."),
        ("25/06", "Pagamento das taxas de protocolo", "Cliente", "cliente", "F4", "base",
         "Taxas SEMMA, municipais e demais custos da submissão."),
        ("26/06", "Protocolo no Processo SEMMA 22760/2020", "Hábilis", "habilis", "F4", "base",
         "Submissão do dossiê completo. Acompanhamento pós-protocolo se inicia."),
    ]

    for prazo, item, resp, resp_kind, fase, camada_kind, obs in itens_chrono:
        row = t.add_row()
        # Prazo
        c = row.cells[0]
        pc = c.paragraphs[0]
        pc.paragraph_format.space_before = Pt(1); pc.paragraph_format.space_after = Pt(1)
        styled_run(pc, prazo, size=8, color=CINZA_9, bold=True, mono=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        # Item
        c = row.cells[1]
        pc = c.paragraphs[0]
        pc.paragraph_format.space_before = Pt(1); pc.paragraph_format.space_after = Pt(1)
        styled_run(pc, item, size=8.5, color=CINZA_9, bold=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        # Responsavel
        c = row.cells[2]
        set_cell_shading(c, RESP_COLORS[resp_kind])
        pc = c.paragraphs[0]
        pc.paragraph_format.space_before = Pt(1); pc.paragraph_format.space_after = Pt(1)
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(pc, resp.upper(), size=7.5, color=BRANCO, bold=True, mono=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Fase
        c = row.cells[3]
        pc = c.paragraphs[0]
        pc.paragraph_format.space_before = Pt(1); pc.paragraph_format.space_after = Pt(1)
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(pc, fase, size=8, color=CINZA_9, bold=True, mono=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Camada
        c = row.cells[4]
        camada_chip(c, camada_kind)
        # Observacao
        c = row.cells[5]
        pc = c.paragraphs[0]
        pc.paragraph_format.space_before = Pt(1); pc.paragraph_format.space_after = Pt(1)
        styled_run(pc, obs, size=8, color=CINZA_7)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # ====================================================================
    # SECAO 06 · BASTIDOR ESTRATEGICO (NOVO)
    # ====================================================================
    heading_section(doc, "06", "Bastidor estratégico", space_before=14)

    body_p(doc,
           "Leituras técnicas que a Hábilis mantém em monitoramento permanente. Não são "
           "obrigações contratuais nem entregáveis. São inteligência de consultoria que "
           "sustenta a condução do processo e antecipa movimentos de órgão. Registradas aqui "
           "para demonstrar profundidade da leitura — não para transformar em escopo.",
           ["monitoramento permanente", "inteligência de consultoria",
            "não para transformar em escopo"], size=10)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.25)
    set_paragraph_shading(p, ROXO_BG)
    set_paragraph_border_left(p, 2.5, "4A148C")
    styled_run(p, "Camada 4 · Bastidor ", size=9.5, color=CINZA_9, bold=True, mono=True)
    styled_run(p, "— itens abaixo não são tarefas do plano. São monitoramentos Hábilis.",
               size=9.5, color=CINZA_7)

    bullet(doc,
           "Contestação pública CONAMA 006/86 · monitoramos a janela de 15-30 dias "
           "pós-publicação (04/05). Caso haja impugnação fundamentada, ativamos a "
           "investigação confirmatória (Camada 2) como resposta técnica, não como escopo "
           "preventivo.",
           ["CONAMA 006/86", "Camada 2"], size=9.5)
    bullet(doc,
           "Passivo potencial em SAO · o laudo 10/2024 não indica vazamento. Monitoramos "
           "via RCA semestral e inspeção 12/05. Se RCA futuro apontar parâmetro fora, a "
           "conduta passa da Camada 1 para Camada 2 (investigação confirmatória).",
           ["Camada 1", "Camada 2"], size=9.5)
    bullet(doc,
           "Reclassificação urbana do imóvel · acompanhamos o PDTU de Guapó. Hoje em "
           "zona rural (confirmado via LAO). Eventual reclassificação para zona de "
           "expansão urbana altera enquadramento futuro — não afeta a renovação atual.",
           ["PDTU de Guapó", "zona rural"], size=9.5)
    bullet(doc,
           "Hipótese de exigência de PGA/PCA integrado (CP 4.14) · análise documental "
           "de 08/05 verifica se a CP demanda peça autônoma ou está satisfeita pelo "
           "conjunto RCA + PAE + PGRS. Decisão técnica interna.",
           ["CP 4.14"], size=9.5)
    bullet(doc,
           "Postura defensiva em fiscalização SEMMA · dossiê de campo e procuração "
           "mantidos atualizados. Em caso de visita surpresa, Hábilis atende in loco "
           "com documentação pronta. Não é entregável — é prontidão operacional.",
           ["prontidão operacional"], size=9.5)
    bullet(doc,
           "Cenário de atraso no Veredas · protocolo 13/05 pode receber exigência "
           "documental complementar. Plano B técnico (memorial simplificado de regime "
           "de captação) pronto sem acionar Camada 2.",
           ["Plano B técnico"], size=9.5)

    # ====================================================================
    # CALLOUT FINAL · PROTEÇÃO DE CAIXA (destaque)
    # ====================================================================
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(0)
    styled_run(p, "PROTEÇÃO DE CAIXA · DISCURSO OFICIAL", size=9, color=VERDE, bold=True, mono=True)

    tb = doc.add_table(rows=1, cols=1)
    c = tb.cell(0, 0)
    set_cell_shading(c, "E8F5E9")
    set_cell_border(c, left=(32, "1B5E20"), top=(8, "1B5E20"),
                    right=(8, "1B5E20"), bottom=(8, "1B5E20"))
    p1 = c.paragraphs[0]
    p1.paragraph_format.space_before = Pt(6); p1.paragraph_format.space_after = Pt(3)
    styled_run(p1, "A condução foi estruturada para evitar custo prematuro. ",
               size=11, color=CINZA_9, bold=True)
    p2 = c.add_paragraph()
    p2.paragraph_format.space_after = Pt(3)
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    styled_run(p2,
               "Camada 1 (Escopo-Base) é fechada, previsível e já contratada — o núcleo "
               "indispensável ao protocolo de 26/06/2026. Camadas 2 (Contingente) e 3 "
               "(Opcional Futura) são forecast controlado: nenhum item é acionado sem "
               "exigência formal do órgão, bloqueio regulatório concreto ou deliberação "
               "explícita do Grupo Z+Z. Camada 4 (Bastidor) é inteligência Hábilis, "
               "sem custo adicional.",
               size=10, color=CINZA_7)
    p3 = c.add_paragraph()
    p3.paragraph_format.space_before = Pt(4); p3.paragraph_format.space_after = Pt(3)
    p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    styled_run(p3,
               "Esta disciplina de acionamento é o que diferencia a condução Hábilis: o plano "
               "técnico é completo (nenhuma leitura foi reduzida), mas o dispêndio é cirúrgico. "
               "O cliente paga pelo protocolo, não pelo preventivo.",
               size=10, color=CINZA_9, bold=True)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ====================================================================
    # ANEXO A1 · COMPLIANCE PARALELO (CAMADA 3)
    # ====================================================================
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(0)
    styled_run(p, "ANEXO | CAMADA 3 · OPCIONAL FUTURA", size=9, color=AZUL, bold=True, mono=True)

    heading_section(doc, "A1", "Compliance Paralelo | 23 itens", space_before=4)

    lede(doc,
         "Frente complementar à renovação da LAO, com cadência própria e aditivo contratual "
         "separado. Itens aqui não integram o dossiê SEMMA — logo, não travam o protocolo de "
         "26/06. Mas falhas contaminam o processo principal (ex.: RAPP IBAMA em atraso impede "
         "certidões que a SEMMA pede; NR-20 irregular bloqueia operação; cadastro ANP "
         "desatualizado fere a CP 3.10; alvará sanitário vencido fere a CP 3.8). Posicionada "
         "como Camada 3: importante, não emergencial, decisão do cliente.",
         ["cadência própria", "aditivo contratual separado", "não travam o protocolo",
          "contaminam o processo principal", "CP 3.10", "CP 3.8", "Camada 3"])

    heading_h3(doc, "Sumário por órgão | 23 itens", space_before=4)
    t = table_hdr(doc, ["Órgão | Norma", "Itens", "Natureza"], [5.5, 1.5, 10.2])
    add_row(t, ["ANP (inclui cadastro de revendedor · CP 3.10 LAO)", "7",
                "Cadastro, revenda, livro de movimentação, relatórios comerciais, comprovação de distribuidora ativa."],
            bold_cols=[0, 1], mono_cols=[1])
    add_row(t, ["IBAMA | CTF/APP + RAPP + TCFA", "4",
                "Cadastro Técnico Federal ativo e Relatório Anual de Atividades Potencialmente Poluidoras entregue."],
            zebra=True, bold_cols=[0, 1], mono_cols=[1])
    add_row(t, ["NR-20 | Segurança do Trabalho", "7",
                "PGR, PCMSO, treinamento NR-20 válido por função, inspeção SESMT e plano de emergência."],
            bold_cols=[0, 1], mono_cols=[1])
    add_row(t, ["PNRS | Resíduos Sólidos (Lei 12.305/10)", "5",
                "Plano de Gerenciamento de Resíduos, MTR, contrato com transportador licenciado e destino final."],
            zebra=True, bold_cols=[0, 1], mono_cols=[1])
    add_row(t, ["Total", "23",
                "Rodando em paralelo a F1-F4. Diagnóstico individual em documento próprio."],
            bold_cols=[0, 1, 2], mono_cols=[1], is_total=True)

    heading_h3(doc, "O que depende do Grupo Z+Z | compliance", space_before=8)
    bullet(doc, "Acesso completo aos cadastros ANP (login/senha ou procuração com poderes específicos) — atualização do cadastro atende à CP 3.10.",
           ["cadastros ANP", "CP 3.10"])
    bullet(doc, "Validação junto à VISA Guapó sobre exigência do alvará sanitário em zona rural — atende à CP 3.8.",
           ["alvará sanitário", "CP 3.8"])
    bullet(doc, "Acesso ao CTF/APP IBAMA (senha ou nomeação de preposto com poderes de consulta e entrega).",
           ["CTF/APP IBAMA"])
    bullet(doc, "Documentos do SESMT ou contrato com serviço terceirizado de segurança do trabalho: PGR, PCMSO, ASO, ficha de EPI.",
           ["PGR", "PCMSO"])
    bullet(doc, "Contratos e MTRs vigentes com transportadores e destinadores de resíduos licenciados (PGRS).",
           ["MTRs"])
    bullet(doc, "Liberação orçamentária para treinamentos NR-20 faltantes e eventual regularização de PGR/PCMSO.",
           ["NR-20"])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_shading(p, AZUL_BG)
    set_paragraph_border_left(p, 2, "0D47A1")
    p.paragraph_format.left_indent = Cm(0.25)
    styled_run(p, "Decisão requerida neste anexo: ", size=9, color=CINZA_9, bold=True)
    styled_run(p,
               "se o Grupo Z+Z deseja contratar o Compliance Paralelo como frente continuada "
               "(aditivo separado da renovação) ou tratá-lo internamente. Recomendação Hábilis: "
               "validar prioritariamente as duas frentes vinculadas a CPs da LAO (ANP CP 3.10 + "
               "alvará CP 3.8) antes do protocolo de 26/06.",
               size=9, color=CINZA_7)

    here = os.path.dirname(os.path.abspath(__file__))
    out = os.path.join(here, "Alinhamento_Executivo_AutoPostoAmerica_v3.docx")
    doc.save(out)
    print(f"OK -> {out}")
    return out


if __name__ == "__main__":
    build()
