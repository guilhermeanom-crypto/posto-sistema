#!/usr/bin/env python3
"""
Mensagens de Cotacao · Auto Posto America.
Catalogo enxuto de 14 pedidos de orcamento prontos para copiar e enviar a prestadores.
Cada bloco: titulo + chip de classificacao + mensagem pronta (copia-cola) + perfil
exigido + tabela compacta para preenchimento.
Uso: python3 build_mensagens_cotacao_v1.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

VERDE = RGBColor(0x1B, 0x5E, 0x20)
VERDE_L = RGBColor(0x4C, 0xAF, 0x50)
VERDE_BG = "E8F5E9"
AMBER_BG = "FFF8E1"
AZUL_BG = "E3F2FD"
CINZA_9 = RGBColor(0x1C, 0x1C, 0x1E)
CINZA_7 = RGBColor(0x3A, 0x3A, 0x3C)
CINZA_5 = RGBColor(0x6C, 0x6C, 0x70)
CINZA_BG = "FAFAFA"
CINZA_MSG = "F4F4F4"
BRANCO = RGBColor(0xFF, 0xFF, 0xFF)
FONT_SANS = "Segoe UI"
FONT_MONO = "Consolas"

CLASSIF_COLORS = {
    "obrig": "1B5E20",
    "prov":  "F57F17",
    "cont":  "F57F17",
    "opc":   "0D47A1",
}
CLASSIF_LABEL = {
    "obrig": "OBRIG. IMEDIATO",
    "prov":  "PROVÁVEL",
    "cont":  "CONTINGENTE",
    "opc":   "OPCIONAL",
}


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for name, val in edges.items():
        if val is None: continue
        size, color = val
        el = OxmlElement(f'w:{name}')
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(size)); el.set(qn('w:color'), color)
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def set_paragraph_shading(p, hex_color):
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    p_pr.append(shd)


def set_paragraph_border_left(p, size_pt, color_hex):
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), str(int(size_pt * 8)))
    left.set(qn('w:space'), '6'); left.set(qn('w:color'), color_hex)
    pbdr.append(left); p_pr.append(pbdr)


def styled_run(p, text, *, bold=False, italic=False, size=None, color=None, mono=False):
    run = p.add_run(text)
    font_name = FONT_MONO if mono else FONT_SANS
    run.font.name = font_name
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.find(qn('w:rFonts'))
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts'); r_pr.append(r_fonts)
    r_fonts.set(qn('w:ascii'), font_name); r_fonts.set(qn('w:hAnsi'), font_name); r_fonts.set(qn('w:cs'), font_name)
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    return run


def field_page(p):
    r = p.add_run()
    r.font.name = FONT_MONO; r.font.size = Pt(8); r.font.color.rgb = CINZA_5
    fld_b = OxmlElement('w:fldChar'); fld_b.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.text = 'PAGE'
    fld_e = OxmlElement('w:fldChar'); fld_e.set(qn('w:fldCharType'), 'end')
    r._r.append(fld_b); r._r.append(it); r._r.append(fld_e)


def field_numpages(p):
    r = p.add_run()
    r.font.name = FONT_MONO; r.font.size = Pt(8); r.font.color.rgb = CINZA_5
    fld_b = OxmlElement('w:fldChar'); fld_b.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.text = 'NUMPAGES'
    fld_e = OxmlElement('w:fldChar'); fld_e.set(qn('w:fldCharType'), 'end')
    r._r.append(fld_b); r._r.append(it); r._r.append(fld_e)


def bloco_cotacao(doc, num, titulo, classif, mensagem, perfil):
    """
    Um bloco compacto por cotacao.
    - num (int): numero sequencial
    - titulo (str): titulo curto
    - classif (str): obrig/prov/cont/opc
    - mensagem (str): texto pronto para copiar e enviar
    - perfil (str): perfil do prestador em uma linha
    """
    # Cabecalho: numero + titulo + chip classificacao
    hdr = doc.add_table(rows=1, cols=2)
    hdr.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Titulo
    c_t = hdr.rows[0].cells[0]
    c_t.width = Cm(13.5)
    pt = c_t.paragraphs[0]
    pt.paragraph_format.space_before = Pt(2); pt.paragraph_format.space_after = Pt(2)
    styled_run(pt, f"{num:02d}  ", size=16, color=VERDE_L, bold=True, mono=True)
    styled_run(pt, titulo, size=13, color=CINZA_9, bold=True)
    c_t.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Chip
    c_c = hdr.rows[0].cells[1]
    c_c.width = Cm(4.2)
    set_cell_shading(c_c, CLASSIF_COLORS[classif])
    pc = c_c.paragraphs[0]
    pc.paragraph_format.space_before = Pt(3); pc.paragraph_format.space_after = Pt(3)
    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(pc, CLASSIF_LABEL[classif], size=8, color=BRANCO, bold=True, mono=True)
    c_c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Mensagem pronta (caixa cinza clara · visualmente "copia-cola")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(1)
    styled_run(p, "MENSAGEM PRONTA · COPIAR E ENVIAR",
               size=7.5, color=CINZA_5, bold=True, mono=True)

    msg_tbl = doc.add_table(rows=1, cols=1)
    c = msg_tbl.cell(0, 0)
    set_cell_shading(c, CINZA_MSG)
    set_cell_border(c, left=(24, "1B5E20"),
                    top=(4, "E0E0E0"), right=(4, "E0E0E0"), bottom=(4, "E0E0E0"))
    paragrafos = mensagem.split("\n\n")
    first = True
    for para in paragrafos:
        if first:
            p_m = c.paragraphs[0]
            first = False
        else:
            p_m = c.add_paragraph()
        p_m.paragraph_format.space_before = Pt(1)
        p_m.paragraph_format.space_after = Pt(3)
        p_m.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        styled_run(p_m, para, size=10.5, color=CINZA_9)

    # Perfil exigido
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
    styled_run(p, "PERFIL EXIGIDO  ", size=7.5, color=CINZA_5, bold=True, mono=True)
    styled_run(p, perfil, size=9.5, color=CINZA_7, italic=True)

    # Tabela compacta preenchimento
    fill = doc.add_table(rows=2, cols=4)
    fill.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ["PRESTADOR", "VALOR (R$)", "PRAZO", "OBSERVAÇÕES"]
    widths = [5.2, 3.0, 2.8, 6.7]
    for i, h in enumerate(headers):
        c = fill.rows[0].cells[i]
        c.width = Cm(widths[i])
        set_cell_shading(c, "1B5E20")
        pp = c.paragraphs[0]
        pp.paragraph_format.space_before = Pt(2); pp.paragraph_format.space_after = Pt(2)
        styled_run(pp, h, size=7.5, color=BRANCO, bold=True, mono=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i in range(4):
        c = fill.rows[1].cells[i]
        c.width = Cm(widths[i])
        pp = c.paragraphs[0]
        pp.paragraph_format.space_before = Pt(6); pp.paragraph_format.space_after = Pt(6)
        styled_run(pp, "", size=10)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # linha em branco como separador
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0); sp.paragraph_format.space_after = Pt(8)


def build():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(1.3); s.bottom_margin = Cm(1.3)
    s.left_margin = Cm(1.6); s.right_margin = Cm(1.6)

    # Header
    h_p = s.header.paragraphs[0]
    styled_run(h_p, "Auto Posto América · Mensagens de Cotação",
               size=8, color=CINZA_5, mono=True)
    h_p.add_run("\t")
    styled_run(h_p, "Hábilis · Abr/2026", size=8, color=VERDE, bold=True, mono=True)

    # Footer
    f_p = s.footer.paragraphs[0]
    styled_run(f_p, "14 cotações · copia-cola direta para prestador",
               size=8, color=CINZA_5, mono=True)
    f_p.add_run("\t")
    styled_run(f_p, "Pagina ", size=8, color=CINZA_5, mono=True)
    field_page(f_p)
    styled_run(f_p, " de ", size=8, color=CINZA_5, mono=True)
    field_numpages(f_p)

    # ====================================================================
    # ABERTURA ENXUTA (nao e capa)
    # ====================================================================
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    styled_run(p, "MENSAGENS DE COTAÇÃO", size=10, color=VERDE, bold=True, mono=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    styled_run(p, "Auto Posto América · Renovação LAO 033/2020",
               size=18, color=CINZA_9, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    styled_run(p,
               "14 cotações a serem levantadas. Cada bloco tem mensagem pronta para copiar "
               "e enviar ao prestador, perfil exigido em uma linha e tabela compacta para "
               "preenchimento do retorno. ",
               size=10.5, color=CINZA_7)
    styled_run(p, "Chip verde = obrigatório imediato · amber = provável/contingente · "
               "azul = opcional.", size=10.5, color=CINZA_9, bold=True)

    # ====================================================================
    # 14 BLOCOS DE COTACAO
    # ====================================================================

    # 01 · MCO
    bloco_cotacao(doc, 1,
                  "Elaboração do MCO · Memorial de Caracterização",
                  "obrig",
                  "Oi, preciso de orçamento para elaboração do Memorial de "
                  "Caracterização do Empreendimento (MCO) de um posto de "
                  "combustível em Guapó-GO, conforme modelo Anexo IX da CEMAM "
                  "029/2018 · Termo de Referência SEMMA.\n\n"
                  "Escopo: MCO completo (identificação, caracterização, processos, "
                  "armazenamento, efluentes, SAO, poços, plantas e croquis) + "
                  "emissão de ART no CREA. Entrega em .docx + .pdf + ART assinada. "
                  "Fornecemos LAO 033/2020, plantas em .dwg (05/2024) e documentação "
                  "de apoio. Inclui uma revisão pós-inspeção de campo.\n\n"
                  "Prazo de entrega: até 28/05/2026. Contratação prevista até 08/05. "
                  "Pode cotar em pacote com os Anexos I e II da CONAMA 273/2000 se "
                  "for o caso — melhor se houver desconto. Me retorna com valor, "
                  "prazo de elaboração a partir da contratação e validade da proposta.",
                  "Eng. Civil ou Ambiental com CREA-GO ativo · experiência em MCO "
                  "para postos de combustível · aceita coordenação CRBio externa.")

    # 02 · Anexos CONAMA 273
    bloco_cotacao(doc, 2,
                  "Anexos I e II da Resolução CONAMA 273/2000 · ART",
                  "obrig",
                  "Oi, preciso de orçamento para elaboração dos Anexos I e II da "
                  "Resolução CONAMA 273/2000 para posto de combustível em Guapó-GO. "
                  "Item 23 do Termo de Referência SEMMA.\n\n"
                  "Escopo: Anexo I (formulário tabular de caracterização) + Anexo II "
                  "(classificação de risco em 3 classes, conforme CONAMA 273) + "
                  "emissão da ART específica. Entrega em .docx + .pdf + ART.\n\n"
                  "Peças se apoiam no MCO (que estou cotando também). Preferência "
                  "por pacote único MCO + Anexos com o mesmo profissional. "
                  "Prazo: até 29/05/2026. Me manda valor em separado e em pacote "
                  "com o MCO, o prazo e a validade da proposta.",
                  "Mesmo profissional do MCO (Eng. Civil/Ambiental CREA-GO ativo).")

    # 03 · Laboratório RCA
    bloco_cotacao(doc, 3,
                  "Análise laboratorial · RCA semestral (efluentes + poços)",
                  "prov",
                  "Oi, preciso de orçamento para coleta e análise laboratorial de um "
                  "posto de combustível em Guapó-GO. Laudo com ART do químico "
                  "responsável. Acreditação INMETRO vigente é requisito.\n\n"
                  "Pontos de coleta: (1) efluente do separador água-óleo SAO, "
                  "(2) 3 poços de monitoramento existentes.\n\n"
                  "Parâmetros no SAO: pH, turbidez, óleos e graxas (O&G), sólidos "
                  "suspensos, DBO5, DQO, OD, surfactantes (MBAS), metais pesados "
                  "(Pb, Zn, Cr, Cd, Ni, Cu, Fe), TPH total e frações.\n\n"
                  "Parâmetros nos poços: BTEX, PAH (16 EPA), TPH, metais.\n\n"
                  "Coleta prevista para 22/05/2026 (feita por técnico do próprio "
                  "laboratório, com cadeia de custódia). Laudo entregue até 03/06. "
                  "Me retorna com valor pacote completo, valor unitário por ponto, "
                  "prazo de emissão do laudo após a coleta e validade da proposta.",
                  "Laboratório com acreditação ABNT NBR ISO/IEC 17025 INMETRO "
                  "vigente no escopo dos parâmetros listados.")

    # 04 · Análise complementar
    bloco_cotacao(doc, 4,
                  "Análise complementar de efluentes e água · apoio à inspeção",
                  "prov",
                  "Oi, preciso de orçamento para coleta e análise complementar de "
                  "efluentes e água subterrânea num posto de combustível em "
                  "Guapó-GO, em até 2 pontos adicionais aos da cotação do RCA.\n\n"
                  "Mesmos parâmetros da análise principal (pH, O&G, DBO, DQO, "
                  "metais, TPH, BTEX, PAH). Objetivo: validar conformidade da SAO "
                  "após inspeção técnica de campo.\n\n"
                  "Coleta em 22/05/2026 (conjugada com a do RCA para otimizar). "
                  "Me retorna custo adicional por ponto e prazo de laudo. Pode "
                  "vir em pacote único com o item anterior.",
                  "Mesmo perfil do RCA · preferencialmente o mesmo laboratório "
                  "(otimiza logística e custo).")

    # 05 · Publicação CONAMA 006
    bloco_cotacao(doc, 5,
                  "Publicação do pedido de licenciamento · CONAMA 006/1986",
                  "obrig",
                  "Oi, preciso de orçamento para publicação de aviso de "
                  "requerimento de licença ambiental para posto de combustível em "
                  "Guapó-GO. Exigência da Resolução CONAMA 006/1986 · Item 6 do "
                  "TR SEMMA.\n\n"
                  "Publicação simultânea em: (1) Diário Oficial do Estado de Goiás "
                  "(DOE/GO) e (2) jornal de circulação regional em Guapó. Texto do "
                  "aviso fornecido no modelo oficial.\n\n"
                  "Entrega: recortes digitalizados + exemplares físicos das duas "
                  "publicações. Data desejada: 04/05/2026 (publicação simultânea). "
                  "Me passa valor DOE/GO + valor do jornal (comparar 2-3 opções de "
                  "jornal) + prazo para publicar após envio do texto.",
                  "Diário Oficial de GO (direto) + jornal com circulação comprovada "
                  "em Guapó-GO · agência de publicações legais pode intermediar.")

    # 06 · Investigação confirmatória
    bloco_cotacao(doc, 6,
                  "Investigação confirmatória · NBR 15515-2/3 + NBR 16209",
                  "cont",
                  "Oi, preciso de orçamento para avaliação preliminar e "
                  "investigação confirmatória em posto de combustível em Guapó-GO. "
                  "Normas: ABNT NBR 15515-2, NBR 15515-3 e NBR 16209. Referências "
                  "CEMAM 029/2018 + CONAMA 420/2009.\n\n"
                  "Escopo: avaliação preliminar (documental e visual) + "
                  "investigação confirmatória (sondagens SPT + coletas de solo e "
                  "água subterrânea — ver parâmetros no item seguinte) + avaliação "
                  "de risco preliminar. Laudo completo com ART de geólogo ou "
                  "engenheiro ambiental. Interpretação cruzada com a LAO.\n\n"
                  "Atenção: execução é contingente. Aciona se SEMMA exigir ou se "
                  "contestação pública na janela CONAMA 006/86. Mesmo assim preciso "
                  "da cotação fechada já. Janela de execução estimada 25/05 a "
                  "15/06/2026. Me manda valor em pacote com as sondagens SPT "
                  "(cotação 07) e valor em separado.",
                  "Empresa especializada em passivo ambiental · geólogo(a) ou "
                  "eng. ambiental com CREA ativo · parceria com laboratório "
                  "INMETRO · portfólio em postos de combustível.")

    # 07 · Sondagens SPT
    bloco_cotacao(doc, 7,
                  "Sondagens SPT · mín. 3 furos + análise cromatográfica",
                  "cont",
                  "Oi, preciso de orçamento para sondagens SPT em posto de "
                  "combustível em Guapó-GO. Componente da investigação "
                  "confirmatória (ABNT NBR 15515-3 + NBR 6484).\n\n"
                  "Escopo: mínimo 3 furos SPT com amostragem contínua até 6m ou até "
                  "atingir o nível freático. 3 amostras de solo por furo + 1 "
                  "amostra de água subterrânea por furo (se atingido). Análise "
                  "química cromatográfica. Laudo com mapa de locação.\n\n"
                  "Parâmetros por amostra: VOC, BTEX, PAH (16 compostos EPA), "
                  "metais pesados (Pb, Cr, Cd, Zn, Cu, Ni, Fe, Mn), TPH total + "
                  "frações C9–C36. Solo: + granulometria e matéria orgânica. "
                  "Água: + pH, condutividade, sólidos.\n\n"
                  "Execução estimada 25/05 se acionada. Laudo até 10/06. Me passa "
                  "valor pacote 3 furos + custo unitário por furo adicional + "
                  "prazo. Locação final dos furos validada por mim antes de executar.",
                  "Empresa de sondagem com equipamento calibrado + laboratório "
                  "INMETRO acoplado · normalmente mesma empresa do item 06.")

    # 08 · Análise de risco
    bloco_cotacao(doc, 8,
                  "Análise de Risco · APP (Preliminar) ou ARQ (Quantitativa)",
                  "cont",
                  "Oi, preciso de orçamento para Análise de Risco em posto de "
                  "combustível em Guapó-GO. Referências: CEMAM 029/2018 · CONAMA "
                  "420/2009 · metodologias ASTM RBCA.\n\n"
                  "Preciso das duas faixas separadas:\n"
                  "(a) APP · Análise de Risco Preliminar — qualitativa, cenários "
                  "de exposição sem modelagem numérica. Entrega estimada 15–30 dias.\n"
                  "(b) ARQ · Análise de Risco Quantitativa — modelagem completa "
                  "com software (RBCA Tool Kit, RISC ou similar), concentrações-"
                  "meta calculadas. Entrega estimada 60–90 dias.\n\n"
                  "Aciona só se SEMMA exigir, tipicamente após resultado da "
                  "investigação confirmatória. Preciso do valor e prazo de cada "
                  "formato já. Se algumas empresas incluem a APP no pacote da "
                  "investigação confirmatória sem custo, confirmar.",
                  "Empresa com eng. ambiental ou químico com formação específica "
                  "em análise de risco · portfólio em áreas avaliadas por "
                  "SEMAD/IBAMA.")

    # 09 · Estanqueidade extra
    bloco_cotacao(doc, 9,
                  "Teste de estanqueidade extra · NBR 16795/2019",
                  "cont",
                  "Oi, preciso de orçamento para teste de estanqueidade em "
                  "tanques subterrâneos e tubulações de produto de posto de "
                  "combustível em Guapó-GO. Norma ABNT NBR 16795/2019.\n\n"
                  "Escopo: teste completo nos tanques e linhas · emissão de laudo "
                  "aprovado/reprovado com ART · se reprovar, apontar ponto de "
                  "vazamento.\n\n"
                  "Atenção: o posto tem laudo vigente de 10/2024 (bienal, próximo "
                  "legal em 10/2026). Este orçamento é backup — só aciono se a "
                  "SEMMA contestar o laudo atual. Cotação fica arquivada. Me manda "
                  "valor + prazo de execução + validade da acreditação INMETRO "
                  "da empresa.",
                  "Empresa com acreditação INMETRO vigente para teste de "
                  "estanqueidade em SASC. Confirmar validade da acreditação na "
                  "data da contratação.")

    # 10 · Adequação SAO
    bloco_cotacao(doc, 10,
                  "Adequação física do sistema SAO · obra civil",
                  "cont",
                  "Oi, preciso de orçamento para eventual adequação do sistema "
                  "separador água-óleo (SAO) de posto de combustível em Guapó-GO. "
                  "Serviço contingente, acionado se diagnóstico de campo em 12/05 "
                  "ou análise de efluentes apontar não-conformidade.\n\n"
                  "Preciso de duas faixas separadas:\n"
                  "(a) Manutenção corretiva simples — limpeza reforçada, troca de "
                  "selos, pequena manutenção.\n"
                  "(b) Obra de adequação estrutural — impermeabilização, troca de "
                  "componentes, ampliação, novo SAO ou adequação de canaletas e "
                  "grelhas.\n\n"
                  "Inclui custo de descarte do material removido como resíduo "
                  "classe I. Janela estimada de execução: 20/05 a 12/06 se "
                  "acionado. Me passa faixa de preço para cada cenário e prazo de "
                  "mobilização.",
                  "Empreiteira com experiência em postos de combustível · "
                  "aceitável prestador regular de manutenção de SAO na região "
                  "de Goiânia.")

    # 11 · Saneamento/fossa
    bloco_cotacao(doc, 11,
                  "Laudo de saneamento ou fossa/sumidouro",
                  "cont",
                  "Oi, preciso de orçamento para laudo técnico de sistema "
                  "sanitário de posto de combustível em Guapó-GO. Depende do "
                  "regime de destinação do efluente sanitário:\n\n"
                  "(a) Se rede pública: só requerimento de anuência à "
                  "concessionária local (sem serviço técnico, apenas taxa).\n"
                  "(b) Se fossa/sumidouro: laudo técnico completo "
                  "(dimensionamento, estanqueidade, distância mínima de poços), "
                  "com ART.\n\n"
                  "Preciso verificar primeiro com o posto qual é o regime — te "
                  "confirmo antes de acionar. Prazo-alvo: 25/05/2026. Me passa "
                  "valor do laudo com ART (cenário b) e, se possível, intermediar "
                  "a anuência (cenário a).",
                  "Eng. civil ou sanitarista com CREA ativo.")

    # 12 · Recontrato Bauer/Limpmil/Ecofenix
    bloco_cotacao(doc, 12,
                  "Recontrato · limpeza de tanques, coleta de resíduos, óleo usado",
                  "prov",
                  "Oi, preciso validar — e, se necessário, recontratar — "
                  "serviços terceirizados de posto de combustível em Guapó-GO. "
                  "Fornecedores atuais: Bauer (limpeza tanques/SAO), Limpmil "
                  "(coleta e transporte de resíduos), Ecofenix (destinação de "
                  "óleo lubrificante usado).\n\n"
                  "Preciso verificar licenças ambientais, cadastro IBAMA "
                  "CTF/APP, autorização ANTT e certificados de destinação dos "
                  "últimos 3 anos de cada um. Se algum estiver irregular, cotar "
                  "substituto regional com mesmas credenciais.\n\n"
                  "Validação até 20/05/2026. Recontrato, se necessário, até "
                  "30/05. Me retorna: status da documentação de cada um + "
                  "cotação de substitutos caso haja irregularidade.",
                  "Prestadores atuais ou substitutos de mesmo porte · licenças "
                  "ambientais e fiscais vigentes · frota própria no caso de "
                  "transporte.")

    # 13 · Topografia
    bloco_cotacao(doc, 13,
                  "Levantamento topográfico complementar · UTM SIRGAS 2000",
                  "opc",
                  "Oi, preciso de orçamento de reserva para levantamento "
                  "topográfico planialtimétrico de posto de combustível em "
                  "Guapó-GO. Opcional · aciona só se o Veredas exigir memorial "
                  "topográfico adicional ou se as plantas atuais (05/2024) não "
                  "tiverem precisão suficiente.\n\n"
                  "Escopo: perímetro do posto + APP + corpo hídrico próximo "
                  "(se houver). Coordenadas UTM SIRGAS 2000. Entrega em .dwg + "
                  "memorial descritivo + ART.\n\n"
                  "Prazo: até 10 dias corridos após acionamento. Me passa valor + "
                  "prazo. Preferência por quem já fez as plantas 05/2024 — me "
                  "confirma se conhece essa empresa.",
                  "Eng. agrimensor(a) ou técnico em topografia com CREA/CFT "
                  "ativo · equipamento GPS/RTK.")

    # 14 · Paisagismo/APP
    bloco_cotacao(doc, 14,
                  "Projeto de adequação paisagística ou PRAD de APP",
                  "opc",
                  "Oi, preciso de orçamento de reserva para projeto de "
                  "adequação paisagística ou Projeto de Recuperação de Área "
                  "Degradada (PRAD) simples para área de APP de posto de "
                  "combustível em Guapó-GO. Opcional · aciona só se inspeção "
                  "12/05 identificar APP comprometida.\n\n"
                  "Escopo do projeto: escolha de espécies nativas do Cerrado + "
                  "plantio + cronograma de manutenção + ART. Só projeto por "
                  "enquanto, não a execução.\n\n"
                  "Prazo de elaboração: 15 dias após acionamento. Me retorna "
                  "valor do projeto + faixa de preço da execução (por m² ou "
                  "por muda) como referência, caso venha a ser contratada.",
                  "Eng. agrônomo(a) ou florestal com CREA ativo · experiência em "
                  "PRAD no Cerrado.")

    here = os.path.dirname(os.path.abspath(__file__))
    out = os.path.join(here, "Mensagens_Cotacao_AutoPostoAmerica_v1.docx")
    doc.save(out)
    print(f"OK -> {out}")
    return out


if __name__ == "__main__":
    build()
