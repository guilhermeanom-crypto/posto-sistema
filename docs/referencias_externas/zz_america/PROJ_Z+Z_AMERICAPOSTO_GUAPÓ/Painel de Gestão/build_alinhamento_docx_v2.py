#!/usr/bin/env python3
"""
Gerador DOCX · Alinhamento Executivo · Auto Posto América.
Peça enxuta (3 paginas + 1 de Compliance Paralelo) para validacao do proprietario.
Uso: python3 build_alinhamento_docx.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Paleta Habilis
VERDE = RGBColor(0x1B, 0x5E, 0x20)
VERDE_L = RGBColor(0x4C, 0xAF, 0x50)
VERDE_BG = "E8F5E9"
AMBER = RGBColor(0xF5, 0x7F, 0x17)
AMBER_BG = "FFF8E1"
VERMELHO = RGBColor(0xB7, 0x1C, 0x1C)
VERMELHO_BG = "FFEBEE"
AZUL = RGBColor(0x0D, 0x47, 0xA1)
AZUL_BG = "E3F2FD"
CINZA_9 = RGBColor(0x1C, 0x1C, 0x1E)
CINZA_7 = RGBColor(0x3A, 0x3A, 0x3C)
CINZA_5 = RGBColor(0x6C, 0x6C, 0x70)
CINZA_BG = "FAFAFA"
BRANCO = RGBColor(0xFF, 0xFF, 0xFF)
FONT_SANS = "Segoe UI"
FONT_MONO = "Consolas"


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for name, val in edges.items():
        if val is None: continue
        size, color = val
        el = OxmlElement(f'w:{name}')
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(size)); el.set(qn('w:color'), color)
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def set_paragraph_shading(p, hex_color):
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    p_pr.append(shd)


def set_paragraph_border_left(p, size_pt, color_hex):
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), str(int(size_pt * 8)))
    left.set(qn('w:space'), '6'); left.set(qn('w:color'), color_hex)
    pbdr.append(left); p_pr.append(pbdr)


def styled_run(p, text, *, bold=False, italic=False, size=None, color=None, mono=False):
    run = p.add_run(text)
    font_name = FONT_MONO if mono else FONT_SANS
    run.font.name = font_name
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.find(qn('w:rFonts'))
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts'); r_pr.append(r_fonts)
    r_fonts.set(qn('w:ascii'), font_name); r_fonts.set(qn('w:hAnsi'), font_name); r_fonts.set(qn('w:cs'), font_name)
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    return run


def mixed_run(p, text, *, size=10, default_color=CINZA_7, bold_parts=None):
    if not bold_parts:
        styled_run(p, text, size=size, color=default_color); return
    remaining = text
    for target in bold_parts:
        idx = remaining.find(target)
        if idx == -1: continue
        if idx > 0: styled_run(p, remaining[:idx], size=size, color=default_color)
        styled_run(p, target, size=size, color=CINZA_9, bold=True)
        remaining = remaining[idx + len(target):]
    if remaining:
        styled_run(p, remaining, size=size, color=default_color)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def heading_section(doc, num, title, space_before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '14')
    bottom.set(qn('w:space'), '3'); bottom.set(qn('w:color'), '1B5E20')
    pbdr.append(bottom); p_pr.append(pbdr)
    styled_run(p, f"{num}  ", size=11, color=VERDE_L, mono=True)
    styled_run(p, title, size=15, color=VERDE, bold=True)


def heading_h3(doc, text, space_before=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    set_paragraph_border_left(p, 2, "1B5E20")
    p.paragraph_format.left_indent = Cm(0.2)
    styled_run(p, text, size=10.5, color=CINZA_9, bold=True)


def body_p(doc, text, bold_parts=None, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    mixed_run(p, text, bold_parts=bold_parts, size=size)


def lede(doc, text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.25)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_shading(p, VERDE_BG)
    set_paragraph_border_left(p, 2.5, "1B5E20")
    mixed_run(p, text, size=10.5, default_color=CINZA_9, bold_parts=bold_parts)


def callout(doc, title, body, kind="verde", bold_body_parts=None):
    bg = {"verde": VERDE_BG, "amber": AMBER_BG, "vermelho": VERMELHO_BG,
          "azul": AZUL_BG, "cinza": CINZA_BG}.get(kind, VERDE_BG)
    border = {"verde": "1B5E20", "amber": "F57F17", "vermelho": "B71C1C",
              "azul": "0D47A1", "cinza": "6C6C70"}.get(kind, "1B5E20")
    t = doc.add_table(rows=1, cols=1)
    c = t.cell(0, 0)
    set_cell_shading(c, bg)
    set_cell_border(c, left=(24, border))
    p1 = c.paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    styled_run(p1, title.upper(), size=9, color=CINZA_9, bold=True, mono=True)
    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    mixed_run(p2, body, size=9.5, default_color=CINZA_7, bold_parts=bold_body_parts)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(2)


def table_hdr(doc, headers, widths_cm=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_shading(c, "1B5E20")
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        styled_run(p, h.upper(), size=8, color=BRANCO, bold=True, mono=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if widths_cm and i < len(widths_cm): c.width = Cm(widths_cm[i])
    return t


def add_row(t, values, zebra=False, bold_cols=None, mono_cols=None, is_total=False, size=9):
    row = t.add_row()
    for i, val in enumerate(values):
        c = row.cells[i]
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
        bold = (bold_cols and i in bold_cols) or is_total
        mono = mono_cols and i in mono_cols
        styled_run(p, val, size=size, color=CINZA_9 if bold else CINZA_7, bold=bold, mono=mono)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        if is_total: set_cell_shading(c, VERDE_BG)
        elif zebra: set_cell_shading(c, CINZA_BG)


def bullet(doc, text, bold_parts=None, size=10):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    mixed_run(p, text, bold_parts=bold_parts, size=size)


# =========== Cores por responsavel (checklist cronologico) ===========
RESP_COLORS = {
    "cliente":  "0D47A1",   # azul
    "habilis":  "1B5E20",   # verde
    "terceiro": "6C6C70",   # cinza
    "misto":    "4A148C",   # roxo (Habilis ou Terceiro)
}


# =========== Checklist tipado por chip ===========
CHIP_STYLE = {
    "DOC":     ("0D47A1", "Documento"),
    "DECISAO": ("1B5E20", "Decisao"),
    "PAGAR":   ("F57F17", "Pagamento"),
    "ACESSO":  ("6C6C70", "Acesso"),
    "ASSINAR": ("2E7D32", "Assinatura"),
}


def set_tbl_borders(t, color="E0E0E0", size=4):
    tbl_pr = t._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(size))
        el.set(qn('w:color'), color)
        borders.append(el)
    tbl_pr.append(borders)


def phase_checklist(doc, title, period, days, items):
    """Header da fase + tabela (chip tipo | descricao)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    set_paragraph_border_left(p, 2.5, "1B5E20")
    p.paragraph_format.left_indent = Cm(0.2)
    styled_run(p, f"{title}  ", size=11, color=CINZA_9, bold=True)
    styled_run(p, f"{period}  ", size=9.5, color=VERDE, mono=True)
    styled_run(p, f"({days} dias)", size=9, color=CINZA_5, mono=True)

    # Contadores por tipo
    counts = {}
    for kind, _ in items:
        counts[kind] = counts.get(kind, 0) + 1
    mini = " | ".join(f"{v} {CHIP_STYLE[k][1].lower()}" for k, v in counts.items())
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(3)
    p2.paragraph_format.left_indent = Cm(0.25)
    styled_run(p2, mini, size=8.5, color=CINZA_5, mono=True)

    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_tbl_borders(t, "F0F0F0", 4)
    for kind, text in items:
        row = t.add_row()
        c_chip, c_body = row.cells[0], row.cells[1]
        color_hex, label = CHIP_STYLE[kind]
        set_cell_shading(c_chip, color_hex)
        pc = c_chip.paragraphs[0]
        pc.paragraph_format.space_before = Pt(2); pc.paragraph_format.space_after = Pt(2)
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(pc, label.upper(), size=7.5, color=BRANCO, bold=True, mono=True)
        c_chip.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c_chip.width = Cm(2.4)

        pb = c_body.paragraphs[0]
        pb.paragraph_format.space_before = Pt(2); pb.paragraph_format.space_after = Pt(2)
        pb.paragraph_format.left_indent = Cm(0.15)
        styled_run(pb, text, size=9.5, color=CINZA_9)
        c_body.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c_body.width = Cm(14.4)


# =========== Matriz visual (cards horizontais) ===========
def matriz_cards(doc, cards, accent_hex="1B5E20"):
    """
    cards: list of dicts { 'title', 'sub', 'stats': [(label, value, color), ...], 'next': str }
    Gera uma tabela N colunas com mini-cards visuais.
    """
    n = len(cards)
    t = doc.add_table(rows=1, cols=n)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, card in enumerate(cards):
        c = t.rows[0].cells[i]
        set_cell_shading(c, "FAFAFA")
        set_cell_border(c, top=(24, accent_hex), left=(4, "E0E0E0"),
                        right=(4, "E0E0E0"), bottom=(4, "E0E0E0"))
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # Título
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(1)
        styled_run(p, card["title"], size=10.5, color=CINZA_9, bold=True)

        # Subtítulo
        p = c.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(5)
        styled_run(p, card["sub"], size=8.5, color=CINZA_5, mono=True)

        # Stats
        for label, val, color in card.get("stats", []):
            p = c.add_paragraph()
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(1)
            styled_run(p, f"{val} ", size=13, color=color, bold=True, mono=True)
            styled_run(p, label, size=7.5, color=CINZA_5, bold=True, mono=True)

        # Próximo
        if card.get("next"):
            p = c.add_paragraph()
            p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(3)
            set_paragraph_border_left(p, 1.5, accent_hex)
            p.paragraph_format.left_indent = Cm(0.15)
            styled_run(p, "PROXIMA ACAO", size=7, color=CINZA_5, bold=True, mono=True)
            p = c.add_paragraph()
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(4)
            set_paragraph_border_left(p, 1.5, accent_hex)
            p.paragraph_format.left_indent = Cm(0.15)
            styled_run(p, card["next"], size=9, color=CINZA_9, bold=True)


def field_page(p):
    r = p.add_run()
    r.font.name = FONT_MONO; r.font.size = Pt(8); r.font.color.rgb = CINZA_5
    fld_b = OxmlElement('w:fldChar'); fld_b.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.text = 'PAGE'
    fld_e = OxmlElement('w:fldChar'); fld_e.set(qn('w:fldCharType'), 'end')
    r._r.append(fld_b); r._r.append(it); r._r.append(fld_e)


def field_numpages(p):
    r = p.add_run()
    r.font.name = FONT_MONO; r.font.size = Pt(8); r.font.color.rgb = CINZA_5
    fld_b = OxmlElement('w:fldChar'); fld_b.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.text = 'NUMPAGES'
    fld_e = OxmlElement('w:fldChar'); fld_e.set(qn('w:fldCharType'), 'end')
    r._r.append(fld_b); r._r.append(it); r._r.append(fld_e)


def build():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(1.5); s.bottom_margin = Cm(1.5)
    s.left_margin = Cm(1.8); s.right_margin = Cm(1.8)

    # Header
    h_p = s.header.paragraphs[0]
    styled_run(h_p, "Auto Posto América | Alinhamento Executivo", size=8, color=CINZA_5, mono=True)
    h_p.add_run("\t")
    styled_run(h_p, "Hábilis | Abr/2026", size=8, color=VERDE, bold=True, mono=True)

    # Footer
    f_p = s.footer.paragraphs[0]
    styled_run(f_p, "Hábilis Regularização Ambiental", size=8, color=CINZA_5, mono=True)
    f_p.add_run("\t")
    styled_run(f_p, "Pagina ", size=8, color=CINZA_5, mono=True)
    field_page(f_p)
    styled_run(f_p, " de ", size=8, color=CINZA_5, mono=True)
    field_numpages(f_p)

    # ====================================================================
    # CAPA
    # ====================================================================
    p = doc.add_paragraph()
    styled_run(p, "H  ", size=16, color=VERDE, bold=True)
    styled_run(p, "Hábilis", size=16, color=VERDE, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(40)
    styled_run(p, "REGULARIZAÇÃO AMBIENTAL", size=9, color=CINZA_5, mono=True)
    for _ in range(4): doc.add_paragraph()

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(p, "PROCESSO DE REGULARIZAÇÃO AMBIENTAL", size=10, color=VERDE, bold=True, mono=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(p, "Auto Posto América", size=30, color=CINZA_9, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(p, "Renovação LAO 033/2020 | SEMMA Guapó | Processo 22760/2020",
               size=13, color=CINZA_7)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    styled_run(p, "Alinhamento executivo v2 para validação do Grupo Z+Z. Resume o plano "
               "técnico em execução pela Hábilis em escopo refinado pós-auditoria documental "
               "da LAO 033/2020: 3 frentes contratadas (Uso do Solo, Renovação LAO, "
               "Condicionantes) com peças de fundamento literal na LAO/TR + 1 frente "
               "paralela (Compliance Paralelo, opcional). Cronograma operacional em 4 fases "
               "(67 dias corridos). Meta acordada ", size=11, color=CINZA_7)
    styled_run(p, "26/06/2026 (180 dias antes) | mínimo legal CP 3.11: 24/08/2026 (120 dias).",
               size=11, color=CINZA_9, bold=True)

    for _ in range(2): doc.add_paragraph()

    tc = doc.add_table(rows=2, cols=4)
    for i, h in enumerate(["DESTINATÁRIO", "EMISSOR", "PROTOCOLO-META", "EMISSÃO"]):
        c = tc.cell(0, i)
        styled_run(c.paragraphs[0], h, size=8, color=CINZA_5, bold=True, mono=True)
        set_cell_border(c, top=(12, "1B5E20"))
    vals = ["Grupo Z+Z", "Hábilis", "26/06/2026", "22/04/2026"]
    for i, v in enumerate(vals):
        c = tc.cell(1, i)
        styled_run(c.paragraphs[0], v, size=10, color=CINZA_9, bold=True, mono=i >= 2)

    add_page_break(doc)

    heading_section(doc, "01", "Panorama", space_before=0)

    lede(doc,
         "A LAO 033/2020 do Auto Posto América vence em 23/12/2026. O mínimo legal de "
         "antecedência (CP 3.11) é 120 dias — piso regulatório efetivo em 24/08/2026. A "
         "Hábilis adota meta operacional 26/06/2026 (180 dias antes), mais defensiva que o "
         "mínimo legal. Esta v2 do alinhamento reflete escopo refinado após auditoria contra "
         "a LAO literal e o TR SEMMA: somente peças com fundamento expresso. Boas práticas "
         "consultivas (PGRS, NR-20, cadastro ANP, alvará sanitário) foram realocadas ao "
         "Compliance Paralelo (anexo). Base documental de 14/04/2026 suporta a execução.",
         ["23/12/2026", "120 dias", "24/08/2026", "26/06/2026", "escopo refinado",
          "fundamento expresso", "Compliance Paralelo (anexo)"])

    heading_h3(doc, "Frentes integradas | escopo refinado v2")
    t = table_hdr(doc, ["Frente", "Status", "Órgão", "Escopo"], [4.0, 2.5, 3.2, 7.5])
    add_row(t, ["Uso do Solo", "Contratada", "Prefeitura Guapó", "Pré-requisito para protocolo da renovação. Certidão de Uso e Ocupação vigente."], bold_cols=[0, 1], mono_cols=[1])
    add_row(t, ["Renovação LAO", "Contratada", "SEMMA (Proc. 22760/2020)", "Dossiê técnico: MCO + Anexos I/II CONAMA 273 + RCA + Publicação CONAMA 006/86 + PAE + PGA/PCA + Estanqueidade (laudo 10/2024) + Relatório Consolidado."], zebra=True, bold_cols=[0, 1], mono_cols=[1])
    add_row(t, ["Condicionantes LAO 033/2020", "Contratada", "SEMMA", "28 CPs mapeadas (10 declaratórias + 18 ações). Recomposição documental + execução das pendentes."], bold_cols=[0, 1], mono_cols=[1])
    add_row(t, ["Compliance Paralelo", "Anexo · Opcional", "ANP | IBAMA | NR-20 | PNRS | VISA", "Cadência própria, fora do caminho crítico do protocolo. Inclui cadastro ANP atualizado (CP 3.10) e alvará sanitário (se aplicável). Detalhe no anexo."], zebra=True, bold_cols=[0, 1], mono_cols=[1])
    add_row(t, ["Total", "3 contratadas + 1 paralela", "4 órgãos/frentes", "Plano único, execução coordenada, escopo enxuto."], bold_cols=[0, 1, 2, 3], mono_cols=[1], is_total=True)

    # ====================================================================
    # SECAO 02 · CRONOGRAMA (fluxo continuo)
    # ====================================================================

    heading_section(doc, "02", "Cronograma | 4 fases em 67 dias", space_before=12)

    body_p(doc,
           "Janela de execução de 67 dias, de 20/04/2026 a 26/06/2026. Cada fase tem "
           "entregáveis próprios e depende de ações do Grupo Z+Z (seção 03). Atrasos em "
           "F1 comprimem F3 e F4 proporcionalmente.",
           ["67 dias", "20/04/2026", "26/06/2026"], size=9.5)

    t = table_hdr(doc, ["Fase", "Período", "Dias", "Entregáveis-chave"], [3.3, 3.4, 1.2, 9.3])
    add_row(t,
            ["F1 | Estruturação", "20/04 -> 08/05", "19",
             "Procuração, dados societários, contratação do terceiro Eng. Civil/Ambiental para ART do MCO/Anexos 273, Uso do Solo protocolado, status outorga consolidado, publicação CONAMA 006/86."],
            bold_cols=[0, 1, 2], mono_cols=[1, 2], size=8.5)
    add_row(t,
            ["F2 | Consolidação", "09/05 -> 30/05", "22",
             "Validação técnica de campo (12/05), protocolo Veredas (13/05), nova coleta RCA (22/05), MCO (28/05), Anexos I/II CONAMA 273 (29/05), Plano das 28 CPs (30/05)."],
            zebra=True, bold_cols=[0, 1, 2], mono_cols=[1, 2], size=8.5)
    add_row(t,
            ["F3 | Execução e Evidência", "31/05 -> 19/06", "20",
             "RCA semestral consolidado (03/06), PAE (05/06), recomposição das CPs parciais (12/06), Relatório Consolidado (19/06)."],
            bold_cols=[0, 1, 2], mono_cols=[1, 2], size=8.5)
    add_row(t,
            ["F4 | Protocolo", "20/06 -> 26/06", "7",
             "Aprovação final, CNDs, checklist final, taxa paga, protocolo no Processo 22760/2020 | SEMMA."],
            zebra=True, bold_cols=[0, 1, 2], mono_cols=[1, 2], size=8.5)

    # ====================================================================
    # SECAO 03 · CHECKLIST CRONOLOGICO (fluxo continuo)
    # ====================================================================

    heading_section(doc, "03", "Checklist cronológico | escopo refinado v2", space_before=12)

    body_p(doc,
           "Itens em ordem cronológica, escopo refinado pós-auditoria contra LAO/TR literal. "
           "Responsável identificado por cor: Cliente em azul, Hábilis em verde e Terceiro "
           "em cinza. MCO e Anexos I/II CONAMA 273 executados por terceiro Eng. "
           "Civil/Ambiental sob coordenação CRBio da Hábilis (modelo padrão · Hábilis não "
           "possui Eng. Civil/Ambiental interno). Itens contingentes e peças de boa prática "
           "consultiva (investigação confirmatória, sondagens SPT, análise de risco, "
           "investigação detalhada, remediação) foram removidos do escopo da renovação. "
           "PGRS, NR-20, cadastro ANP e alvará sanitário tratados no anexo.",
           ["Cliente", "Hábilis", "Terceiro", "escopo refinado",
            "removidos do escopo da renovação", "anexo"], size=9.5)

    # Legenda de responsaveis (3 cores: sem mais 'misto')
    legenda = doc.add_table(rows=1, cols=3)
    legenda.alignment = WD_TABLE_ALIGNMENT.LEFT
    legenda_data = [
        ("Cliente", "cliente"), ("Hábilis", "habilis"), ("Terceiro", "terceiro"),
    ]
    for i, (label, kind) in enumerate(legenda_data):
        c = legenda.rows[0].cells[i]
        set_cell_shading(c, RESP_COLORS[kind])
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(p, label.upper(), size=8, color=BRANCO, bold=True, mono=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Tabela cronologica (sem checkbox, versao 'proposto')
    t = table_hdr(doc, ["Prazo", "Item", "Responsável", "Fase", "Observação"],
                  [1.5, 5.2, 2.8, 0.9, 7.0])

    itens_chrono = [
        ("22/04", "Validação do AVCB Cercon 161726", "Hábilis", "habilis", "F1",
         "Certificado CBMGO 161726, emitido 07/11/2025, válido até 20/10/2026. Em mãos. Atende CP 3.10."),
        ("22/04", "Coleta de RG/CPF dos sócios", "Cliente", "cliente", "F1",
         "RG e CPF de Sérgio Marques e José Guilherme. Exigência dos dois processos (Uso do Solo + LAO)."),
        ("23/04", "Procuração pública para a Hábilis", "Cliente", "cliente", "F1",
         "Firma reconhecida. Poderes para SEMMA, Prefeitura, SEMAD/GO e Veredas. Substitui a procuração A3 de 2023."),
        ("24/04", "Cartão CNPJ + Comprovante de Endereço", "Cliente", "cliente", "F1",
         "CNPJ ativo conferindo CNAE/endereço. Comprovante de endereço dos sócios nos últimos 3 meses."),
        ("25/04", "Verificação formal da Certidão de Uso do Solo", "Cliente", "cliente", "F1",
         "Confirmar existência, número, data e disponibilidade. Apoio da Hábilis."),
        ("27/04", "Enquadramento do imóvel + CAR", "Hábilis", "habilis", "F1",
         "Zona rural confirmada via LAO. CAR obrigatório. ART de Uso do Solo via Eng. Agronômica (CREA)."),
        ("27/04", "ART específica para Uso do Solo", "Hábilis", "habilis", "F1",
         "ART distinta da Renovação LAO. Assinada pela Hábilis via RT Eng. Agronômica."),
        ("28/04", "Diagnóstico de outorga no Web Outorga", "Hábilis", "habilis", "F1",
         "Consulta ao Sistema Web Outorga para histórico da DURH002737 (processos legados ainda não migrados ao Veredas). Protocolo novo será no Veredas em 13/05. DURH002737 vencida 09/12/2023."),
        ("30/04", "Protocolo do requerimento da Certidão", "Hábilis", "habilis", "F1",
         "Somente se verificação indicar ausência ou desatualização."),
        ("04/05", "Publicação do pedido de licenciamento (CONAMA 006/86)", "Hábilis", "habilis", "F1",
         "DOE/GO + jornal de grande circulação. Item 6 do TR SEMMA. Aguardar 30 dias para contestação pública."),
        ("05/05", "Contratação do terceiro Eng. Civil/Ambiental", "Cliente", "cliente", "F1",
         "Para ART do MCO (28/05) + Anexos I/II CONAMA 273 (29/05). Modelo padrão Hábilis: subcontratação sob coordenação CRBio. Indicação Hábilis. Faixa R$ 8-15k + ART."),
        ("08/05", "Análise documental consolidada", "Hábilis", "habilis", "F1",
         "Fechamento da triagem do lote de 14/04 e identificação de lacunas."),
        ("10/05", "Mapeamento documental das CPs parciais/pendentes", "Hábilis", "habilis", "F1",
         "Plano de recomposição por CP (documental, física, contratação)."),
        ("12/05", "Validação técnica prévia em campo · 5 eixos", "Hábilis", "habilis", "F2",
         "Diagnóstico interno (NÃO entregável formal): implantação/APP/RL/CAR, armazenamento e tubulações, efluentes/SAO + poços, resíduos, segurança e rotina. Subsidia produção das peças."),
        ("12/05", "Validação das plantas em campo", "Hábilis", "habilis", "F2",
         "Plantas 05/2024 já em .dwg vs configuração atual. Quadro de Áreas + UTM SIRGAS 2000 (CP 3.12)."),
        ("12/05", "Levantamento hídrico: define regime Veredas", "Hábilis", "habilis", "F2",
         "Rede pública: dispensa simplificada. Captação direta: processo completo com memorial técnico."),
        ("13/05", "Protocolo novo de outorga via Veredas", "Hábilis", "habilis", "F2",
         "Dia útil seguinte à validação 12/05. Gera certidão de tramitação que compõe o dossiê SEMMA. IN 15/2026 §15."),
        ("15/05", "Contratação do laboratório acreditado", "Cliente", "cliente", "F2",
         "Laboratório com acreditação INMETRO vigente para o RCA semestral (CP 4.4). Indicação Hábilis."),
        ("18/05", "NFs e certificados dos equipamentos", "Cliente", "cliente", "F2",
         "NFs de tanques/tubulações + certificados CONAMA 319/2002 e INMETRO 37/109/110/111/009. Subsidia MCO 28/05."),
        ("20/05", "Contratos + NFs semestrais (Bauer, Limpmil, Ecofenix)", "Cliente", "cliente", "F2",
         "CP 4.5/4.6/4.7. Validar ANP + licenciamento + autorização de transporte. Certificados de coleta de óleo dos últimos 3 anos."),
        ("22/05", "Nova coleta de efluentes · RCA", "Terceiro", "terceiro", "F2",
         "Só após adequação SAO (diagnóstico 12/05). Laboratório indicado pela Hábilis. ART do químico."),
        ("28/05", "Atualização do MCO · Anexo IX CEMAM 029/2018", "Terceiro", "terceiro", "F2",
         "Peça central do dossiê. Eng. Civil/Ambiental subcontratado sob coordenação CRBio. ART do profissional contratado."),
        ("29/05", "Anexos I e II da Resolução CONAMA 273/2000", "Terceiro", "terceiro", "F2",
         "Anexo I tabular + Anexo II classifica risco. Item 23 TR. ART própria. Mesmo terceiro do MCO."),
        ("30/05", "Plano de atendimento das 28 condicionantes", "Hábilis", "habilis", "F2",
         "Documento interno: responsável, evidência e prazo por condicionante."),
        ("03/06", "RCA · Relatório de Controle Ambiental semestral", "Hábilis", "habilis", "F3",
         "Anexo XI CEMAM 029/2018. CP 4.4. Análises SAO: pH, turbidez, O&G, sólidos, DBO, DQO, OD, MBAS, metais, TPH, BTEX/PAH em poços. ART RT Ambiental (CRBio)."),
        ("05/06", "PAE · Plano de Atendimento a Emergências", "Hábilis", "habilis", "F3",
         "CP 4.9. Cenários de vazamento, incêndio, derramamento. ART RT Ambiental (CRBio)."),
        ("12/06", "Recomposição documental das CPs parciais", "Hábilis", "habilis", "F3",
         "Conforme Plano de Atendimento de 30/05."),
        ("19/06", "Relatório técnico consolidado", "Hábilis", "habilis", "F3",
         "Peça integradora: MCO, RCA, PAE, plantas, plano de CPs, evidências. ART RT coordenador CRBio. Laudos específicos mantêm ART dos RTs próprios."),
        ("22/06", "Aprovação formal do relatório", "Cliente", "cliente", "F4",
         "Aprovação escrita do Grupo Z+Z antes do protocolo."),
        ("23/06", "CND Municipal + CND Estadual", "Cliente", "cliente", "F4",
         "Certidões Negativas de Débitos emitidas próximas ao protocolo. Validade 90 dias."),
        ("24/06", "Checklist final cruzado", "Hábilis", "habilis", "F4",
         "Conferência de peças, assinaturas, ARTs e taxas antes da submissão."),
        ("25/06", "Pagamento das taxas de protocolo (DUAM)", "Cliente", "cliente", "F4",
         "Taxas SEMMA, municipais e demais custos da submissão."),
        ("26/06", "Protocolo no Processo SEMMA 22760/2020", "Hábilis", "habilis", "F4",
         "Meta acordada (180d antes do vencimento). Mínimo legal: 24/08/2026 (CP 3.11). Acompanhamento pós-protocolo se inicia."),
    ]

    for prazo, item, resp, resp_kind, fase, obs in itens_chrono:
        row = t.add_row()
        # Prazo
        c = row.cells[0]
        pc = c.paragraphs[0]
        pc.paragraph_format.space_before = Pt(1); pc.paragraph_format.space_after = Pt(1)
        styled_run(pc, prazo, size=8, color=CINZA_9, bold=True, mono=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        # Item
        c = row.cells[1]
        pc = c.paragraphs[0]
        pc.paragraph_format.space_before = Pt(1); pc.paragraph_format.space_after = Pt(1)
        styled_run(pc, item, size=8.5, color=CINZA_9, bold=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        # Responsavel (celula colorida)
        c = row.cells[2]
        set_cell_shading(c, RESP_COLORS[resp_kind])
        pc = c.paragraphs[0]
        pc.paragraph_format.space_before = Pt(1); pc.paragraph_format.space_after = Pt(1)
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(pc, resp.upper(), size=7.5, color=BRANCO, bold=True, mono=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Fase
        c = row.cells[3]
        pc = c.paragraphs[0]
        pc.paragraph_format.space_before = Pt(1); pc.paragraph_format.space_after = Pt(1)
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(pc, fase, size=8, color=CINZA_9, bold=True, mono=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Observacao
        c = row.cells[4]
        pc = c.paragraphs[0]
        pc.paragraph_format.space_before = Pt(1); pc.paragraph_format.space_after = Pt(1)
        styled_run(pc, obs, size=8, color=CINZA_7)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # ====================================================================
    # ANEXO A1 · COMPLIANCE PARALELO (fluxo continuo)
    # ====================================================================

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(0)
    styled_run(p, "ANEXO | CADÊNCIA PRÓPRIA", size=9, color=VERDE, bold=True, mono=True)

    heading_section(doc, "A1", "Compliance Paralelo | frente opcional", space_before=4)

    lede(doc,
         "Frente complementar à renovação da LAO, com cadência própria, fora do caminho "
         "crítico do protocolo de 26/06. Itens aqui não integram o dossiê SEMMA, mas falhas "
         "podem contaminar o processo principal (ex.: cadastro ANP desatualizado fere a CP "
         "3.10; alvará sanitário vencido fere a CP 3.8; RAPP IBAMA atrasado dificulta CNDs; "
         "NR-20 irregular bloqueia operação). Tratada como frente opcional permanente, com "
         "aditivo contratual à parte. Custo anual estimado R$ 20k–60k.",
         ["cadência própria", "fora do caminho crítico do protocolo",
          "frente opcional permanente", "R$ 20k–60k"])

    heading_h3(doc, "Sumário por órgão", space_before=4)
    t = table_hdr(doc, ["Órgão | Norma", "Vínculo", "Natureza"], [5.5, 2.3, 9.4])
    add_row(t, ["ANP (inclui cadastro de revendedor atualizado)", "CP 3.10 LAO", "Cadastro, revenda, livro de movimentação, relatórios comerciais, comprovação de distribuidora ativa."], bold_cols=[0, 1], mono_cols=[1])
    add_row(t, ["IBAMA | CTF/APP + RAPP + TCFA", "Federal", "Cadastro Técnico Federal ativo e Relatório Anual de Atividades Potencialmente Poluidoras entregue."], zebra=True, bold_cols=[0, 1], mono_cols=[1])
    add_row(t, ["NR-20 | Segurança do Trabalho", "MTb", "PGR, PCMSO, treinamento NR-20 válido por função, inspeção SESMT e plano de emergência."], bold_cols=[0, 1], mono_cols=[1])
    add_row(t, ["PNRS | PGRS · Resíduos Sólidos (Lei 12.305/10)", "Federal", "Plano de Gerenciamento de Resíduos, MTR, contrato com transportador licenciado e destino final."], zebra=True, bold_cols=[0, 1], mono_cols=[1])
    add_row(t, ["Alvará Sanitário municipal (se aplicável)", "CP 3.8 LAO", "Validar exigência em zona rural junto à VISA Guapó. Se exigível, emissão/renovação."], bold_cols=[0, 1], mono_cols=[1])
    add_row(t, ["Monitoramentos pós-LAO", "Permanente", "Estanqueidade bienal (próximo 10/2026), RCA semestral contínuo, certificados de coleta."], zebra=True, bold_cols=[0, 1], mono_cols=[1])
    add_row(t, ["Total", "5 frentes + monitoramentos", "Rodando em paralelo. Custo anual R$ 20k–60k. Diagnóstico individual em documento próprio."], bold_cols=[0, 1, 2], mono_cols=[1], is_total=True)

    heading_h3(doc, "O que depende do Grupo Z+Z | compliance", space_before=8)
    bullet(doc, "Acesso completo aos cadastros ANP (login/senha ou procuração com poderes específicos) — atualização do cadastro de revendedor atende a CP 3.10 da LAO.", ["cadastros ANP", "CP 3.10"])
    bullet(doc, "Validação junto à VISA Guapó sobre exigência do alvará sanitário em zona rural (atende a CP 3.8 LAO).", ["alvará sanitário", "CP 3.8"])
    bullet(doc, "Acesso ao CTF/APP IBAMA (senha ou nomeação de preposto com poderes de consulta e entrega).", ["CTF/APP IBAMA"])
    bullet(doc, "Documentos do SESMT ou contrato com serviço terceirizado de segurança do trabalho: PGR, PCMSO, ASO, ficha de EPI.", ["PGR", "PCMSO"])
    bullet(doc, "Contratos e MTRs vigentes com transportadores e destinadores de resíduos licenciados (PGRS conforme PNRS).", ["MTRs", "PGRS"])
    bullet(doc, "Liberação orçamentária para treinamentos NR-20 faltantes e regularizações pontuais.", ["NR-20"])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_shading(p, AZUL_BG)
    set_paragraph_border_left(p, 2, "0D47A1")
    p.paragraph_format.left_indent = Cm(0.25)
    styled_run(p, "Decisão requerida neste anexo: ", size=9, color=CINZA_9, bold=True)
    styled_run(p, "se o Grupo Z+Z deseja contratar o Compliance Paralelo como frente continuada (aditivo separado da renovação) ou tratá-lo internamente. Recomendação Hábilis: validar prioritariamente as duas frentes vinculadas a CPs da LAO (ANP CP 3.10 + alvará CP 3.8) antes do protocolo de 26/06.",
               size=9, color=CINZA_7)

    here = os.path.dirname(os.path.abspath(__file__))
    out = os.path.join(here, "Alinhamento_Executivo_AutoPostoAmerica_v2.docx")
    doc.save(out)
    print(f"OK -> {out}")
    return out


if __name__ == "__main__":
    build()
